"""文档生成模块"""

import tempfile
import cv2
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from typing import List, Tuple
import os
import re

from .models import LicenseData, ExtractedImage, TextBlock
from .exceptions import DocumentGenerationError


class DocumentGenerator:
    """文档生成器，基于模板填充字段值"""
    
    # 字段标签映射（英文标签 -> 中文字段名）
    FIELD_MAPPING = {
        "License No": "证号",
        "License No.": "证号",
        "Name": "姓名",
        "Sex": "性别",
        "Nationality": "国籍",
        "Address": "住址",
        "Date of Birth": "出生日期",
        "Date of First Issue": "初次领证日期",
        "Class": "准驾车型",
        "Valid Period": "有效期限",
        "Valid From": "有效期起始日期",
        "Valid For": "有效期限",  # 旧版驾驶证使用 Valid For
        "Issuing Authority": "发证机关",
        "File No": "档案编号",
        "File No.": "档案编号",
        "Record": "记录"
    }
    
    # 新版模板路径（只有 Valid Period）
    NEW_VERSION_TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "jsz_translated", "模版 - 新版.docx")
    # 旧版模板路径（有 Valid From 和 Valid For）
    OLD_VERSION_TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "jsz_translated", "模版 - 旧版.docx")
    
    def __init__(self, template_path: str = None):
        """
        初始化文档生成器
        
        Args:
            template_path: 模板文档路径，如果为None则根据驾驶证版本自动选择
        """
        # 如果指定了模板路径，使用指定的路径
        if template_path is not None:
            self.template_path = template_path
            self.logger = logging.getLogger(__name__)
            
            # 检查模板是否存在
            if not os.path.exists(self.template_path):
                raise DocumentGenerationError(f"模板文件不存在: {self.template_path}")
        else:
            # 不指定模板路径时，将在 generate_document 中根据版本选择
            self.template_path = None
            self.logger = logging.getLogger(__name__)
    
    def _save_extracted_images(self, images: List[ExtractedImage]) -> None:
        """
        保存提取的图像到临时文件
        
        Args:
            images: 提取的图像列表
        """
        for img in images:
            # 创建临时文件
            temp_file = tempfile.NamedTemporaryFile(
                delete=False,
                suffix='.png'
            )
            temp_path = temp_file.name
            temp_file.close()
            
            # 保存图像
            cv2.imwrite(temp_path, img.image_data)
            img.temp_path = temp_path
    
    def _remove_cell_borders(self, cell) -> None:
        """
        移除单元格边框
        
        Args:
            cell: docx 表格单元格对象
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tcPr.append(tcBorders)
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = tcBorders.find(qn(f'w:{border_name}'))
            if border is None:
                border = parse_xml(r'<w:%s xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="none"/>' % border_name)
                tcBorders.append(border)
            else:
                border.set(qn('w:val'), 'none')

    
    def _get_english_label(self, field_name: str) -> str:
        """
        获取字段的英文标签
        
        Args:
            field_name: 中文字段名
            
        Returns:
            英文标签
        """
        labels = {
            "姓名": "Name",
            "性别": "Sex",
            "国籍": "Nationality",
            "出生日期": "Date of Birth",
            "住址": "Address",
            "证号": "License No.",
            "准驾车型": "Class",
            "有效期起始日期": "Valid From",
            "有效起始日期": "Valid From",  # OCR 变体
            "有效期限": "Valid Period",
            "发证机关": "Issuing Authority",
            "初次领证日期": "Date of First Issue",
            "档案编号": "File No.",
            "记录": "Record"
        }
        return labels.get(field_name, field_name)
    
    def _get_english_label_for_old_version(self, field_name: str) -> str:
        """
        获取旧版驾驶证字段的英文标签
        
        旧版驾驶证使用 Valid For 而不是 Valid Period
        
        Args:
            field_name: 中文字段名
            
        Returns:
            英文标签
        """
        labels = {
            "姓名": "Name",
            "性别": "Sex",
            "国籍": "Nationality",
            "出生日期": "Date of Birth",
            "住址": "Address",
            "证号": "License No.",
            "准驾车型": "Class",
            "有效期起始日期": "Valid From",
            "有效起始日期": "Valid From",  # OCR 变体
            "有效期限": "Valid For",  # 旧版使用 Valid For
            "发证机关": "Issuing Authority",
            "初次领证日期": "Date of First Issue",
            "档案编号": "File No.",
            "记录": "Record"
        }
        return labels.get(field_name, field_name)
    
    def _add_text_to_layout(
        self,
        layout_table,
        text: str,
        position: Tuple[int, int],
        img_width: int,
        img_height: int,
        scale: float
    ) -> None:
        """
        在布局表格中添加文字
        
        Args:
            layout_table: 布局表格
            text: 文字内容
            position: 位置坐标 (x, y)
            img_width: 原图宽度
            img_height: 原图高度
            scale: 缩放比例
        """
        x, y = position
        
        # 计算在表格中的位置（行列索引）
        col_index = int((x / img_width) * 50)
        row_index = int((y / img_height) * 50)
        
        # 确保索引在范围内
        col_index = min(col_index, 49)
        row_index = min(row_index, 49)
        
        # 获取单元格
        cell = layout_table.rows[row_index].cells[col_index]
        
        # 添加文字
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _add_image_to_layout(
        self,
        layout_table,
        img: ExtractedImage,
        img_width: int,
        img_height: int,
        scale: float
    ) -> None:
        """
        在布局表格中添加图像
        
        Args:
            layout_table: 布局表格
            img: 提取的图像
            img_width: 原图宽度
            img_height: 原图高度
            scale: 缩放比例
        """
        x, y = img.position
        w, h = img.size
        
        # 计算在表格中的位置
        col_index = int((x / img_width) * 50)
        row_index = int((y / img_height) * 50)
        
        # 确保索引在范围内
        col_index = min(col_index, 49)
        row_index = min(row_index, 49)
        
        # 获取单元格
        cell = layout_table.rows[row_index].cells[col_index]
        
        # 添加图像
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(img.temp_path, width=Inches(w / 96 * scale))

    def _remove_missing_pages(
        self,
        doc: Document,
        has_main: bool,
        has_duplicate: bool,
        has_legend: bool
    ) -> None:
        """
        根据检测到的页面类型删除模板中不存在的页面
        
        新版模板默认包含三个页面：主页、副页、准驾车型代号规定页
        根据实际检测结果删除不存在的页面
        
        Args:
            doc: Document 对象
            has_main: 是否有主页
            has_duplicate: 是否有副页
            has_legend: 是否有准驾车型代号规定页
        """
        print(f"\n[页面处理] 检测结果: 主页={has_main}, 副页={has_duplicate}, 准驾车型代号规定页={has_legend}")
        
        # 收集需要删除的页面标记
        pages_to_remove = []
        
        if not has_duplicate:
            pages_to_remove.append('duplicate')
            print("[页面处理] 将删除副页")
        
        if not has_legend:
            pages_to_remove.append('legend')
            print("[页面处理] 将删除准驾车型代号规定页")
        
        if not pages_to_remove:
            print("[页面处理] 所有页面都存在，无需删除")
            return
        
        # 遍历文档中的所有表格，找到并删除对应页面的内容
        # 新版模板结构：每个页面通常是一个独立的表格或一组表格
        
        # 标记要删除的元素
        elements_to_remove = []
        
        # 页面标识文本
        duplicate_markers = ["Duplicate of Driving License", "副页"]
        legend_markers = ["Legend for Class of Vehicles", "准驾车型代号规定"]
        
        # 遍历文档的所有元素（段落和表格）
        current_page_type = 'main'  # 当前正在处理的页面类型
        
        for element in doc.element.body:
            # 检查是否是段落
            if element.tag.endswith('p'):
                text = element.text if element.text else ''
                # 获取段落的完整文本（包括子元素）
                for child in element.iter():
                    if child.text:
                        text += child.text
                    if child.tail:
                        text += child.tail
                
                # 检测页面类型变化
                if any(marker in text for marker in duplicate_markers):
                    current_page_type = 'duplicate'
                elif any(marker in text for marker in legend_markers):
                    current_page_type = 'legend'
                
                # 如果当前页面需要删除，标记该元素
                if current_page_type == 'duplicate' and 'duplicate' in pages_to_remove:
                    elements_to_remove.append(element)
                elif current_page_type == 'legend' and 'legend' in pages_to_remove:
                    elements_to_remove.append(element)
            
            # 检查是否是表格
            elif element.tag.endswith('tbl'):
                # 获取表格中的所有文本
                table_text = ''
                for cell in element.iter():
                    if cell.text:
                        table_text += cell.text
                    if cell.tail:
                        table_text += cell.tail
                
                # 检测页面类型变化
                if any(marker in table_text for marker in duplicate_markers):
                    current_page_type = 'duplicate'
                elif any(marker in table_text for marker in legend_markers):
                    current_page_type = 'legend'
                
                # 如果当前页面需要删除，标记该元素
                if current_page_type == 'duplicate' and 'duplicate' in pages_to_remove:
                    elements_to_remove.append(element)
                elif current_page_type == 'legend' and 'legend' in pages_to_remove:
                    elements_to_remove.append(element)
        
        # 删除标记的元素
        for element in elements_to_remove:
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        
        print(f"[页面处理] 已删除 {len(elements_to_remove)} 个元素")

    def generate_document(
        self,
        license_data: LicenseData,
        output_path: str
    ) -> None:
        """
        基于模板生成 DOCX 翻译文档
        
        Args:
            license_data: 驾驶证数据
            output_path: 输出文件路径
            
        Raises:
            DocumentGenerationError: 文档生成失败
        """
        try:
            # 根据驾驶证版本选择模板
            if self.template_path is None:
                if license_data.is_old_version:
                    template_path = self.OLD_VERSION_TEMPLATE
                    print(f"[模板选择] 检测到旧版驾驶证（有 Valid From 字段），使用旧版模板")
                else:
                    template_path = self.NEW_VERSION_TEMPLATE
                    print(f"[模板选择] 检测到新版驾驶证（只有 Valid Period 字段），使用新版模板")
                
                # 检查模板是否存在
                if not os.path.exists(template_path):
                    raise DocumentGenerationError(f"模板文件不存在: {template_path}")
            else:
                template_path = self.template_path
            
            # 加载模板文档
            doc = Document(template_path)
            
            # 根据检测到的页面类型删除不存在的页面（仅新版模板）
            if not license_data.is_old_version:
                self._remove_missing_pages(
                    doc, 
                    has_main=license_data.has_main_page,
                    has_duplicate=license_data.has_duplicate,
                    has_legend=license_data.has_legend_page
                )
            
            # 创建字段值字典（英文标签 -> 翻译值）
            field_values = self._build_field_values_dict(license_data.fields, license_data.is_old_version)
            
            # 在文档中查找并替换字段标签
            self._fill_template_fields(doc, field_values, license_data.has_duplicate)
            
            # 保存提取的图像到临时文件
            self._save_extracted_images(license_data.images)
            
            # 替换模板中的图片和印章
            self._fill_template_images(doc, license_data.images)
            
            # 添加印章文字到左边方框
            if license_data.seal_texts:
                self._add_seal_text_to_left_box(doc, license_data.seal_texts)
            
            # 添加条形码数字到占位符
            if license_data.barcode_number:
                self._fill_barcode_placeholder(doc, license_data.barcode_number, license_data.is_old_version)
            
            self.logger.info(f"条形码数字: {license_data.barcode_number if license_data.barcode_number else '未识别'}")
            
            # 保存文档
            doc.save(output_path)
            
        except Exception as e:
            raise DocumentGenerationError(f"文档生成失败: {str(e)}")
    
    def _build_field_values_dict(self, fields: List, is_old_version: bool = False) -> dict:
        """
        构建字段值字典（英文标签 -> 翻译值）
        
        Args:
            fields: 字段列表
            is_old_version: 是否是旧版驾驶证
            
        Returns:
            字段值字典
        """
        field_values = {}
        
        for field in fields:
            # 根据版本获取英文标签
            if is_old_version:
                english_label = self._get_english_label_for_old_version(field.field_name)
            else:
                english_label = self._get_english_label(field.field_name)
            
            if english_label:
                # 使用翻译值，如果没有则使用原值
                value = field.translated_value if field.translated_value else field.field_value
                field_values[english_label] = value
        
        print(f"识别到 {len(field_values)} 个字段")
        for key, value in field_values.items():
            print(f"  {key}: {value}")
        
        return field_values
    
    def _fill_template_fields(self, doc: Document, field_values: dict, has_duplicate: bool = False) -> None:
        """
        在模板中填充字段值
        
        Args:
            doc: Document 对象
            field_values: 字段值字典
            has_duplicate: 是否有副页
        """
        print("\n" + "="*60)
        if has_duplicate:
            print("填充字段（正页+副页）")
        else:
            print("填充字段（仅正页）")
        print("="*60)
        
        # 收集正页和副页的段落对象
        first_page_paragraphs = []
        duplicate_page_paragraphs = []
        found_duplicate = False
        
        # 1. 文档级别的段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                if "Duplicate of Driving License" in text:
                    found_duplicate = True
                
                if not found_duplicate:
                    first_page_paragraphs.append(('doc_para', paragraph))
                elif has_duplicate:
                    duplicate_page_paragraphs.append(('doc_para', paragraph))
        
        # 2. 表格中的段落
        if doc.tables:
            for table_idx, table in enumerate(doc.tables):
                found_duplicate_in_table = False
                
                def traverse_table(table, level=0, is_duplicate_page=False):
                    nonlocal found_duplicate_in_table
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                text = paragraph.text.strip()
                                if "Duplicate of Driving License" in text:
                                    found_duplicate_in_table = True
                                    is_duplicate_page = True
                                
                                if text:
                                    if not is_duplicate_page:
                                        first_page_paragraphs.append((f'table_L{level}', paragraph))
                                    elif has_duplicate:
                                        duplicate_page_paragraphs.append((f'table_L{level}', paragraph))
                            
                            # 嵌套表格
                            for nested_table in cell.tables:
                                traverse_table(nested_table, level + 1, is_duplicate_page)
                
                traverse_table(table)
        
        # 填充正页字段
        for source, para in first_page_paragraphs:
            self._fill_paragraph_fields(para, field_values)
        
        # 填充副页字段（如果有）
        if has_duplicate and duplicate_page_paragraphs:
            print("填充副页字段...")
            
            # 副页需要的字段：License No., Name, File No., Record
            # 同时添加带点号和不带点号的版本，确保能匹配模板
            duplicate_field_values = {}
            
            # License No
            license_no = field_values.get("License No.", field_values.get("License No", ""))
            duplicate_field_values["License No."] = license_no
            duplicate_field_values["License No"] = license_no
            
            # Name
            name = field_values.get("Name", "")
            duplicate_field_values["Name"] = name
            
            # File No - 关键修复：同时添加带点号和不带点号的版本
            file_no = field_values.get("File No.", field_values.get("File No", ""))
            duplicate_field_values["File No."] = file_no
            duplicate_field_values["File No"] = file_no
            
            # Record
            record = field_values.get("Record", "")
            duplicate_field_values["Record"] = record
            
            # Valid From - 副页也需要这个字段
            valid_from = field_values.get("Valid From", "")
            duplicate_field_values["Valid From"] = valid_from
            
            print(f"副页字段值:")
            for key, value in duplicate_field_values.items():
                print(f"  {key}: {value}")
            print("-"*60)
            
            for source, para in duplicate_page_paragraphs:
                self._fill_paragraph_fields(para, duplicate_field_values)
    
    def _fill_paragraph_fields(self, paragraph, field_values: dict) -> None:
        """
        在段落中填充字段值（处理同一行多个字段的情况，保留未识别的字段标签）
        
        Args:
            paragraph: 段落对象
            field_values: 字段值字典
        """
        text = paragraph.text.strip()
        
        if not text:
            return
        
        original_text = text
        
        import re
        
        # 找出文本中所有的字段标签及其位置
        field_positions = []
        
        # 按标签长度降序排序，确保长标签优先匹配（如 "Valid From" 优先于 "Valid"）
        sorted_labels = sorted(field_values.keys(), key=len, reverse=True)
        
        for label in sorted_labels:
            # 生成标签的变体（带点号和不带点号）
            label_variants = [label]
            if label.endswith('.'):
                # 如果标签以点号结尾，也尝试不带点号的版本
                label_variants.append(label[:-1])
            
            for label_variant in label_variants:
                # 精确匹配：标签后面必须跟着冒号，使用单词边界
                # 使用 \b 确保是完整单词匹配
                pattern = r'(?<![a-zA-Z])' + re.escape(label_variant) + r'\s*[:：]\s*'
                for match in re.finditer(pattern, text, re.IGNORECASE):
                    start_pos = match.start()
                    end_pos = match.end()
                    
                    # 检查这个位置是否已经被更长的标签占用
                    already_covered = False
                    for fp in field_positions:
                        # 如果当前匹配的范围与已有匹配重叠，跳过
                        if not (end_pos <= fp['start'] or start_pos >= fp['end']):
                            already_covered = True
                            break
                    
                    if not already_covered:
                        field_positions.append({
                            'label': label_variant,  # 使用实际匹配到的变体
                            'start': match.start(),
                            'end': match.end(),
                            'match_text': match.group(),
                            'value': field_values[label]  # 使用原始标签的值
                        })
        
        if not field_positions:
            return
        
        # 按位置排序
        field_positions.sort(key=lambda x: x['start'])
        
        # 构建新文本
        new_text = ""
        last_pos = 0
        
        for i, field_info in enumerate(field_positions):
            # 添加从上次位置到当前标签开始的文本
            new_text += text[last_pos:field_info['start']]
            
            # 添加字段标签和冒号
            new_text += field_info['label'] + ": "
            
            # 添加字段值
            new_text += field_info['value']
            
            # 处理标签后面的内容
            match_end = field_info['end']
            
            # 如果有下一个字段
            if i + 1 < len(field_positions):
                next_start = field_positions[i + 1]['start']
                # 检查两个字段之间是否有其他未识别的字段标签
                between_text = text[match_end:next_start]
                # 查找其他字段标签（单词 + 冒号）
                if re.search(r'[A-Z][a-zA-Z\s\.]*[:：]', between_text):
                    # 保留其他字段标签
                    new_text += between_text
                else:
                    # 没有其他字段，只保留空格
                    space_count = max(5, len(between_text.strip()) + 3)
                    new_text += " " * space_count
                last_pos = next_start
            else:
                # 最后一个字段
                # 检查后面是否还有其他字段标签（未被识别的）
                remaining = text[match_end:]
                # 查找是否有其他字段标签模式（单词 + 冒号）
                other_field_match = re.search(r'\s*([A-Z][a-zA-Z\s\.]+)[:：]', remaining)
                if other_field_match:
                    # 保留从匹配开始的字段标签部分
                    # 如果匹配开始位置>0，说明前面有空格，保留它
                    if other_field_match.start() > 0:
                        new_text += remaining[other_field_match.start():]
                    else:
                        # 没有空格，添加空格
                        new_text += "   " + remaining[other_field_match.start():]
                else:
                    # 没有其他字段，不添加任何内容
                    pass
                last_pos = len(text)
        
        # 如果文本发生了变化，更新段落
        if new_text.strip() != original_text.strip():
            print(f"[OK] 填充段落")
            print(f"  原文: {original_text}")
            print(f"  新文: {new_text}")
            print("-" * 60)
            
            # 清空段落并重新添加文本
            paragraph.clear()
            run = paragraph.add_run(new_text)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _fill_template_images(self, doc: Document, images: List[ExtractedImage]) -> None:
        """
        在模板中填充图片
        
        查找 [PHOTO] 占位符并替换为相片
        
        Args:
            doc: Document 对象
            images: 图片列表
        """
        if not images:
            return
        
        # 查找照片
        photo = None
        for img in images:
            if img.image_type == "photo":
                photo = img
                break
        
        # 替换 [PHOTO] 占位符
        if photo:
            self._replace_placeholder_in_doc(doc, "[PHOTO]", photo, "photo")
    
    def _replace_placeholder_in_doc(
        self,
        doc: Document,
        placeholder: str,
        image: ExtractedImage,
        image_type: str
    ) -> None:
        """
        在文档中查找占位符并替换为图片
        
        Args:
            doc: Document 对象
            placeholder: 占位符文本（如 "[SEAL]" 或 "[PHOTO]"）
            image: 图片对象
            image_type: 图片类型（"seal" 或 "photo"）
        """
        if not image or not image.temp_path:
            return
        
        # 遍历所有段落
        for paragraph in doc.paragraphs:
            if placeholder.lower() in paragraph.text.lower():
                self._replace_placeholder_in_paragraph(paragraph, placeholder, image, image_type)
                return
        
        # 遍历所有表格
        for table in doc.tables:
            if self._replace_placeholder_in_table(table, placeholder, image, image_type):
                return
    
    def _replace_placeholder_in_table(
        self,
        table,
        placeholder: str,
        image: ExtractedImage,
        image_type: str
    ) -> bool:
        """
        在表格中查找并替换占位符
        
        Returns:
            是否找到并替换了占位符
        """
        for row in table.rows:
            for cell in row.cells:
                # 先检查嵌套表格（优先处理嵌套表格中的占位符）
                for nested_table in cell.tables:
                    if self._replace_placeholder_in_table(nested_table, placeholder, image, image_type):
                        return True
                
                # 再检查单元格的段落
                for paragraph in cell.paragraphs:
                    if placeholder.lower() in paragraph.text.lower():
                        # 调整单元格设置以容纳图片
                        self._adjust_cell_for_image(cell, row)
                        self._replace_placeholder_in_paragraph(paragraph, placeholder, image, image_type)
                        return True
        
        return False
    
    def _adjust_cell_for_image(self, cell, row) -> None:
        """
        调整单元格设置以容纳图片
        
        Args:
            cell: 单元格对象
            row: 行对象
        """
        try:
            # 设置单元格垂直对齐为顶部
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            vAlign = tcPr.find(qn('w:vAlign'))
            if vAlign is None:
                vAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="top"/>')
                tcPr.append(vAlign)
            else:
                vAlign.set(qn('w:val'), 'top')
            
            # 设置行高为至少 3.5 厘米（约 1984 twips）
            tr = row._element
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = parse_xml(r'<w:trPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                tr.insert(0, trPr)
            
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is None:
                trHeight = parse_xml(r'<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="1984" w:hRule="atLeast"/>')
                trPr.append(trHeight)
            else:
                trHeight.set(qn('w:val'), '1984')
                trHeight.set(qn('w:hRule'), 'atLeast')
        except Exception as e:
            print(f"[警告] 调整单元格设置时出错: {str(e)}")
    
    def _replace_placeholder_in_paragraph(
        self,
        paragraph,
        placeholder: str,
        image: ExtractedImage,
        image_type: str
    ) -> None:
        """
        在段落中替换占位符为图片
        
        Args:
            paragraph: 段落对象
            placeholder: 占位符文本
            image: 图片对象
            image_type: 图片类型
        """
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn as qn_func
        
        # 清空段落
        paragraph.clear()
        
        # 添加图片（只处理相片）
        run = paragraph.add_run()
        
        # 相片大小 - 调整为适合方框的尺寸（约1.1英寸宽）
        img_width_inch = 1.1
        run.add_picture(image.temp_path, width=Inches(img_width_inch))
        
        # 设置段落居中对齐
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 尝试调整包含图片的单元格和行
        try:
            # 获取段落所在的单元格元素
            para_element = paragraph._element
            tc = para_element.getparent()
            while tc is not None and tc.tag != qn('w:tc'):
                tc = tc.getparent()
            
            if tc is not None:
                # 设置单元格垂直居中对齐
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = parse_xml(r'<w:tcPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    tc.insert(0, tcPr)
                
                vAlign = tcPr.find(qn('w:vAlign'))
                if vAlign is None:
                    vAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
                    tcPr.append(vAlign)
                else:
                    vAlign.set(qn('w:val'), 'center')
                
                # 获取单元格所在的行
                row = tc.getparent()
                
                # 设置行高
                trPr = row.find(qn('w:trPr'))
                if trPr is None:
                    trPr = parse_xml(r'<w:trPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    row.insert(0, trPr)
                
                # 设置行高为至少 4 厘米（约 2268 twips）
                trHeight = trPr.find(qn('w:trHeight'))
                if trHeight is None:
                    trHeight = parse_xml(r'<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="2268" w:hRule="atLeast"/>')
                    trPr.append(trHeight)
                else:
                    trHeight.set(qn('w:val'), '2268')
                    trHeight.set(qn('w:hRule'), 'atLeast')
        except Exception as e:
            print(f"[警告] 调整图片单元格时出错: {str(e)}")
        
        print(f"[OK] {placeholder} 已替换为相片（宽度: {img_width_inch} 英寸，垂直居中）")
    
    def _add_seal_text_to_left_box(self, doc: Document, seal_texts: List[str]) -> None:
        """
        在 [SEAL] 占位符位置添加印章文字翻译（红色）
        
        Args:
            doc: Document 对象
            seal_texts: 印章文字列表（已翻译）
        """
        if not seal_texts:
            return
        
        # 查找并替换 [SEAL] 占位符
        replaced = self._replace_seal_placeholder_with_text(doc, "[SEAL]", seal_texts)
        
        if not replaced:
            print("[警告] 未找到 [SEAL] 占位符，无法添加印章文字")
        else:
            print(f"[OK] 印章文字已添加（红色）：{', '.join(seal_texts)}")
    
    def _replace_seal_placeholder_with_text(
        self,
        doc: Document,
        placeholder: str,
        seal_texts: List[str]
    ) -> bool:
        """
        查找 [SEAL] 占位符并替换为红色文字
        
        Returns:
            是否成功替换
        """
        # 遍历所有段落
        for paragraph in doc.paragraphs:
            if placeholder.lower() in paragraph.text.lower():
                self._replace_seal_in_paragraph(paragraph, seal_texts)
                return True
        
        # 遍历所有表格
        for table in doc.tables:
            if self._replace_seal_in_table(table, placeholder, seal_texts):
                return True
        
        return False
    
    def _replace_seal_in_table(
        self,
        table,
        placeholder: str,
        seal_texts: List[str]
    ) -> bool:
        """
        在表格中查找并替换 [SEAL] 占位符
        """
        for row in table.rows:
            for cell in row.cells:
                # 检查单元格的段落
                for paragraph in cell.paragraphs:
                    if placeholder.lower() in paragraph.text.lower():
                        self._replace_seal_in_paragraph(paragraph, seal_texts)
                        return True
                
                # 检查嵌套表格
                for nested_table in cell.tables:
                    if self._replace_seal_in_table(nested_table, placeholder, seal_texts):
                        return True
        
        return False
    
    def _replace_seal_in_paragraph(
        self,
        paragraph,
        seal_texts: List[str]
    ) -> None:
        """
        在段落中替换 [SEAL] 为红色印章文字
        """
        from docx.shared import RGBColor
        
        # 清空段落
        paragraph.clear()
        
        # 设置段落居中对齐
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加印章文字，红色，居中
        for i, text in enumerate(seal_texts):
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 设置红色
            run.font.color.rgb = RGBColor(255, 0, 0)
            
            # 除了最后一个，都添加换行
            if i < len(seal_texts) - 1:
                paragraph.add_run("\n")

    def _fill_barcode_placeholder(self, doc: Document, barcode_number: str, is_old_version: bool = False) -> None:
        """
        在 *[BARCODE]* 占位符位置填充条形码数字
        
        Args:
            doc: Document 对象
            barcode_number: 条形码数字
            is_old_version: 是否是旧版驾驶证（旧版条形码右对齐）
        """
        if not barcode_number:
            return
        
        # 尝试多种占位符格式
        placeholders = ["*[BARCODE]*", "[BARCODE]", "*BARCODE*", "BARCODE"]
        replaced = False
        
        for placeholder in placeholders:
            if self._replace_barcode_placeholder(doc, placeholder, barcode_number, is_old_version):
                print(f"[OK] 条形码数字已填充（占位符: {placeholder}）：{barcode_number}")
                replaced = True
                break
        
        if not replaced:
            print(f"[警告] 未找到条形码占位符（尝试了: {', '.join(placeholders)}）")
    
    def _replace_barcode_placeholder(
        self,
        doc: Document,
        placeholder: str,
        barcode_number: str,
        is_old_version: bool = False
    ) -> bool:
        """
        查找 [BARCODE] 占位符并替换为条形码数字
        
        Returns:
            是否成功替换
        """
        # 遍历所有段落
        for paragraph in doc.paragraphs:
            if placeholder.lower() in paragraph.text.lower():
                self._replace_barcode_in_paragraph(paragraph, barcode_number, is_old_version)
                return True
        
        # 遍历所有表格
        for table in doc.tables:
            if self._replace_barcode_in_table(table, placeholder, barcode_number, is_old_version):
                return True
        
        return False
    
    def _replace_barcode_in_table(
        self,
        table,
        placeholder: str,
        barcode_number: str,
        is_old_version: bool = False
    ) -> bool:
        """
        在表格中查找并替换 [BARCODE] 占位符
        """
        for row in table.rows:
            for cell in row.cells:
                # 检查单元格的段落
                for paragraph in cell.paragraphs:
                    if placeholder.lower() in paragraph.text.lower():
                        self._replace_barcode_in_paragraph(paragraph, barcode_number, is_old_version)
                        return True
                
                # 检查嵌套表格
                for nested_table in cell.tables:
                    if self._replace_barcode_in_table(nested_table, placeholder, barcode_number, is_old_version):
                        return True
        
        return False
    
    def _replace_barcode_in_paragraph(
        self,
        paragraph,
        barcode_number: str,
        is_old_version: bool = False
    ) -> None:
        """
        在段落中替换 [BARCODE] 为条形码数字（保留星号格式）
        
        Args:
            paragraph: 段落对象
            barcode_number: 条形码数字
            is_old_version: 是否是旧版驾驶证（旧版条形码右对齐）
        """
        # 获取原始文本，检查是否有星号格式
        original_text = paragraph.text
        has_star_format = '*[BARCODE]*' in original_text or '*BARCODE*' in original_text
        
        # 清空段落
        paragraph.clear()
        
        # 设置段落对齐方式：旧版右对齐，新版居中
        if is_old_version:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            print(f"[条形码] 旧版驾驶证，设置右对齐")
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"[条形码] 新版驾驶证，设置居中对齐")
        
        # 添加条形码数字（如果原文有星号格式，保留星号）
        if has_star_format:
            formatted_barcode = f"*{barcode_number}*"
        else:
            formatted_barcode = barcode_number
        
        run = paragraph.add_run(formatted_barcode)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
