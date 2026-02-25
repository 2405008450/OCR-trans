"""诊断 DOCX 文件结构"""
import os, sys, glob
from docx import Document
from lxml import etree

# 找到最新上传的文件
upload_dir = "uploads/alignment"
files = glob.glob(os.path.join(upload_dir, "*_original.docx"))
if not files:
    print("没有找到上传的 docx 文件")
    sys.exit(1)

fpath = max(files, key=os.path.getmtime)
print(f"文件: {fpath} ({os.path.getsize(fpath)} bytes)")

doc = Document(fpath)

print(f"\n=== doc.paragraphs: {len(doc.paragraphs)} ===")
for i, p in enumerate(doc.paragraphs[:5]):
    print(f"  [{i}] '{p.text[:100]}'")

print(f"\n=== doc.tables: {len(doc.tables)} ===")

print(f"\n=== doc.element.body 子元素 ===")
if hasattr(doc, 'element') and hasattr(doc.element, 'body'):
    for i, child in enumerate(doc.element.body.iterchildren()):
        tag = child.tag
        short_tag = tag.split('}')[-1] if '}' in tag else tag
        # 获取子元素数量
        sub_count = len(list(child.iterchildren()))
        # 获取文本内容预览
        texts = []
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for t in child.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t.text:
                texts.append(t.text)
        text_preview = ''.join(texts)[:100]
        print(f"  [{i}] tag={short_tag} (full: {tag[:80]}), 子元素={sub_count}, 文本='{text_preview}'")

        # 如果是 sdt, 展开一层
        if short_tag == 'sdt':
            for j, sub in enumerate(child.iterchildren()):
                sub_tag = sub.tag.split('}')[-1] if '}' in sub.tag else sub.tag
                sub_sub_count = len(list(sub.iterchildren()))
                sub_texts = []
                for t in sub.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    if t.text:
                        sub_texts.append(t.text)
                sub_text = ''.join(sub_texts)[:100]
                print(f"    [{i}.{j}] tag={sub_tag}, 子元素={sub_sub_count}, 文本='{sub_text}'")

                # 再展开 sdtContent
                if sub_tag == 'sdtContent':
                    for k, subsub in enumerate(sub.iterchildren()):
                        ss_tag = subsub.tag.split('}')[-1] if '}' in subsub.tag else subsub.tag
                        ss_texts = []
                        for t in subsub.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                            if t.text:
                                ss_texts.append(t.text)
                        ss_text = ''.join(ss_texts)[:80]
                        print(f"      [{i}.{j}.{k}] tag={ss_tag}, 文本='{ss_text}'")

print(f"\n=== 整个 body XML 前 2000 字符 ===")
body_xml = etree.tostring(doc.element.body, encoding='unicode')
print(body_xml[:2000])
