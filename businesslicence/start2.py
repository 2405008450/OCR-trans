"""图片翻译系统启动文件

功能：
1. 使用 Google AI Studio (Gemini API) 识别图片中的字段标签和字段值
2. 翻译识别结果（中文->英文）
3. 根据字段标签匹配，填充到Word模板中

直接运行: python start2.py
右键运行: 右键点击此文件 -> 运行Python文件
"""

import os
import sys
import base64
import json
import re
from pathlib import Path
from datetime import datetime
import tkinter as tk
from tkinter import ttk, messagebox

# 获取脚本所在目录（项目根目录）
SCRIPT_DIR = Path(__file__).parent.absolute()

# 确保项目根目录在 Python 路径中（对于右键运行很重要）
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

# ==================== 配置区域 ====================

# Google AI Studio API 配置
GOOGLE_API_KEY = "AIzaSyBgpxxxxxxxxxC"  # 请替换为你的 Google AI Studio API Key
MODEL = "models/gemini-3-flash-preview"  # 可用:  models/gemini-2.5-pro, models/gemini-3.1-pro-preview


# 固定翻译映射（字段标签：中文 -> 英文模板标签）
FIXED_LABEL_TRANSLATIONS = {
    "名称": "business name",
    "名 称": "business name",  # 带空格的版本
    "类型": "ownership",
    "类 型": "ownership",  # 带空格的版本
    "商事主体类型": "ownership",  # 珠海等地区用这个
    "法定代表人": "legal representative",
    "负责人": "legal representative",  # 有些执照用负责人
    "经营范围": "operations",
    "注册资本": "registered capital",
    "注册资金": "registered capital",  # 另一种写法
    "成立日期": "date of incorporation",
    "住所": "domicile",
    "住 所": "domicile",  # 带空格的版本
    "住址": "domicile",
    "经营场所": "domicile",  # 个体户用这个
    "登记机关": "registration authority",
    "核准日期": "date",
    "重要提示": "important notice",  # 额外字段
    "营业期限": "business term",  # 额外字段
    "组成形式": "organization form",  # 额外字段
    "经营期限": "term of operation",  # 八字段模板专用
}

# 需要排除的字段（这些不是表格内的字段，不应该填充到模板中）
EXCLUDED_FIELDS = {
    "统一社会信用代码",
    "unified social credit code",
    "注册号",
    "registration number",
    "执照编号",
    "license number",
    "副本",
    "duplicate",
    "duplicate information",
    "duplicate identifier",
    "正本",
    "original",
    "营业执照",
    "business license",
    # 排除文档标题和监制等非核心字段
    "文档标题",
    "document title",
    "title",  # 排除 Title 标签
    "监制",
    "监制单位",
    "supervision",
    "supervised by",
    "supervising authority",
    "made under the supervision",
    "国家企业信用信息公示系统网址",
    "website of national enterprise credit information publicity system",
    "national enterprise credit information publicity system website",
    "国家市场监督管理总局监制",
    "supervised and produced by the state administration for market regulation",
    "state administration for market regulation",
    # 登记机关通过印章 [Seal] 处理，不作为字段填充
    "登记机关",
    "registration authority",
    # 副本信息
    "副本信息",
}

# 允许添加到模板的额外字段（白名单）
ALLOWED_EXTRA_FIELDS = {
    "important notice",  # 重要提示
    "business term",     # 营业期限
    "organization form", # 组成形式
}

# 新增占位符定义
# [CreditCode] - 统一社会信用代码
# [QRText] - 二维码右侧说明文字

# 模板中的标准字段顺序（用于动态填充）
TEMPLATE_STANDARD_FIELDS = [
    "business name",
    "ownership",
    "legal representative", 
    "operations",
    "registered capital",
    "date of incorporation",
    "domicile",
]

# 模板标签别名映射（处理模板中可能的不同写法）
TEMPLATE_LABEL_ALIASES = {
    "business name": ["business name", "name", "company name"],
    "ownership": ["ownership", "type", "company type", "type of business entity"],
    "legal representative": ["legal representative", "representative", "person in charge"],
    "operations": ["operations", "business scope", "scope"],
    "registered capital": ["registered capital", "capital"],
    "date of incorporation": ["date of incorporation", "incorporation date", "established", "date of establishment", "establishment date", "date established"],
    "domicile": ["domicile", "address", "location", "business address"],
    "term of operation": ["term of operation", "business term", "operating period", "operation term"],
    "important notice": ["important notice", "notice", "important"],
}

# 模板布局配置 - 定义左右栏的字段标签位置
# 注意：这只是默认布局，实际会根据原图的布局动态调整
# 这些是模板中红框区域内的字段标签占位符
TEMPLATE_LAYOUT = {
    "left_column": [
        "business name",      # 第1行左
        "ownership",          # 第2行左
        "legal representative", # 第3行左
        "operations",         # 第4行左
    ],
    "right_column": [
        "registered capital",    # 第1行右
        "date of incorporation", # 第2行右
        "domicile",             # 第3行右
    ],
    # 底部区域（不在红框内，单独处理）
    "bottom": []
}

# 是否根据原图布局动态调整模板（True = 按原图布局，False = 按模板固定布局）
USE_SOURCE_IMAGE_LAYOUT = True

# 字段优先级配置 - 当某些字段缺失时，如何重新分配
# 优先级越高的字段越重要，应该优先显示
FIELD_PRIORITY = {
    "business name": 1,
    "ownership": 2,
    "legal representative": 3,
    "registered capital": 4,
    "date of incorporation": 5,
    "domicile": 6,
    "operations": 7,
}

# 额外字段的英文标签映射（用于显示在模板中）
EXTRA_FIELD_LABELS = {
    "important notice": "Important Notice",
    "business term": "Business Term",
    "organization form": "Organization Form",
    "business scope note": "Business Scope Note",
}

# 输入图片路径；留空表示运行时必须显式传入
INPUT_IMAGE = ""

TEMPLATE_DIR = SCRIPT_DIR / "template"

# Word模板路径
TEMPLATE_PATH = str(TEMPLATE_DIR / "模板.docx")
# 八字段模板路径（包含经营期限字段）
TEMPLATE_PATH_8 = str(TEMPLATE_DIR / "模板8.docx")
# 竖版模板路径（用于竖版营业执照）
TEMPLATE_PATH_VERTICAL = str(TEMPLATE_DIR / "模板竖.docx")

# 输出目录
OUTPUT_DIR = str(SCRIPT_DIR.parent / "outputs" / "business_licence")

# ==================================================


def show_company_name_dialog(ai_translated_name: str, original_cn_name: str) -> str:
    """显示公司名称翻译选择弹窗
    
    Args:
        ai_translated_name: AI翻译的英文公司名称
        original_cn_name: 原始中文公司名称
        
    Returns:
        用户选择或输入的英文公司名称
    """
    result = {"name": ai_translated_name}  # 默认使用AI翻译
    
    def on_confirm():
        if choice_var.get() == "ai":
            result["name"] = ai_translated_name
        else:
            manual_name = manual_entry.get().strip()
            if not manual_name:
                messagebox.showwarning("提示", "请输入公司英文名称")
                return
            result["name"] = manual_name
        dialog.destroy()
    
    def on_choice_change():
        if choice_var.get() == "manual":
            manual_entry.config(state='normal')
            manual_entry.focus_set()
        else:
            manual_entry.config(state='disabled')
    
    # 创建弹窗
    dialog = tk.Tk()
    dialog.title("公司名称翻译确认")
    dialog.geometry("550x420")
    dialog.resizable(False, False)
    
    # 居中显示
    dialog.update_idletasks()
    x = (dialog.winfo_screenwidth() - 550) // 2
    y = (dialog.winfo_screenheight() - 420) // 2
    dialog.geometry(f"550x420+{x}+{y}")
    
    # 主框架
    main_frame = ttk.Frame(dialog, padding="20")
    main_frame.pack(fill=tk.BOTH, expand=True)
    
    # 标题
    title_label = ttk.Label(main_frame, text="请确认公司名称翻译", font=('Microsoft YaHei', 14, 'bold'))
    title_label.pack(pady=(0, 15))
    
    # 原始中文名称
    ttk.Label(main_frame, text="中文名称:", font=('Microsoft YaHei', 10)).pack(anchor=tk.W)
    cn_text = tk.Text(main_frame, height=2, width=60, font=('Microsoft YaHei', 10), wrap=tk.WORD, bg='#f0f0f0')
    cn_text.insert('1.0', original_cn_name)
    cn_text.config(state='disabled')
    cn_text.pack(fill=tk.X, pady=(2, 10))
    
    # 分隔线
    ttk.Separator(main_frame, orient='horizontal').pack(fill=tk.X, pady=10)
    
    # 选项变量
    choice_var = tk.StringVar(value="ai")
    
    # AI翻译选项
    ai_radio = ttk.Radiobutton(main_frame, text="使用AI翻译", variable=choice_var, value="ai", command=on_choice_change)
    ai_radio.pack(anchor=tk.W, pady=(5, 2))
    
    # AI翻译结果显示框
    ai_text = tk.Text(main_frame, height=2, width=60, font=('Microsoft YaHei', 10), wrap=tk.WORD, fg='green', bg='#f5fff5')
    ai_text.insert('1.0', ai_translated_name)
    ai_text.config(state='disabled')
    ai_text.pack(fill=tk.X, pady=(0, 10), padx=(20, 0))
    
    # 手动输入选项
    manual_radio = ttk.Radiobutton(main_frame, text="手动输入公司英文名称", variable=choice_var, value="manual", command=on_choice_change)
    manual_radio.pack(anchor=tk.W, pady=(5, 2))
    
    # 手动输入框
    manual_entry = ttk.Entry(main_frame, width=60, font=('Microsoft YaHei', 10))
    manual_entry.pack(fill=tk.X, pady=(0, 10), padx=(20, 0))
    manual_entry.config(state='disabled')
    
    # 提示文字
    hint_label = ttk.Label(main_frame, text="提示: 某些公司有固定的英文译名，请根据实际情况选择", 
                          font=('Microsoft YaHei', 9), foreground='gray')
    hint_label.pack(pady=(5, 10))
    
    # 确认按钮
    confirm_btn = ttk.Button(main_frame, text="确  认", command=on_confirm, width=15)
    confirm_btn.pack(pady=10)
    
    # 绑定回车键
    dialog.bind('<Return>', lambda e: on_confirm())
    
    # 运行弹窗
    dialog.mainloop()
    
    return result["name"]


def get_image_orientation(image_path: str) -> str:
    """根据图片尺寸判断是横版还是竖版
    
    Args:
        image_path: 图片文件路径
        
    Returns:
        "horizontal" 表示横版（宽 > 高）
        "vertical" 表示竖版（高 >= 宽）
    """
    from PIL import Image
    
    try:
        with Image.open(image_path) as img:
            width, height = img.size
            
            if width > height:
                print(f"图片尺寸: {width}x{height}，判定为横版")
                return "horizontal"
            else:
                print(f"图片尺寸: {width}x{height}，判定为竖版")
                return "vertical"
    except Exception as e:
        print(f"无法读取图片尺寸，默认使用横版: {e}")
        return "horizontal"


def encode_image(image_path: str) -> str:
    """将图片编码为base64"""
    with open(image_path, "rb") as f:
        return base64.standard_b64encode(f.read()).decode("utf-8")


def get_image_mime_type(image_path: str) -> str:
    """根据文件扩展名获取MIME类型"""
    ext = Path(image_path).suffix.lower()
    mime_types = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }
    return mime_types.get(ext, "image/jpeg")


def extract_and_translate(image_path: str) -> dict:
    """使用Gemini从图片中提取字段并翻译"""
    import requests
    
    image_data = encode_image(image_path)
    mime_type = get_image_mime_type(image_path)
    
    prompt = """请仔细识别这张图片中的所有文字内容，提取出字段标签和对应的字段值，并将中文翻译成英文。

请以JSON格式返回结果，格式如下：
{
    "document_type": "文档类型（如：Business License, Invoice, ID Card等）",
    "layout": "文档布局（single_column 或 two_column）",
    "fields": [
        {
            "label_cn": "中文字段标签",
            "label_en": "English Field Label", 
            "value_cn": "中文字段值",
            "value_en": "English Field Value",
            "position": "字段在原图中的位置（left 表示左栏，right 表示右栏，full_width 表示跨越整行，如果是单栏布局则都填 left）",
            "row": 行号（从1开始，同一行的字段行号相同）,
            "importance": "字段重要性（primary 表示主要字段如名称/类型/法人/经营范围/注册资本/成立日期/住所，secondary 表示次要字段）"
        }
    ],
    "seal_text": {
        "organization_cn": "印章中的机构名称（中文）",
        "organization_en": "Organization name in seal (English, use standard government agency translation format)",
        "date_cn": "印章覆盖的日期（中文，通常在印章下方或被印章覆盖的日期）",
        "date_en": "Date covered by seal (English)"
    },
    "credit_code": {
        "code": "统一社会信用代码的编号（如：91440000MA4XXXXXX）",
        "full_text_cn": "完整的信用代码文字（包括'统一社会信用代码'标签和编号）",
        "full_text_en": "Full credit code text in English (e.g., 'Unified Social Credit Code: 91440000MA4XXXXXX')"
    },
    "registration_no": {
        "no": "注册号/编号（如果文档中有单独的注册号或编号字段，提取其值；如果没有则为空字符串）",
        "full_text_en": "Registration No.: XXXXXXXX（完整的英文格式，如果没有则为空字符串）"
    },
    "qr_code": {
        "exists": true或false（图片中是否存在二维码）, 
        "bbox": [x1, y1, x2, y2]（二维码的边界框坐标，像素值，左上角为原点。如果不存在则为null）,
        "position_description": "二维码在图片中的位置描述（如：右上角、左下角等）"
    },
    "qr_text": {
        "text_cn": "二维码旁边的说明文字（中文）",
        "text_en": "QR code description text (English translation)"
    },
    "duplicate_number": {
        "exists": true或false（图片中是否存在副本编号，如"副本 1-1"、"第一副本"等）,
        "number_cn": "副本编号的中文原文（如：1-1、第一副本等）",
        "number_en": "副本编号的英文格式（如：1-1）"
    }
}

注意：
1. 字段标签是指表单中的标题或提示文字（如"名称"、"地址"、"注册号"等）
2. 字段值是指对应标签后面的具体内容
3. 请识别所有可见的字段
4. 如果某个字段值为空，请用空字符串表示
5. 翻译要准确、专业
6. 只返回JSON，不要有其他解释文字
7. 特别注意识别红色印章中的文字（通常是圆形或椭圆形的公章），提取机构名称
8. 印章覆盖的日期通常在印章附近或被印章部分覆盖，格式如"2016年04月13日"，翻译成英文格式如"April 13, 2016"
9. 如果没有印章或印章文字不清晰，seal_text中的字段可以为空字符串
16. 重要：印章中的机构名称翻译必须使用标准政府机构翻译格式，例如：
    - "XX市XX区市场监督管理局" -> "Administration for Market Regulation of XX District, XX City"
    - "XX市市场监督管理局" -> "Administration for Market Regulation of XX City"
    - "XX省XX市工商行政管理局" -> "Administration for Industry and Commerce of XX City, XX Province"
    - 翻译时保持地名的拼音形式（如 Foshan, Zhuhai, Guangzhou 等），机构名称使用标准英文翻译
    - 严禁使用 "Municipality" 翻译"市"，必须使用 "City"
    - 严禁使用 "Bureau" 翻译"局"，必须使用 "Administration"
    - 正确格式示例：佛山市市场监督管理局 -> "Administration for Market Regulation of Foshan City"
    - 错误格式示例：佛山市市场监督管理局 -> "Foshan Municipality Market Supervision Bureau"（禁止）
10. 重要：请判断文档是单栏还是双栏布局，并标注每个字段的位置（left/right/full_width）和行号
11. 对于双栏布局，左栏和右栏同一水平位置的字段应该有相同的行号
12. 对于跨越整行的字段（如经营范围通常占据整行），position 应标记为 full_width
13. 特别注意识别"统一社会信用代码"（通常在文档左上角），提取完整的代码编号
13. 如果文档中有单独的"注册号"或"编号"字段（不同于统一社会信用代码），请提取到 registration_no 中
14. 特别注意识别二维码旁边的说明文字（通常在文档右上角，描述扫码用途），并翻译成英文。重要：只识别图片中实际存在的文字，如果二维码旁边没有任何说明文字，qr_text 的 text_cn 和 text_en 必须返回空字符串，不要猜测或推断可能存在的文字
15. 重要：请识别图片中的二维码位置，返回其边界框坐标（像素值）。坐标格式为[x1, y1, x2, y2]，其中(x1,y1)是左上角，(x2,y2)是右下角
17. 特别注意识别营业执照上的副本编号（通常在文档右上角或标题附近，如"副本"、"副本 1-1"、"第一副本"等）。如果存在副本编号，提取其编号格式（如1-1）；如果只写"副本"没有具体编号，或者是"正本"，则 duplicate_number.exists 为 false"""

    print("正在使用 Google AI Studio (Gemini) 识别图片并翻译...")
    import requests
    
    # Google AI Studio REST API
    api_url = f"https://generativelanguage.googleapis.com/v1beta/{MODEL}:generateContent?key={GOOGLE_API_KEY}"
    
    payload = {
        "contents": [
            {
                "parts": [
                    {"text": prompt},
                    {
                        "inline_data": {
                            "mime_type": mime_type,
                            "data": image_data
                        }
                    }
                ]
            }
        ],
        "generationConfig": {
            "temperature": 0
        }
    }
    
    try:
        response = requests.post(
            api_url,
            json=payload,
            timeout=180,
            proxies={"http": PROXY, "https": PROXY} if PROXY else None,
        )
    except requests.exceptions.RequestException as e:
        print(f"请求异常: {e}")
        return {"fields": [], "document_type": "Unknown"}
    
    if response.status_code != 200:
        print(f"API请求失败: {response.status_code}")
        print(response.text)
        return {"fields": [], "document_type": "Unknown"}
    
    result = response.json()
    try:
        result_text = result["candidates"][0]["content"]["parts"][0]["text"]
    except (KeyError, IndexError) as e:
        print(f"解析响应失败: {e}")
        print(f"响应内容: {result}")
        return {"fields": [], "document_type": "Unknown"}
    print(f"\n识别和翻译结果:\n{result_text}\n")
    
    # 解析JSON
    try:
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', result_text)
        if json_match:
            json_str = json_match.group(1)
        else:
            json_str = result_text
        parsed_data = json.loads(json_str)
        
        # 确保 seal_text 字段存在
        if "seal_text" not in parsed_data:
            parsed_data["seal_text"] = {
                "organization_cn": "",
                "organization_en": "",
                "date_cn": "",
                "date_en": ""
            }
            print("注意: API未返回印章信息，seal_text 为空")
        
        # 确保 credit_code 字段存在
        if "credit_code" not in parsed_data:
            parsed_data["credit_code"] = {
                "code": "",
                "full_text_cn": "",
                "full_text_en": ""
            }
            print("注意: API未返回信用代码信息，credit_code 为空")
        
        # 确保 registration_no 字段存在
        if "registration_no" not in parsed_data:
            parsed_data["registration_no"] = {
                "no": "",
                "full_text_en": ""
            }
            print("注意: API未返回注册号信息，registration_no 为空")
        
        # 确保 qr_text 字段存在
        if "qr_text" not in parsed_data:
            parsed_data["qr_text"] = {
                "text_cn": "",
                "text_en": ""
            }
            print("注意: API未返回二维码说明文字，qr_text 为空")
        
        # 确保 qr_code 字段存在
        if "qr_code" not in parsed_data:
            parsed_data["qr_code"] = {
                "exists": False,
                "bbox": None,
                "position_description": ""
            }
            print("注意: API未返回二维码位置信息，qr_code 为空")
        
        # 确保 duplicate_number 字段存在
        if "duplicate_number" not in parsed_data:
            parsed_data["duplicate_number"] = {
                "exists": False,
                "number_cn": "",
                "number_en": ""
            }
            print("注意: API未返回副本编号信息，duplicate_number 为空")
        
        # 处理营业执照编号中的"外"字，翻译为"W."
        def replace_wai_in_code(code_str):
            """将编号中的'外'字替换为'W.'"""
            if code_str:
                return code_str.replace("外", "W.")
            return code_str
        
        # 处理 credit_code 中的编号
        if "credit_code" in parsed_data:
            if parsed_data["credit_code"].get("code"):
                parsed_data["credit_code"]["code"] = replace_wai_in_code(parsed_data["credit_code"]["code"])
            if parsed_data["credit_code"].get("full_text_en"):
                parsed_data["credit_code"]["full_text_en"] = replace_wai_in_code(parsed_data["credit_code"]["full_text_en"])
        
        # 处理 registration_no 中的编号
        if "registration_no" in parsed_data:
            if parsed_data["registration_no"].get("no"):
                parsed_data["registration_no"]["no"] = replace_wai_in_code(parsed_data["registration_no"]["no"])
            if parsed_data["registration_no"].get("full_text_en"):
                parsed_data["registration_no"]["full_text_en"] = replace_wai_in_code(parsed_data["registration_no"]["full_text_en"])
        
        return parsed_data
    except json.JSONDecodeError as e:
        print(f"JSON解析失败: {e}")
        return {"fields": [], "document_type": "Unknown", "seal_text": {}}


def fill_vertical_template(template_path: str, output_path: str, data: dict, image_path: str = None):
    """填充竖版营业执照模板
    
    竖版模板结构：
    - 外层表格只有1行1列
    - 嵌套表格有10行3列
    - 行0-7是字段行：列0是标签，列1-2是值（合并单元格）
    - 行8是印章区域
    - 行9是底部信息
    """
    from docx import Document
    from docx.shared import RGBColor, Pt
    import re
    
    print(f"正在填充竖版模板: {template_path}")
    
    doc = Document(template_path)
    fields = data.get("fields", [])
    seal_text = data.get("seal_text", {})
    
    # 获取印章文字和日期
    seal_org_en = seal_text.get("organization_en", "").strip()
    seal_date_en = seal_text.get("date_en", "").strip()
    
    # 获取信用代码信息
    credit_code_data = data.get("credit_code", {})
    credit_code = credit_code_data.get("code", "").strip()
    credit_code_full_en = credit_code_data.get("full_text_en", "").strip()
    credit_code_value = credit_code if credit_code else credit_code_full_en
    
    # 获取注册号信息
    registration_no_data = data.get("registration_no", {})
    registration_no = registration_no_data.get("no", "").strip()
    registration_no_full_en = registration_no_data.get("full_text_en", "").strip()
    registration_no_value = registration_no if registration_no else registration_no_full_en
    
    # 获取副本编号信息
    duplicate_number_data = data.get("duplicate_number", {})
    duplicate_number_exists = duplicate_number_data.get("exists", False)
    duplicate_number_en = duplicate_number_data.get("number_en", "").strip()
    
    replacements_made = 0
    
    # 构建字段值映射（标签 -> 英文值）
    field_value_map = {}
    for field in fields:
        label_cn = field.get("label_cn", "").strip().replace(" ", "").replace("　", "")
        label_en = field.get("label_en", "").strip().lower()
        value_en = field.get("value_en", "").strip()
        
        if not value_en:
            continue
        
        # 通过中文标签映射到标准英文标签
        standard_label = None
        if label_cn in FIXED_LABEL_TRANSLATIONS:
            standard_label = FIXED_LABEL_TRANSLATIONS[label_cn]
        else:
            for k, v in FIXED_LABEL_TRANSLATIONS.items():
                if k.replace(" ", "").replace("　", "") == label_cn:
                    standard_label = v
                    break
        
        if standard_label:
            # 查找主标签名
            main_label = standard_label
            for ml, aliases in TEMPLATE_LABEL_ALIASES.items():
                if standard_label == ml or standard_label in aliases:
                    main_label = ml
                    break
            field_value_map[main_label] = value_en
            print(f"  字段映射: {label_cn} -> {main_label} = {value_en[:40]}...")
    
    # 参照横版模板处理：检查是否有重要提示字段，如果有则替换到经营范围的位置
    important_notice_value = None
    important_notice_display = "Important Notice"  # 重要提示的显示标签
    has_important_notice = False
    for label in list(field_value_map.keys()):
        if label.lower() == "important notice":
            important_notice_value = field_value_map[label]
            has_important_notice = True
            # 将重要提示映射为operations，这样它会填充到经营范围的位置
            field_value_map["operations"] = important_notice_value
            # 删除原来的important notice
            del field_value_map[label]
            print(f"  重要提示将替换到经营范围位置: {important_notice_value[:40]}...")
            break
    
    print(f"\n识别到 {len(field_value_map)} 个有效字段")
    print(f"印章机构: {seal_org_en}")
    print(f"印章日期: {seal_date_en}")
    print(f"信用代码: {credit_code_value}")
    print(f"注册号: {registration_no_value}")
    
    # 竖版模板的标签映射（模板中的英文标签 -> 主标签）
    vertical_label_map = {
        "business name": "business name",
        "ownership": "ownership",
        "domicile": "domicile",
        "legal representative": "legal representative",
        "registered capital": "registered capital",
        "date of incorporation": "date of incorporation",
        "term of operation": "term of operation",
        "operations": "operations",
        "important notice": "important notice",  # 重要提示
    }
    
    # 辅助函数：更新单元格内容并保留格式
    def update_cell_content(cell, new_text):
        if cell.paragraphs and cell.paragraphs[0].runs:
            original_para = cell.paragraphs[0]
            original_run = original_para.runs[0]
            original_font = original_run.font
            original_alignment = original_para.alignment
            
            original_para.clear()
            new_run = original_para.add_run(new_text)
            
            if original_font.name:
                new_run.font.name = original_font.name
            if original_font.size:
                new_run.font.size = original_font.size
            if original_font.bold is not None:
                new_run.font.bold = original_font.bold
            
            original_para.alignment = original_alignment
        else:
            cell.text = new_text
    
    # 辅助函数：处理段落中的占位符（处理分散在多个run中的占位符）
    def process_paragraph_placeholder(para, placeholder_pattern, replacement_text, is_seal=False):
        """处理段落中的占位符，支持跨run的占位符"""
        nonlocal replacements_made
        
        para_text = para.text
        if not re.search(placeholder_pattern, para_text, re.IGNORECASE):
            return False
        
        # 保存原有格式
        original_alignment = para.alignment
        original_font_name = None
        original_font_size = None
        if para.runs:
            original_font_name = para.runs[0].font.name
            original_font_size = para.runs[0].font.size
        
        # 替换文本
        new_text = re.sub(placeholder_pattern, replacement_text, para_text, flags=re.IGNORECASE)
        
        # 清空段落并重新填充
        para.clear()
        
        if is_seal and replacement_text:
            # 印章需要特殊格式（红色斜体）
            run = para.add_run(new_text.replace(replacement_text, ''))
            if original_font_name:
                run.font.name = original_font_name
            if original_font_size:
                run.font.size = original_font_size
            
            seal_run = para.add_run(replacement_text)
            seal_run.italic = True
            seal_run.font.color.rgb = RGBColor(255, 0, 0)
            if original_font_size:
                seal_run.font.size = original_font_size
        else:
            run = para.add_run(new_text)
            if original_font_name:
                run.font.name = original_font_name
            if original_font_size:
                run.font.size = original_font_size
        
        para.alignment = original_alignment
        replacements_made += 1
        return True
    
    # 处理所有段落中的占位符
    def process_all_placeholders_in_para(para):
        para_text = para.text
        
        # 检查各种占位符
        has_credit_code = re.search(r'\[credit\s*code\s*\]', para_text, re.IGNORECASE)
        has_seal = re.search(r'\[seal\s*\]', para_text, re.IGNORECASE)
        has_date = re.search(r'\[date\s*\]', para_text, re.IGNORECASE)
        has_no = re.search(r'\[no\.?\s*\]', para_text, re.IGNORECASE)
        has_duplicate = re.search(r'\[duplicate\s*number\s*\]', para_text, re.IGNORECASE)
        
        if not any([has_credit_code, has_seal, has_date, has_no, has_duplicate]):
            return
        
        # 保存原有格式
        original_alignment = para.alignment
        original_font_name = None
        original_font_size = None
        if para.runs:
            original_font_name = para.runs[0].font.name
            original_font_size = para.runs[0].font.size
        
        new_text = para_text
        seal_replacement = None
        
        # 处理 [Credit Code]
        if has_credit_code and credit_code_value:
            new_text = re.sub(r'\[credit\s*code\s*\]', credit_code_value, new_text, flags=re.IGNORECASE)
            print(f"填充信用代码: [Credit Code] -> {credit_code_value}")
            nonlocal replacements_made
            replacements_made += 1
        
        # 处理 [No.]
        if has_no:
            if registration_no_value:
                no_text = f"No. {registration_no_value}"
                new_text = re.sub(r'\[no\.?\s*\]', no_text, new_text, flags=re.IGNORECASE)
                print(f"填充注册号: [No.] -> {no_text}")
                replacements_made += 1
            else:
                new_text = re.sub(r'\s*\[no\.?\s*\]\s*', '', new_text, flags=re.IGNORECASE)
                print(f"删除空的注册号占位符")
        
        # 处理 [Duplicate number]
        if has_duplicate:
            if duplicate_number_exists and duplicate_number_en:
                dup_text = f"(Duplicate number: {duplicate_number_en})"
                new_text = re.sub(r'\[duplicate\s*number\s*\]', dup_text, new_text, flags=re.IGNORECASE)
                print(f"填充副本号: [Duplicate number] -> {dup_text}")
                replacements_made += 1
            else:
                new_text = re.sub(r'\s*\[duplicate\s*number\s*\]\s*', '', new_text, flags=re.IGNORECASE)
                print(f"删除空的副本号占位符")
        
        # 处理 [Date]
        if has_date and seal_date_en:
            new_text = re.sub(r'\[date\s*\]', seal_date_en, new_text, flags=re.IGNORECASE)
            print(f"填充日期: [Date] -> {seal_date_en}")
            replacements_made += 1
        
        # 处理 [Seal] - 需要特殊格式
        if has_seal and seal_org_en:
            seal_replacement = f"(Seal of {seal_org_en})"
            seal_placeholder = "___SEAL___"
            new_text = re.sub(r'\[seal\s*\]', seal_placeholder, new_text, flags=re.IGNORECASE)
        
        # 清空段落并重新填充
        para.clear()
        
        if seal_replacement and seal_placeholder in new_text:
            parts = new_text.split(seal_placeholder)
            for idx, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    if original_font_name:
                        run.font.name = original_font_name
                    if original_font_size:
                        run.font.size = original_font_size
                if idx < len(parts) - 1:
                    seal_run = para.add_run(seal_replacement)
                    seal_run.italic = True
                    seal_run.font.color.rgb = RGBColor(255, 0, 0)
                    if original_font_size:
                        seal_run.font.size = original_font_size
            print(f"填充印章: [Seal] -> {seal_replacement}")
            replacements_made += 1
        else:
            run = para.add_run(new_text)
            if original_font_name:
                run.font.name = original_font_name
            if original_font_size:
                run.font.size = original_font_size
        
        para.alignment = original_alignment
    
    # 处理文本框中的占位符（[No.] 可能在文本框中）
    print("\n扫描文本框中的占位符...")
    from docx.oxml.ns import qn as docx_qn
    
    for element in doc._element.iter():
        if element.tag.endswith('txbxContent'):
            for p_elem in element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                text_parts = []
                t_elements = p_elem.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                for t_elem in t_elements:
                    if t_elem.text:
                        text_parts.append(t_elem.text)
                para_text = ''.join(text_parts)
                
                if not para_text.strip():
                    continue
                
                # 检查 [No.] 占位符（支持各种变体）
                has_no = re.search(r'\[No\.?\s*\]', para_text, re.IGNORECASE)
                if has_no:
                    if registration_no_value:
                        # 有注册号，替换占位符
                        new_text = re.sub(r'\[No\.?\s*\]', f"No. {registration_no_value}", para_text, flags=re.IGNORECASE)
                        # 清空所有文本元素，只在第一个中填入新文本
                        first = True
                        for t in t_elements:
                            if first:
                                t.text = new_text
                                first = False
                            else:
                                t.text = ''
                        print(f"  填充文本框注册号: [No.] -> No. {registration_no_value}")
                        replacements_made += 1
                    else:
                        # 没有注册号，删除整个占位符
                        new_text = re.sub(r'\s*\[No\.?\s*\]\s*', '', para_text, flags=re.IGNORECASE)
                        # 清空所有文本元素，只在第一个中填入新文本
                        first = True
                        for t in t_elements:
                            if first:
                                t.text = new_text
                                first = False
                            else:
                                t.text = ''
                        print(f"  删除文本框中空的注册号占位符")
    
    # 处理外层表格中的占位符（如 [Credit Code]）
    print("\n处理外层表格中的占位符...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_all_placeholders_in_para(para)
    
    # 处理嵌套表格中的字段
    print("\n处理嵌套表格中的字段...")
    filled_fields = set()
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 检查是否有嵌套表格
                nested_tables = cell.tables
                for nested_table in nested_tables:
                    print(f"  找到嵌套表格: {len(nested_table.rows)}行 x {len(nested_table.columns)}列")
                    
                    for row_idx, nested_row in enumerate(nested_table.rows):
                        # 去重单元格
                        seen_cells = []
                        for c in nested_row.cells:
                            if c not in seen_cells:
                                seen_cells.append(c)
                        
                        if len(seen_cells) < 2:
                            continue
                        
                        label_cell = seen_cells[0]
                        value_cell = seen_cells[1]  # 值单元格（可能是合并的）
                        
                        label_text = label_cell.text.strip().rstrip(':').rstrip('：').strip().lower()
                        
                        # 处理所有单元格中的占位符（印章、日期等）
                        for c in seen_cells:
                            for para in c.paragraphs:
                                process_all_placeholders_in_para(para)
                        
                        if not label_text:
                            continue
                        
                        # 查找匹配的主标签
                        matched_label = None
                        for template_label, main_label in vertical_label_map.items():
                            if label_text == template_label or label_text.startswith(template_label.split()[0]):
                                matched_label = main_label
                                break
                        
                        # 也检查别名
                        if not matched_label:
                            for main_label, aliases in TEMPLATE_LABEL_ALIASES.items():
                                if label_text in aliases or any(label_text.startswith(a.split()[0]) for a in aliases):
                                    matched_label = main_label
                                    break
                        
                        if matched_label and matched_label in field_value_map:
                            value = field_value_map[matched_label]
                            if matched_label not in filled_fields:
                                # 如果是operations且有重要提示，需要同时更新标签
                                if matched_label == "operations" and has_important_notice:
                                    update_cell_content(label_cell, f"{important_notice_display}:")
                                    print(f"    行{row_idx}: 标签替换为 '{important_notice_display}:'")
                                update_cell_content(value_cell, value)
                                filled_fields.add(matched_label)
                                print(f"    行{row_idx}: {label_text} -> {value[:40]}...")
                                replacements_made += 1
                        elif matched_label and matched_label not in field_value_map:
                            # 字段不存在，清空标签和值
                            update_cell_content(label_cell, "")
                            update_cell_content(value_cell, "")
                            print(f"    行{row_idx}: 清空不存在的字段 '{label_text}'")
    
    # 检查未填充的字段
    unfilled = set(field_value_map.keys()) - filled_fields
    if unfilled:
        print(f"\n警告: 以下字段未能填充: {unfilled}")
    
    print(f"\n竖版模板填充了 {replacements_made} 个字段")
    
    doc.save(output_path)
    print(f"\n文档已保存到: {output_path}")
    return replacements_made


def fill_template(template_path: str, output_path: str, data: dict, image_path: str = None):
    """将识别和翻译的字段填充到Word模板中
    
    模板格式：表格中标签在左边单元格，值填充到右边单元格
    特殊占位符：
    - [Seal]: 印章中的机构名称翻译
    - [Date]: 印章覆盖的日期翻译
    - [CreditCode]: 统一社会信用代码
    - [QRText]: 二维码右侧说明文字
    
    智能填充逻辑：
    1. 根据原图布局（单栏/双栏）智能分配字段到模板的左右栏
    2. 按照字段标签匹配模板中的位置
    3. 对于不同排版的营业执照，自动调整字段位置
    """
    from docx import Document
    from docx.shared import RGBColor, Inches, Pt
    import cv2
    import numpy as np
    import tempfile
    import os
    
    print(f"正在填充模板: {template_path}")
    
    doc = Document(template_path)
    fields = data.get("fields", [])
    seal_text = data.get("seal_text", {})
    layout = data.get("layout", "single_column")
    
    print(f"文档布局: {layout}")
    
    # ==================== 智能布局映射 ====================
    # 根据原图的字段位置，智能分配到模板的左右栏
    # 动态调整：删除缺失字段的标签，添加新字段
    
    def map_fields_to_template_layout(fields, source_layout):
        """
        根据原图字段位置，智能映射到模板布局
        
        智能处理逻辑：
        1. 识别原图中存在的字段及其位置（左/右栏）
        2. 如果 USE_SOURCE_IMAGE_LAYOUT=True，按原图的位置来排版
        3. 如果 USE_SOURCE_IMAGE_LAYOUT=False，按模板固定布局排版
        4. 删除模板中不存在的字段标签
        5. 添加原图中有但模板中没有的新字段
        
        返回: {
            "left_column": [(label, value, display_label), ...],
            "right_column": [(label, value, display_label), ...],
            "bottom": [(label, value, display_label), ...]
        }, field_value_map, field_display_label_map
        """
        # 首先提取所有有效字段及其标准标签
        # 使用白名单机制：只有在 FIXED_LABEL_TRANSLATIONS 中定义的标签才能被填充
        valid_fields = []
        for field in fields:
            label_cn = field.get("label_cn", "").strip()
            label_en = field.get("label_en", "").strip()
            value_en = field.get("value_en", "").strip()
            position = field.get("position", "left")  # 原图中的位置：left/right/full_width
            row = field.get("row", 0)
            
            if not value_en:
                continue
            
            # 白名单机制：只接受在 FIXED_LABEL_TRANSLATIONS 中定义的字段标签
            # 先尝试精确匹配，再尝试去空格匹配
            standard_label = None
            label_cn_normalized = label_cn.replace(" ", "").replace("　", "")  # 去除空格
            
            if label_cn in FIXED_LABEL_TRANSLATIONS:
                standard_label = FIXED_LABEL_TRANSLATIONS[label_cn]
            else:
                # 找到去空格后匹配的键
                for k, v in FIXED_LABEL_TRANSLATIONS.items():
                    if k.replace(" ", "").replace("　", "") == label_cn_normalized:
                        standard_label = v
                        break
            
            # 如果中文标签没有匹配，跳过这个字段（白名单机制）
            if not standard_label:
                print(f"    跳过非白名单字段: {label_cn} ({label_en})")
                continue
            
            # 检查是否是需要排除的字段
            if standard_label.lower() in [ex.lower() for ex in EXCLUDED_FIELDS]:
                print(f"    跳过排除字段: {label_cn} ({standard_label})")
                continue
            
            # 查找主标签名
            main_label = None
            for ml, aliases in TEMPLATE_LABEL_ALIASES.items():
                if standard_label == ml or standard_label in aliases:
                    main_label = ml
                    break
            
            # 如果没有匹配到主标签，使用标准标签
            if not main_label:
                main_label = standard_label
            
            valid_fields.append({
                "main_label": main_label,
                "value": value_en,
                "position": position,  # 保留原图位置信息
                "row": row,
                "label_cn": label_cn,
                "label_en": label_en,  # 保留原始英文标签用于显示
            })
        
        # 创建字段值映射、显示标签映射和位置映射
        field_value_map = {}
        field_display_label_map = {}
        field_position_map = {}  # 记录每个字段在原图中的位置
        field_row_map = {}  # 记录每个字段在原图中的行号
        
        for f in valid_fields:
            field_value_map[f["main_label"]] = f["value"]
            field_position_map[f["main_label"]] = f["position"]
            field_row_map[f["main_label"]] = f["row"]
            
            # 显示标签：优先使用模板标准标签，否则使用原始英文标签
            if f["main_label"] in TEMPLATE_LABEL_ALIASES:
                # 标准字段，使用 Title Case
                field_display_label_map[f["main_label"]] = f["main_label"].title().replace("Of", "of")
            elif f["main_label"] in EXTRA_FIELD_LABELS:
                field_display_label_map[f["main_label"]] = EXTRA_FIELD_LABELS[f["main_label"]]
            else:
                # 新字段，使用原始英文标签
                field_display_label_map[f["main_label"]] = f["label_en"].rstrip(':').rstrip('：').strip()
        
        # 打印识别到的字段及其原图位置
        print(f"\n  从原图识别到 {len(field_value_map)} 个有效字段:")
        for label, value in field_value_map.items():
            display = field_display_label_map.get(label, label)
            pos = field_position_map.get(label, "unknown")
            row = field_row_map.get(label, 0)
            print(f"    - {display}: {value[:40]}... [原图位置: {pos}, 行: {row}]")
        
        # 确定哪些是标准字段，哪些是额外字段
        all_standard_fields = set(TEMPLATE_LAYOUT["left_column"] + 
                                   TEMPLATE_LAYOUT["right_column"] + 
                                   TEMPLATE_LAYOUT["bottom"])
        
        present_standard_fields = {f for f in field_value_map.keys() if f in all_standard_fields}
        extra_fields = {f for f in field_value_map.keys() if f not in all_standard_fields}
        missing_standard_fields = all_standard_fields - present_standard_fields
        
        if missing_standard_fields:
            print(f"\n  缺失的标准字段（将从模板删除）: {', '.join(missing_standard_fields)}")
        if extra_fields:
            print(f"  额外的字段（将添加到模板）: {', '.join(extra_fields)}")
        
        # 根据配置决定布局方式
        result = {
            "left_column": [],
            "right_column": [],
            "bottom": []
        }
        
        # 检查是否有重要提示字段，如果有则替换到经营范围的位置
        important_notice_value = None
        important_notice_display = None
        for label in list(field_value_map.keys()):
            if label.lower() == "important notice":
                important_notice_value = field_value_map[label]
                important_notice_display = field_display_label_map.get(label, "Important Notice")
                # 将重要提示映射为operations，这样它会填充到经营范围的位置
                field_value_map["operations"] = important_notice_value
                field_display_label_map["operations"] = important_notice_display
                field_position_map["operations"] = field_position_map.get(label, "left")
                field_row_map["operations"] = field_row_map.get(label, 0)
                # 删除原来的important notice
                del field_value_map[label]
                if label in field_display_label_map:
                    del field_display_label_map[label]
                if label in field_position_map:
                    del field_position_map[label]
                if label in field_row_map:
                    del field_row_map[label]
                print(f"  重要提示将替换到经营范围位置: {important_notice_display}")
                break
        
        if USE_SOURCE_IMAGE_LAYOUT:
            # ========== 按原图布局排版 ==========
            print(f"\n  使用原图布局模式 (USE_SOURCE_IMAGE_LAYOUT=True)")
            
            # 按原图位置分组字段
            left_fields = []
            right_fields = []
            full_width_fields = []
            bottom_fields = []
            
            for label in field_value_map.keys():
                pos = field_position_map.get(label, "left")
                row = field_row_map.get(label, 0)
                display_label = field_display_label_map.get(label, label.title())
                value = field_value_map[label]
                
                # 按原图位置分配
                if pos == "right":
                    right_fields.append((label, value, display_label, row))
                elif pos == "full_width":
                    full_width_fields.append((label, value, display_label, row))
                else:  # left 或其他
                    left_fields.append((label, value, display_label, row))
            
            # 按行号排序
            left_fields.sort(key=lambda x: x[3])
            right_fields.sort(key=lambda x: x[3])
            full_width_fields.sort(key=lambda x: x[3])
            bottom_fields.sort(key=lambda x: x[3])
            
            # 转换为结果格式（去掉行号）
            result["left_column"] = [(f[0], f[1], f[2]) for f in left_fields]
            result["right_column"] = [(f[0], f[1], f[2]) for f in right_fields]
            result["bottom"] = [(f[0], f[1], f[2]) for f in bottom_fields]
            
            # full_width 字段添加到左栏末尾（它们会跨越整行）
            for f in full_width_fields:
                result["left_column"].append((f[0], f[1], f[2]))
            
            # 打印布局结果
            print(f"\n  按原图布局分配:")
            for label, value, display in result["left_column"]:
                print(f"    左栏: {display} = {value[:30]}...")
            for label, value, display in result["right_column"]:
                print(f"    右栏: {display} = {value[:30]}...")
            for label, value, display in result["bottom"]:
                print(f"    底部: {display} = {value[:30]}...")
        else:
            # ========== 按模板固定布局排版 ==========
            print(f"\n  使用模板固定布局模式 (USE_SOURCE_IMAGE_LAYOUT=False)")
            
            # 按模板定义的顺序填充左栏（只填充存在的字段）
            for template_label in TEMPLATE_LAYOUT["left_column"]:
                if template_label in field_value_map:
                    display_label = field_display_label_map.get(template_label, template_label.title())
                    result["left_column"].append((template_label, field_value_map[template_label], display_label))
                    print(f"  左栏: {display_label} = {field_value_map[template_label][:30]}...")
            
            # 按模板定义的顺序填充右栏（只填充存在的字段）
            for template_label in TEMPLATE_LAYOUT["right_column"]:
                if template_label in field_value_map:
                    display_label = field_display_label_map.get(template_label, template_label.title())
                    result["right_column"].append((template_label, field_value_map[template_label], display_label))
                    print(f"  右栏: {display_label} = {field_value_map[template_label][:30]}...")
            
            # 底部字段
            for template_label in TEMPLATE_LAYOUT["bottom"]:
                if template_label in field_value_map:
                    display_label = field_display_label_map.get(template_label, template_label.title())
                    result["bottom"].append((template_label, field_value_map[template_label], display_label))
                    print(f"  底部: {display_label} = {field_value_map[template_label][:30]}...")
            
            # 添加额外字段到左栏末尾（如重要提示等）
            for extra_label in extra_fields:
                display_label = field_display_label_map.get(extra_label, extra_label.title())
                result["left_column"].append((extra_label, field_value_map[extra_label], display_label))
                print(f"  左栏(额外): {display_label} = {field_value_map[extra_label][:30]}...")
        
        return result, field_value_map, field_display_label_map
    
    print("\n开始智能布局映射...")
    mapped_layout, field_values, display_labels = map_fields_to_template_layout(fields, layout)
    
    print(f"\n识别到 {len(field_values)} 个字段:")
    for label, value in field_values.items():
        print(f"  {label}: {value[:50]}...")
    
    # 获取印章文字和日期
    seal_org_en = seal_text.get("organization_en", "").strip()
    seal_date_en = seal_text.get("date_en", "").strip()
    
    # 获取信用代码信息
    credit_code_data = data.get("credit_code", {})
    credit_code = credit_code_data.get("code", "").strip()
    credit_code_full_en = credit_code_data.get("full_text_en", "").strip()
    
    # 优先使用纯编号，避免标签重复
    credit_code_value = credit_code if credit_code else credit_code_full_en
    
    # 获取二维码说明文字
    qr_text_data = data.get("qr_text", {})
    qr_text_en = qr_text_data.get("text_en", "").strip()
    
    # 获取副本编号信息
    duplicate_number_data = data.get("duplicate_number", {})
    duplicate_number_exists = duplicate_number_data.get("exists", False)
    duplicate_number_en = duplicate_number_data.get("number_en", "").strip()
    
    # 获取注册号信息
    registration_no_data = data.get("registration_no", {})
    registration_no = registration_no_data.get("no", "").strip()
    registration_no_full_en = registration_no_data.get("full_text_en", "").strip()
    registration_no_value = registration_no if registration_no else registration_no_full_en
    
    # 如果 registration_no 为空，尝试从 fields 中查找
    if not registration_no_value:
        for field in fields:
            label_cn = field.get("label_cn", "").strip().lower()
            label_en = field.get("label_en", "").strip().lower()
            value_en = field.get("value_en", "").strip()
            
            if any(keyword in label_cn for keyword in ["注册号", "编号", "执照编号"]) or \
               any(keyword in label_en for keyword in ["registration no", "registration number", "license no"]):
                registration_no_value = value_en
                print(f"从 fields 中找到注册号: {label_cn} -> {value_en}")
                break
    
    print(f"\n印章机构名称: {seal_org_en}")
    print(f"印章日期: {seal_date_en}")
    print(f"信用代码: {credit_code_value}")
    print(f"注册号: {registration_no_value}")
    print(f"二维码说明文字: {qr_text_en[:50]}..." if len(qr_text_en) > 50 else f"二维码说明文字: {qr_text_en}")
    print(f"副本编号: {duplicate_number_en if duplicate_number_exists else '无'}")
    
    replacements_made = 0
    
    # 先处理 [Seal] 和 [Date] 占位符（在填充字段之前）
    # 处理单元格中的所有段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 遍历单元格中的每个段落
                for para in cell.paragraphs:
                    para_text = para.text
                    has_seal = re.search(r'\[seal\s*\]', para_text, re.IGNORECASE)
                    has_date = re.search(r'\[date\s*\]', para_text, re.IGNORECASE)
                    
                    if has_seal or has_date:
                        print(f"找到占位符段落: '{para_text}'")
                        new_text = para_text
                        
                        # 保存原有格式
                        original_alignment = para.alignment
                        original_font_name = None
                        original_font_size = None
                        if para.runs:
                            original_font_name = para.runs[0].font.name
                            original_font_size = para.runs[0].font.size
                        
                        if has_date and seal_date_en:
                            new_text = re.sub(r'\[date\s*\]', seal_date_en, new_text, flags=re.IGNORECASE)
                            print(f"填充日期: [Date] -> {seal_date_en}")
                            replacements_made += 1
                        
                        seal_placeholder = "___SEAL_PLACEHOLDER___"
                        if has_seal and seal_org_en:
                            new_text = re.sub(r'\[seal\s*\]', seal_placeholder, new_text, flags=re.IGNORECASE)
                        
                        # 清空段落并重新填充
                        para.clear()
                        
                        if seal_placeholder in new_text:
                            parts = new_text.split(seal_placeholder)
                            for idx, part in enumerate(parts):
                                if part:
                                    run = para.add_run(part)
                                    if original_font_name:
                                        run.font.name = original_font_name
                                    if original_font_size:
                                        run.font.size = original_font_size
                                if idx < len(parts) - 1:
                                    seal_run = para.add_run(f"(Seal of {seal_org_en})")
                                    seal_run.italic = True
                                    seal_run.font.color.rgb = RGBColor(255, 0, 0)
                                    if original_font_size:
                                        seal_run.font.size = original_font_size
                            print(f"填充印章: [Seal] -> (Seal of {seal_org_en})")
                            replacements_made += 1
                        else:
                            run = para.add_run(new_text)
                            if original_font_name:
                                run.font.name = original_font_name
                            if original_font_size:
                                run.font.size = original_font_size
                        
                        # 恢复对齐方式
                        para.alignment = original_alignment
    
    # 处理文档段落中的占位符（表格外的段落）
    for para in doc.paragraphs:
        has_seal = re.search(r'\[seal\s*\]', para.text, re.IGNORECASE)
        has_date = re.search(r'\[date\s*\]', para.text, re.IGNORECASE)
        
        if has_seal or has_date:
            original_text = para.text
            new_text = original_text
            
            # 保存原有格式
            original_alignment = para.alignment
            original_font_name = None
            original_font_size = None
            if para.runs:
                original_font_name = para.runs[0].font.name
                original_font_size = para.runs[0].font.size
            
            if has_date and seal_date_en:
                new_text = re.sub(r'\[date\s*\]', seal_date_en, new_text, flags=re.IGNORECASE)
                print(f"填充日期(段落): [Date] -> {seal_date_en}")
                replacements_made += 1
            
            seal_placeholder = "___SEAL_PLACEHOLDER___"
            if has_seal and seal_org_en:
                new_text = re.sub(r'\[seal\s*\]', seal_placeholder, new_text, flags=re.IGNORECASE)
            
            para.clear()
            
            if seal_placeholder in new_text:
                parts = new_text.split(seal_placeholder)
                for idx, part in enumerate(parts):
                    if part:
                        run = para.add_run(part)
                        if original_font_name:
                            run.font.name = original_font_name
                        if original_font_size:
                            run.font.size = original_font_size
                    if idx < len(parts) - 1:
                        seal_run = para.add_run(f"(Seal of {seal_org_en})")
                        seal_run.italic = True
                        seal_run.font.color.rgb = RGBColor(255, 0, 0)
                        if original_font_size:
                            seal_run.font.size = original_font_size
                print(f"填充印章(段落): [Seal] -> (Seal of {seal_org_en})")
                replacements_made += 1
            else:
                run = para.add_run(new_text)
                if original_font_name:
                    run.font.name = original_font_name
                if original_font_size:
                    run.font.size = original_font_size
            
            # 恢复对齐方式
            para.alignment = original_alignment
    
    # 处理 [CreditCode]、[QRCode]、[QRText] 占位符
    
    # 辅助函数：处理段落中的占位符
    def process_paragraph_placeholders(para, credit_code_value, qr_text_en, registration_no_value, duplicate_number_exists, duplicate_number_en):
        """处理段落中的 [CreditCode]、[QRText]、[No.]、[Duplicate number] 占位符"""
        nonlocal replacements_made
        para_text = para.text
        
        # 匹配各种变体: [CreditCode], [Credit Code], [credit code] 等
        has_credit_code = re.search(r'\[credit\s*code\s*\]', para_text, re.IGNORECASE)
        has_qr_text = re.search(r'\[qr\s*text\s*\]', para_text, re.IGNORECASE)
        # 匹配 [No.] 或 [No] 占位符
        has_no = re.search(r'\[no\.?\s*\]', para_text, re.IGNORECASE)
        # 匹配 [Duplicate number] 或 [Duplicate Number] 占位符
        has_duplicate_number = re.search(r'\[duplicate\s*number\s*\]', para_text, re.IGNORECASE)
        
        if not (has_credit_code or has_qr_text or has_no or has_duplicate_number):
            return False
        
        print(f"找到新占位符段落: '{para_text}'")
        
        # 保存原有格式
        original_alignment = para.alignment
        original_font_name = None
        original_font_size = None
        if para.runs:
            original_font_name = para.runs[0].font.name
            original_font_size = para.runs[0].font.size
        
        new_text = para_text
        
        # 处理 [Credit Code] / [CreditCode]
        if has_credit_code and credit_code_value:
            new_text = re.sub(r'\[credit\s*code\s*\]', credit_code_value, new_text, flags=re.IGNORECASE)
            print(f"填充信用代码: [Credit Code] -> {credit_code_value}")
            replacements_made += 1
        
        # 处理 [QRText] / [QR Text] - 如果有值则替换，没有则删除占位符
        if has_qr_text:
            if qr_text_en:
                new_text = re.sub(r'\[qr\s*text\s*\]', qr_text_en, new_text, flags=re.IGNORECASE)
                print(f"填充二维码说明: [QRText] -> {qr_text_en[:50]}...")
                replacements_made += 1
            else:
                # 没有二维码说明文字，删除占位符（包括可能的前后空格）
                new_text = re.sub(r'\s*\[qr\s*text\s*\]\s*', '', new_text, flags=re.IGNORECASE)
                print(f"删除空的二维码说明占位符: [QRText]")
        
        # 处理 [No.] / [No] - 如果有值则替换，没有则删除占位符
        if has_no:
            if registration_no_value:
                # 输出格式: No. XXXXXXXX
                no_text = f"No. {registration_no_value}"
                new_text = re.sub(r'\[no\.?\s*\]', no_text, new_text, flags=re.IGNORECASE)
                print(f"填充注册号: [No.] -> {no_text}")
                replacements_made += 1
            else:
                # 没有注册号，删除占位符（包括可能的前后空格）
                new_text = re.sub(r'\s*\[no\.?\s*\]\s*', '', new_text, flags=re.IGNORECASE)
                print(f"删除空的注册号占位符: [No.]")
        
        # 处理 [Duplicate number] - 如果有副本号则替换，没有则删除占位符
        if has_duplicate_number:
            if duplicate_number_exists and duplicate_number_en:
                # 输出格式: (Duplicate number: 1-1)
                dup_text = f"(Duplicate number: {duplicate_number_en})"
                new_text = re.sub(r'\[duplicate\s*number\s*\]', dup_text, new_text, flags=re.IGNORECASE)
                print(f"填充副本号: [Duplicate number] -> {dup_text}")
                replacements_made += 1
            else:
                # 没有副本号，删除占位符（包括可能的前后空格）
                new_text = re.sub(r'\s*\[duplicate\s*number\s*\]\s*', '', new_text, flags=re.IGNORECASE)
                print(f"删除空的副本号占位符: [Duplicate number]")
        
        # 没有 [QRCode]，直接替换文本
        para.clear()
        run = para.add_run(new_text)
        if original_font_name:
            run.font.name = original_font_name
        if original_font_size:
            run.font.size = original_font_size
        para.alignment = original_alignment
        
        return True
    
    print(f"\n开始处理新占位符...")
    print(f"  credit_code_value: '{credit_code_value}'")
    print(f"  qr_text_en: '{qr_text_en[:80]}...'" if len(qr_text_en) > 80 else f"  qr_text_en: '{qr_text_en}'")
    
    # 处理表格中的新占位符
    print("\n扫描表格中的占位符...")
    for table_idx, table in enumerate(doc.tables):
        print(f"  表格 {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    # 打印所有非空单元格
                    display_text = cell_text[:60] + "..." if len(cell_text) > 60 else cell_text
                    display_text = display_text.replace('\n', ' ')
                    if '[' in cell_text and ']' in cell_text:
                        print(f"    [{row_idx},{cell_idx}] 占位符: '{display_text}'")
                for para in cell.paragraphs:
                    process_paragraph_placeholders(para, credit_code_value, qr_text_en, registration_no_value, duplicate_number_exists, duplicate_number_en)
    
    # 处理文档段落中的新占位符
    print("\n扫描文档段落中的占位符...")
    for para in doc.paragraphs:
        if para.text.strip():
            # 打印所有非空段落
            print(f"  段落内容: '{para.text[:100]}'" if len(para.text) > 100 else f"  段落内容: '{para.text}'")
            if '[' in para.text and ']' in para.text:
                print(f"    ^ 发现可能的占位符!")
        process_paragraph_placeholders(para, credit_code_value, qr_text_en, registration_no_value, duplicate_number_exists, duplicate_number_en)
    
    # 处理文本框中的占位符（[QRText] 通常在文本框中）
    print("\n扫描文本框中的占位符...")
    from docx.oxml.ns import qn as docx_qn
    
    # 遍历文档中的所有文本框
    for element in doc._element.iter():
        if element.tag.endswith('txbxContent'):
            # 找到文本框内容
            for p_elem in element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                # 获取段落文本
                text_parts = []
                for t_elem in p_elem.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    if t_elem.text:
                        text_parts.append(t_elem.text)
                para_text = ''.join(text_parts)
                
                if not para_text.strip():
                    continue
                
                print(f"  文本框内容: '{para_text}'")
                
                # 检查并替换占位符
                has_qr_text = re.search(r'\[qr\s*text\s*\]', para_text, re.IGNORECASE)
                has_credit_code = re.search(r'\[credit\s*code\s*\]', para_text, re.IGNORECASE)
                has_no = re.search(r'\[no\.?\s*\]', para_text, re.IGNORECASE)
                has_duplicate_number = re.search(r'\[duplicate\s*number\s*\]', para_text, re.IGNORECASE)
                
                if has_credit_code and credit_code_value:
                    # 替换文本框中的 [Credit Code]
                    for t_elem in p_elem.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t_elem.text and re.search(r'\[credit\s*code\s*\]', t_elem.text, re.IGNORECASE):
                            t_elem.text = re.sub(r'\[credit\s*code\s*\]', credit_code_value, t_elem.text, flags=re.IGNORECASE)
                            print(f"  填充文本框信用代码: [Credit Code] -> {credit_code_value}")
                            replacements_made += 1
                
                if has_no:
                    # 替换文本框中的 [No.] - 如果有值则替换，没有则删除
                    for t_elem in p_elem.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t_elem.text and re.search(r'\[no\.?\s*\]', t_elem.text, re.IGNORECASE):
                            if registration_no_value:
                                # 输出格式: No. XXXXXXXX
                                no_text = f"No. {registration_no_value}"
                                t_elem.text = re.sub(r'\[no\.?\s*\]', no_text, t_elem.text, flags=re.IGNORECASE)
                                print(f"  填充文本框注册号: [No.] -> {no_text}")
                                replacements_made += 1
                            else:
                                # 没有注册号，删除占位符
                                t_elem.text = re.sub(r'\s*\[no\.?\s*\]\s*', '', t_elem.text, flags=re.IGNORECASE)
                                print(f"  删除文本框中空的注册号占位符: [No.]")
                
                if has_duplicate_number:
                    # 替换文本框中的 [Duplicate number] - 如果有值则替换，没有则删除
                    for t_elem in p_elem.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t_elem.text and re.search(r'\[duplicate\s*number\s*\]', t_elem.text, re.IGNORECASE):
                            if duplicate_number_exists and duplicate_number_en:
                                # 输出格式: (Duplicate number: 1-1)
                                dup_text = f"(Duplicate number: {duplicate_number_en})"
                                t_elem.text = re.sub(r'\[duplicate\s*number\s*\]', dup_text, t_elem.text, flags=re.IGNORECASE)
                                print(f"  填充文本框副本号: [Duplicate number] -> {dup_text}")
                                replacements_made += 1
                            else:
                                # 没有副本号，删除占位符
                                t_elem.text = re.sub(r'\s*\[duplicate\s*number\s*\]\s*', '', t_elem.text, flags=re.IGNORECASE)
                                print(f"  删除文本框中空的副本号占位符: [Duplicate number]")
                
                if has_qr_text:
                    # 替换文本框中的 [QRText]
                    for t_elem in p_elem.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t_elem.text and re.search(r'\[qr\s*text\s*\]', t_elem.text, re.IGNORECASE):
                            if qr_text_en:
                                t_elem.text = re.sub(r'\[qr\s*text\s*\]', qr_text_en, t_elem.text, flags=re.IGNORECASE)
                                print(f"  填充文本框二维码说明: [QRText] -> {qr_text_en[:50]}...")
                                replacements_made += 1
                            else:
                                # 没有二维码说明文字，删除占位符
                                t_elem.text = re.sub(r'\s*\[qr\s*text\s*\]\s*', '', t_elem.text, flags=re.IGNORECASE)
                                print(f"  删除文本框中空的二维码说明占位符: [QRText]")
    
    # 遍历模板中的表格，动态调整字段
    # 根据 mapped_layout 中的左右栏分配来填充
    # 如果 USE_SOURCE_IMAGE_LAYOUT=True，会按原图的位置来排版
    
    filled_fields = set()
    
    # 确定哪些标准字段存在
    all_standard_fields = set(TEMPLATE_LAYOUT["left_column"] + 
                               TEMPLATE_LAYOUT["right_column"] + 
                               TEMPLATE_LAYOUT["bottom"])
    present_fields = set(field_values.keys())
    missing_fields = all_standard_fields - present_fields
    
    # 从 mapped_layout 获取左右栏字段列表
    left_column_fields = mapped_layout.get("left_column", [])
    right_column_fields = mapped_layout.get("right_column", [])
    bottom_fields = mapped_layout.get("bottom", [])
    
    # 创建字段到位置的映射（用于快速查找）
    field_to_column = {}
    for label, value, display in left_column_fields:
        field_to_column[label] = ("left", value, display)
    for label, value, display in right_column_fields:
        field_to_column[label] = ("right", value, display)
    for label, value, display in bottom_fields:
        field_to_column[label] = ("bottom", value, display)
    
    print(f"\n开始处理模板表格...")
    print(f"  存在的字段: {present_fields}")
    print(f"  缺失的字段: {missing_fields}")
    print(f"  左栏字段: {[f[0] for f in left_column_fields]}")
    print(f"  右栏字段: {[f[0] for f in right_column_fields]}")
    print(f"  底部字段: {[f[0] for f in bottom_fields]}")
    
    # 辅助函数：更新单元格内容并保留格式
    def update_cell_content(cell, new_text, is_label=False):
        """更新单元格内容，保留原有格式"""
        if cell.paragraphs and cell.paragraphs[0].runs:
            original_para = cell.paragraphs[0]
            original_run = original_para.runs[0]
            original_font = original_run.font
            original_alignment = original_para.alignment
            
            original_para.clear()
            new_run = original_para.add_run(new_text)
            
            if original_font.name:
                new_run.font.name = original_font.name
            if original_font.size:
                new_run.font.size = original_font.size
            if original_font.bold is not None:
                new_run.font.bold = original_font.bold
            if original_font.italic is not None:
                new_run.font.italic = original_font.italic
            
            original_para.alignment = original_alignment
        else:
            cell.text = new_text
    
    # 遍历表格，按行处理
    # 模板结构：每行有4个单元格 [左标签, 左值, 右标签, 右值]
    # 
    # 新策略：
    # 1. 先收集模板中所有的标签位置
    # 2. 根据 mapped_layout 中的字段，按行填充
    # 3. 对于额外字段，找空位填充
    
    # 记录每行的单元格信息
    row_cells_info = []  # [(row_idx, left_label_cell, left_value_cell, right_label_cell, right_value_cell, left_matched, right_matched), ...]
    
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            # 去重单元格（Word表格中合并的单元格会重复出现）
            seen_cells = []
            for cell in row.cells:
                if cell not in seen_cells:
                    seen_cells.append(cell)
            
            # 检查是否是标准的4列布局（左标签-左值-右标签-右值）
            if len(seen_cells) >= 4:
                left_label_cell = seen_cells[0]
                left_value_cell = seen_cells[1]
                right_label_cell = seen_cells[2]
                right_value_cell = seen_cells[3]
                
                # 获取当前行的左右标签
                left_label_text = left_label_cell.text.strip().rstrip(':').rstrip('：').strip().lower()
                right_label_text = right_label_cell.text.strip().rstrip(':').rstrip('：').strip().lower()
                
                # 查找匹配的主标签
                left_matched = None
                right_matched = None
                
                for main_label, aliases in TEMPLATE_LABEL_ALIASES.items():
                    if left_label_text == main_label or left_label_text in aliases:
                        left_matched = main_label
                    if right_label_text == main_label or right_label_text in aliases:
                        right_matched = main_label
                
                row_cells_info.append({
                    'row_idx': row_idx,
                    'left_label_cell': left_label_cell,
                    'left_value_cell': left_value_cell,
                    'right_label_cell': right_label_cell,
                    'right_value_cell': right_value_cell,
                    'left_matched': left_matched,
                    'right_matched': right_matched,
                    'left_label_text': left_label_text,
                    'right_label_text': right_label_text,
                    'cols': 4
                })
                
                print(f"\n  行 {row_idx}: 左标签='{left_label_text}'({left_matched}), 右标签='{right_label_text}'({right_matched})")
            
            elif len(seen_cells) >= 2:
                # 2列布局（标签-值），用于底部字段等
                label_cell = seen_cells[0]
                value_cell = seen_cells[1]
                
                cell_text = label_cell.text.strip().rstrip(':').rstrip('：').strip().lower()
                
                matched_main_label = None
                for main_label, aliases in TEMPLATE_LABEL_ALIASES.items():
                    if cell_text == main_label or cell_text in aliases:
                        matched_main_label = main_label
                        break
                
                row_cells_info.append({
                    'row_idx': row_idx,
                    'label_cell': label_cell,
                    'value_cell': value_cell,
                    'matched': matched_main_label,
                    'label_text': cell_text,
                    'cols': 2
                })
                
                print(f"  行 {row_idx} (2列): 标签='{cell_text}'({matched_main_label})")
    
    # 第一遍：填充标准字段（按原图位置）
    print(f"\n  === 第一遍：填充标准字段 ===")
    
    for row_info in row_cells_info:
        if row_info['cols'] == 4:
            left_matched = row_info['left_matched']
            right_matched = row_info['right_matched']
            left_label_text = row_info.get('left_label_text', '')
            right_label_text = row_info.get('right_label_text', '')
            row_processed = False  # 标记这一行是否已经完全处理（交换情况）
            
            # 处理未匹配的模板标签（如 Title: 等不需要的标签）
            # 如果左侧标签不为空但未匹配到任何字段，清空它
            if left_label_text and not left_matched:
                update_cell_content(row_info['left_label_cell'], "")
                update_cell_content(row_info['left_value_cell'], "")
                print(f"    左侧清空未匹配标签: '{left_label_text}'")
            
            # 如果右侧标签不为空但未匹配到任何字段，清空它
            if right_label_text and not right_matched:
                update_cell_content(row_info['right_label_cell'], "")
                update_cell_content(row_info['right_value_cell'], "")
                print(f"    右侧清空未匹配标签: '{right_label_text}'")
            
            # 处理左侧
            if left_matched and left_matched not in filled_fields:
                if left_matched in field_to_column:
                    column, value, display = field_to_column[left_matched]
                    if column == "left":
                        # 字段应该在左边，直接填充（同时更新标签）
                        update_cell_content(row_info['left_label_cell'], f"{display}:")
                        update_cell_content(row_info['left_value_cell'], value)
                        filled_fields.add(left_matched)
                        print(f"    左侧填充: {display} = {value[:30]}...")
                    elif column == "right":
                        # 字段应该在右边，检查是否需要交换
                        if right_matched and right_matched in field_to_column and right_matched not in filled_fields:
                            r_column, r_value, r_display = field_to_column[right_matched]
                            if r_column == "left":
                                # 需要交换
                                update_cell_content(row_info['left_label_cell'], f"{r_display}:")
                                update_cell_content(row_info['left_value_cell'], r_value)
                                filled_fields.add(right_matched)
                                print(f"    左侧交换填充: {r_display} = {r_value[:30]}...")
                                
                                update_cell_content(row_info['right_label_cell'], f"{display}:")
                                update_cell_content(row_info['right_value_cell'], value)
                                filled_fields.add(left_matched)
                                print(f"    右侧交换填充: {display} = {value[:30]}...")
                                row_processed = True  # 整行已处理
                        
                        if not row_processed:
                            # 检查右边是否已经被填充（不要覆盖已填充的内容）
                            if right_matched and right_matched in filled_fields:
                                # 右边已经填充了，清空左边，不移动
                                update_cell_content(row_info['left_label_cell'], "")
                                update_cell_content(row_info['left_value_cell'], "")
                                print(f"    左侧清空: {left_matched} 应在右边但右边已被占用")
                            else:
                                # 不需要交换，移动到右边
                                update_cell_content(row_info['left_label_cell'], "")
                                update_cell_content(row_info['left_value_cell'], "")
                                update_cell_content(row_info['right_label_cell'], f"{display}:")
                                update_cell_content(row_info['right_value_cell'], value)
                                filled_fields.add(left_matched)
                                print(f"    移动到右侧: {display} = {value[:30]}...")
                            row_processed = True
                else:
                    # 字段不存在，清空
                    update_cell_content(row_info['left_label_cell'], "")
                    update_cell_content(row_info['left_value_cell'], "")
                    print(f"    左侧清空: {left_matched} 不存在")
            
            # 处理右侧（如果这一行还没完全处理）
            if not row_processed and right_matched and right_matched not in filled_fields:
                if right_matched in field_to_column:
                    column, value, display = field_to_column[right_matched]
                    if column == "right":
                        # 字段应该在右边，直接填充（同时更新标签）
                        update_cell_content(row_info['right_label_cell'], f"{display}:")
                        update_cell_content(row_info['right_value_cell'], value)
                        filled_fields.add(right_matched)
                        print(f"    右侧填充: {display} = {value[:30]}...")
                    elif column == "left":
                        # 检查左边是否空
                        if not row_info['left_label_cell'].text.strip():
                            update_cell_content(row_info['left_label_cell'], f"{display}:")
                            update_cell_content(row_info['left_value_cell'], value)
                            update_cell_content(row_info['right_label_cell'], "")
                            update_cell_content(row_info['right_value_cell'], "")
                            filled_fields.add(right_matched)
                            print(f"    移动到左侧: {display} = {value[:30]}...")
                        else:
                            # 保持在右侧（同时更新标签）
                            update_cell_content(row_info['right_label_cell'], f"{display}:")
                            update_cell_content(row_info['right_value_cell'], value)
                            filled_fields.add(right_matched)
                            print(f"    保持右侧: {display} = {value[:30]}...")
                else:
                    update_cell_content(row_info['right_label_cell'], "")
                    update_cell_content(row_info['right_value_cell'], "")
                    print(f"    右侧清空: {right_matched} 不存在")
        
        elif row_info['cols'] == 2:
            matched = row_info['matched']
            if matched and matched not in filled_fields:
                if matched in field_to_column:
                    column, value, display = field_to_column[matched]
                    # 同时更新标签和值
                    update_cell_content(row_info['label_cell'], f"{display}:")
                    update_cell_content(row_info['value_cell'], value)
                    filled_fields.add(matched)
                    print(f"    2列填充: {display} = {value[:30]}...")
                else:
                    update_cell_content(row_info['label_cell'], "")
                    update_cell_content(row_info['value_cell'], "")
                    print(f"    2列清空: {matched} 不存在")
    
    # 第二遍：填充未填充的字段到空位
    # 根据字段在原图中的位置，优先填充到对应的位置
    # 左侧字段只填充到左侧空位，右侧字段只填充到右侧空位
    print(f"\n  === 第二遍：填充剩余字段到空位 ===")
    
    unfilled = [label for label in field_to_column.keys() if label not in filled_fields]
    print(f"  未填充的字段: {unfilled}")
    
    # 分离左侧字段和右侧字段
    unfilled_left = [label for label in unfilled if field_to_column[label][0] == "left"]
    unfilled_right = [label for label in unfilled if field_to_column[label][0] == "right"]
    
    print(f"  未填充的左侧字段: {unfilled_left}")
    print(f"  未填充的右侧字段: {unfilled_right}")
    
    # 第一轮：严格按位置填充（左侧字段填左侧，右侧字段填右侧）
    for row_info in row_cells_info:
        if not unfilled_left and not unfilled_right:
            break
        
        if row_info['cols'] == 4:
            # 左侧空位只填充左侧字段
            if not row_info['left_label_cell'].text.strip() and unfilled_left:
                label = unfilled_left.pop(0)
                column, value, display = field_to_column[label]
                update_cell_content(row_info['left_label_cell'], f"{display}:")
                update_cell_content(row_info['left_value_cell'], value)
                filled_fields.add(label)
                print(f"    填充到空左侧: {display} = {value[:30]}...")
            
            # 右侧空位只填充右侧字段
            if not row_info['right_label_cell'].text.strip() and unfilled_right:
                label = unfilled_right.pop(0)
                column, value, display = field_to_column[label]
                update_cell_content(row_info['right_label_cell'], f"{display}:")
                update_cell_content(row_info['right_value_cell'], value)
                filled_fields.add(label)
                print(f"    填充到空右侧: {display} = {value[:30]}...")
        
        elif row_info['cols'] == 2:
            if not row_info['label_cell'].text.strip():
                # 2列布局，优先填充左侧字段
                if unfilled_left:
                    label = unfilled_left.pop(0)
                    column, value, display = field_to_column[label]
                    update_cell_content(row_info['label_cell'], f"{display}:")
                    update_cell_content(row_info['value_cell'], value)
                    filled_fields.add(label)
                    print(f"    填充到空2列: {display} = {value[:30]}...")
    
    # 第二轮：如果还有未填充的字段，交叉填充（左侧字段填右侧空位，或右侧字段填左侧空位）
    if unfilled_left or unfilled_right:
        print(f"  === 第二轮：交叉填充剩余字段 ===")
        print(f"  剩余左侧字段: {unfilled_left}")
        print(f"  剩余右侧字段: {unfilled_right}")
        
        for row_info in row_cells_info:
            if not unfilled_left and not unfilled_right:
                break
            
            if row_info['cols'] == 4:
                # 左侧空位填充右侧字段
                if not row_info['left_label_cell'].text.strip() and unfilled_right:
                    label = unfilled_right.pop(0)
                    column, value, display = field_to_column[label]
                    update_cell_content(row_info['left_label_cell'], f"{display}:")
                    update_cell_content(row_info['left_value_cell'], value)
                    filled_fields.add(label)
                    print(f"    填充到空左侧(右侧字段): {display} = {value[:30]}...")
                
                # 右侧空位填充左侧字段
                if not row_info['right_label_cell'].text.strip() and unfilled_left:
                    label = unfilled_left.pop(0)
                    column, value, display = field_to_column[label]
                    update_cell_content(row_info['right_label_cell'], f"{display}:")
                    update_cell_content(row_info['right_value_cell'], value)
                    filled_fields.add(label)
                    print(f"    填充到空右侧(左侧字段): {display} = {value[:30]}...")
            
            elif row_info['cols'] == 2:
                if not row_info['label_cell'].text.strip() and unfilled_right:
                    label = unfilled_right.pop(0)
                    column, value, display = field_to_column[label]
                    update_cell_content(row_info['label_cell'], f"{display}:")
                    update_cell_content(row_info['value_cell'], value)
                    filled_fields.add(label)
                    print(f"    填充到空2列(右侧字段): {display} = {value[:30]}...")
    
    # 检查是否有未填充的字段
    unfilled_fields = set(field_values.keys()) - filled_fields
    if unfilled_fields:
        print(f"\n  警告: 以下字段未能填充到模板: {unfilled_fields}")
    
    # 更新填充计数
    replacements_made += len(filled_fields)
    
    print(f"\n模板中填充了 {replacements_made} 个字段")
    
    doc.save(output_path)
    print(f"\n文档已保存到: {output_path}")
    print(f"共填充 {replacements_made} 处")
    return replacements_made


def select_template(data, image_path: str = None):
    """根据图片尺寸和识别到的字段选择合适的模板
    
    首先根据图片尺寸判断横版/竖版：
    - 竖版：使用竖版模板
    - 横版：根据字段判断使用默认模板还是八字段模板
    
    Args:
        data: 识别到的数据
        image_path: 图片路径，用于判断横竖版
    """
    # 首先判断横竖版
    if image_path:
        orientation = get_image_orientation(image_path)
        if orientation == "vertical":
            print(f"竖版营业执照，使用竖版模板: {TEMPLATE_PATH_VERTICAL}")
            return TEMPLATE_PATH_VERTICAL
    
    # 横版：根据字段选择模板
    fields = data.get("fields", [])
    
    # 检查是否有经营期限字段
    has_term_of_operation = False
    for field in fields:
        label_cn = field.get("label_cn", "").strip().replace(" ", "").replace("　", "")
        if label_cn in ["经营期限", "营业期限"]:
            has_term_of_operation = True
            break
    
    if has_term_of_operation:
        print(f"检测到经营期限字段，使用八字段模板: {TEMPLATE_PATH_8}")
        return TEMPLATE_PATH_8
    else:
        print(f"使用默认模板: {TEMPLATE_PATH}")
        return TEMPLATE_PATH


def main():
    """主函数"""
    
    print("=" * 60)
    print("Gemini 文档识别翻译填充工具")
    print("=" * 60)
    
    # 处理命令行参数
    if len(sys.argv) > 1:
        input_image = sys.argv[1]
        if not Path(input_image).is_absolute():
            input_image = str(SCRIPT_DIR / input_image)
    else:
        input_image = INPUT_IMAGE
    
    print(f"输入图片: {input_image}")
    print(f"默认模板: {TEMPLATE_PATH}")
    print(f"八字段模板: {TEMPLATE_PATH_8}")
    print(f"竖版模板: {TEMPLATE_PATH_VERTICAL}")
    print("=" * 60)

    if not input_image:
        print("错误: 未提供输入图片路径。请通过命令行传入，例如: python start2.py your_image.jpg")
        sys.exit(1)

    # 检查文件
    if not Path(input_image).exists():
        print(f"错误: 输入图片不存在: {input_image}")
        sys.exit(1)

    # 生成输出文件路径
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    input_path = Path(input_image)
    output_filename = f"{input_path.stem}_translated_{timestamp}.docx"
    output_path = str(Path(OUTPUT_DIR) / output_filename)

    # 确保输出目录存在
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
    print(f"输出文件: {output_path}")

    # 检查默认模板是否存在（八字段模板在选择时检查）
    if not Path(TEMPLATE_PATH).exists():
        print(f"错误: 默认模板文件不存在: {TEMPLATE_PATH}")
        sys.exit(1)
    
    # 识别并翻译
    data = extract_and_translate(input_image)
    
    # 查找公司名称字段，弹窗让用户确认翻译
    business_name_cn = ""
    business_name_en = ""
    business_name_field_idx = -1
    
    for idx, field in enumerate(data.get("fields", [])):
        label_cn = field.get("label_cn", "").strip().replace(" ", "").replace("　", "")
        if label_cn in ["名称", "名　称", "公司名称", "企业名称"]:
            business_name_cn = field.get("value_cn", "").strip()
            business_name_en = field.get("value_en", "").strip()
            business_name_field_idx = idx
            break
    
    # 如果找到公司名称，弹窗让用户确认
    if business_name_cn and business_name_en and business_name_field_idx >= 0:
        print("\n" + "=" * 60)
        print("公司名称翻译确认")
        print("=" * 60)
        print(f"中文名称: {business_name_cn}")
        print(f"AI翻译: {business_name_en}")
        
        # 显示弹窗
        confirmed_name = show_company_name_dialog(business_name_en, business_name_cn)
        
        # 更新数据
        if confirmed_name != business_name_en:
            print(f"用户选择手动输入: {confirmed_name}")
            data["fields"][business_name_field_idx]["value_en"] = confirmed_name
        else:
            print(f"用户确认使用AI翻译: {confirmed_name}")
        print()
    
    # 打印识别结果
    print("\n" + "=" * 60)
    print("识别到的字段:")
    print("=" * 60)
    for field in data.get("fields", []):
        print(f"  {field.get('label_cn', '')} ({field.get('label_en', '')})")
        print(f"    中文值: {field.get('value_cn', '')}")
        print(f"    英文值: {field.get('value_en', '')}")
        print()
    
    # 打印印章信息
    seal_text = data.get("seal_text", {})
    if seal_text:
        print("=" * 60)
        print("印章信息:")
        print("=" * 60)
        print(f"  机构名称(中文): {seal_text.get('organization_cn', '')}")
        print(f"  机构名称(英文): {seal_text.get('organization_en', '')}")
        print(f"  印章日期(中文): {seal_text.get('date_cn', '')}")
        print(f"  印章日期(英文): {seal_text.get('date_en', '')}")
        print()
    
    # 打印信用代码信息
    credit_code = data.get("credit_code", {})
    if credit_code:
        print("=" * 60)
        print("信用代码信息:")
        print("=" * 60)
        print(f"  代码: {credit_code.get('code', '')}")
        print(f"  完整英文: {credit_code.get('full_text_en', '')}")
        print()
    
    # 打印二维码说明文字
    qr_text = data.get("qr_text", {})
    if qr_text:
        print("=" * 60)
        print("二维码说明文字:")
        print("=" * 60)
        print(f"  中文: {qr_text.get('text_cn', '')}")
        print(f"  英文: {qr_text.get('text_en', '')}")
        print()
    
    # 打印副本编号信息
    duplicate_number = data.get("duplicate_number", {})
    if duplicate_number:
        print("=" * 60)
        print("副本编号信息:")
        print("=" * 60)
        print(f"  存在副本号: {duplicate_number.get('exists', False)}")
        print(f"  中文: {duplicate_number.get('number_cn', '')}")
        print(f"  英文: {duplicate_number.get('number_en', '')}")
        print()
    
    # 根据识别结果选择合适的模板
    selected_template = select_template(data, input_image)
    
    # 检查选中的模板是否存在
    if not Path(selected_template).exists():
        print(f"警告: 选中的模板不存在: {selected_template}")
        print(f"回退到默认模板: {TEMPLATE_PATH}")
        selected_template = TEMPLATE_PATH
    
    # 根据模板类型选择填充函数
    if selected_template == TEMPLATE_PATH_VERTICAL:
        # 竖版模板使用专门的填充函数
        fill_vertical_template(selected_template, output_path, data, input_image)
    else:
        # 横版模板使用原有的填充函数
        fill_template(selected_template, output_path, data, input_image)
    
    print("\n" + "=" * 60)
    print("处理完成!")
    print("=" * 60)


if __name__ == '__main__':
    main()
