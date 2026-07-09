import zipfile
import sys
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.service.chat_preserve_docx_service import render_chat_preserve_docx_from_layout


def test_render_chat_preserve_docx_keeps_editable_text_and_images(tmp_path):
    image_path = tmp_path / "chat.png"
    image = Image.new("RGB", (420, 260), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((24, 36, 72, 84), fill=(46, 92, 180))
    draw.ellipse((260, 132, 292, 164), fill=(255, 196, 0))
    draw.text((88, 42), "Fabian", fill=(120, 120, 120))
    draw.text((88, 66), "Hello with sticker", fill=(0, 0, 0))
    image.save(image_path)

    layout = {
        "messages": [
            {
                "sender": "Fabian",
                "time": "08-16 12:28",
                "text": "Hello with sticker",
                "message_bbox": [18, 28, 380, 176],
                "avatar_bbox": [24, 36, 72, 84],
                "visuals": [
                    {"type": "emoji", "bbox": [260, 132, 292, 164], "alt": ""}
                ],
            }
        ]
    }

    output_docx = tmp_path / "chat.docx"
    result = render_chat_preserve_docx_from_layout(
        image_path=image_path,
        layout=layout,
        output_docx_path=output_docx,
        assets_dir=tmp_path / "assets",
    )

    assert output_docx.exists()
    assert result.asset_count == 2
    assert "Hello with sticker" in result.raw_text
    assert result.layout["render"]["debug_overlays"]
    assert Path(result.layout["render"]["debug_overlays"][0]).exists()

    document = Document(output_docx)
    cell_text = "\n".join(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    assert "Fabian" in cell_text
    assert "Hello with sticker" in cell_text

    with zipfile.ZipFile(output_docx) as docx_zip:
        media_files = [name for name in docx_zip.namelist() if name.startswith("word/media/")]
    assert len(media_files) >= 2


def test_render_chat_preserve_refines_sloppy_asset_boxes(tmp_path):
    image_path = tmp_path / "chat_sloppy.png"
    image = Image.new("RGB", (360, 220), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((30, 40, 78, 88), fill=(20, 90, 180))
    draw.rectangle((230, 124, 270, 164), fill=(255, 210, 0))
    draw.text((94, 46), "Pepe", fill=(120, 120, 120))
    draw.text((94, 70), "Sticker below", fill=(0, 0, 0))
    image.save(image_path)

    layout = {
        "messages": [
            {
                "sender": "Pepe",
                "time": "08-16 12:30",
                "text": "Sticker below",
                "message_bbox": [18, 30, 330, 178],
                "avatar_bbox": [20, 30, 90, 100],
                "visuals": [
                    {"type": "sticker", "bbox": [212, 110, 286, 180], "alt": ""}
                ],
            }
        ]
    }

    result = render_chat_preserve_docx_from_layout(
        image_path=image_path,
        layout=layout,
        output_docx_path=tmp_path / "chat_sloppy.docx",
        assets_dir=tmp_path / "assets_sloppy",
    )

    message = result.layout["pages"][0]["layout"]["messages"][0]
    avatar_refined = message["avatar_bbox_refined"]
    visual_refined = message["visuals"][0]["bbox_refined"]

    assert avatar_refined != [20, 30, 90, 100]
    assert avatar_refined[0] >= 24
    assert avatar_refined[1] >= 34
    assert avatar_refined[2] <= 84
    assert avatar_refined[3] <= 94
    assert visual_refined != [212, 110, 286, 180]
    assert visual_refined[0] >= 224
    assert visual_refined[1] >= 118
    assert visual_refined[2] <= 276
    assert visual_refined[3] <= 170


def test_render_chat_preserve_uses_local_avatar_candidate_when_model_box_drifts(tmp_path):
    image_path = tmp_path / "chat_drifted_avatar.png"
    image = Image.new("RGB", (480, 1000), (245, 245, 245))
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 480, 96), fill=(255, 255, 255))
    draw.line((36, 42, 18, 58, 36, 74), fill=(20, 20, 20), width=4)
    draw.rectangle((0, 96, 480, 1000), fill=(58, 191, 205))
    draw.rounded_rectangle((112, 196, 430, 326), radius=8, fill=(255, 255, 255))

    for y in range(210, 274):
        for x in range(34, 98):
            value = (x * 17 + y * 11) % 180
            image.putpixel((x, y), (45 + value // 3, 60 + value // 2, 90 + value))

    draw.text((112, 160), "Sender-STS", fill=(255, 255, 255))
    draw.text((132, 218), "Morning check", fill=(0, 0, 0))
    image.save(image_path)

    layout = {
        "messages": [
            {
                "sender": "Sender-STS",
                "time": "2023-10-31 07:17",
                "text": "Morning check",
                "message_bbox": [20, 150, 440, 340],
                "avatar_bbox": [18, 24, 78, 84],
                "visuals": [],
            }
        ]
    }

    result = render_chat_preserve_docx_from_layout(
        image_path=image_path,
        layout=layout,
        output_docx_path=tmp_path / "chat_drifted_avatar.docx",
        assets_dir=tmp_path / "assets_drifted_avatar",
    )

    message = result.layout["pages"][0]["layout"]["messages"][0]
    avatar_refined = message["avatar_bbox_refined"]

    assert message["avatar_bbox_source"] == "local_candidate_detection"
    assert message["avatar_bbox_model"] == [18, 24, 78, 84]
    assert message["avatar_bbox_detected"][1] >= 180
    assert avatar_refined[0] >= 20
    assert avatar_refined[1] >= 190
    assert avatar_refined[2] <= 112
    assert avatar_refined[3] <= 288


def test_render_chat_preserve_filters_text_visuals_and_detects_right_avatar(tmp_path):
    image_path = tmp_path / "two_sided_chat.png"
    image = Image.new("RGB", (374, 887), (238, 238, 238))
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 374, 110), fill=(245, 245, 245))

    def avatar_box(box, seed):
        left, top, right, bottom = box
        draw.rounded_rectangle(box, radius=4, fill=(230, 230, 230))
        for y in range(top + 2, bottom - 2):
            for x in range(left + 2, right - 2):
                value = (x * 19 + y * 23 + seed) % 190
                image.putpixel((x, y), (45 + value // 4, 35 + value // 3, 70 + value // 2))

    avatar_box((8, 136, 52, 179), 3)
    avatar_box((8, 567, 52, 610), 11)
    draw.rounded_rectangle((48, 128, 190, 160), radius=4, fill=(255, 255, 255))
    draw.rounded_rectangle((64, 545, 320, 620), radius=4, fill=(255, 255, 255))

    draw.rounded_rectangle((58, 640, 322, 735), radius=4, fill=(142, 217, 93))
    draw.polygon([(300, 661), (322, 650), (322, 759), (300, 759)], fill=(142, 217, 93))
    draw.rounded_rectangle((328, 661, 365, 698), radius=4, fill=(28, 42, 66))
    draw.line((329, 695, 364, 672), fill=(230, 210, 150), width=5)
    draw.text((70, 668), "text area", fill=(20, 80, 30))
    image.save(image_path)

    layout = {
        "messages": [
            {
                "sender": "left",
                "time": "",
                "text": "left message",
                "message_bbox": [48, 128, 190, 160],
                "avatar_bbox": [10, 120, 39, 149],
                "visuals": [],
            },
            {
                "sender": "left",
                "time": "",
                "text": "file card",
                "message_bbox": [64, 545, 320, 620],
                "avatar_bbox": [10, 464, 39, 493],
                "visuals": [],
            },
            {
                "sender": "",
                "time": "",
                "text": "right message",
                "message_bbox": [58, 640, 322, 735],
                "avatar_bbox": [335, 545, 364, 574],
                "visuals": [
                    {"type": "image", "bbox": [264, 669, 308, 713], "alt": ""}
                ],
            },
        ]
    }

    result = render_chat_preserve_docx_from_layout(
        image_path=image_path,
        layout=layout,
        output_docx_path=tmp_path / "two_sided_chat.docx",
        assets_dir=tmp_path / "assets_two_sided",
    )

    messages = result.layout["pages"][0]["layout"]["messages"]

    left_avatar = messages[1]["avatar_bbox_refined"]
    assert left_avatar[0] in {8, 9}
    assert left_avatar[1] in {566, 567}
    assert left_avatar[2] in {52, 53}
    assert left_avatar[3] in {610, 611}
    right_avatar = messages[2]["avatar_bbox_refined"]
    assert right_avatar[0] in {327, 328}
    assert right_avatar[1] == 661
    assert right_avatar[2] in {365, 366}
    assert right_avatar[3] in {698, 699}
    assert messages[2]["visuals"] == []
