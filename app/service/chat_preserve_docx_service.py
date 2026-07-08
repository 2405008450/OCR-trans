import copy
import json
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageSequence

from app.service.gemini_service import generate_vision_html


CHAT_LAYOUT_SYSTEM_PROMPT = """You analyze screenshots of chat software and return strict JSON only.

Return one JSON object and no markdown/code fences/explanations.

Image coordinates:
- Use absolute pixel coordinates relative to the uploaded image.
- Every bbox must be [left, top, right, bottom].
- If a coordinate is uncertain, use null for that bbox instead of guessing.

Goal:
- Extract chat messages in reading order.
- Preserve editable text.
- Identify non-text visuals that must be cropped later: user avatars, image emoji, stickers, and inline image reactions.
- Ignore decorative divider lines and ordinary app chrome/icons unless they are part of a message.

Schema:
{
  "messages": [
    {
      "sender": "visible sender name or empty string",
      "time": "visible timestamp or empty string",
      "text": "message body text, preserving line breaks",
      "message_bbox": [left, top, right, bottom] or null,
      "avatar_bbox": [left, top, right, bottom] or null,
      "visuals": [
        {
          "type": "emoji|sticker|image",
          "bbox": [left, top, right, bottom],
          "alt": "short visible label or empty string"
        }
      ]
    }
  ]
}

Rules:
- Do not invent messages.
- Unicode emoji inside editable text should remain in text and should not be listed in visuals.
- Only list visuals that are actual image-based emoji/stickers/images or avatars.
- For repeated messages by the same sender, repeat the sender if visible; otherwise leave it empty.
"""


CHAT_LOCAL_ASSET_SYSTEM_PROMPT = """You locate avatar and image-based emoji/sticker assets inside one cropped chat message.

Return strict JSON only. Use pixel coordinates relative to the cropped image.

Schema:
{
  "avatar_bbox": [left, top, right, bottom] or null,
  "visuals": [
    {"type": "emoji|sticker|image", "bbox": [left, top, right, bottom], "alt": ""}
  ]
}

Rules:
- Do not include editable text as visuals.
- Unicode emoji that appears as normal text should not be listed.
- Only locate the sender avatar and image-based emoji/sticker/image content.
- If unsure, return null or an empty list instead of guessing.
"""


LOCAL_ASSET_RECHECK_MAX_MESSAGES = 25
AVATAR_CANDIDATE_MAX_PER_PAGE = 80


@dataclass
class ChatPreserveResult:
    raw_text: str
    total_pages: int
    asset_count: int
    fallback_count: int
    layout: dict[str, Any]


@dataclass
class ChatPage:
    page_no: int
    image: Image.Image
    layout: dict[str, Any]


StatusCallback = Callable[[str], None]


def convert_chat_screenshot_to_docx(
    *,
    input_path: str | Path,
    output_docx_path: str | Path,
    layout_json_path: str | Path,
    assets_dir: str | Path,
    model: str,
    gemini_route: str,
    status_callback: StatusCallback | None = None,
) -> ChatPreserveResult:
    """将聊天截图/PDF转换为保留头像和图片表情的可编辑 Word。"""
    input_file = Path(input_path)
    output_docx = Path(output_docx_path)
    layout_json = Path(layout_json_path)
    asset_root = Path(assets_dir)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    layout_json.parent.mkdir(parents=True, exist_ok=True)
    asset_root.mkdir(parents=True, exist_ok=True)

    pages = _load_input_pages(input_file)
    chat_pages: list[ChatPage] = []
    full_layout: dict[str, Any] = {"mode": "chat_preserve", "pages": []}

    for page_index, image in enumerate(pages, start=1):
        _emit(status_callback, f"正在分析聊天截图第 {page_index}/{len(pages)} 页")
        layout = _analyze_chat_page(
            image=image,
            page_no=page_index,
            model=model,
            gemini_route=gemini_route,
        )
        layout = _refine_layout_with_local_asset_pass(
            image=image,
            layout=layout,
            page_no=page_index,
            model=model,
            gemini_route=gemini_route,
            status_callback=status_callback,
        )
        chat_pages.append(ChatPage(page_no=page_index, image=image, layout=layout))

    render_result = render_chat_preserve_docx_from_pages(
        pages=chat_pages,
        output_docx_path=output_docx,
        assets_dir=asset_root,
    )
    full_layout["pages"] = [
        {
            "page_no": page.page_no,
            "image_width": page.image.width,
            "image_height": page.image.height,
            "layout": page.layout,
        }
        for page in chat_pages
    ]
    full_layout["render"] = {
        "asset_count": render_result.asset_count,
        "fallback_count": render_result.fallback_count,
        "debug_overlays": [str(path).replace("\\", "/") for path in render_result.debug_overlay_paths],
    }
    layout_json.write_text(
        json.dumps(full_layout, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    return ChatPreserveResult(
        raw_text=render_result.raw_text,
        total_pages=len(chat_pages),
        asset_count=render_result.asset_count,
        fallback_count=render_result.fallback_count,
        layout=full_layout,
    )


def render_chat_preserve_docx_from_layout(
    *,
    image_path: str | Path,
    layout: dict[str, Any],
    output_docx_path: str | Path,
    assets_dir: str | Path,
) -> ChatPreserveResult:
    """测试和离线调试入口：用已知布局直接生成 DOCX。"""
    image = _open_image_as_rgb(Path(image_path))
    page = ChatPage(page_no=1, image=image, layout=layout)
    render_result = render_chat_preserve_docx_from_pages(
        pages=[page],
        output_docx_path=output_docx_path,
        assets_dir=assets_dir,
    )
    return ChatPreserveResult(
        raw_text=render_result.raw_text,
        total_pages=1,
        asset_count=render_result.asset_count,
        fallback_count=render_result.fallback_count,
        layout={
            "mode": "chat_preserve",
            "pages": [{"page_no": 1, "layout": page.layout}],
            "render": {
                "asset_count": render_result.asset_count,
                "fallback_count": render_result.fallback_count,
                "debug_overlays": [str(path).replace("\\", "/") for path in render_result.debug_overlay_paths],
            },
        },
    )


@dataclass
class _RenderResult:
    raw_text: str
    asset_count: int
    fallback_count: int
    debug_overlay_paths: list[Path]


def render_chat_preserve_docx_from_pages(
    *,
    pages: Iterable[ChatPage],
    output_docx_path: str | Path,
    assets_dir: str | Path,
) -> _RenderResult:
    document = Document()
    _setup_document(document)
    asset_root = Path(assets_dir)
    asset_root.mkdir(parents=True, exist_ok=True)

    raw_sections: list[str] = []
    asset_count = 0
    fallback_count = 0
    debug_overlay_paths: list[Path] = []
    page_list = list(pages)

    for page_index, page in enumerate(page_list, start=1):
        if page_index > 1:
            document.add_page_break()
        page.layout = _refine_layout_assets(page.image, page.layout)
        overlay_path = asset_root / f"debug_overlay_page_{page.page_no:03d}.png"
        _save_debug_overlay(page.image, page.layout, overlay_path)
        debug_overlay_paths.append(overlay_path)
        messages = _normalize_messages(page.layout)
        raw_sections.append(f"--- Page {page.page_no} ---")

        if not messages:
            fallback_path = _crop_to_file(
                image=page.image,
                bbox=(0, 0, page.image.width, page.image.height),
                output_path=asset_root / f"page_{page.page_no:03d}_full.png",
            )
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(fallback_path), width=Inches(6.2))
            asset_count += 1
            fallback_count += 1
            raw_sections.append("")
            continue

        for message_index, message in enumerate(messages, start=1):
            render_state = _render_message(
                document=document,
                image=page.image,
                message=message,
                page_no=page.page_no,
                message_index=message_index,
                assets_dir=asset_root,
            )
            asset_count += render_state["asset_count"]
            fallback_count += render_state["fallback_count"]
            raw_sections.append(_message_to_raw_text(message))

    document.save(str(output_docx_path))
    raw_text = "\n\n".join(section for section in raw_sections if section is not None).strip()
    return _RenderResult(
        raw_text=raw_text,
        asset_count=asset_count,
        fallback_count=fallback_count,
        debug_overlay_paths=debug_overlay_paths,
    )


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)


def _render_message(
    *,
    document: Document,
    image: Image.Image,
    message: dict[str, Any],
    page_no: int,
    message_index: int,
    assets_dir: Path,
) -> dict[str, int]:
    asset_count = 0
    fallback_count = 0
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, "FFFFFF")
    _set_cell_width(table.rows[0].cells[0], Inches(0.55))
    _set_cell_width(table.rows[0].cells[1], Inches(5.45))
    _set_cell_width(table.rows[0].cells[2], Inches(1.05))

    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    avatar_bbox = _coerce_bbox(message.get("avatar_bbox_refined") or message.get("avatar_bbox"), image.size)
    if avatar_bbox:
        avatar_path = _crop_to_file(
            image=image,
            bbox=avatar_bbox,
            output_path=assets_dir / f"p{page_no:03d}_m{message_index:03d}_avatar.png",
            padding=2,
        )
        paragraph = table.rows[0].cells[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(avatar_path), width=Inches(0.38))
        asset_count += 1

    body_cell = table.rows[0].cells[1]
    header_para = body_cell.paragraphs[0]
    header_para.paragraph_format.space_after = Pt(1)
    sender = _clean_text(message.get("sender"))
    if sender:
        sender_run = header_para.add_run(sender)
        sender_run.bold = True
        sender_run.font.size = Pt(9)
        sender_run.font.color.rgb = RGBColor(140, 145, 152)

    text = _clean_text(message.get("text"))
    if text:
        text_para = body_cell.add_paragraph()
        text_para.paragraph_format.space_after = Pt(2)
        _add_multiline_text(text_para, text)

    valid_visuals = 0
    invalid_visuals = 0
    for visual_index, visual in enumerate(_normalize_visuals(message), start=1):
        bbox = _coerce_bbox(visual.get("bbox_refined") or visual.get("bbox"), image.size)
        if not bbox:
            invalid_visuals += 1
            continue
        visual_path = _crop_to_file(
            image=image,
            bbox=bbox,
            output_path=assets_dir / f"p{page_no:03d}_m{message_index:03d}_visual{visual_index:02d}.png",
            padding=1,
        )
        para = body_cell.add_paragraph()
        width_inch = _visual_width_inches(bbox)
        para.add_run().add_picture(str(visual_path), width=Inches(width_inch))
        valid_visuals += 1
        asset_count += 1

    needs_fallback = _needs_message_fallback(message, text, avatar_bbox, valid_visuals, invalid_visuals, image.size)
    if needs_fallback:
        message_bbox = _coerce_bbox(message.get("message_bbox"), image.size)
        if message_bbox:
            fallback_path = _crop_to_file(
                image=image,
                bbox=message_bbox,
                output_path=assets_dir / f"p{page_no:03d}_m{message_index:03d}_fallback.png",
                padding=4,
            )
            para = body_cell.add_paragraph()
            para.paragraph_format.space_before = Pt(2)
            para.add_run().add_picture(str(fallback_path), width=Inches(4.8))
            asset_count += 1
            fallback_count += 1

    time_text = _clean_text(message.get("time"))
    if time_text:
        time_para = table.rows[0].cells[2].paragraphs[0]
        time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        time_run = time_para.add_run(time_text)
        time_run.font.size = Pt(8.5)
        time_run.font.color.rgb = RGBColor(140, 145, 152)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return {"asset_count": asset_count, "fallback_count": fallback_count}


def _needs_message_fallback(
    message: dict[str, Any],
    text: str,
    avatar_bbox: tuple[int, int, int, int] | None,
    valid_visuals: int,
    invalid_visuals: int,
    image_size: tuple[int, int],
) -> bool:
    if invalid_visuals:
        return True
    if message.get("avatar_bbox") and not avatar_bbox:
        return True
    if not text and valid_visuals == 0:
        return _coerce_bbox(message.get("message_bbox"), image_size) is not None
    return False


def _message_to_raw_text(message: dict[str, Any]) -> str:
    pieces = []
    sender = _clean_text(message.get("sender"))
    time_text = _clean_text(message.get("time"))
    text = _clean_text(message.get("text"))
    header = " | ".join(item for item in (sender, time_text) if item)
    if header:
        pieces.append(header)
    if text:
        pieces.append(text)
    return "\n".join(pieces)


def _add_multiline_text(paragraph, text: str) -> None:
    lines = text.splitlines() or [text]
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        paragraph.add_run(line)


def _normalize_messages(layout: dict[str, Any]) -> list[dict[str, Any]]:
    messages = layout.get("messages") if isinstance(layout, dict) else None
    if not isinstance(messages, list):
        return []
    return [item for item in messages if isinstance(item, dict)]


def _normalize_visuals(message: dict[str, Any]) -> list[dict[str, Any]]:
    visuals = message.get("visuals")
    if not isinstance(visuals, list):
        return []
    return [item for item in visuals if isinstance(item, dict)]


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\r\n", "\n").replace("\r", "\n").strip()


def _visual_width_inches(bbox: tuple[int, int, int, int]) -> float:
    width = max(bbox[2] - bbox[0], 1)
    height = max(bbox[3] - bbox[1], 1)
    if width <= 48 and height <= 48:
        return 0.28
    if width <= 96 and height <= 96:
        return 0.55
    return min(max(width / 180, 0.75), 2.35)


def _refine_layout_assets(image: Image.Image, layout: dict[str, Any]) -> dict[str, Any]:
    refined = copy.deepcopy(layout if isinstance(layout, dict) else {})
    messages = _normalize_messages(refined)
    avatar_candidates = _detect_chat_avatar_candidates(image)
    assigned_avatars = _assign_detected_avatar_candidates(
        messages=messages,
        candidates=avatar_candidates,
        image_size=image.size,
    )

    for message in messages:
        message_bbox = _coerce_bbox(message.get("message_bbox"), image.size)
        avatar_bbox = _coerce_bbox(message.get("avatar_bbox"), image.size)
        if avatar_bbox:
            original_avatar_bbox = _coerce_bbox(message.get("avatar_bbox_model"), image.size) or avatar_bbox
            message["avatar_bbox_original"] = _bbox_to_list(original_avatar_bbox)
            if message.get("avatar_bbox_source") == "local_candidate_detection":
                message["avatar_bbox_refined"] = _bbox_to_list(avatar_bbox)
                message["avatar_bbox_refine_method"] = "local_candidate"
            else:
                snapped = _refine_asset_bbox(
                    image=image,
                    bbox=avatar_bbox,
                    kind="avatar",
                    bounds=message_bbox,
                )
                message["avatar_bbox_refined"] = _bbox_to_list(snapped or avatar_bbox)
                message["avatar_bbox_refine_method"] = "edge_snap" if snapped else "original"

        filtered_visuals = []
        for visual in _normalize_visuals(message):
            visual_bbox = _coerce_bbox(visual.get("bbox"), image.size)
            if not visual_bbox:
                continue
            if not _should_keep_chat_visual(
                image=image,
                visual=visual,
                visual_bbox=visual_bbox,
                avatar_bbox=avatar_bbox,
                message_bbox=message_bbox,
            ):
                visual["_filtered_out"] = True
                visual["filter_reason"] = "non_preserved_chat_visual"
                continue
            visual["bbox_original"] = _bbox_to_list(visual_bbox)
            snapped = _refine_asset_bbox(
                image=image,
                bbox=visual_bbox,
                kind=str(visual.get("type") or "image"),
                bounds=message_bbox,
            )
            visual["bbox_refined"] = _bbox_to_list(snapped or visual_bbox)
            visual["bbox_refine_method"] = "edge_snap" if snapped else "original"
            filtered_visuals.append(visual)
        if message.get("visuals") is not None:
            message["visuals"] = filtered_visuals

    refined["_asset_refine"] = {
        "method": "local_avatar_candidates_then_asset_refine",
        "image_width": image.width,
        "image_height": image.height,
        "avatar_candidate_count": len(avatar_candidates),
        "avatar_candidate_assigned": assigned_avatars,
    }
    if avatar_candidates:
        refined["_avatar_candidate_detection"] = {
            "method": "local_margin_cv_candidates",
            "candidates": [
                {
                    "bbox": _bbox_to_list(candidate["bbox"]),
                    "side": candidate["side"],
                    "score": round(float(candidate["score"]), 3),
                }
                for candidate in avatar_candidates[:AVATAR_CANDIDATE_MAX_PER_PAGE]
            ],
        }
    return refined


def _detect_chat_avatar_candidates(image: Image.Image) -> list[dict[str, Any]]:
    try:
        import cv2
        import numpy as np

        arr = np.asarray(image.convert("RGB"))
        if arr.size == 0:
            return []

        height, width = arr.shape[:2]
        if width < 80 or height < 80:
            return []

        mobile_like = height / max(width, 1) > 1.8
        top_skip = max(64, int(width * 0.18)) if mobile_like else 0
        bottom_skip = max(48, int(width * 0.12)) if mobile_like else 0
        y_limit = max(top_skip + 1, height - bottom_skip)
        band_width = min(width, max(56, int(width * 0.16)))
        min_size = max(18, int(width * 0.065))
        max_size = min(220, max(56, int(width * 0.20)))
        bands = [
            ("left", 0, band_width),
            ("right", max(0, width - band_width), width),
        ]

        candidates: list[dict[str, Any]] = []
        for side, band_left, band_right in bands:
            crop = arr[top_skip:y_limit, band_left:band_right]
            if crop.size == 0:
                continue
            gray = cv2.cvtColor(crop, cv2.COLOR_RGB2GRAY)
            edges = cv2.Canny(gray, 45, 140)
            kernel = np.ones((5, 5), dtype="uint8")
            mask = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, kernel, iterations=2)
            mask = cv2.dilate(mask, kernel, iterations=1)
            contours, _hierarchy = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

            for contour in contours:
                x, y, candidate_width, candidate_height = cv2.boundingRect(contour)
                bbox = (
                    int(band_left + x),
                    int(top_skip + y),
                    int(band_left + x + candidate_width),
                    int(top_skip + y + candidate_height),
                )
                bbox = _trim_avatar_candidate_to_outer_visual(arr, bbox, side=side) or bbox
                candidate_width = bbox[2] - bbox[0]
                candidate_height = bbox[3] - bbox[1]
                if candidate_width < min_size or candidate_height < min_size:
                    continue
                if candidate_width > max_size or candidate_height > max_size:
                    continue
                aspect = candidate_width / max(candidate_height, 1)
                if not 0.65 <= aspect <= 1.45:
                    continue
                bbox = _normalize_avatar_candidate_bbox(bbox, image_size=(width, height))
                if not bbox:
                    continue
                score = _score_avatar_candidate(arr, bbox, side=side)
                if score < 22:
                    continue
                candidates.append(
                    {
                        "bbox": bbox,
                        "side": side,
                        "score": score,
                    }
                )

        return _dedupe_avatar_candidates(candidates)[:AVATAR_CANDIDATE_MAX_PER_PAGE]
    except Exception:
        return []


def _should_keep_chat_visual(
    *,
    image: Image.Image,
    visual: dict[str, Any],
    visual_bbox: tuple[int, int, int, int],
    avatar_bbox: tuple[int, int, int, int] | None,
    message_bbox: tuple[int, int, int, int] | None,
) -> bool:
    kind = _clean_text(visual.get("type")).lower()
    alt = _clean_text(visual.get("alt")).lower()
    label = f"{kind} {alt}"
    if "avatar" in label or "头像" in label:
        return False
    if any(marker in label for marker in ("icon", "pdf", "file", "文件", "图标")):
        return False
    if avatar_bbox and (
        _bbox_iou(visual_bbox, avatar_bbox) > 0.08
        or _bbox_center_distance(visual_bbox, avatar_bbox)
        < max(visual_bbox[2] - visual_bbox[0], visual_bbox[3] - visual_bbox[1], avatar_bbox[2] - avatar_bbox[0], avatar_bbox[3] - avatar_bbox[1])
    ):
        return False

    width = visual_bbox[2] - visual_bbox[0]
    height = visual_bbox[3] - visual_bbox[1]
    area = width * height
    if width < 16 or height < 16:
        return False
    if kind == "image" and area < 2500:
        return False
    if kind not in {"emoji", "sticker", "image"} and area < 4096:
        return False
    if message_bbox and kind == "image" and area < _bbox_area(message_bbox) * 0.04:
        return False
    return True


def _trim_avatar_candidate_to_outer_visual(
    arr: Any,
    bbox: tuple[int, int, int, int],
    *,
    side: str,
) -> tuple[int, int, int, int] | None:
    try:
        import cv2
        import numpy as np

        image_height, image_width = arr.shape[:2]
        left, top, right, bottom = bbox
        width = right - left
        height = bottom - top
        if side != "right" or image_width >= 600 or height <= width * 1.08:
            return None

        crop = arr[top:bottom, left:right]
        if crop.size == 0:
            return None

        focus_start = max(0, int(crop.shape[1] * 0.42))
        focus = crop[:, focus_start:]
        hsv = cv2.cvtColor(focus, cv2.COLOR_RGB2HSV)
        mask = (((hsv[:, :, 1] > 28) & (hsv[:, :, 2] < 248))).astype("uint8") * 255
        mask = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, np.ones((3, 3), dtype="uint8"), iterations=1)
        count, _labels, stats, _centroids = cv2.connectedComponentsWithStats(mask, 8)

        best = None
        best_area = 0
        for label in range(1, count):
            x, y, component_width, component_height, area = stats[label]
            if area < 24:
                continue
            aspect = component_width / max(component_height, 1)
            if not 0.55 <= aspect <= 1.8:
                continue
            if int(area) > best_area:
                best_area = int(area)
                best = (
                    left + focus_start + int(x),
                    top + int(y),
                    left + focus_start + int(x + component_width),
                    top + int(y + component_height),
                )
        if best and side == "right":
            best_width = best[2] - best[0]
            best_height = best[3] - best[1]
            size = max(best_width, best_height, int(round(image_width * 0.10)))
            best = (
                max(0, best[2] - size),
                best[1],
                best[2],
                min(image_height, best[1] + size),
            )
        return best
    except Exception:
        return None


def _normalize_avatar_candidate_bbox(
    bbox: tuple[int, int, int, int],
    *,
    image_size: tuple[int, int],
) -> tuple[int, int, int, int] | None:
    left, top, right, bottom = bbox
    width = right - left
    height = bottom - top
    if width <= 0 or height <= 0:
        return None

    if height > width * 1.18:
        bottom = top + width
    elif width > height * 1.18:
        right = left + height

    return _intersect_bbox((left, top, right, bottom), (0, 0, image_size[0], image_size[1]))


def _score_avatar_candidate(
    arr: Any,
    bbox: tuple[int, int, int, int],
    *,
    side: str,
) -> float:
    try:
        import cv2
        import numpy as np

        height, width = arr.shape[:2]
        left, top, right, bottom = bbox
        patch = arr[top:bottom, left:right]
        if patch.size == 0:
            return 0.0
        patch_height, patch_width = patch.shape[:2]
        gray = cv2.cvtColor(patch, cv2.COLOR_RGB2GRAY)
        edges = cv2.Canny(gray, 45, 140)
        edge_density = float(np.count_nonzero(edges) / max(patch_width * patch_height, 1))
        gray_std = float(gray.std())
        color_std = float(patch.reshape(-1, 3).std(axis=0).mean())
        size_balance = 1.0 - min(abs(patch_width - patch_height) / max(patch_width, patch_height, 1), 1.0)

        ring_padding = max(5, int(max(patch_width, patch_height) * 0.12))
        ring_bbox = _expand_bbox(bbox, (width, height), padding=ring_padding)
        ring = arr[ring_bbox[1] : ring_bbox[3], ring_bbox[0] : ring_bbox[2]]
        center_mask = np.ones(ring.shape[:2], dtype=bool)
        inner_left = left - ring_bbox[0]
        inner_top = top - ring_bbox[1]
        inner_right = right - ring_bbox[0]
        inner_bottom = bottom - ring_bbox[1]
        center_mask[inner_top:inner_bottom, inner_left:inner_right] = False
        ring_pixels = ring[center_mask]
        ring_contrast = 0.0
        if ring_pixels.size:
            ring_contrast = float(np.linalg.norm(patch.reshape(-1, 3).mean(axis=0) - ring_pixels.reshape(-1, 3).mean(axis=0)))

        margin_bonus = 0.0
        center_x = (left + right) / 2
        if side == "left":
            margin_bonus = max(0.0, 1.0 - center_x / max(width * 0.28, 1)) * 8
        else:
            margin_bonus = max(0.0, 1.0 - (width - center_x) / max(width * 0.28, 1)) * 8

        return (
            gray_std * 0.38
            + color_std * 0.32
            + edge_density * 105
            + ring_contrast * 0.12
            + size_balance * 12
            + margin_bonus
        )
    except Exception:
        return 0.0


def _dedupe_avatar_candidates(candidates: list[dict[str, Any]]) -> list[dict[str, Any]]:
    selected: list[dict[str, Any]] = []
    for candidate in sorted(candidates, key=lambda item: float(item.get("score", 0)), reverse=True):
        bbox = candidate["bbox"]
        duplicate = False
        for existing in selected:
            existing_bbox = existing["bbox"]
            if _bbox_iou(bbox, existing_bbox) > 0.18:
                duplicate = True
                break
            if _bbox_center_distance(bbox, existing_bbox) < max(bbox[2] - bbox[0], bbox[3] - bbox[1]) * 0.55:
                duplicate = True
                break
        if not duplicate:
            selected.append(candidate)
    return sorted(selected, key=lambda item: (item["bbox"][1], item["bbox"][0]))


def _assign_detected_avatar_candidates(
    *,
    messages: list[dict[str, Any]],
    candidates: list[dict[str, Any]],
    image_size: tuple[int, int],
) -> int:
    if not messages or not candidates:
        return 0

    remaining = list(candidates)
    assigned = 0
    for message in messages:
        if not _message_accepts_detected_avatar(message):
            continue
        message_bbox = _coerce_bbox(message.get("message_bbox"), image_size)
        existing_bbox = _coerce_bbox(message.get("avatar_bbox"), image_size)
        preferred_side = _preferred_avatar_side(
            message=message,
            existing_bbox=existing_bbox,
            message_bbox=message_bbox,
            image_size=image_size,
        )
        match_index = _select_avatar_candidate_for_message(
            message_bbox=message_bbox,
            existing_bbox=existing_bbox,
            candidates=remaining,
            image_size=image_size,
            preferred_side=preferred_side,
        )
        if match_index is None:
            match_index = _select_next_avatar_candidate_by_sequence(
                candidates=remaining,
                preferred_side=preferred_side,
            )
        if match_index is None:
            continue

        candidate = remaining.pop(match_index)
        candidate_bbox = candidate["bbox"]
        if message.get("avatar_bbox") is not None and message.get("avatar_bbox_model") is None:
            message["avatar_bbox_model"] = message.get("avatar_bbox")
        message["avatar_bbox"] = _bbox_to_list(candidate_bbox)
        message["avatar_bbox_detected"] = _bbox_to_list(candidate_bbox)
        message["avatar_bbox_source"] = "local_candidate_detection"
        message["avatar_candidate_score"] = round(float(candidate["score"]), 3)
        message["avatar_candidate_side"] = candidate["side"]
        assigned += 1
    return assigned


def _message_accepts_detected_avatar(message: dict[str, Any]) -> bool:
    if message.get("avatar_bbox") is not None:
        return True
    return bool(_clean_text(message.get("sender")))


def _preferred_avatar_side(
    *,
    message: dict[str, Any],
    existing_bbox: tuple[int, int, int, int] | None,
    message_bbox: tuple[int, int, int, int] | None,
    image_size: tuple[int, int],
) -> str | None:
    image_mid_x = image_size[0] / 2
    if existing_bbox:
        center_x = (existing_bbox[0] + existing_bbox[2]) / 2
        return "left" if center_x < image_mid_x else "right"
    if message_bbox:
        center_x = (message_bbox[0] + message_bbox[2]) / 2
        if center_x > image_mid_x * 1.12:
            return "right"
    if _clean_text(message.get("sender")):
        return "left"
    return None


def _select_next_avatar_candidate_by_sequence(
    *,
    candidates: list[dict[str, Any]],
    preferred_side: str | None,
) -> int | None:
    if not candidates:
        return None
    if preferred_side:
        for index, candidate in enumerate(candidates):
            if candidate.get("side") == preferred_side:
                return index
    return 0


def _select_avatar_candidate_for_message(
    *,
    message_bbox: tuple[int, int, int, int] | None,
    existing_bbox: tuple[int, int, int, int] | None,
    candidates: list[dict[str, Any]],
    image_size: tuple[int, int],
    preferred_side: str | None,
) -> int | None:
    best_index = None
    best_score = None
    has_preferred_candidate = bool(preferred_side and any(candidate.get("side") == preferred_side for candidate in candidates))
    for index, candidate in enumerate(candidates):
        bbox = candidate["bbox"]
        candidate_score = float(candidate["score"])
        score = candidate_score

        if preferred_side and candidate.get("side") != preferred_side:
            if has_preferred_candidate:
                continue
            score -= 70.0
        elif preferred_side:
            score += 35.0

        if message_bbox:
            tolerance = max(36, bbox[3] - bbox[1], int(image_size[1] * 0.015))
            if not _bbox_y_matches(bbox, message_bbox, tolerance=tolerance):
                continue
            y_overlap = _bbox_y_overlap(bbox, message_bbox)
            score += min(45.0, y_overlap / max(bbox[3] - bbox[1], 1) * 45.0)
            candidate_center_y = (bbox[1] + bbox[3]) / 2
            message_center_y = (message_bbox[1] + message_bbox[3]) / 2
            score -= min(18.0, abs(candidate_center_y - message_center_y) * 0.025)
        elif existing_bbox:
            distance = _bbox_center_distance(bbox, existing_bbox)
            if distance > max(image_size) * 0.35:
                continue
            score -= min(22.0, distance * 0.04)
        else:
            continue

        if existing_bbox:
            existing_distance = _bbox_center_distance(bbox, existing_bbox)
            score -= min(10.0, existing_distance * 0.015)

        if best_score is None or score > best_score:
            best_score = score
            best_index = index
    return best_index


def _refine_asset_bbox(
    *,
    image: Image.Image,
    bbox: tuple[int, int, int, int],
    kind: str,
    bounds: tuple[int, int, int, int] | None = None,
) -> tuple[int, int, int, int] | None:
    if not _is_reasonable_asset_bbox(bbox, kind):
        return None

    padding = max(6, int(max(bbox[2] - bbox[0], bbox[3] - bbox[1]) * 0.28))
    search_bbox = _expand_bbox(bbox, image.size, padding=padding)
    if bounds:
        expanded_bounds = _expand_bbox(bounds, image.size, padding=4)
        search_bbox = _intersect_bbox(search_bbox, expanded_bounds)
    if not search_bbox:
        return None

    crop = image.crop(search_bbox)
    target_rect = (
        bbox[0] - search_bbox[0],
        bbox[1] - search_bbox[1],
        bbox[2] - search_bbox[0],
        bbox[3] - search_bbox[1],
    )
    local_refined = _detect_foreground_bbox(crop, target_rect)
    if not local_refined:
        return None

    refined = _offset_bbox(local_refined, search_bbox[0], search_bbox[1])
    refined = _intersect_bbox(refined, bounds or (0, 0, image.width, image.height))
    if not refined:
        return None

    refined = _normalize_refined_bbox(refined, original=bbox, kind=kind, image_size=image.size, bounds=bounds)
    if not refined or not _is_reasonable_asset_bbox(refined, kind):
        return None
    if _bbox_area(refined) < max(16, _bbox_area(bbox) * 0.18):
        return None
    if _bbox_area(refined) > _bbox_area(bbox) * 3.0:
        return None
    return refined


def _detect_foreground_bbox(
    crop: Image.Image,
    target_rect: tuple[int, int, int, int],
) -> tuple[int, int, int, int] | None:
    try:
        import cv2
        import numpy as np

        arr = np.asarray(crop.convert("RGB"))
        if arr.size == 0:
            return None

        border_pixels = np.concatenate(
            [
                arr[0, :, :],
                arr[-1, :, :],
                arr[:, 0, :],
                arr[:, -1, :],
            ],
            axis=0,
        )
        bg = np.median(border_pixels, axis=0)
        diff = np.linalg.norm(arr.astype(np.float32) - bg.astype(np.float32), axis=2)
        gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
        edges = cv2.Canny(gray, 50, 150)
        mask = ((diff > 18) | (edges > 0)).astype("uint8") * 255
        kernel = np.ones((3, 3), dtype="uint8")
        mask = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, kernel, iterations=1)
        mask = cv2.dilate(mask, kernel, iterations=1)

        count, labels, stats, _centroids = cv2.connectedComponentsWithStats(mask, 8)
        best_rect = None
        best_score = None
        crop_area = crop.width * crop.height
        for label in range(1, count):
            x, y, w, h, area = stats[label]
            if area < 12 or area > crop_area * 0.92:
                continue
            rect = (int(x), int(y), int(x + w), int(y + h))
            overlap = _bbox_area(_intersect_bbox(rect, target_rect))
            center_distance = _bbox_center_distance(rect, target_rect)
            score = (overlap * 3.0) + (area * 0.02) - center_distance
            if best_score is None or score > best_score:
                best_score = score
                best_rect = rect
        return best_rect
    except Exception:
        return _detect_foreground_bbox_pillow(crop)


def _detect_foreground_bbox_pillow(crop: Image.Image) -> tuple[int, int, int, int] | None:
    image = crop.convert("RGB")
    pixels = image.load()
    width, height = image.size
    if width <= 0 or height <= 0:
        return None

    border_samples = []
    for x in range(width):
        border_samples.append(pixels[x, 0])
        border_samples.append(pixels[x, height - 1])
    for y in range(height):
        border_samples.append(pixels[0, y])
        border_samples.append(pixels[width - 1, y])
    bg = tuple(sorted(channel)[len(channel) // 2] for channel in zip(*border_samples))

    xs = []
    ys = []
    for y in range(height):
        for x in range(width):
            pixel = pixels[x, y]
            diff = sum(abs(pixel[i] - bg[i]) for i in range(3))
            if diff > 36:
                xs.append(x)
                ys.append(y)
    if not xs or not ys:
        return None
    return min(xs), min(ys), max(xs) + 1, max(ys) + 1


def _normalize_refined_bbox(
    refined: tuple[int, int, int, int],
    *,
    original: tuple[int, int, int, int],
    kind: str,
    image_size: tuple[int, int],
    bounds: tuple[int, int, int, int] | None,
) -> tuple[int, int, int, int] | None:
    lower_kind = kind.lower()
    if lower_kind == "avatar":
        original_w = original[2] - original[0]
        original_h = original[3] - original[1]
        refined_w = refined[2] - refined[0]
        refined_h = refined[3] - refined[1]
        size = max(refined_w, refined_h, int(max(original_w, original_h) * 0.78))
        center_x = (refined[0] + refined[2]) // 2
        center_y = (refined[1] + refined[3]) // 2
        half = max(size // 2, 1)
        refined = (center_x - half, center_y - half, center_x + half, center_y + half)

    page_bounds = bounds or (0, 0, image_size[0], image_size[1])
    return _intersect_bbox(refined, page_bounds)


def _is_reasonable_asset_bbox(bbox: tuple[int, int, int, int], kind: str) -> bool:
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    if width <= 0 or height <= 0:
        return False
    aspect = width / max(height, 1)
    lower_kind = kind.lower()
    if lower_kind == "avatar":
        return 12 <= width <= 220 and 12 <= height <= 220 and 0.55 <= aspect <= 1.8
    return width >= 8 and height >= 8 and 0.08 <= aspect <= 12


def _save_debug_overlay(image: Image.Image, layout: dict[str, Any], output_path: Path) -> None:
    overlay = image.convert("RGB").copy()
    draw = ImageDraw.Draw(overlay)
    for index, message in enumerate(_normalize_messages(layout), start=1):
        message_bbox = _coerce_bbox(message.get("message_bbox"), overlay.size)
        if message_bbox:
            _draw_bbox(draw, message_bbox, "#2F80ED", f"m{index}")

        avatar_original = _coerce_bbox(message.get("avatar_bbox_original"), overlay.size)
        avatar_detected = _coerce_bbox(message.get("avatar_bbox_detected"), overlay.size)
        avatar_refined = _coerce_bbox(message.get("avatar_bbox_refined") or message.get("avatar_bbox"), overlay.size)
        if avatar_original:
            _draw_bbox(draw, avatar_original, "#EB5757", "avatar original")
        if avatar_detected and avatar_detected != avatar_refined:
            _draw_bbox(draw, avatar_detected, "#F2C94C", "avatar detected")
        if avatar_refined:
            _draw_bbox(draw, avatar_refined, "#27AE60", "avatar refined")

        for visual_index, visual in enumerate(_normalize_visuals(message), start=1):
            visual_original = _coerce_bbox(visual.get("bbox_original"), overlay.size)
            visual_refined = _coerce_bbox(visual.get("bbox_refined") or visual.get("bbox"), overlay.size)
            if visual_original:
                _draw_bbox(draw, visual_original, "#F2994A", f"v{visual_index} original")
            if visual_refined:
                _draw_bbox(draw, visual_refined, "#9B51E0", f"v{visual_index} refined")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    overlay.save(output_path, format="PNG")


def _draw_bbox(draw: ImageDraw.ImageDraw, bbox: tuple[int, int, int, int], color: str, label: str) -> None:
    for inset in range(2):
        draw.rectangle(
            (bbox[0] - inset, bbox[1] - inset, bbox[2] + inset, bbox[3] + inset),
            outline=color,
        )
    draw.text((bbox[0] + 2, max(0, bbox[1] - 12)), label, fill=color)


def _set_table_borders(table, color: str) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_width(cell, width) -> None:
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def _load_input_pages(input_file: Path) -> list[Image.Image]:
    if input_file.suffix.lower() == ".pdf":
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError("请先安装 PyMuPDF: pip install pymupdf") from exc

        pages: list[Image.Image] = []
        document = fitz.open(str(input_file))
        try:
            for page in document:
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                image = Image.open(BytesIO(pix.tobytes("png"))).convert("RGB")
                pages.append(image)
        finally:
            close = getattr(document, "close", None)
            if callable(close):
                close()
        return pages

    return [_open_image_as_rgb(input_file)]


def _open_image_as_rgb(image_path: Path) -> Image.Image:
    with Image.open(image_path) as image:
        first_frame = next(ImageSequence.Iterator(image))
        return first_frame.convert("RGB")


def _analyze_chat_page(
    *,
    image: Image.Image,
    page_no: int,
    model: str,
    gemini_route: str,
) -> dict[str, Any]:
    image_bytes = BytesIO()
    image.save(image_bytes, format="PNG")
    user_prompt = (
        f"Analyze this chat screenshot page {page_no}. "
        f"The image size is width={image.width}px, height={image.height}px. "
        "Return strict JSON using the requested schema."
    )
    response = generate_vision_html(
        system_prompt=CHAT_LAYOUT_SYSTEM_PROMPT,
        image_bytes=image_bytes.getvalue(),
        mime_type="image/png",
        model=model,
        route=gemini_route,
        user_prompt=user_prompt,
        temperature=0,
    )
    parsed = _parse_json_object(response)
    if not isinstance(parsed, dict):
        raise RuntimeError("聊天截图布局识别结果不是 JSON 对象")
    return parsed


def _refine_layout_with_local_asset_pass(
    *,
    image: Image.Image,
    layout: dict[str, Any],
    page_no: int,
    model: str,
    gemini_route: str,
    status_callback: StatusCallback | None = None,
) -> dict[str, Any]:
    refined = copy.deepcopy(layout if isinstance(layout, dict) else {})
    messages = _normalize_messages(refined)
    checked = 0
    failed = 0
    skipped = 0

    for message_index, message in enumerate(messages, start=1):
        if checked >= LOCAL_ASSET_RECHECK_MAX_MESSAGES:
            skipped += 1
            continue

        message_bbox = _coerce_bbox(message.get("message_bbox"), image.size)
        if not message_bbox:
            skipped += 1
            continue

        checked += 1
        _emit(
            status_callback,
            f"正在局部复核第 {page_no} 页第 {message_index} 条消息的头像/表情位置",
        )
        try:
            local_layout = _analyze_message_assets_local(
                image=image,
                message_bbox=message_bbox,
                page_no=page_no,
                message_index=message_index,
                model=model,
                gemini_route=gemini_route,
            )
        except Exception as exc:
            failed += 1
            message["_local_asset_error"] = f"{exc.__class__.__name__}: {exc}"
            continue

        _merge_local_asset_layout(
            message=message,
            local_layout=local_layout,
            message_bbox=message_bbox,
        )
        message["_local_asset_checked"] = True

    refined["_local_asset_pass"] = {
        "checked": checked,
        "failed": failed,
        "skipped": skipped,
        "max_messages": LOCAL_ASSET_RECHECK_MAX_MESSAGES,
    }
    return refined


def _analyze_message_assets_local(
    *,
    image: Image.Image,
    message_bbox: tuple[int, int, int, int],
    page_no: int,
    message_index: int,
    model: str,
    gemini_route: str,
) -> dict[str, Any]:
    crop_bbox = _expand_bbox(message_bbox, image.size, padding=2)
    crop = image.crop(crop_bbox)
    image_bytes = BytesIO()
    crop.save(image_bytes, format="PNG")
    user_prompt = (
        f"Analyze cropped message {message_index} on page {page_no}. "
        f"The cropped image size is width={crop.width}px, height={crop.height}px. "
        "Return strict JSON using the requested schema."
    )
    response = generate_vision_html(
        system_prompt=CHAT_LOCAL_ASSET_SYSTEM_PROMPT,
        image_bytes=image_bytes.getvalue(),
        mime_type="image/png",
        model=model,
        route=gemini_route,
        user_prompt=user_prompt,
        temperature=0,
    )
    parsed = _parse_json_object(response)
    if not isinstance(parsed, dict):
        raise RuntimeError("局部头像/表情识别结果不是 JSON 对象")
    parsed["_crop_bbox"] = _bbox_to_list(crop_bbox)
    return parsed


def _merge_local_asset_layout(
    *,
    message: dict[str, Any],
    local_layout: dict[str, Any],
    message_bbox: tuple[int, int, int, int],
) -> None:
    crop_bbox = _coerce_bbox(local_layout.get("_crop_bbox"), (10**9, 10**9)) or message_bbox
    crop_size = (crop_bbox[2] - crop_bbox[0], crop_bbox[3] - crop_bbox[1])

    local_avatar = _coerce_bbox(local_layout.get("avatar_bbox"), crop_size)
    if local_avatar:
        global_avatar = _offset_bbox(local_avatar, crop_bbox[0], crop_bbox[1])
        if message.get("avatar_bbox") is not None:
            message["avatar_bbox_model"] = message.get("avatar_bbox")
        message["avatar_bbox"] = _bbox_to_list(global_avatar)
        message["avatar_bbox_source"] = "local_asset_pass"

    local_visuals = []
    for visual in _normalize_visuals(local_layout):
        local_bbox = _coerce_bbox(visual.get("bbox"), crop_size)
        if not local_bbox:
            continue
        global_bbox = _offset_bbox(local_bbox, crop_bbox[0], crop_bbox[1])
        local_visual = {
            "type": visual.get("type") or "image",
            "bbox": _bbox_to_list(global_bbox),
            "alt": _clean_text(visual.get("alt")),
            "bbox_source": "local_asset_pass",
        }
        local_visuals.append(local_visual)

    if local_visuals:
        if message.get("visuals"):
            message["visuals_model"] = message.get("visuals")
        message["visuals"] = local_visuals


def _parse_json_object(text: str) -> Any:
    normalized = (text or "").strip()
    normalized = re.sub(r"^```(?:json)?\s*", "", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"\s*```$", "", normalized)
    try:
        return json.loads(normalized)
    except json.JSONDecodeError:
        pass

    start = normalized.find("{")
    end = normalized.rfind("}")
    if start >= 0 and end > start:
        return json.loads(normalized[start : end + 1])
    raise ValueError("未能从聊天截图布局识别结果中解析 JSON")


def _coerce_bbox(value: Any, image_size: tuple[int, int]) -> tuple[int, int, int, int] | None:
    width, height = image_size
    raw: list[Any] | None = None
    if isinstance(value, (list, tuple)) and len(value) == 4:
        raw = list(value)
    elif isinstance(value, dict):
        if all(key in value for key in ("left", "top", "right", "bottom")):
            raw = [value["left"], value["top"], value["right"], value["bottom"]]
        elif all(key in value for key in ("x", "y", "width", "height")):
            raw = [value["x"], value["y"], value["x"] + value["width"], value["y"] + value["height"]]
    if raw is None:
        return None

    try:
        coords = [float(item) for item in raw]
    except (TypeError, ValueError):
        return None
    if all(0 <= item <= 1 for item in coords):
        coords = [coords[0] * width, coords[1] * height, coords[2] * width, coords[3] * height]

    left, top, right, bottom = coords
    if right <= left or bottom <= top:
        right = left + max(right, 0)
        bottom = top + max(bottom, 0)

    left_i = max(0, min(width, int(round(left))))
    top_i = max(0, min(height, int(round(top))))
    right_i = max(0, min(width, int(round(right))))
    bottom_i = max(0, min(height, int(round(bottom))))

    if right_i - left_i < 4 or bottom_i - top_i < 4:
        return None
    return left_i, top_i, right_i, bottom_i


def _bbox_to_list(bbox: tuple[int, int, int, int]) -> list[int]:
    return [int(bbox[0]), int(bbox[1]), int(bbox[2]), int(bbox[3])]


def _offset_bbox(
    bbox: tuple[int, int, int, int],
    offset_x: int,
    offset_y: int,
) -> tuple[int, int, int, int]:
    return (
        bbox[0] + offset_x,
        bbox[1] + offset_y,
        bbox[2] + offset_x,
        bbox[3] + offset_y,
    )


def _expand_bbox(
    bbox: tuple[int, int, int, int],
    image_size: tuple[int, int],
    *,
    padding: int,
) -> tuple[int, int, int, int]:
    width, height = image_size
    return (
        max(0, bbox[0] - padding),
        max(0, bbox[1] - padding),
        min(width, bbox[2] + padding),
        min(height, bbox[3] + padding),
    )


def _intersect_bbox(
    first: tuple[int, int, int, int] | None,
    second: tuple[int, int, int, int] | None,
) -> tuple[int, int, int, int] | None:
    if not first or not second:
        return None
    left = max(first[0], second[0])
    top = max(first[1], second[1])
    right = min(first[2], second[2])
    bottom = min(first[3], second[3])
    if right - left < 1 or bottom - top < 1:
        return None
    return left, top, right, bottom


def _bbox_y_overlap(
    first: tuple[int, int, int, int],
    second: tuple[int, int, int, int],
) -> int:
    return max(0, min(first[3], second[3]) - max(first[1], second[1]))


def _bbox_y_matches(
    first: tuple[int, int, int, int],
    second: tuple[int, int, int, int],
    *,
    tolerance: int,
) -> bool:
    if _bbox_y_overlap(first, second) > 0:
        return True
    first_center_y = (first[1] + first[3]) / 2
    return second[1] - tolerance <= first_center_y <= second[3] + tolerance


def _bbox_area(bbox: tuple[int, int, int, int] | None) -> int:
    if not bbox:
        return 0
    return max(0, bbox[2] - bbox[0]) * max(0, bbox[3] - bbox[1])


def _bbox_iou(
    first: tuple[int, int, int, int],
    second: tuple[int, int, int, int],
) -> float:
    intersection = _bbox_area(_intersect_bbox(first, second))
    if intersection <= 0:
        return 0.0
    union = _bbox_area(first) + _bbox_area(second) - intersection
    return intersection / max(union, 1)


def _bbox_center_distance(
    first: tuple[int, int, int, int],
    second: tuple[int, int, int, int],
) -> float:
    first_x = (first[0] + first[2]) / 2
    first_y = (first[1] + first[3]) / 2
    second_x = (second[0] + second[2]) / 2
    second_y = (second[1] + second[3]) / 2
    return ((first_x - second_x) ** 2 + (first_y - second_y) ** 2) ** 0.5


def _crop_to_file(
    *,
    image: Image.Image,
    bbox: tuple[int, int, int, int],
    output_path: Path,
    padding: int = 0,
) -> Path:
    left, top, right, bottom = bbox
    left = max(0, left - padding)
    top = max(0, top - padding)
    right = min(image.width, right + padding)
    bottom = min(image.height, bottom + padding)
    crop = image.crop((left, top, right, bottom))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    crop.save(output_path, format="PNG")
    return output_path


def _emit(callback: StatusCallback | None, message: str) -> None:
    if callback:
        callback(message)
