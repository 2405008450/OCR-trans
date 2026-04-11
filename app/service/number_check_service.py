import asyncio
import json
import logging
import os
import re
import shutil
import sys
import threading
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

from docx import Document
from fastapi import UploadFile

from app.core.config import settings
from app.core.file_naming import build_user_visible_filename, ensure_unique_path
from app.service.gemini_service import (
    GEMINI_ROUTE_OPENROUTER,
    ensure_gemini_route_configured,
    resolve_model_for_route,
)
from app.service.libreoffice_service import convert_doc_to_docx_via_libreoffice


_task_progress: Dict[str, Dict[str, Any]] = {}
logger = logging.getLogger("app.number_check")

REPO_ROOT = Path(__file__).resolve().parents[2]
NUMBER_CHECK_LATEST_ROOT = REPO_ROOT / "专检" / "数检"
NUMBER_CHECK_MAIN_FILE = NUMBER_CHECK_LATEST_ROOT / "main_zhongying.py"
NUMBER_CHECK_MODE_DOUBLE = "double"
NUMBER_CHECK_MODE_SINGLE = "single"
SECTION_DIR_NAMES = {"正文": "zhengwen", "页眉": "yemei", "页脚": "yejiao"}
SECTION_REPORT_KEYS = {"正文": "body_json", "页眉": "header_json", "页脚": "footer_json"}
SECTION_COUNT_KEYS = {"正文": "body_issues", "页眉": "header_issues", "页脚": "footer_issues"}


def _specialized_output_dir(section_dir: str, output_dir: Path, folder: str = "output_json") -> Path:
    return output_dir / section_dir / folder


def _normalize_number_check_mode(mode: Optional[str]) -> str:
    candidate = (mode or NUMBER_CHECK_MODE_DOUBLE).strip().lower()
    if candidate not in {NUMBER_CHECK_MODE_DOUBLE, NUMBER_CHECK_MODE_SINGLE}:
        raise ValueError(f"不支持的数字专检模式: {mode}")
    return candidate


def get_number_check_default_mode() -> str:
    """读取最新版主脚本里的 BILINGUAL_MODE，作为网页默认模式。"""
    try:
        content = NUMBER_CHECK_MAIN_FILE.read_text(encoding="utf-8")
    except Exception:
        return NUMBER_CHECK_MODE_DOUBLE

    matched = re.search(r"^\s*BILINGUAL_MODE\s*=\s*(True|False)\s*$", content, re.MULTILINE)
    if not matched:
        return NUMBER_CHECK_MODE_DOUBLE
    return NUMBER_CHECK_MODE_SINGLE if matched.group(1) == "True" else NUMBER_CHECK_MODE_DOUBLE


def _patch_v2_replace_helper(fix_replace_docx_module) -> None:
    def _safe_replace_and_add_comment_in_paragraph(
        paragraph,
        pattern: str,
        old_value: str,
        new_value: str,
        comment_manager,
        reason: str = "",
        anchor_pattern: str = None,
    ) -> bool:
        runs = list(paragraph.runs)
        if not runs:
            return False

        full_text = "".join(r.text or "" for r in runs)
        flags = re.IGNORECASE | re.DOTALL
        if anchor_pattern and not re.search(anchor_pattern, full_text, flags=flags):
            return False

        match = re.search(pattern, full_text, flags=flags)
        if not match:
            return False

        start, end = match.span()
        if start == end:
            return False

        spans = []
        cursor = 0
        for run in runs:
            text = run.text or ""
            spans.append((run, cursor, cursor + len(text)))
            cursor += len(text)

        hit = [(run, s, e) for run, s, e in spans if start < e and end > s]
        if not hit:
            return False

        first_run, fs, _ = hit[0]
        last_run, ls, _ = hit[-1]
        prefix = first_run.text[: max(0, start - fs)]
        suffix = last_run.text[max(0, end - ls):]

        for run, _, _ in hit:
            run.text = ""

        parent = first_run._element.getparent()
        insert_pos = parent.index(first_run._element)

        if prefix:
            prefix_run = paragraph.add_run(prefix)
            parent.remove(prefix_run._element)
            parent.insert(insert_pos, prefix_run._element)
            insert_pos += 1

        new_run = paragraph.add_run(new_value)
        parent.remove(new_run._element)
        parent.insert(insert_pos, new_run._element)
        insert_pos += 1

        new_run.bold = first_run.bold
        new_run.italic = first_run.italic
        new_run.font.name = first_run.font.name
        new_run.font.size = first_run.font.size

        if suffix:
            suffix_run = paragraph.add_run(suffix)
            parent.remove(suffix_run._element)
            parent.insert(insert_pos, suffix_run._element)

        comment_text = (
            f"【修改建议】\n"
            f"原值: {old_value}\n"
            f"新值: {new_value}\n"
            f"修改理由: {reason or ''}"
        )
        return comment_manager.add_comment_to_run(new_run, comment_text)

    fix_replace_docx_module.replace_and_add_comment_in_paragraph = _safe_replace_and_add_comment_in_paragraph

# 专检模块导入锁：防止并发任务修改 sys.path / sys.modules 时互相干扰
_specialist_import_lock = threading.Lock()

# ── 客户端日志脱敏规则（服务器终端日志不受影响）────────────────────────────
# 规则按顺序依次替换，(pattern, replacement) 格式
_CLIENT_LOG_SANITIZE_RULES: List[tuple] = [
    # [config] 整行：隐藏模型名和路线，仅保留结构标记
    (re.compile(r'\[config\]\s+route=\S+,\s*model=\S+'), '[config] 任务已初始化'),
    (re.compile(r'\[config\]\s+.*'), '[config] 任务已初始化'),
    # LLM 对比行：路线= / 模型= 后面的值
    (re.compile(r'(路线=)\S+'), r'\1***'),
    (re.compile(r'(模型=)\S+'), r'\1***'),
    # 英文 route= / model= 后面的值（不区分大小写）
    (re.compile(r'(?i)(route=)\S+'), r'\1***'),
    (re.compile(r'(?i)(model=)\S+'), r'\1***'),
    # gemini 系列模型名（如 gemini-3-flash-preview、gemini-3.1-pro-preview 等）
    (re.compile(r'(?i)gemini[-/][\w.\-]+'), '[AI模型]'),
    # openrouter / google 作为路线标识
    (re.compile(r'\bopenrouter\b', re.IGNORECASE), '[路线]'),
    (re.compile(r'\bgoogle\b', re.IGNORECASE), '[路线]'),
]


def _sanitize_client_log(line: str) -> str:
    """对即将写入客户端 stream_log 的单行内容进行脱敏。"""
    for pattern, replacement in _CLIENT_LOG_SANITIZE_RULES:
        line = pattern.sub(replacement, line)
    return line

NUMBER_CHECK_MODELS: Dict[str, Dict[str, str]] = {
    "gemini-3-flash-preview": {
        "label": "快速版V2",
        "description": "速度更快，适合常规数字核对场景。",
    },
    "gemini-3.1-pro-preview": {
        "label": "增强版V2",
        "description": "推理更强，适合复杂编号和上下文判断场景。",
    },
}

NUMBER_CHECK_MODEL_ALIASES: Dict[str, str] = {
    "gemini-3-flash-preview": "gemini-3-flash-preview",
    "google/gemini-3-flash-preview": "gemini-3-flash-preview",
    "gemini 3 flash preview": "gemini-3-flash-preview",
    "gemini-3.1-pro-preview": "gemini-3.1-pro-preview",
    "google/gemini-3.1-pro-preview": "gemini-3.1-pro-preview",
    "gemini 3.1 pro preview": "gemini-3.1-pro-preview",
}


def get_number_check_models() -> Dict[str, Dict[str, str]]:
    return NUMBER_CHECK_MODELS


def normalize_number_check_model(model_name: Optional[str]) -> str:
    candidate = (model_name or "gemini-3.1-pro-preview").strip()
    key = NUMBER_CHECK_MODEL_ALIASES.get(candidate.lower(), candidate)
    if key not in NUMBER_CHECK_MODELS:
        raise ValueError(f"不支持的数字专检模型: {model_name}")
    return key


def _append_stream_log(task_id: str, message: str) -> None:
    if task_id not in _task_progress:
        return
    line = _sanitize_client_log((message or "").strip())
    if not line:
        return
    current = _task_progress[task_id].get("stream_log", "")
    lines = current.splitlines() if current else []
    if lines and lines[-1] == line:
        return
    combined = f"{current}\n{line}" if current else line
    _task_progress[task_id]["stream_log"] = combined[-50000:]


def _emit_log(task_id: str, message: str, level: str = "info") -> None:
    _append_stream_log(task_id, message)
    log_line = f"[number-check][{task_id}] {message}"
    if level == "error":
        logger.error(log_line)
    elif level == "warning":
        logger.warning(log_line)
    else:
        logger.info(log_line)


def _init_task_progress(task_id: str, total_steps: int = 6) -> None:
    _task_progress[task_id] = {
        "status": "running",
        "current_step": 0,
        "total_steps": total_steps,
        "message": "初始化任务...",
        "progress": 0,
        "details": [],
        "stream_log": "",
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat(),
    }


def _update_progress(
    task_id: str,
    current_step: int,
    total_steps: int,
    message: str,
    details: Optional[List[str]] = None,
) -> None:
    if task_id not in _task_progress:
        return

    progress = int((current_step / total_steps) * 100)
    detail_list = details or []
    _task_progress[task_id].update(
        {
            "current_step": current_step,
            "total_steps": total_steps,
            "message": message,
            "progress": progress,
            "details": detail_list,
            "updated_at": datetime.now().isoformat(),
        }
    )
    _emit_log(task_id, f"[{progress:>3}%] {message}")
    for item in detail_list:
        _emit_log(task_id, f"  - {item}")


def _complete_task(task_id: str, result: Optional[Dict[str, Any]] = None, error: Optional[str] = None) -> None:
    if task_id not in _task_progress:
        return

    _task_progress[task_id].update(
        {
            "status": "done" if not error else "failed",
            "progress": 100,
            "message": "处理完成" if not error else f"处理失败: {error}",
            "result": result,
            "error": error,
            "updated_at": datetime.now().isoformat(),
        }
    )
    _emit_log(task_id, "[done] 处理完成" if not error else f"[error] {error}", level="error" if error else "info")


def _get_task_progress(task_id: str) -> Optional[Dict[str, Any]]:
    return _task_progress.get(task_id)


def _cleanup_task(task_id: str) -> None:
    if task_id in _task_progress:
        del _task_progress[task_id]


def _normalize_path(path: str) -> str:
    return path.replace("\\", "/") if path else path


def _validate_docx(file: UploadFile, label: str) -> None:
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in {".docx", ".doc"}:
        raise ValueError(f"{label} 必须是 .docx 或 .doc 文件，当前为 {file.filename}")


def _validate_single_file(file: UploadFile, label: str) -> None:
    ext = os.path.splitext(file.filename or "")[1].lower()
    allowed = {".docx", ".doc", ".pdf", ".xlsx", ".pptx"}
    if ext not in allowed:
        raise ValueError(f"{label} 仅支持 {', '.join(sorted(allowed))}，当前为 {file.filename}")


def _supports_revised_doc_output(ext: str) -> bool:
    return ext.lower() in {".docx", ".doc"}


def _convert_doc_input_if_needed(source: Path, work_dir: Path, role: str) -> Path:
    if source.suffix.lower() != ".doc":
        return source

    work_dir.mkdir(parents=True, exist_ok=True)
    target = work_dir / f"{role}.docx"
    return Path(convert_doc_to_docx_via_libreoffice(source, target))


def _prepare_specialized_import_path() -> None:
    """
    将最新版专检目录加入 sys.path，并清理同名包缓存，避免命名空间污染。
    """
    specialized_root = NUMBER_CHECK_LATEST_ROOT
    if not specialized_root.exists():
        raise FileNotFoundError(f"未找到数字专检依赖目录: {specialized_root}")

    specialized_root_str = str(specialized_root)
    package_roots = ("llm", "parsers", "replace", "utils", "revise", "backup_copy", "divide")

    for package_name in package_roots:
        mod = sys.modules.get(package_name)
        if mod is None:
            continue

        package_path: Optional[Path] = None
        try:
            spec = getattr(mod, "__spec__", None)
            if spec and spec.submodule_search_locations:
                package_path = Path(list(spec.submodule_search_locations)[0])
            elif hasattr(mod, "__path__"):
                package_path = Path(list(mod.__path__)[0])
            elif hasattr(mod, "__file__") and mod.__file__:
                package_path = Path(mod.__file__).parent
        except Exception:
            package_path = None

        is_correct = False
        if package_path is not None:
            try:
                package_path.relative_to(specialized_root)
                is_correct = True
            except ValueError:
                is_correct = False

        if is_correct:
            continue

        stale = [key for key in list(sys.modules.keys()) if key == package_name or key.startswith(f"{package_name}.")]
        for key in stale:
            del sys.modules[key]
        if stale:
            logger.info(f"[import] 清除了 {len(stale)} 个 {package_name}.* 缓存")

    if specialized_root_str not in sys.path:
        sys.path.insert(0, specialized_root_str)


def _count_report_errors(report_file: Path) -> int:
    if not report_file.exists():
        return 0
    try:
        payload = json.loads(report_file.read_text(encoding="utf-8"))
    except Exception:
        return 0
    return len(payload) if isinstance(payload, list) else 0


def _copy_report_outputs(report_paths: Dict[str, str], output_dir: Path, folder_name: str) -> tuple[Dict[str, str], Dict[str, int]]:
    reports: Dict[str, str] = {}
    counts: Dict[str, int] = {value: 0 for value in SECTION_COUNT_KEYS.values()}

    for section_name, source_path in (report_paths or {}).items():
        if not source_path:
            continue
        source = Path(source_path)
        if not source.exists():
            continue

        section_dir = SECTION_DIR_NAMES.get(section_name, section_name)
        target_dir = output_dir / section_dir / "output_json"
        target_dir.mkdir(parents=True, exist_ok=True)
        target = target_dir / source.name
        try:
            same_file = source.resolve() == target.resolve()
        except Exception:
            same_file = str(source) == str(target)
        if not same_file:
            shutil.copy2(source, target)

        report_key = SECTION_REPORT_KEYS.get(section_name, f"{section_dir}_json")
        count_key = SECTION_COUNT_KEYS.get(section_name, f"{section_dir}_issues")
        reports[report_key] = f"outputs/number_check/{folder_name}/{section_dir}/output_json/{target.name}"
        counts[count_key] = _count_report_errors(target)

    return reports, counts


def _extract_quoted_segments(text: str) -> List[str]:
    if not text:
        return []

    patterns = [
        r"“([^”]+)”",
        r'"([^"]+)"',
        r"‘([^’]+)’",
        r"'([^']+)'",
    ]
    result: List[str] = []
    for pattern in patterns:
        result.extend(item.strip() for item in re.findall(pattern, text) if item and item.strip())
    return result


def _cleanup_replacement_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    return text.strip()


def _apply_replace_pair(base_text: str, old_fragment: str, new_fragment: str) -> str:
    if not base_text or not old_fragment:
        return ""
    if old_fragment in base_text:
        return _cleanup_replacement_text(base_text.replace(old_fragment, new_fragment, 1))
    return ""


def _apply_delete_fragment(base_text: str, fragment: str) -> str:
    if not base_text or not fragment:
        return ""

    direct = _apply_replace_pair(base_text, fragment, "")
    if direct:
        return direct

    candidates = [
        f" {fragment}",
        f"{fragment} ",
        f" {fragment} ",
    ]
    for candidate in candidates:
        direct = _apply_replace_pair(base_text, candidate, " ")
        if direct:
            return direct

    return ""


def _apply_insert_fragment(base_text: str, anchor: str, fragment: str, position: str) -> str:
    if not base_text or not anchor or not fragment or anchor not in base_text:
        return ""
    if position == "after":
        replaced = base_text.replace(anchor, f"{anchor} {fragment}", 1)
    else:
        replaced = base_text.replace(anchor, f"{fragment} {anchor}", 1)
    return _cleanup_replacement_text(replaced)


def _derive_new_text_from_suggestion(translated_text: str, suggestion: str) -> str:
    translated_text = (translated_text or "").strip()
    suggestion = (suggestion or "").strip()
    if not translated_text or not suggestion:
        return ""

    quoted = _extract_quoted_segments(suggestion)

    pair_patterns = [
        r"将“([^”]+)”修改为“([^”]+)”",
        r'将"([^"]+)"修改为"([^"]+)"',
        r"把“([^”]+)”改为“([^”]+)”",
        r'把"([^"]+)"改为"([^"]+)"',
        r"把“([^”]+)”改成“([^”]+)”",
        r'把"([^"]+)"改成"([^"]+)"',
        r"由“([^”]+)”改为“([^”]+)”",
        r'由"([^"]+)"改为"([^"]+)"',
    ]
    for pattern in pair_patterns:
        match = re.search(pattern, suggestion)
        if match:
            old_fragment = match.group(1).strip()
            new_fragment = match.group(2).strip()
            replaced = _apply_replace_pair(translated_text, old_fragment, new_fragment)
            return replaced or new_fragment

    delete_patterns = [
        r"删除多译的“([^”]+)”",
        r'删除多译的"([^"]+)"',
        r"删除“([^”]+)”",
        r'删除"([^"]+)"',
        r"删去“([^”]+)”",
        r'删去"([^"]+)"',
        r"去掉“([^”]+)”",
        r'去掉"([^"]+)"',
    ]
    for pattern in delete_patterns:
        match = re.search(pattern, suggestion)
        if match:
            fragment = match.group(1).strip()
            deleted = _apply_delete_fragment(translated_text, fragment)
            return deleted or translated_text

    add_patterns = [
        (r"在“([^”]+)”后(?:添加|增加|补充)“([^”]+)”", "after"),
        (r'在"([^"]+)"后(?:添加|增加|补充)"([^"]+)"', "after"),
        (r"在“([^”]+)”前(?:添加|增加|补充)“([^”]+)”", "before"),
        (r'在"([^"]+)"前(?:添加|增加|补充)"([^"]+)"', "before"),
    ]
    for pattern, position in add_patterns:
        match = re.search(pattern, suggestion)
        if match:
            anchor = match.group(1).strip()
            fragment = match.group(2).strip()
            inserted = _apply_insert_fragment(translated_text, anchor, fragment, position)
            if inserted:
                return inserted

    standalone_patterns = [
        r"修改为“([^”]+)”",
        r'修改为"([^"]+)"',
        r"改为“([^”]+)”",
        r'改为"([^"]+)"',
        r"建议修改为“([^”]+)”",
        r'建议修改为"([^"]+)"',
        r"建议改为“([^”]+)”",
        r'建议改为"([^"]+)"',
        r"建议补充译文，例如修改为“([^”]+)”",
        r'建议补充译文，例如修改为"([^"]+)"',
        r"建议译为“([^”]+)”",
        r'建议译为"([^"]+)"',
        r"应修改为“([^”]+)”",
        r'应修改为"([^"]+)"',
        r"应为“([^”]+)”",
        r'应为"([^"]+)"',
    ]
    for pattern in standalone_patterns:
        match = re.search(pattern, suggestion)
        if match:
            candidate = match.group(1).strip()
            if not candidate:
                continue
            if len(candidate) >= max(6, int(len(translated_text) * 0.6)):
                return candidate

    if len(quoted) >= 2 and quoted[-2] in translated_text:
        replaced = _apply_replace_pair(translated_text, quoted[-2], quoted[-1])
        if replaced:
            return replaced

    if len(quoted) == 1 and len(quoted[0]) >= max(6, int(len(translated_text) * 0.6)):
        return quoted[0]

    return ""


def _normalize_error_item(error: Dict[str, Any], index: int) -> Dict[str, Any]:
    normalized = dict(error or {})

    error_type = (
        normalized.get("错误类型")
        or normalized.get("error_type")
        or normalized.get("type")
        or ""
    )
    original_text = (
        normalized.get("原文数值")
        or normalized.get("original_text")
        or normalized.get("原文")
        or ""
    )
    translated_text = (
        normalized.get("译文数值")
        or normalized.get("translated_text")
        or normalized.get("trans_text")
        or normalized.get("译文")
        or ""
    )
    suggestion_text = (
        normalized.get("译文修改建议值")
        or normalized.get("correction_suggestion")
        or normalized.get("suggestion")
        or ""
    )
    direct_reason = (
        normalized.get("修改理由")
        or normalized.get("reason")
        or normalized.get("correction_suggestion")
        or ""
    )
    translated_context = (
        normalized.get("译文上下文")
        or normalized.get("translated_context")
        or translated_text
    )
    original_context = (
        normalized.get("原文上下文")
        or normalized.get("original_context")
        or original_text
    )
    anchor_text = (
        normalized.get("替换锚点")
        or normalized.get("anchor_text")
        or translated_text
    )

    derived_new_text = (normalized.get("译文修改建议值") or "").strip()
    if not derived_new_text and translated_text and suggestion_text:
        derived_new_text = _derive_new_text_from_suggestion(translated_text, suggestion_text)

    normalized["错误编号"] = str(
        normalized.get("错误编号")
        or normalized.get("error_id")
        or normalized.get("id")
        or index
    )
    normalized["错误类型"] = str(error_type).strip()
    normalized["原文数值"] = str(original_text).strip()
    normalized["译文数值"] = str(translated_text).strip()
    normalized["译文修改建议值"] = str(derived_new_text or suggestion_text).strip()
    normalized["修改理由"] = str(direct_reason).strip()
    normalized["原文上下文"] = str(original_context).strip()
    normalized["译文上下文"] = str(translated_context).strip()
    normalized["替换锚点"] = str(anchor_text).strip()
    return normalized


def _normalize_error_list(errors: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return [_normalize_error_item(error, idx) for idx, error in enumerate(errors or [], 1)]


def _apply_all_fixes(
    doc: Document,
    errors: List[Dict[str, Any]],
    label: str,
    region: str,
    replace_func,
    processor,
    task_id: Optional[str] = None,
    progress_callback: Optional[Callable[[int, int, str, List[str]], None]] = None,
) -> Dict[str, int]:
    if not errors:
        if task_id:
            _emit_log(task_id, f"[fix] {label} 无需修复")
        return {"success": 0, "failed": 0, "skipped": 0}

    success = 0
    failed = 0
    skipped = 0
    total = len(errors)
    if task_id:
        _emit_log(task_id, f"[fix] >>> 正在修复 {label}（共 {total} 条）...")

    for idx, raw_error in enumerate(errors, 1):
        error = _normalize_error_item(raw_error, idx)
        old = (error.get("译文数值") or "").strip()
        new = (error.get("译文修改建议值") or "").strip()
        reason = (error.get("修改理由") or "").strip()
        context = error.get("译文上下文", "") or ""
        anchor = error.get("替换锚点", "") or ""

        if not old or not new:
            skipped += 1
            if task_id:
                _emit_log(task_id, f"[fix]   [{idx}/{total}] 跳过: 缺少【译文数值】或【译文修改建议值】", level="warning")
        else:
            ok, strategy = replace_func(
                doc,
                old,
                new,
                reason,
                processor,
                context=context,
                anchor_text=anchor,
                region=region,
            )
            if ok:
                success += 1
                if task_id:
                    _emit_log(task_id, f"[fix]   [{idx}/{total}] ✓ '{old}' → '{new}'  策略={strategy}  理由={reason}")
            else:
                failed += 1
                if task_id:
                    _emit_log(task_id, f"[fix]   [{idx}/{total}] ✗ 未匹配到 '{old}'", level="warning")

        if progress_callback and (idx % max(1, total // 10) == 0 or idx == total):
            details = [
                f"{label}: 成功 {success}，失败 {failed}，跳过 {skipped}",
                f"进度: {idx}/{total}",
            ]
            progress_callback(idx, total, f"正在修复{label} ({idx}/{total})", details)

    rate = success / (success + failed) if (success + failed) else 0
    if task_id:
        _emit_log(
            task_id,
            f"[fix] --- {label} 修复统计: 成功={success} 失败={failed} 跳过={skipped} 成功率={rate:.0%} ---",
        )
    return {"success": success, "failed": failed, "skipped": skipped}


def _run_number_check_sync(
    task_id: str,
    original_path: Path,
    translated_path: Path,
    output_dir: Path,
    report_dir: Path,
    gemini_route: str,
    model_name: str,
    original_filename: str,
    translated_filename: str,
) -> Dict[str, Any]:
    """双文件模式：新版数检的同步执行链路。"""
    import time as _time

    effective_route = GEMINI_ROUTE_OPENROUTER
    resolved_model_name = resolve_model_for_route(model_name, effective_route)
    converted_dir = output_dir / "converted_inputs"

    if original_path.suffix.lower() == ".doc" or translated_path.suffix.lower() == ".doc":
        _update_progress(task_id, 2, 7, "检测到 .doc 文件，正在转换为 .docx...")
        original_path = _convert_doc_input_if_needed(original_path, converted_dir, "original")
        translated_path = _convert_doc_input_if_needed(translated_path, converted_dir, "translated")

    with _specialist_import_lock:
        _prepare_specialized_import_path()
        if not settings.OPENROUTER_API_KEY:
            raise ValueError("未配置 OPENROUTER_API_KEY 或 OPENAI_API_KEY，无法执行数字专检。")

        os.environ["OPENROUTER_API_KEY"] = settings.OPENROUTER_API_KEY
        os.environ["OPENAI_API_KEY"] = settings.OPENROUTER_API_KEY
        os.environ["API_KEY"] = settings.OPENROUTER_API_KEY
        os.environ["NUMBER_CHECK_MODEL"] = resolved_model_name
        if settings.OPENROUTER_BASE_URL:
            os.environ["OPENROUTER_BASE_URL"] = settings.OPENROUTER_BASE_URL
            os.environ["OPENAI_BASE_URL"] = settings.OPENROUTER_BASE_URL
            os.environ["BASE_URL"] = settings.OPENROUTER_BASE_URL

        _update_progress(task_id, 2, 7, "正在加载最新版数检模块...")
        from llm.check import Match
        from parsers.word.body_extractor import extract_body_text
        from parsers.word.footer_extractor import extract_footers
        from parsers.word.header_extractor import extract_headers
        from backup_copy.backup_manager import ensure_backup_copy
        from replace.replace_revision import replace_and_revise_in_docx
        from revise.revision import RevisionManager
        from parsers.json.clean_json import extract_and_parse, parse_json_content
        from utils.json_files import write_json_with_timestamp

    if gemini_route != effective_route:
        _emit_log(task_id, f"[double] 双文件模式固定走 route={effective_route}，已忽略 route={gemini_route}", level="warning")

    _update_progress(task_id, 3, 7, "正在提取文档文本...")

    original_body = extract_body_text(str(original_path))
    translated_body = extract_body_text(str(translated_path))
    original_header_raw = extract_headers(str(original_path))
    translated_header_raw = extract_headers(str(translated_path))
    original_footer_raw = extract_footers(str(original_path))
    translated_footer_raw = extract_footers(str(translated_path))

    original_header = "\n".join(original_header_raw) if isinstance(original_header_raw, list) else (original_header_raw or "")
    translated_header = "\n".join(translated_header_raw) if isinstance(translated_header_raw, list) else (translated_header_raw or "")
    original_footer = "\n".join(original_footer_raw) if isinstance(original_footer_raw, list) else (original_footer_raw or "")
    translated_footer = "\n".join(translated_footer_raw) if isinstance(translated_footer_raw, list) else (translated_footer_raw or "")

    _emit_log(task_id, f"[double] route={effective_route}, model={resolved_model_name}")

    _update_progress(task_id, 4, 7, "正在对比数值差异...")
    matcher = Match(model_name=resolved_model_name)
    report_output_dirs = {
        "正文": _specialized_output_dir("zhengwen", output_dir),
        "页眉": _specialized_output_dir("yemei", output_dir),
        "页脚": _specialized_output_dir("yejiao", output_dir),
    }
    parts = [
        ("正文", original_body, translated_body, report_output_dirs["正文"]),
        ("页眉", original_header, translated_header, report_output_dirs["页眉"]),
        ("页脚", original_footer, translated_footer, report_output_dirs["页脚"]),
    ]

    report_paths: Dict[str, str] = {}
    for idx, (name, original_text, translated_text, output_subdir) in enumerate(parts, 1):
        output_subdir.mkdir(parents=True, exist_ok=True)
        _update_progress(task_id, 4, 7, f"正在对比{name} ({idx}/3)...")
        if original_text and translated_text:
            t0 = _time.time()
            raw_result = matcher.compare_texts(original_text, translated_text)
            result = parse_json_content(raw_result) if raw_result else []
            elapsed = _time.time() - t0
            _emit_log(task_id, f"[compare] {name}检查完成，耗时 {elapsed:.1f}s")
        else:
            result = []
            _emit_log(task_id, f"[skip] {name}原文或译文为空，跳过模型比对", level="warning")
        _, json_path = write_json_with_timestamp(result, str(output_subdir))
        report_paths[name] = json_path

    _update_progress(task_id, 5, 7, "正在加载检查报告...")
    body_errors = _normalize_error_list(extract_and_parse(report_paths.get("正文")))
    header_errors = _normalize_error_list(extract_and_parse(report_paths.get("页眉")))
    footer_errors = _normalize_error_list(extract_and_parse(report_paths.get("页脚")))

    _emit_log(task_id, f"[report] 正文 {len(body_errors)} 条，页眉 {len(header_errors)} 条，页脚 {len(footer_errors)} 条")

    _update_progress(task_id, 6, 7, "正在生成修订版文档...")
    backup_copy_path = ensure_backup_copy(str(translated_path))
    doc = Document(backup_copy_path)
    revision_manager = RevisionManager(doc, author="翻译校对")

    def progress_callback(current: int, total: int, message: str, details: List[str]) -> None:
        _update_progress(task_id, 6, 7, message, details)

    body_stat = _apply_all_fixes(doc, body_errors, "正文", "body", replace_and_revise_in_docx, revision_manager, task_id=task_id, progress_callback=progress_callback)
    header_stat = _apply_all_fixes(doc, header_errors, "页眉", "header", replace_and_revise_in_docx, revision_manager, task_id=task_id, progress_callback=progress_callback)
    footer_stat = _apply_all_fixes(doc, footer_errors, "页脚", "footer", replace_and_revise_in_docx, revision_manager, task_id=task_id, progress_callback=progress_callback)

    doc.save(backup_copy_path)
    final_doc_path = ensure_unique_path(
        output_dir / build_user_visible_filename(translated_filename, suffix="corrected", ext=".docx")
    )
    shutil.copy2(backup_copy_path, final_doc_path)

    copied_reports, report_counts = _copy_report_outputs(report_paths, output_dir, output_dir.name)
    total_success = body_stat["success"] + header_stat["success"] + footer_stat["success"]
    total_failed = body_stat["failed"] + header_stat["failed"] + footer_stat["failed"]
    total_skipped = body_stat["skipped"] + header_stat["skipped"] + footer_stat["skipped"]

    result = {
        "task_id": task_id,
        "mode": NUMBER_CHECK_MODE_DOUBLE,
        "original_filename": original_filename,
        "translated_filename": translated_filename,
        "model_name": resolved_model_name,
        "corrected_docx": f"outputs/number_check/{output_dir.name}/{final_doc_path.name}",
        "reports": copied_reports,
        "report_counts": report_counts,
        "stats": {
            "success": total_success,
            "failed": total_failed,
            "skipped": total_skipped,
        },
        "summary": "双文件模式已完成检查并生成修订版文档。",
    }
    _complete_task(task_id, result=result)
    return result


def _run_number_check_single_sync(
    task_id: str,
    single_path: Path,
    output_dir: Path,
    gemini_route: str,
    model_name: str,
    single_filename: str,
) -> Dict[str, Any]:
    """单文件模式：docx 生成修订版，其它格式生成报告。"""
    effective_route = GEMINI_ROUTE_OPENROUTER
    resolved_model_name = resolve_model_for_route(model_name, effective_route)
    if single_path.suffix.lower() == ".doc":
        _update_progress(task_id, 2, 6, "检测到 .doc 文件，正在转换为 .docx...")
        single_path = _convert_doc_input_if_needed(single_path, output_dir / "converted_inputs", "single")

    total_steps = 6 if single_path.suffix.lower() == ".docx" else 5

    with _specialist_import_lock:
        _prepare_specialized_import_path()
        if not settings.OPENROUTER_API_KEY:
            raise ValueError("单文件模式需要配置 OPENROUTER_API_KEY 或 OPENAI_API_KEY。")

        os.environ["OPENROUTER_API_KEY"] = settings.OPENROUTER_API_KEY
        os.environ["OPENAI_API_KEY"] = settings.OPENROUTER_API_KEY
        os.environ["API_KEY"] = settings.OPENROUTER_API_KEY
        os.environ["NUMBER_CHECK_SINGLE_MODEL"] = resolved_model_name
        if settings.OPENROUTER_BASE_URL:
            os.environ["OPENROUTER_BASE_URL"] = settings.OPENROUTER_BASE_URL
            os.environ["OPENAI_BASE_URL"] = settings.OPENROUTER_BASE_URL
            os.environ["BASE_URL"] = settings.OPENROUTER_BASE_URL

        _update_progress(task_id, 2, total_steps, "正在加载最新版单文件数检模块...")
        from llm.check_2 import run_bilingual_comparison
        from parsers.json.clean_json import extract_and_parse
        from backup_copy.backup_manager import ensure_backup_copy
        from replace.replace_revision import replace_and_revise_in_docx
        from revise.revision import RevisionManager

    if gemini_route != effective_route:
        _emit_log(task_id, f"[single] 单文件模式固定走 route={effective_route}，已忽略 route={gemini_route}", level="warning")

    _update_progress(task_id, 3, total_steps, "正在执行单文件双语对照检查...")
    _emit_log(task_id, f"[single] route={effective_route}, model={resolved_model_name}")
    report_output_dirs = {
        "正文": str(_specialized_output_dir("zhengwen", output_dir)),
        "页眉": str(_specialized_output_dir("yemei", output_dir)),
        "页脚": str(_specialized_output_dir("yejiao", output_dir)),
    }
    report_paths = run_bilingual_comparison(
        str(single_path),
        output_dirs=report_output_dirs,
        report_prefix="文本对比结果",
        model_name=resolved_model_name,
    )
    if report_paths is None:
        raise RuntimeError("单文件双语对照检查失败")

    _update_progress(task_id, 4, total_steps, "正在整理检查报告...")
    copied_reports, report_counts = _copy_report_outputs(report_paths, output_dir, output_dir.name)
    body_errors = _normalize_error_list(extract_and_parse(report_paths.get("正文")))
    header_errors = _normalize_error_list(extract_and_parse(report_paths.get("页眉")))
    footer_errors = _normalize_error_list(extract_and_parse(report_paths.get("页脚")))
    total_issues = sum(report_counts.values())

    result: Dict[str, Any] = {
        "task_id": task_id,
        "mode": NUMBER_CHECK_MODE_SINGLE,
        "single_filename": single_filename,
        "model_name": resolved_model_name,
        "reports": copied_reports,
        "report_counts": report_counts,
        "stats": {
            "total_issues": total_issues,
            "body_issues": report_counts.get("body_issues", 0),
            "header_issues": report_counts.get("header_issues", 0),
            "footer_issues": report_counts.get("footer_issues", 0),
        },
    }

    if single_path.suffix.lower() == ".docx":
        _update_progress(task_id, 5, total_steps, "正在生成单文件修订版文档...")
        backup_copy_path = ensure_backup_copy(str(single_path))
        doc = Document(backup_copy_path)
        revision_manager = RevisionManager(doc, author="翻译校对")

        body_stat = _apply_all_fixes(doc, body_errors, "正文", "body", replace_and_revise_in_docx, revision_manager, task_id=task_id)
        header_stat = _apply_all_fixes(doc, header_errors, "页眉", "header", replace_and_revise_in_docx, revision_manager, task_id=task_id)
        footer_stat = _apply_all_fixes(doc, footer_errors, "页脚", "footer", replace_and_revise_in_docx, revision_manager, task_id=task_id)

        doc.save(backup_copy_path)
        final_doc_path = ensure_unique_path(
            output_dir / build_user_visible_filename(single_filename, suffix="corrected", ext=".docx")
        )
        shutil.copy2(backup_copy_path, final_doc_path)

        result["corrected_docx"] = f"outputs/number_check/{output_dir.name}/{final_doc_path.name}"
        result["fix_stats"] = {
            "success": body_stat["success"] + header_stat["success"] + footer_stat["success"],
            "failed": body_stat["failed"] + header_stat["failed"] + footer_stat["failed"],
            "skipped": body_stat["skipped"] + header_stat["skipped"] + footer_stat["skipped"],
        }
        result["summary"] = "单文件 DOCX 模式已生成检查报告，并输出修订版文档。"
    else:
        result["summary"] = "单文件模式已生成检查报告；当前文件格式暂不支持自动修订。"

    _complete_task(task_id, result=result)
    return result


async def run_number_check_task(
    original_file: Optional[UploadFile] = None,
    translated_file: Optional[UploadFile] = None,
    single_file: Optional[UploadFile] = None,
    mode: str = NUMBER_CHECK_MODE_DOUBLE,
    task_id: str = "",
    display_no: Optional[str] = None,
    gemini_route: str = "openrouter",
    model_name: str = "gemini-3.1-pro-preview",
) -> Dict[str, Any]:
    normalized_mode = _normalize_number_check_mode(mode)
    model_name = normalize_number_check_model(model_name)

    if normalized_mode == NUMBER_CHECK_MODE_SINGLE:
        if single_file is None:
            raise ValueError("单文件模式缺少待检查文件")
        _validate_single_file(single_file, "单文件")
        gemini_route = ensure_gemini_route_configured(GEMINI_ROUTE_OPENROUTER)
    else:
        if original_file is None or translated_file is None:
            raise ValueError("双文件模式缺少原文或译文文件")
        _validate_docx(original_file, "原文")
        _validate_docx(translated_file, "译文")
        gemini_route = ensure_gemini_route_configured(GEMINI_ROUTE_OPENROUTER)

    if not task_id:
        task_id = str(uuid.uuid4())

    single_ext = os.path.splitext(single_file.filename or "")[1].lower() if single_file else ""
    total_steps = 6 if normalized_mode == NUMBER_CHECK_MODE_SINGLE and _supports_revised_doc_output(single_ext) else (5 if normalized_mode == NUMBER_CHECK_MODE_SINGLE else 7)
    _init_task_progress(task_id, total_steps=total_steps)
    _emit_log(task_id, f"[config] mode={normalized_mode}, route={gemini_route}, model={model_name}")

    folder_name = display_no or task_id
    upload_dir = Path(settings.UPLOAD_DIR) / "number_check" / folder_name
    output_dir = Path(settings.OUTPUT_DIR) / "number_check" / folder_name
    upload_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    _update_progress(task_id, 1, total_steps, "正在保存上传文件...")

    try:
        if normalized_mode == NUMBER_CHECK_MODE_SINGLE:
            single_path = upload_dir / f"single{single_ext or '.docx'}"
            with open(single_path, "wb") as f:
                f.write(await single_file.read())
            return await asyncio.to_thread(
                _run_number_check_single_sync,
                task_id,
                single_path,
                output_dir,
                gemini_route,
                model_name,
                single_file.filename or single_path.name,
            )

        original_ext = os.path.splitext(original_file.filename or "")[1].lower() or ".docx"
        translated_ext = os.path.splitext(translated_file.filename or "")[1].lower() or ".docx"
        original_path = upload_dir / f"original{original_ext}"
        translated_path = upload_dir / f"translated{translated_ext}"
        with open(original_path, "wb") as f:
            f.write(await original_file.read())
        with open(translated_path, "wb") as f:
            f.write(await translated_file.read())

        return await asyncio.to_thread(
            _run_number_check_sync,
            task_id,
            original_path,
            translated_path,
            output_dir,
            output_dir / "reports",
            gemini_route,
            model_name,
            original_file.filename or "original.docx",
            translated_file.filename or "translated.docx",
        )
    except Exception as exc:
        _emit_log(task_id, f"[error] 任务失败: {type(exc).__name__}: {exc}", level="error")
        _complete_task(task_id, error=str(exc))
        raise







