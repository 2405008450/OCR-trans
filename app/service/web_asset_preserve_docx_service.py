import copy
import json
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw

from app.service.chat_preserve_docx_service import (
    _bbox_area,
    _bbox_to_list,
    _clean_text,
    _coerce_bbox,
    _crop_to_file,
    _expand_bbox,
    _intersect_bbox,
    _load_input_pages,
    _offset_bbox,
    _parse_json_object,
    _refine_asset_bbox,
)
from app.service.gemini_service import generate_vision_html


WEB_LAYOUT_SYSTEM_PROMPT = """You analyze screenshots of web pages and return strict JSON only.

Return one JSON object and no markdown/code fences/explanations.

Image coordinates:
- Use absolute pixel coordinates relative to the uploaded image.
- Every bbox must be [left, top, right, bottom].
- If coordinates are uncertain, set bbox to null instead of guessing.

Goal:
- Reconstruct the page in reading order.
- Keep text editable.
- Preserve only important image assets that carry content or evidence.

Preserve image assets when they are:
- product photos, charts, graphs, dashboards, diagrams, flowcharts, maps, QR codes, screenshots inside the page, evidence photos, primary logos, or content illustrations.

Ignore image assets when they are:
- decorative backgrounds, tiny navigation icons, button icons, social icons, arrows, separators, generic avatars, stock decoration, or repeated chrome.

Schema:
{
  "blocks": [
    {
      "type": "text",
      "order": 1,
      "text": "visible text, preserving line breaks",
      "bbox": [left, top, right, bottom] or null,
      "role": "heading|paragraph|caption|metadata|button|other"
    },
    {
      "type": "image",
      "order": 2,
      "bbox": [left, top, right, bottom],
      "asset_type": "product|chart|diagram|map|qr|screenshot|logo|photo|illustration|other",
      "importance": "high|medium|low",
      "reason": "why this image should be preserved",
      "caption": "nearby caption or empty string"
    }
  ]
}

Rules:
- Do not include decorative or low-value images unless they are central to understanding the page.
- Do not include ordinary editable text as images.
- Keep ordering top-to-bottom and left-to-right.
- Prefer fewer, high-value image blocks over many small icons.
"""


WEB_LOCAL_ASSET_SYSTEM_PROMPT = """You locate one important web-page image asset inside a cropped region.

Return strict JSON only, using coordinates relative to the cropped image.

Schema:
{
  "bbox": [left, top, right, bottom] or null,
  "asset_type": "product|chart|diagram|map|qr|screenshot|logo|photo|illustration|other",
  "importance": "high|medium|low",
  "reason": "short reason"
}

Rules:
- Locate the main important image asset only.
- Ignore nearby text, icons, buttons, and decorative background.
- Return null bbox if the crop does not contain a meaningful image asset.
"""


WEB_LOCAL_ASSET_RECHECK_MAX_BLOCKS = 8
WEB_RECOMMENDED_MODEL = "anthropic/claude-sonnet-5"
PRESERVED_IMPORTANCE = {"high", "medium"}


@dataclass
class WebAssetPreserveResult:
    raw_text: str
    total_pages: int
    asset_count: int
    fallback_count: int
    layout: dict[str, Any]


@dataclass
class WebPage:
    page_no: int
    image: Image.Image
    layout: dict[str, Any]


@dataclass
class _RenderResult:
    raw_text: str
    asset_count: int
    fallback_count: int
    debug_overlay_paths: list[Path]


StatusCallback = Callable[[str], None]


def convert_web_screenshot_to_docx(
    *,
    input_path: str | Path,
    output_docx_path: str | Path,
    layout_json_path: str | Path,
    assets_dir: str | Path,
    model: str,
    gemini_route: str,
    status_callback: StatusCallback | None = None,
) -> WebAssetPreserveResult:
    """将网页截图/PDF转换为保留重要图片的可编辑 Word。"""
    input_file = Path(input_path)
    output_docx = Path(output_docx_path)
    layout_json = Path(layout_json_path)
    asset_root = Path(assets_dir)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    layout_json.parent.mkdir(parents=True, exist_ok=True)
    asset_root.mkdir(parents=True, exist_ok=True)

    pages = _load_input_pages(input_file)
    web_pages: list[WebPage] = []
    full_layout: dict[str, Any] = {"mode": "web_asset_preserve", "pages": []}

    for page_index, image in enumerate(pages, start=1):
        _emit(status_callback, f"正在分析网页截图第 {page_index}/{len(pages)} 页")
        layout = _analyze_web_page(
            image=image,
            page_no=page_index,
            model=model,
            gemini_route=gemini_route,
        )
        layout = _refine_web_layout_with_local_asset_pass(
            image=image,
            layout=layout,
            page_no=page_index,
            model=model,
            gemini_route=gemini_route,
            status_callback=status_callback,
        )
        web_pages.append(WebPage(page_no=page_index, image=image, layout=layout))

    render_result = render_web_asset_preserve_docx_from_pages(
        pages=web_pages,
        output_docx_path=output_docx,
        assets_dir=asset_root,
    )
    full_layout["pages"] = [
        {
            "page_no": page.page_no,
            "image_width": page.image.width,
            "image_height": page.image.height,
            "layout": page.layout,
        }
        for page in web_pages
    ]
    full_layout["render"] = {
        "asset_count": render_result.asset_count,
        "fallback_count": render_result.fallback_count,
        "debug_overlays": [str(path).replace("\\", "/") for path in render_result.debug_overlay_paths],
    }
    layout_json.write_text(json.dumps(full_layout, ensure_ascii=False, indent=2), encoding="utf-8")

    return WebAssetPreserveResult(
        raw_text=render_result.raw_text,
        total_pages=len(web_pages),
        asset_count=render_result.asset_count,
        fallback_count=render_result.fallback_count,
        layout=full_layout,
    )


def render_web_asset_preserve_docx_from_layout(
    *,
    image_path: str | Path,
    layout: dict[str, Any],
    output_docx_path: str | Path,
    assets_dir: str | Path,
) -> WebAssetPreserveResult:
    image = Image.open(image_path).convert("RGB")
    page = WebPage(page_no=1, image=image, layout=layout)
    render_result = render_web_asset_preserve_docx_from_pages(
        pages=[page],
        output_docx_path=output_docx_path,
        assets_dir=assets_dir,
    )
    full_layout = {
        "mode": "web_asset_preserve",
        "pages": [{"page_no": 1, "layout": page.layout}],
        "render": {
            "asset_count": render_result.asset_count,
            "fallback_count": render_result.fallback_count,
            "debug_overlays": [str(path).replace("\\", "/") for path in render_result.debug_overlay_paths],
        },
    }
    return WebAssetPreserveResult(
        raw_text=render_result.raw_text,
        total_pages=1,
        asset_count=render_result.asset_count,
        fallback_count=render_result.fallback_count,
        layout=full_layout,
    )


def render_web_asset_preserve_docx_from_pages(
    *,
    pages: Iterable[WebPage],
    output_docx_path: str | Path,
    assets_dir: str | Path,
) -> _RenderResult:
    document = Document()
    _setup_document(document)
    asset_root = Path(assets_dir)
    asset_root.mkdir(parents=True, exist_ok=True)

    raw_sections: list[str] = []
    asset_count = 0
    fallback_count = 0
    debug_overlay_paths: list[Path] = []

    for page_index, page in enumerate(list(pages), start=1):
        if page_index > 1:
            document.add_page_break()

        page.layout = _refine_web_layout_assets(page.image, page.layout)
        overlay_path = asset_root / f"debug_overlay_page_{page.page_no:03d}.png"
        _save_web_debug_overlay(page.image, page.layout, overlay_path)
        debug_overlay_paths.append(overlay_path)

        blocks = _sorted_blocks(page.layout)
        raw_sections.append(f"--- Page {page.page_no} ---")
        if not blocks:
            fallback_path = _crop_to_file(
                image=page.image,
                bbox=(0, 0, page.image.width, page.image.height),
                output_path=asset_root / f"page_{page.page_no:03d}_full.png",
            )
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(fallback_path), width=Inches(6.2))
            asset_count += 1
            fallback_count += 1
            continue

        for block_index, block in enumerate(blocks, start=1):
            if _block_type(block) == "image":
                rendered = _render_image_block(
                    document=document,
                    image=page.image,
                    block=block,
                    page_no=page.page_no,
                    block_index=block_index,
                    assets_dir=asset_root,
                )
                asset_count += rendered["asset_count"]
                fallback_count += rendered["fallback_count"]
                raw_text = _image_block_to_raw_text(block)
                if raw_text:
                    raw_sections.append(raw_text)
                continue

            text = _clean_text(block.get("text"))
            if text:
                _render_text_block(document, block, text)
                raw_sections.append(text)

    document.save(str(output_docx_path))
    return _RenderResult(
        raw_text="\n\n".join(item for item in raw_sections if item).strip(),
        asset_count=asset_count,
        fallback_count=fallback_count,
        debug_overlay_paths=debug_overlay_paths,
    )


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)


def _render_text_block(document: Document, block: dict[str, Any], text: str) -> None:
    role = _clean_text(block.get("role")).lower()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    lines = text.splitlines() or [text]
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        if role == "heading":
            run.bold = True
            run.font.size = Pt(14)
        elif role in {"caption", "metadata"}:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(105, 112, 122)


def _render_image_block(
    *,
    document: Document,
    image: Image.Image,
    block: dict[str, Any],
    page_no: int,
    block_index: int,
    assets_dir: Path,
) -> dict[str, int]:
    bbox = _coerce_bbox(block.get("bbox_refined") or block.get("bbox"), image.size)
    if not bbox:
        return {"asset_count": 0, "fallback_count": 1}

    crop_path = _crop_to_file(
        image=image,
        bbox=bbox,
        output_path=assets_dir / f"p{page_no:03d}_b{block_index:03d}_web_asset.png",
        padding=1,
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(crop_path), width=Inches(_image_width_inches(bbox)))

    caption = _clean_text(block.get("caption"))
    reason = _clean_text(block.get("reason"))
    if caption or reason:
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.add_run(caption or reason)
        caption_run.font.size = Pt(8.5)
        caption_run.font.color.rgb = RGBColor(105, 112, 122)

    return {"asset_count": 1, "fallback_count": 0}


def _image_width_inches(bbox: tuple[int, int, int, int]) -> float:
    width = max(bbox[2] - bbox[0], 1)
    return min(max(width / 150, 1.0), 6.2)


def _analyze_web_page(
    *,
    image: Image.Image,
    page_no: int,
    model: str,
    gemini_route: str,
) -> dict[str, Any]:
    image_bytes = BytesIO()
    image.save(image_bytes, format="PNG")
    user_prompt = (
        f"Analyze this web-page screenshot page {page_no}. "
        f"The image size is width={image.width}px, height={image.height}px. "
        "Return strict JSON using the requested schema."
    )
    response = generate_vision_html(
        system_prompt=WEB_LAYOUT_SYSTEM_PROMPT,
        image_bytes=image_bytes.getvalue(),
        mime_type="image/png",
        model=model,
        route=gemini_route,
        user_prompt=user_prompt,
        temperature=0,
    )
    parsed = _parse_json_object(response)
    if not isinstance(parsed, dict):
        raise RuntimeError("网页截图布局识别结果不是 JSON 对象")
    return parsed


def _refine_web_layout_with_local_asset_pass(
    *,
    image: Image.Image,
    layout: dict[str, Any],
    page_no: int,
    model: str,
    gemini_route: str,
    status_callback: StatusCallback | None = None,
) -> dict[str, Any]:
    refined = _refine_web_layout_assets(image, layout)
    checked = 0
    failed = 0
    skipped = 0

    for block_index, block in enumerate(_image_blocks(refined), start=1):
        if checked >= WEB_LOCAL_ASSET_RECHECK_MAX_BLOCKS:
            skipped += 1
            continue
        bbox = _coerce_bbox(block.get("bbox"), image.size)
        refined_bbox = _coerce_bbox(block.get("bbox_refined"), image.size)
        if not bbox:
            skipped += 1
            continue
        high_value = _importance(block) == "high" or _bbox_area(bbox) >= image.width * image.height * 0.08
        needs_local = high_value and (
            not refined_bbox
            or block.get("bbox_refine_method") == "original"
            or _bbox_area(refined_bbox) < _bbox_area(bbox) * 0.25
        )
        if not needs_local:
            skipped += 1
            continue

        checked += 1
        _emit(status_callback, f"正在局部复核第 {page_no} 页第 {block_index} 个网页图片块")
        try:
            local_layout = _analyze_web_asset_local(
                image=image,
                bbox=bbox,
                page_no=page_no,
                block_index=block_index,
                model=model,
                gemini_route=gemini_route,
            )
        except Exception as exc:
            failed += 1
            block["_local_asset_error"] = f"{exc.__class__.__name__}: {exc}"
            continue

        local_bbox = _coerce_bbox(local_layout.get("bbox"), _local_crop_size(local_layout, bbox))
        crop_bbox = _coerce_bbox(local_layout.get("_crop_bbox"), (10**9, 10**9)) or bbox
        if local_bbox:
            global_bbox = _offset_bbox(local_bbox, crop_bbox[0], crop_bbox[1])
            block["bbox_model"] = block.get("bbox")
            block["bbox"] = _bbox_to_list(global_bbox)
            block["bbox_source"] = "local_asset_pass"
            block["asset_type"] = local_layout.get("asset_type") or block.get("asset_type") or "image"
            block["importance"] = local_layout.get("importance") or block.get("importance") or "high"
            if local_layout.get("reason"):
                block["reason"] = local_layout.get("reason")

    refined = _refine_web_layout_assets(image, refined)
    refined["_local_asset_pass"] = {
        "checked": checked,
        "failed": failed,
        "skipped": skipped,
        "max_blocks": WEB_LOCAL_ASSET_RECHECK_MAX_BLOCKS,
    }
    return refined


def _analyze_web_asset_local(
    *,
    image: Image.Image,
    bbox: tuple[int, int, int, int],
    page_no: int,
    block_index: int,
    model: str,
    gemini_route: str,
) -> dict[str, Any]:
    crop_bbox = _expand_bbox(bbox, image.size, padding=max(8, int(max(bbox[2] - bbox[0], bbox[3] - bbox[1]) * 0.08)))
    crop = image.crop(crop_bbox)
    image_bytes = BytesIO()
    crop.save(image_bytes, format="PNG")
    user_prompt = (
        f"Analyze cropped web asset candidate {block_index} on page {page_no}. "
        f"The cropped image size is width={crop.width}px, height={crop.height}px. "
        "Return strict JSON using the requested schema."
    )
    response = generate_vision_html(
        system_prompt=WEB_LOCAL_ASSET_SYSTEM_PROMPT,
        image_bytes=image_bytes.getvalue(),
        mime_type="image/png",
        model=model,
        route=gemini_route,
        user_prompt=user_prompt,
        temperature=0,
    )
    parsed = _parse_json_object(response)
    if not isinstance(parsed, dict):
        raise RuntimeError("网页局部图片识别结果不是 JSON 对象")
    parsed["_crop_bbox"] = _bbox_to_list(crop_bbox)
    return parsed


def _refine_web_layout_assets(image: Image.Image, layout: dict[str, Any]) -> dict[str, Any]:
    refined = copy.deepcopy(layout if isinstance(layout, dict) else {})
    for block in _image_blocks(refined):
        bbox = _coerce_bbox(block.get("bbox"), image.size)
        if not bbox:
            continue
        block["bbox_original"] = _bbox_to_list(bbox)
        snapped = _refine_asset_bbox(image=image, bbox=bbox, kind=str(block.get("asset_type") or "image"))
        block["bbox_refined"] = _bbox_to_list(snapped or bbox)
        block["bbox_refine_method"] = "edge_snap" if snapped else "original"
    refined["_asset_refine"] = {
        "method": "edge_snap_with_optional_local_asset_pass",
        "image_width": image.width,
        "image_height": image.height,
    }
    return refined


def _sorted_blocks(layout: dict[str, Any]) -> list[dict[str, Any]]:
    blocks = layout.get("blocks") if isinstance(layout, dict) else None
    if not isinstance(blocks, list):
        return []
    return sorted(
        [block for block in blocks if isinstance(block, dict) and _should_render_block(block)],
        key=_block_sort_key,
    )


def _image_blocks(layout: dict[str, Any]) -> list[dict[str, Any]]:
    blocks = layout.get("blocks") if isinstance(layout, dict) else None
    if not isinstance(blocks, list):
        return []
    return [block for block in blocks if isinstance(block, dict) and _block_type(block) == "image" and _should_preserve_image(block)]


def _should_render_block(block: dict[str, Any]) -> bool:
    if _block_type(block) == "image":
        return _should_preserve_image(block)
    return bool(_clean_text(block.get("text")))


def _should_preserve_image(block: dict[str, Any]) -> bool:
    return _importance(block) in PRESERVED_IMPORTANCE and _coerce_bbox(block.get("bbox"), (10**9, 10**9)) is not None


def _block_sort_key(block: dict[str, Any]) -> tuple[float, float, int]:
    order = block.get("order")
    try:
        order_value = int(order)
    except (TypeError, ValueError):
        order_value = 999999
    bbox = block.get("bbox_refined") or block.get("bbox")
    if isinstance(bbox, (list, tuple)) and len(bbox) == 4:
        try:
            return order_value, float(bbox[1]), float(bbox[0])
        except (TypeError, ValueError):
            pass
    return order_value, 999999.0, 999999.0


def _block_type(block: dict[str, Any]) -> str:
    return _clean_text(block.get("type")).lower()


def _importance(block: dict[str, Any]) -> str:
    return _clean_text(block.get("importance")).lower() or "medium"


def _image_block_to_raw_text(block: dict[str, Any]) -> str:
    pieces = []
    caption = _clean_text(block.get("caption"))
    reason = _clean_text(block.get("reason"))
    asset_type = _clean_text(block.get("asset_type"))
    if caption:
        pieces.append(caption)
    if reason:
        pieces.append(f"[保留图片: {asset_type or 'image'} - {reason}]")
    return "\n".join(pieces)


def _local_crop_size(local_layout: dict[str, Any], fallback_bbox: tuple[int, int, int, int]) -> tuple[int, int]:
    crop_bbox = _coerce_bbox(local_layout.get("_crop_bbox"), (10**9, 10**9))
    if crop_bbox:
        return crop_bbox[2] - crop_bbox[0], crop_bbox[3] - crop_bbox[1]
    return fallback_bbox[2] - fallback_bbox[0], fallback_bbox[3] - fallback_bbox[1]


def _save_web_debug_overlay(image: Image.Image, layout: dict[str, Any], output_path: Path) -> None:
    overlay = image.convert("RGB").copy()
    draw = ImageDraw.Draw(overlay)
    for index, block in enumerate(_sorted_blocks(layout), start=1):
        if _block_type(block) == "image":
            original = _coerce_bbox(block.get("bbox_original") or block.get("bbox"), overlay.size)
            refined = _coerce_bbox(block.get("bbox_refined") or block.get("bbox"), overlay.size)
            if original:
                _draw_bbox(draw, original, "#F2994A", f"img{index} original")
            if refined:
                _draw_bbox(draw, refined, "#9B51E0", f"img{index} refined")
        else:
            bbox = _coerce_bbox(block.get("bbox"), overlay.size)
            if bbox:
                _draw_bbox(draw, bbox, "#2F80ED", f"text{index}")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    overlay.save(output_path, format="PNG")


def _draw_bbox(draw: ImageDraw.ImageDraw, bbox: tuple[int, int, int, int], color: str, label: str) -> None:
    for inset in range(2):
        draw.rectangle(
            (bbox[0] - inset, bbox[1] - inset, bbox[2] + inset, bbox[3] + inset),
            outline=color,
        )
    draw.text((bbox[0] + 2, max(0, bbox[1] - 12)), label, fill=color)


def _emit(callback: StatusCallback | None, message: str) -> None:
    if callback:
        callback(message)
