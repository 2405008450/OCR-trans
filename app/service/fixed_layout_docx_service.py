# -*- coding: utf-8 -*-
"""
将浏览器已经排版完成的 HTML 转换为固定布局 DOCX。

这里不让 OCR/视觉模型猜测 bbox，而是启动本机 Chrome/Edge，通过 DevTools
协议读取 DOM 的实际渲染坐标，再把每一行文字写入 Word 的绝对定位文本框。
非文字内容（图片、底色、边框等）会作为页面背景保留。
"""

from __future__ import annotations

import base64
import json
import os
import shutil
import socket
import subprocess
import tempfile
import time
import urllib.request
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree
from websockets.sync.client import connect


CSS_PX_PER_INCH = 96.0
POINTS_PER_INCH = 72.0
PX_TO_PT = POINTS_PER_INCH / CSS_PX_PER_INCH
NS_VML = "urn:schemas-microsoft-com:vml"
NS_OFFICE = "urn:schemas-microsoft-com:office:office"

# A4 纵向在 96 DPI 下的 CSS 像素尺寸。
DEFAULT_PAGE_WIDTH_PX = 210 / 25.4 * CSS_PX_PER_INCH
DEFAULT_PAGE_HEIGHT_PX = 297 / 25.4 * CSS_PX_PER_INCH


class FixedLayoutConversionError(RuntimeError):
    """固定布局转换无法完成。"""


@dataclass(frozen=True)
class BrowserTextLine:
    page_index: int
    text: str
    x: float
    y: float
    width: float
    height: float
    font_family: str
    font_size: float
    font_weight: int
    italic: bool
    color: str
    line_height: float
    direction: str


@dataclass(frozen=True)
class BrowserLayout:
    page_width_px: float
    page_height_px: float
    pages: list[list[BrowserTextLine]]
    background_images: list[Path]


def convert_html_to_fixed_layout_docx(
    html_text: str,
    output_path: str | Path,
    *,
    html_output_path: str | Path | None = None,
    browser_path: str | Path | None = None,
    debug_layout_path: str | Path | None = None,
    page_width_px: float = DEFAULT_PAGE_WIDTH_PX,
    page_height_px: float = DEFAULT_PAGE_HEIGHT_PX,
) -> str:
    """
    使用浏览器实测布局生成固定版式 DOCX。

    Args:
        html_text: 已完成翻译和排版的完整 HTML。
        output_path: 输出 DOCX 路径。
        html_output_path: 可选，保存本次输入 HTML，便于问题复现。
        browser_path: 可选，Chrome/Edge 可执行文件。
        debug_layout_path: 可选，保存浏览器实测出的坐标 JSON。
        page_width_px/page_height_px: Word 页面对应的浏览器画布尺寸。
    """
    output_file = Path(output_path).resolve()
    output_file.parent.mkdir(parents=True, exist_ok=True)

    if html_output_path:
        html_file = Path(html_output_path).resolve()
        html_file.parent.mkdir(parents=True, exist_ok=True)
        html_file.write_text(html_text, encoding="utf-8")
        remove_html_after = False
    else:
        temp_html = tempfile.NamedTemporaryFile(
            mode="w",
            suffix=".html",
            encoding="utf-8",
            delete=False,
            dir=output_file.parent,
        )
        try:
            temp_html.write(html_text)
        finally:
            temp_html.close()
        html_file = Path(temp_html.name)
        remove_html_after = True

    background_dir = Path(
        tempfile.mkdtemp(prefix=".fixed-layout-backgrounds-", dir=output_file.parent)
    )
    try:
        layout = _extract_browser_layout(
            html_file=html_file,
            background_dir=background_dir,
            browser_path=browser_path,
            page_width_px=page_width_px,
            page_height_px=page_height_px,
        )
        if debug_layout_path:
            _write_layout_debug(layout, Path(debug_layout_path))
        _build_fixed_layout_docx(layout, output_file)
    finally:
        shutil.rmtree(background_dir, ignore_errors=True)
        if remove_html_after:
            html_file.unlink(missing_ok=True)

    if not output_file.exists() or output_file.stat().st_size == 0:
        raise FixedLayoutConversionError("固定布局 DOCX 未生成")
    return str(output_file)


def resolve_browser_path(configured_path: str | Path | None = None) -> str:
    """寻找可用于无头渲染的 Chrome 或 Edge。"""
    candidates = [
        str(configured_path or "").strip(),
        os.getenv("FIXED_LAYOUT_BROWSER_PATH", "").strip(),
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        "google-chrome",
        "chromium",
        "chromium-browser",
        "msedge",
    ]
    checked: set[str] = set()
    for candidate in candidates:
        if not candidate or candidate in checked:
            continue
        checked.add(candidate)
        candidate_path = Path(candidate)
        if candidate_path.is_absolute() or any(sep in candidate for sep in ("/", "\\")):
            if candidate_path.exists():
                return str(candidate_path)
            continue
        resolved = shutil.which(candidate)
        if resolved:
            return resolved
    raise FileNotFoundError(
        "未找到 Chrome/Edge。请安装浏览器，或设置 FIXED_LAYOUT_BROWSER_PATH。"
    )


def _find_free_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


class _CdpClient:
    def __init__(self, websocket_url: str):
        self._socket = connect(websocket_url, open_timeout=10, close_timeout=3)
        self._next_id = 0

    def close(self) -> None:
        self._socket.close()

    def call(self, method: str, params: dict[str, Any] | None = None) -> dict[str, Any]:
        self._next_id += 1
        request_id = self._next_id
        self._socket.send(
            json.dumps(
                {"id": request_id, "method": method, "params": params or {}},
                ensure_ascii=False,
            )
        )
        while True:
            payload = json.loads(self._socket.recv())
            if payload.get("id") != request_id:
                continue
            if "error" in payload:
                raise FixedLayoutConversionError(
                    f"浏览器命令 {method} 执行失败: {payload['error']}"
                )
            return payload.get("result") or {}

    def evaluate(self, expression: str) -> Any:
        result = self.call(
            "Runtime.evaluate",
            {
                "expression": expression,
                "awaitPromise": True,
                "returnByValue": True,
                "userGesture": True,
            },
        )
        exception = result.get("exceptionDetails")
        if exception:
            raise FixedLayoutConversionError(f"浏览器脚本执行失败: {exception}")
        return (result.get("result") or {}).get("value")


def _extract_browser_layout(
    *,
    html_file: Path,
    background_dir: Path,
    browser_path: str | Path | None,
    page_width_px: float,
    page_height_px: float,
) -> BrowserLayout:
    browser = resolve_browser_path(browser_path)
    port = _find_free_port()
    with tempfile.TemporaryDirectory(
        prefix="fixed-layout-browser-",
        ignore_cleanup_errors=True,
    ) as profile_dir:
        command = [
            browser,
            "--headless=new",
            "--disable-gpu",
            "--hide-scrollbars",
            "--no-first-run",
            "--no-default-browser-check",
            "--allow-file-access-from-files",
            "--force-device-scale-factor=1",
            f"--remote-debugging-port={port}",
            f"--user-data-dir={profile_dir}",
            f"--window-size={int(page_width_px)},{int(page_height_px)}",
            html_file.as_uri(),
        ]
        creation_flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
        process = subprocess.Popen(
            command,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            creationflags=creation_flags,
        )
        client: _CdpClient | None = None
        try:
            websocket_url = _wait_for_page_websocket(port, html_file)
            client = _CdpClient(websocket_url)
            client.call("Page.enable")
            client.call("Runtime.enable")
            client.call(
                "Emulation.setDeviceMetricsOverride",
                {
                    "width": int(round(page_width_px)),
                    "height": int(round(page_height_px)),
                    "deviceScaleFactor": 1,
                    "mobile": False,
                },
            )
            client.evaluate(
                "document.fonts && document.fonts.ready "
                "? document.fonts.ready.then(() => true) : Promise.resolve(true)"
            )
            raw_layout = client.evaluate(
                _browser_layout_script(page_width_px, page_height_px)
            )
            if not isinstance(raw_layout, dict):
                raise FixedLayoutConversionError("浏览器未返回有效的 DOM 布局数据")

            pages = _parse_browser_pages(raw_layout)
            background_images = _capture_page_backgrounds(
                client=client,
                background_dir=background_dir,
                page_count=len(pages),
                page_width_px=page_width_px,
                page_height_px=page_height_px,
            )
            return BrowserLayout(
                page_width_px=page_width_px,
                page_height_px=page_height_px,
                pages=pages,
                background_images=background_images,
            )
        finally:
            if client is not None:
                try:
                    client.call("Browser.close")
                except Exception:
                    pass
                try:
                    client.close()
                except Exception:
                    pass
            if process.poll() is None:
                process.terminate()
            try:
                process.wait(timeout=5)
            except subprocess.TimeoutExpired:
                process.kill()
                process.wait(timeout=5)


def _wait_for_page_websocket(port: int, html_file: Path) -> str:
    endpoint = f"http://127.0.0.1:{port}/json/list"
    deadline = time.monotonic() + 15
    last_error: Exception | None = None
    while time.monotonic() < deadline:
        try:
            with urllib.request.urlopen(endpoint, timeout=1) as response:
                pages = json.loads(response.read().decode("utf-8"))
            html_uri = html_file.as_uri()
            for page in pages:
                if page.get("type") != "page":
                    continue
                if page.get("url") == html_uri and page.get("webSocketDebuggerUrl"):
                    return str(page["webSocketDebuggerUrl"])
            for page in pages:
                if page.get("type") == "page" and page.get("webSocketDebuggerUrl"):
                    return str(page["webSocketDebuggerUrl"])
        except Exception as exc:
            last_error = exc
        time.sleep(0.1)
    raise FixedLayoutConversionError(f"浏览器启动超时: {last_error}")


def _browser_layout_script(page_width_px: float, page_height_px: float) -> str:
    width = json.dumps(float(page_width_px))
    height = json.dumps(float(page_height_px))
    return f"""
(async () => {{
  const PAGE_WIDTH = {width};
  const PAGE_HEIGHT = {height};
  if (document.fonts && document.fonts.ready) await document.fonts.ready;

  const originalBodyStyle = getComputedStyle(document.body);
  const inset = {{
    top: parseFloat(originalBodyStyle.marginTop) + parseFloat(originalBodyStyle.paddingTop),
    right: parseFloat(originalBodyStyle.marginRight) + parseFloat(originalBodyStyle.paddingRight),
    bottom: parseFloat(originalBodyStyle.marginBottom) + parseFloat(originalBodyStyle.paddingBottom),
    left: parseFloat(originalBodyStyle.marginLeft) + parseFloat(originalBodyStyle.paddingLeft)
  }};
  const originalNodes = Array.from(document.body.childNodes);
  const style = document.createElement("style");
  style.textContent = `
    html, body {{ margin: 0 !important; padding: 0 !important; width: ${{PAGE_WIDTH}}px !important;
      background: white !important; }}
    .codex-fixed-page {{ position: relative; display: block; box-sizing: border-box;
      width: ${{PAGE_WIDTH}}px; height: ${{PAGE_HEIGHT}}px; overflow: hidden; background: white; }}
  `;
  document.head.appendChild(style);
  document.body.replaceChildren();

  const pages = [];
  function newPage() {{
    const page = document.createElement("section");
    page.className = "codex-fixed-page";
    page.style.padding = `${{inset.top}}px ${{inset.right}}px ${{inset.bottom}}px ${{inset.left}}px`;
    document.body.appendChild(page);
    pages.push(page);
    return page;
  }}
  function isPageBreak(node) {{
    if (!(node instanceof Element)) return false;
    const own = (node.getAttribute("style") || "").toLowerCase();
    const computed = getComputedStyle(node);
    return own.includes("page-break-before: always")
      || own.includes("break-before: page")
      || computed.pageBreakBefore === "always"
      || computed.breakBefore === "page";
  }}

  let currentPage = newPage();
  for (const node of originalNodes) {{
    if (isPageBreak(node)) {{
      currentPage = newPage();
      continue;
    }}
    currentPage.appendChild(node);
  }}
  await new Promise(resolve => requestAnimationFrame(() => requestAnimationFrame(resolve)));

  function numericWeight(value) {{
    if (value === "bold" || value === "bolder") return 700;
    const parsed = parseInt(value, 10);
    return Number.isFinite(parsed) ? parsed : 400;
  }}
  function visibleParent(element) {{
    let current = element;
    while (current && current !== document.documentElement) {{
      const cs = getComputedStyle(current);
      if (cs.display === "none" || cs.visibility === "hidden" || parseFloat(cs.opacity) === 0) return false;
      current = current.parentElement;
    }}
    return true;
  }}

  const resultPages = pages.map((page, pageIndex) => {{
    const pageRect = page.getBoundingClientRect();
    const lines = [];
    const walker = document.createTreeWalker(page, NodeFilter.SHOW_TEXT);
    let node;
    while ((node = walker.nextNode())) {{
      const parent = node.parentElement;
      if (!parent || ["SCRIPT", "STYLE", "NOSCRIPT"].includes(parent.tagName)) continue;
      if (!visibleParent(parent) || !node.nodeValue || !node.nodeValue.trim()) continue;
      const cs = getComputedStyle(parent);
      const chars = [];
      for (let index = 0; index < node.nodeValue.length; index++) {{
        const range = document.createRange();
        range.setStart(node, index);
        range.setEnd(node, index + 1);
        const rect = range.getBoundingClientRect();
        if (rect.width <= 0 || rect.height <= 0) continue;
        chars.push({{ ch: node.nodeValue[index], rect }});
      }}
      const groups = [];
      for (const item of chars) {{
        let group = groups.find(g => Math.abs(g.top - item.rect.top) <= 1.5);
        if (!group) {{
          group = {{ top: item.rect.top, items: [] }};
          groups.push(group);
        }}
        group.items.push(item);
      }}
      for (const group of groups) {{
        group.items.sort((a, b) => a.rect.left - b.rect.left);
        while (group.items.length && /\\s/.test(group.items[0].ch)) group.items.shift();
        while (group.items.length && /\\s/.test(group.items[group.items.length - 1].ch)) group.items.pop();
        if (!group.items.length) continue;
        const text = group.items.map(item => item.ch).join("").replace(/\\s+/g, " ").trim();
        if (!text) continue;
        const left = Math.min(...group.items.map(item => item.rect.left));
        const right = Math.max(...group.items.map(item => item.rect.right));
        const top = Math.min(...group.items.map(item => item.rect.top));
        const bottom = Math.max(...group.items.map(item => item.rect.bottom));
        const parsedLineHeight = parseFloat(cs.lineHeight);
        lines.push({{
          page_index: pageIndex,
          text,
          x: left - pageRect.left,
          y: top - pageRect.top,
          width: right - left,
          height: bottom - top,
          font_family: cs.fontFamily || "Arial",
          font_size: parseFloat(cs.fontSize) || 16,
          font_weight: numericWeight(cs.fontWeight),
          italic: cs.fontStyle === "italic" || cs.fontStyle === "oblique",
          color: cs.color || "rgb(0, 0, 0)",
          line_height: Number.isFinite(parsedLineHeight) ? parsedLineHeight : (bottom - top) * 1.15,
          direction: cs.direction || "ltr"
        }});
      }}
    }}
    return lines;
  }});

  // 只隐藏文字节点，图片、边框、底色仍保留，供后续截取页面背景。
  for (const page of pages) {{
    const walker = document.createTreeWalker(page, NodeFilter.SHOW_TEXT);
    const textNodes = [];
    let node;
    while ((node = walker.nextNode())) {{
      if (node.nodeValue && node.nodeValue.trim()) textNodes.push(node);
    }}
    for (const textNode of textNodes) {{
      const span = document.createElement("span");
      span.className = "codex-hidden-text";
      span.textContent = textNode.nodeValue;
      textNode.replaceWith(span);
    }}
  }}
  const hideStyle = document.createElement("style");
  hideStyle.textContent = ".codex-hidden-text {{ color: transparent !important; -webkit-text-fill-color: transparent !important; text-shadow: none !important; }}";
  document.head.appendChild(hideStyle);
  await new Promise(resolve => requestAnimationFrame(() => requestAnimationFrame(resolve)));

  return {{
    page_width_px: PAGE_WIDTH,
    page_height_px: PAGE_HEIGHT,
    pages: resultPages
  }};
}})()
"""


def _parse_browser_pages(raw_layout: dict[str, Any]) -> list[list[BrowserTextLine]]:
    raw_pages = raw_layout.get("pages")
    if not isinstance(raw_pages, list) or not raw_pages:
        raise FixedLayoutConversionError("HTML 中没有可输出的页面")
    pages: list[list[BrowserTextLine]] = []
    for page_index, raw_page in enumerate(raw_pages):
        lines: list[BrowserTextLine] = []
        if isinstance(raw_page, list):
            for raw_line in raw_page:
                if not isinstance(raw_line, dict) or not str(raw_line.get("text") or "").strip():
                    continue
                lines.append(
                    BrowserTextLine(
                        page_index=page_index,
                        text=str(raw_line["text"]),
                        x=float(raw_line.get("x") or 0),
                        y=float(raw_line.get("y") or 0),
                        width=max(float(raw_line.get("width") or 1), 1),
                        height=max(float(raw_line.get("height") or 1), 1),
                        font_family=str(raw_line.get("font_family") or "Arial"),
                        font_size=max(float(raw_line.get("font_size") or 16), 1),
                        font_weight=int(raw_line.get("font_weight") or 400),
                        italic=bool(raw_line.get("italic")),
                        color=str(raw_line.get("color") or "rgb(0, 0, 0)"),
                        line_height=max(float(raw_line.get("line_height") or 16), 1),
                        direction=str(raw_line.get("direction") or "ltr"),
                    )
                )
        pages.append(lines)
    return pages


def _capture_page_backgrounds(
    *,
    client: _CdpClient,
    background_dir: Path,
    page_count: int,
    page_width_px: float,
    page_height_px: float,
) -> list[Path]:
    paths: list[Path] = []
    for page_index in range(page_count):
        result = client.call(
            "Page.captureScreenshot",
            {
                "format": "png",
                "fromSurface": True,
                "captureBeyondViewport": True,
                "clip": {
                    "x": 0,
                    "y": page_index * page_height_px,
                    "width": page_width_px,
                    "height": page_height_px,
                    "scale": 1,
                },
            },
        )
        encoded = result.get("data")
        if not encoded:
            raise FixedLayoutConversionError(f"第 {page_index + 1} 页背景截图失败")
        path = background_dir / f"page-{page_index + 1}.png"
        path.write_bytes(base64.b64decode(encoded))
        paths.append(path)
    return paths


def _build_fixed_layout_docx(layout: BrowserLayout, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Pt(layout.page_width_px * PX_TO_PT)
    section.page_height = Pt(layout.page_height_px * PX_TO_PT)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    anchor = document.add_paragraph()
    for page_index, lines in enumerate(layout.pages):
        if page_index:
            anchor = document.add_paragraph()
        _minimize_anchor_paragraph(anchor)
        if page_index < len(layout.background_images):
            _append_background_shape(
                anchor,
                layout.background_images[page_index],
                layout.page_width_px * PX_TO_PT,
                layout.page_height_px * PX_TO_PT,
                page_index,
            )
        for line_index, line in enumerate(lines):
            _append_textbox(
                anchor,
                line,
                page_width_px=layout.page_width_px,
                shape_id=f"p{page_index + 1}t{line_index + 1}",
            )
        if page_index < len(layout.pages) - 1:
            anchor.add_run().add_break(WD_BREAK.PAGE)

    document.save(output_path)


def _minimize_anchor_paragraph(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "20")
    spacing.set(qn("w:lineRule"), "exact")


def _append_background_shape(
    paragraph,
    image_path: Path,
    width_pt: float,
    height_pt: float,
    page_index: int,
) -> None:
    r_id, _image = paragraph.part.get_or_add_image(str(image_path))
    run = OxmlElement("w:r")
    pict = OxmlElement("w:pict")
    shape = etree.Element(f"{{{NS_VML}}}shape", nsmap={"v": NS_VML, "o": NS_OFFICE})
    shape.set("id", f"FixedLayoutBackground{page_index + 1}")
    shape.set(
        "style",
        _shape_style(
            x_pt=0,
            y_pt=0,
            width_pt=width_pt,
            height_pt=height_pt,
            z_index=-251654144 + page_index,
        ),
    )
    shape.set("stroked", "f")
    shape.set("filled", "f")
    image_data = etree.Element(f"{{{NS_VML}}}imagedata")
    image_data.set(qn("r:id"), r_id)
    image_data.set(f"{{{NS_OFFICE}}}title", "")
    shape.append(image_data)
    pict.append(shape)
    run.append(pict)
    paragraph._p.append(run)


def _append_textbox(
    paragraph,
    line: BrowserTextLine,
    *,
    page_width_px: float,
    shape_id: str,
) -> None:
    # Word 与 Chromium 的字体回退和字距不完全一致，因此给文本框保留余量。
    # 靠近右边缘时同步向左扩展，避免右对齐字段被页面裁掉。
    width_px = max(line.width * 1.18 + 6, line.font_size)
    x_px = line.x
    if x_px + width_px > page_width_px - 2:
        x_px = max(0, page_width_px - 2 - width_px)
    y_px = max(0, line.y - max(line.font_size * 0.10, 1))
    width_pt = width_px * PX_TO_PT
    height_pt = max(line.line_height * 1.55, line.height * 1.55) * PX_TO_PT

    run = OxmlElement("w:r")
    pict = OxmlElement("w:pict")
    shape = etree.Element(f"{{{NS_VML}}}shape", nsmap={"v": NS_VML})
    shape.set("id", shape_id)
    shape.set(
        "style",
        _shape_style(
            x_pt=x_px * PX_TO_PT,
            y_pt=y_px * PX_TO_PT,
            width_pt=width_pt,
            height_pt=height_pt,
            z_index=1000 + line.page_index,
        ),
    )
    shape.set("stroked", "f")
    shape.set("filled", "f")

    textbox = etree.Element(f"{{{NS_VML}}}textbox")
    textbox.set("inset", "0,0,0,0")
    text_content = OxmlElement("w:txbxContent")
    text_paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), str(max(int(round(line.line_height * PX_TO_PT * 20)), 20)))
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)
    if line.direction.lower() == "rtl":
        p_pr.append(OxmlElement("w:bidi"))
    text_paragraph.append(p_pr)

    text_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    font_name = _primary_font_name(line.font_family)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), font_name)
    r_pr.append(fonts)
    half_points = max(int(round(line.font_size * PX_TO_PT * 2)), 2)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(half_points))
    r_pr.append(size)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(half_points))
    r_pr.append(size_cs)
    if line.font_weight >= 600:
        r_pr.append(OxmlElement("w:b"))
        r_pr.append(OxmlElement("w:bCs"))
    if line.italic:
        r_pr.append(OxmlElement("w:i"))
        r_pr.append(OxmlElement("w:iCs"))
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _rgb_to_hex(line.color))
    r_pr.append(color)
    text_run.append(r_pr)
    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = line.text
    text_run.append(text)
    text_paragraph.append(text_run)
    text_content.append(text_paragraph)
    textbox.append(text_content)
    shape.append(textbox)
    pict.append(shape)
    run.append(pict)
    paragraph._p.append(run)


def _shape_style(
    *,
    x_pt: float,
    y_pt: float,
    width_pt: float,
    height_pt: float,
    z_index: int,
) -> str:
    return (
        "position:absolute;"
        f"margin-left:{x_pt:.3f}pt;"
        f"margin-top:{y_pt:.3f}pt;"
        f"width:{width_pt:.3f}pt;"
        f"height:{height_pt:.3f}pt;"
        f"z-index:{z_index};"
        "visibility:visible;"
        "mso-wrap-style:none;"
        "mso-position-horizontal-relative:page;"
        "mso-position-vertical-relative:page"
    )


def _primary_font_name(font_family: str) -> str:
    first = (font_family or "Arial").split(",", 1)[0].strip().strip("\"'")
    return first or "Arial"


def _rgb_to_hex(color: str) -> str:
    value = (color or "").strip().lower()
    if value.startswith("#"):
        compact = value[1:]
        if len(compact) == 3:
            return "".join(ch * 2 for ch in compact).upper()
        if len(compact) >= 6:
            return compact[:6].upper()
    if value.startswith("rgb"):
        start = value.find("(")
        end = value.find(")", start + 1)
        if start >= 0 and end > start:
            parts = [part.strip() for part in value[start + 1 : end].split(",")[:3]]
            try:
                return "".join(f"{max(0, min(int(float(part)), 255)):02X}" for part in parts)
            except ValueError:
                pass
    return "000000"


def _write_layout_debug(layout: BrowserLayout, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "coordinate_source": "browser_dom",
        "page_width_px": layout.page_width_px,
        "page_height_px": layout.page_height_px,
        "pages": [[asdict(line) for line in page] for page in layout.pages],
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def iter_layout_text(layout: BrowserLayout) -> Iterable[str]:
    """测试和诊断辅助：按页面顺序迭代实测文字。"""
    for page in layout.pages:
        for line in page:
            yield line.text
