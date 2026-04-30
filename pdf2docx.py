"""
OCR 图片 → 混合 HTML/Markdown → 格式化 Word 文档 (一步到位)
"""

import base64
import re
import sys
import os
import time
from pathlib import Path
from typing import Any

from app.core.config import settings
from app.service.gemini_service import (
    GEMINI_ROUTE_GOOGLE,
    GEMINI_ROUTE_OPENROUTER,
    ensure_gemini_route_configured,
    generate_vision_html,
    normalize_gemini_route,
)
from app.service.libreoffice_service import (
    LIBREOFFICE_PATH,
    convert_to_docx_via_libreoffice,
    resolve_libreoffice_path,
)

# ============================================================
# 依赖检查与导入
# ============================================================
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml
except ImportError:
    print("请安装 python-docx: pip install python-docx")
    sys.exit(1)

try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except ImportError:
    print("请安装 BeautifulSoup: pip install beautifulsoup4")
    sys.exit(1)


# ============================================================
# 1. LLM OCR 调用
# ============================================================
SYS_PROMPT = """Convert the document into simple, Microsoft Word-safe HTML.

Return only a complete HTML document. Do not include explanations. Do not use markdown. Do not wrap the output in code fences.

The output must be a full HTML document with:
- <!DOCTYPE html>
- <html>
- <head>
- <meta charset="utf-8">
- <title>Document</title>
- <body>

Important:
- The HTML will be opened by Microsoft Word, not a browser.
- Prioritize Word compatibility over visual fidelity.
- Use simple HTML only.

Rules:
- Include all visible text content.
- Preserve reading order.
- Do not summarize or rewrite.
- Do not invent missing text.
- If a page has no readable text and no essential stamp/seal, return a full HTML document with an empty <body></body>.
- Replace any sequence of 3 or more repeated characters used as visual separators or fill lines (e.g. -----, _____, ....., =====) with a single short placeholder: ___
- For blank form fields shown as long lines (e.g. "Name: _____________"), keep only one short underscore placeholder: "Name: ___"
- Do NOT reproduce decorative divider lines or page-wide rules; omit them entirely.

Allowed tags:
- <html>, <head>, <meta>, <title>, <body>
- <h1> to <h6>
- <p>
- <br>
- <strong>, <em>, <u>, <sub>, <sup>
- <ul>, <ol>, <li>
- <table>, <thead>, <tbody>, <tr>, <th>, <td>
- <hr>

Formatting rules:
- Prefer <p> over nested <div>.
- Before writing HTML, inspect the page as a 2D layout, not only as linear text.
- Use tables aggressively for structured alignment, even when the source document has no visible table borders.
- If two or more text blocks sit on the same visual row or horizontal band but are separated left/right, place them in the same table row with separate cells.
- This rule applies to sparse forms, metadata headers, label/value pairs, signatures, names, dates, IDs, addresses, and any side-by-side content blocks.
- Do NOT flatten left/right blocks from the same row into separate stacked paragraphs.
- If one side contains multiple lines, keep them inside the same cell with <br> instead of splitting them into separate paragraphs.
- If you are unsure whether a layout should be paragraphs or a table, prefer a borderless table that preserves the horizontal relationship.
- For alignment-only layouts, use simple borderless tables such as:
  <table style="width:100%; border:none;"><tr><td style="width:50%; border:none; vertical-align:top;">...</td><td style="width:50%; border:none; vertical-align:top;">...</td></tr></table>
- Do not invent columns that are not present, but preserve every obvious horizontal grouping that appears in the source.
- Keep styles minimal and inline only when necessary.
- Only use simple inline styles such as:
  text-align, font-size, font-family, font-weight, margin, text-indent, width, vertical-align, border
- Do not use:
  flex, grid, float, position, transform, rgba(), opacity, negative margins, external CSS, JavaScript, SVG, canvas
- Avoid unnecessary nesting.

Images / non-text visuals:
- Extract only readable text and layout structure.
- For charts, diagrams, logos, icons, photos, illustrations, handwriting, or other graphics, keep only the readable text that appears inside them.
- Do NOT describe the visual itself.
- Never output placeholders or descriptions such as [Image: ...], [Blank Page], [Chart], [Icon], [Logo], [Photo], or similar.
- If a chart / icon / logo / illustration contains no readable text, omit it entirely.
- Do not rely on external image paths unless explicitly provided and required.

Output must be valid HTML that Microsoft Word can open directly.
Visual color preservation:
- Preserve important visual colors from the document.
- If text is clearly colored (e.g., red titles or red stamps), reflect that using simple HTML inline styles.

Seals / stamps:
- If a stamp or official seal is present, represent it as a paragraph.
- Use red color to reflect the stamp.
- Prefer simple styles compatible with Microsoft Word.
- Do not use rgba() or opacity.
- Example:
  <p style="color:#C00000; font-weight:bold;">[Official Seal]</p>"""

OCR_FALLBACK_MODEL_ORDER = (
    "google/gemini-3.1-pro-preview",
    "google/gemini-3-flash-preview",
    "gemini-3.1-flash-lite-preview",
)
OCR_BLANK_HINTS = {
    "blank",
    "blankpage",
    "emptypage",
    "nocontent",
    "空白",
    "空白页",
    "空白页面",
    "无内容",
    "本页为空白",
}


class OCRIncompleteResultError(RuntimeError):
    pass


def _strip_optional_code_fence(text: str) -> str:
    normalized = (text or "").strip()
    normalized = re.sub(r"^```(?:html|markdown)?\s*\n?", "", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"\n?```\s*$", "", normalized)
    return normalized.strip()


def _emit_ocr_status(message: str, status_callback=None) -> None:
    text = (message or "").strip()
    if not text:
        return
    print(text, flush=True)
    if status_callback:
        status_callback(text)


def _ocr_exception_message(exc: Exception) -> str:
    if isinstance(exc, OCRIncompleteResultError):
        return "OCR 输出疑似被截断"
    return f"OCR 请求失败（{exc.__class__.__name__}）"


def _route_display_name(route: str) -> str:
    return {
        GEMINI_ROUTE_GOOGLE: "Google Vertex",
        GEMINI_ROUTE_OPENROUTER: "OpenRouter",
    }.get(route, route)


def _route_is_available(route: str) -> bool:
    try:
        ensure_gemini_route_configured(route)
        return True
    except Exception:
        return False


def _build_ocr_route_candidates(gemini_route: str) -> list[str]:
    normalized = normalize_gemini_route(gemini_route)
    candidates: list[str] = []
    for route in (
        normalized,
        GEMINI_ROUTE_GOOGLE if normalized != GEMINI_ROUTE_GOOGLE else GEMINI_ROUTE_OPENROUTER,
    ):
        if route not in candidates and _route_is_available(route):
            candidates.append(route)
    if not candidates and _route_is_available(normalized):
        candidates.append(normalized)
    return candidates or [normalized]


def _build_ocr_model_candidates(model: str) -> list[str]:
    if model in OCR_FALLBACK_MODEL_ORDER:
        start_index = OCR_FALLBACK_MODEL_ORDER.index(model)
        return list(OCR_FALLBACK_MODEL_ORDER[start_index:])

    candidates = [model]
    for fallback_model in OCR_FALLBACK_MODEL_ORDER[1:]:
        if fallback_model not in candidates:
            candidates.append(fallback_model)
    return candidates


def _build_ocr_attempt_plan(model: str, gemini_route: str, retries: int) -> list[dict[str, Any]]:
    routes = _build_ocr_route_candidates(gemini_route)
    models = _build_ocr_model_candidates(model)
    plan: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()

    for model_index, candidate_model in enumerate(models):
        for route_index, candidate_route in enumerate(routes):
            stage_key = (candidate_route, candidate_model)
            if stage_key in seen:
                continue
            seen.add(stage_key)
            stage_retries = retries if model_index == 0 and route_index == 0 else 1
            plan.append(
                {
                    "route": candidate_route,
                    "model": candidate_model,
                    "retries": max(stage_retries, 1),
                }
            )
    return plan


def _is_blank_ocr_result(text: str) -> bool:
    if not (text or "").strip():
        return True

    normalized = _strip_optional_code_fence(text)
    normalized = re.sub(r"<page_break\s*/>", " ", normalized, flags=re.IGNORECASE)

    soup = BeautifulSoup(normalized, "html.parser")
    body = soup.find("body")
    visible_text = " ".join(body.stripped_strings) if body else " ".join(soup.stripped_strings)
    visible_text = visible_text.replace("\xa0", " ").strip()
    if not visible_text:
        return True

    visible_text = re.sub(r"[\s\-_.,;:!?'\"`~|/\\()\[\]{}<>+=]+", "", visible_text)
    if not visible_text:
        return True

    return visible_text.lower() in OCR_BLANK_HINTS


def _is_likely_truncated_ocr_result(text: str) -> bool:
    normalized = _strip_optional_code_fence(text)
    if not normalized:
        return False

    lower = normalized.lower()
    looks_like_full_html = any(marker in lower for marker in ("<!doctype html", "<html", "<body"))
    if not looks_like_full_html:
        return False

    if "<body" in lower and "</body>" not in lower:
        return True
    if "<html" in lower and "</html>" not in lower:
        return True

    return False


def _ocr_single_image(
    img_b64: str,
    mime_type: str,
    model: str,
    gemini_route: str = GEMINI_ROUTE_GOOGLE,
    retries: int = 3,
    status_callback=None,
) -> str:
    """对单张图片调用 OCR，失败时按“原路线 -> 备用路线 -> 轻量模型”逐级降级。"""
    image_bytes = base64.standard_b64decode(img_b64)
    attempt_plan = _build_ocr_attempt_plan(model, gemini_route, retries)

    for stage_index, stage in enumerate(attempt_plan):
        route = stage["route"]
        candidate_model = stage["model"]
        stage_retries = stage["retries"]

        if stage_index > 0:
            _emit_ocr_status(
                f"↘ OCR 降级到 {_route_display_name(route)} / {candidate_model}",
                status_callback,
            )

        for attempt in range(stage_retries):
            try:
                response_text = generate_vision_html(
                    system_prompt=SYS_PROMPT,
                    image_bytes=image_bytes,
                    mime_type=mime_type,
                    model=candidate_model,
                    route=route,
                    temperature=0,
                )
                if _is_likely_truncated_ocr_result(response_text):
                    raise OCRIncompleteResultError("OCR 输出疑似被截断：HTML 未完整闭合")
                print(response_text, end="", flush=True)
                if stage_index > 0:
                    _emit_ocr_status(
                        f"✅ OCR 已恢复：{_route_display_name(route)} / {candidate_model}",
                        status_callback,
                    )
                return response_text
            except Exception as exc:
                retry_same_stage = attempt < stage_retries - 1 and not isinstance(
                    exc,
                    OCRIncompleteResultError,
                )
                if retry_same_stage:
                    wait = 3 * (attempt + 1)
                    _emit_ocr_status(
                        f"⚠️ {_ocr_exception_message(exc)}，{wait} 秒后重试 "
                        f"[{attempt + 1}/{stage_retries}]...",
                        status_callback,
                    )
                    time.sleep(wait)
                    continue

                if stage_index < len(attempt_plan) - 1:
                    next_stage = attempt_plan[stage_index + 1]
                    _emit_ocr_status(
                        f"⚠️ {_ocr_exception_message(exc)}，准备切换到 "
                        f"{_route_display_name(next_stage['route'])} / {next_stage['model']}",
                        status_callback,
                    )
                    break

                raise RuntimeError(
                    f"OCR 失败，已用尽重试与降级策略：{exc}"
                ) from exc


def ocr_file(
    file_path: str,
    api_key: str = "",
    model: str = "google/gemini-3.1-pro-preview",
    gemini_route: str = GEMINI_ROUTE_GOOGLE,
    page_progress_callback=None,
    return_metadata: bool = False,
    ocr_status_callback=None,
) -> str | dict[str, Any]:
    """调用 LLM 对图片或 PDF 进行 OCR，返回混合 HTML/Markdown 文本。
    PDF 会逐页渲染为图片后分页发送，避免大文件直传超限。
    page_progress_callback(current_page, total_pages) 可选，用于报告逐页进度。
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        try:
            import fitz  # PyMuPDF
        except ImportError:
            print("请先安装 PyMuPDF: pip install pymupdf")
            sys.exit(1)

        doc = fitz.open(file_path)
        try:
            total = len(doc)
            all_results = []
            blank_pages: list[int] = []

            for i, page in enumerate(doc):
                page_no = i + 1
                print(f"\n🔄 正在处理第 {page_no}/{total} 页...")
                if page_progress_callback:
                    page_progress_callback(page_no, total)

                def page_status(message: str, current_page: int = page_no, total_pages: int = total):
                    if ocr_status_callback:
                        ocr_status_callback(f"第 {current_page}/{total_pages} 页：{message}")

                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                img_b64 = base64.standard_b64encode(pix.tobytes("jpeg", jpg_quality=85)).decode("utf-8")
                text = _ocr_single_image(
                    img_b64,
                    "image/jpeg",
                    model,
                    gemini_route=gemini_route,
                    status_callback=page_status,
                )
                if _is_blank_ocr_result(text):
                    blank_pages.append(page_no)
                    _emit_ocr_status(
                        f"ℹ️ 第 {page_no}/{total} 页 OCR 输出为空，已计为空白页",
                        ocr_status_callback,
                    )
                all_results.append(text)
                if i < total - 1:
                    time.sleep(1)
        finally:
            close = getattr(doc, "close", None)
            if callable(close):
                close()

        print("\n✅ PDF OCR 完成")
        joined_text = "\n\n<page_break/>\n\n".join(all_results)
        if return_metadata:
            return {
                "text": joined_text,
                "total_pages": total,
                "blank_page_count": len(blank_pages),
                "blank_pages": blank_pages,
            }
        return joined_text

    else:
        mime_map = {
            ".png":  "image/png",
            ".jpg":  "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif":  "image/gif",
            ".webp": "image/webp",
            ".bmp":  "image/bmp",
        }
        mime_type = mime_map.get(ext, "image/png")

        with open(file_path, "rb") as f:
            img_b64 = base64.standard_b64encode(f.read()).decode("utf-8")

        print("🔄 正在调用 LLM 进行 OCR...")
        result = _ocr_single_image(
            img_b64,
            mime_type,
            model,
            gemini_route=gemini_route,
            status_callback=ocr_status_callback,
        )
        print("\n✅ OCR 完成")
        if return_metadata:
            is_blank = _is_blank_ocr_result(result)
            blank_pages = [1] if is_blank else []
            return {
                "text": result,
                "total_pages": 1,
                "blank_page_count": len(blank_pages),
                "blank_pages": blank_pages,
            }
        return result


# ============================================================
# 2. 混合格式解析与 DOCX 渲染引擎
# ============================================================
def normalize_to_word_html(raw_text: str, title: str = "Document") -> str:
    """Normalize OCR/translation output into full HTML for LibreOffice DOCX conversion."""
    text = raw_text.strip()
    text = re.sub(r"^```(?:html|markdown)?\s*\n?", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\n?```\s*$", "", text)

    page_break_pattern = re.compile(r"\s*<page_break\s*/>\s*", flags=re.IGNORECASE)
    parts = [part.strip() for part in page_break_pattern.split(text) if part.strip()]
    if not parts:
        parts = [""]

    body_segments = []
    for part in parts:
        soup = BeautifulSoup(part, "html.parser")
        body = soup.find("body")
        if body:
            segment = "".join(str(child) for child in body.contents).strip()
        else:
            segment = part
        normalized_segment, _table_specs = _normalize_html_tables_for_word(segment)
        body_segments.append(normalized_segment)

    page_break_html = '<p style="page-break-before: always;"></p>'
    body_html = f"\n{page_break_html}\n".join(body_segments)

    return (
        "<!DOCTYPE html>\n"
        "<html>\n"
        "<head>\n"
        '  <meta charset="utf-8">\n'
        f"  <title>{title}</title>\n"
        "</head>\n"
        "<body>\n"
        f"{body_html}\n"
        "</body>\n"
        "</html>\n"
    )


def _parse_inline_style(style_text: str) -> dict[str, str]:
    style_map: dict[str, str] = {}
    for fragment in style_text.split(";"):
        if ":" not in fragment:
            continue
        key, value = fragment.split(":", 1)
        key = key.strip().lower()
        value = value.strip()
        if key and value:
            style_map[key] = value
    return style_map


def _serialize_inline_style(style_map: dict[str, str]) -> str:
    return "; ".join(f"{key}: {value}" for key, value in style_map.items())


def _extract_pct_width(value: str | None) -> int | None:
    if not value:
        return None
    match = re.fullmatch(r"\s*(\d{1,3})\s*%\s*", value)
    if not match:
        return None
    return max(1, min(int(match.group(1)), 100))


def _extract_table_alignment(table: Tag) -> str:
    align = (table.get("align") or "").strip().lower()
    if align in {"left", "center", "right"}:
        return align

    current: Tag | None = table
    while isinstance(current, Tag):
        style_map = _parse_inline_style(current.get("style", ""))
        text_align = style_map.get("text-align", "").strip().lower()
        if text_align in {"left", "center", "right"}:
            return text_align
        current = current.parent if isinstance(current.parent, Tag) else None

    return "center"


def _iter_table_rows(table: Tag) -> list[Tag]:
    rows = table.find_all("tr", recursive=False)
    if rows:
        return rows

    collected: list[Tag] = []
    for section in table.find_all(["thead", "tbody", "tfoot"], recursive=False):
        collected.extend(section.find_all("tr", recursive=False))
    return collected


def _extract_table_column_widths_pct(table: Tag) -> list[float]:
    for row in _iter_table_rows(table):
        widths: list[float] = []
        cells = row.find_all(["td", "th"], recursive=False)
        if not cells:
            continue

        for cell in cells:
            cell_style = _parse_inline_style(cell.get("style", ""))
            width_pct = _extract_pct_width(cell_style.get("width")) or _extract_pct_width(cell.get("width"))
            if width_pct is None:
                widths = []
                break

            colspan_raw = (cell.get("colspan") or "1").strip()
            try:
                colspan = max(int(colspan_raw), 1)
            except ValueError:
                colspan = 1

            widths.extend([width_pct / colspan] * colspan)

        if widths:
            return widths

    return []


def _normalize_html_tables_for_word(html_fragment: str) -> tuple[str, list[dict[str, Any]]]:
    soup = BeautifulSoup(html_fragment, "html.parser")
    table_specs: list[dict[str, Any]] = []

    for table in soup.find_all("table"):
        table_style = _parse_inline_style(table.get("style", ""))
        width_pct = _extract_pct_width(table_style.get("width")) or _extract_pct_width(table.get("width"))
        if width_pct is None:
            width_pct = 100

        alignment = _extract_table_alignment(table)
        border_value = table_style.get("border", "").strip().lower()
        borderless = border_value == "none" or (table.get("border") or "").strip() == "0"

        table_style["width"] = f"{width_pct}%"
        table_style.setdefault("border-collapse", "collapse")
        table_style.setdefault("table-layout", "fixed")
        table_style.setdefault("margin-left", "auto")
        table_style.setdefault("margin-right", "auto")
        if borderless:
            table_style["border"] = "none"
            table["border"] = "0"

        table["width"] = f"{width_pct}%"
        table["align"] = alignment
        table["cellpadding"] = table.get("cellpadding", "0")
        table["cellspacing"] = table.get("cellspacing", "0")
        table["style"] = _serialize_inline_style(table_style)

        for cell in table.find_all(["td", "th"]):
            cell_style = _parse_inline_style(cell.get("style", ""))
            cell_style.setdefault("vertical-align", "top")
            if borderless:
                cell_style.setdefault("border", "none")
            if cell_style:
                cell["style"] = _serialize_inline_style(cell_style)

        table_specs.append(
            {
                "align": alignment,
                "width_pct": width_pct,
                "column_widths_pct": _extract_table_column_widths_pct(table),
            }
        )

    if soup.body:
        return "".join(str(child) for child in soup.body.contents).strip(), table_specs
    return str(soup).strip(), table_specs


def _extract_table_layout_specs(html_text: str) -> list[dict[str, Any]]:
    _normalized_html, table_specs = _normalize_html_tables_for_word(html_text)
    return table_specs


def _get_or_add_xml_child(parent, tag_name: str):
    child = parent.find(qn(tag_name))
    if child is None:
        child = OxmlElement(tag_name)
        parent.append(child)
    return child


def _set_table_preferred_width_pct(table, width_pct: int):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    tbl_w = _get_or_add_xml_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(max(1, min(width_pct, 100)) * 50))


def _set_table_alignment(table, alignment: str):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    jc = _get_or_add_xml_child(tbl_pr, "w:jc")
    jc.set(qn("w:val"), alignment if alignment in {"left", "center", "right"} else "center")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is not None:
        tbl_pr.remove(tbl_ind)


def _enable_table_autofit(table):
    if hasattr(table, "autofit"):
        table.autofit = True
    if hasattr(table, "allow_autofit"):
        table.allow_autofit = True

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    tbl_layout = _get_or_add_xml_child(tbl_pr, "w:tblLayout")
    tbl_layout.set(qn("w:type"), "autofit")


def _set_table_grid_widths(table, widths_twips: list[int]):
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        insert_at = 1 if table._tbl.tblPr is not None else 0
        table._tbl.insert(insert_at, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)

    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(max(int(width), 1)))
        tbl_grid.append(grid_col)


def _set_cell_width_twips(cell, width_twips: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(max(int(width_twips), 1)))


def _get_cell_grid_span(cell) -> int:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return 1

    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is None:
        return 1

    try:
        return max(int(grid_span.get(qn("w:val"))), 1)
    except (TypeError, ValueError):
        return 1


def _apply_column_widths_to_table(table, column_widths_pct: list[float], content_width_twips: int):
    if not column_widths_pct or content_width_twips <= 0:
        return

    total_pct = sum(column_widths_pct)
    if total_pct <= 0:
        return

    normalized = [max(width, 0.0) / total_pct for width in column_widths_pct]
    widths_twips = [max(int(round(content_width_twips * ratio)), 1) for ratio in normalized]
    delta = content_width_twips - sum(widths_twips)
    if widths_twips and delta:
        widths_twips[-1] += delta

    _set_table_grid_widths(table, widths_twips)

    for row in table.rows:
        col_idx = 0
        for cell in row.cells:
            span = _get_cell_grid_span(cell)
            cell_width = sum(widths_twips[col_idx:col_idx + span]) or widths_twips[min(col_idx, len(widths_twips) - 1)]
            _set_cell_width_twips(cell, cell_width)
            col_idx += span


def _postprocess_docx_tables(docx_path: str, table_specs: list[dict[str, Any]]):
    if not table_specs:
        return

    document = Document(docx_path)
    if not document.tables:
        return

    section = document.sections[0]
    content_width_twips = int((section.page_width - section.left_margin - section.right_margin) / 635)

    for index, table in enumerate(document.tables):
        spec = table_specs[index] if index < len(table_specs) else {}
        _set_table_preferred_width_pct(table, int(spec.get("width_pct", 100)))
        _set_table_alignment(table, str(spec.get("align", "center")))
        _enable_table_autofit(table)

        column_widths_pct = spec.get("column_widths_pct") or []
        if isinstance(column_widths_pct, list) and len(column_widths_pct) == len(table.columns):
            _apply_column_widths_to_table(table, [float(item) for item in column_widths_pct], content_width_twips)

    document.save(docx_path)


def convert_html_to_docx_via_libreoffice(
    html_text: str,
    output_path: str,
    html_output_path: str | None = None,
    libreoffice_path: str = LIBREOFFICE_PATH,
) -> str:
    """将 HTML 写盘后通过 LibreOffice 转为 DOCX。"""
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    html_file = Path(html_output_path) if html_output_path else output_file.with_suffix(".html")
    html_file.parent.mkdir(parents=True, exist_ok=True)
    html_file.write_text(html_text, encoding="utf-8")

    expected_docx = html_file.with_suffix(".docx")
    if expected_docx.exists():
        expected_docx.unlink()
    if output_file.exists() and output_file != expected_docx:
        output_file.unlink()

    docx_file = convert_to_docx_via_libreoffice(
        input_path=html_file,
        output_path=output_file,
        libreoffice_path=libreoffice_path,
    )

    try:
        _postprocess_docx_tables(str(docx_file), _extract_table_layout_specs(html_text))
    except Exception as exc:
        print(f"[warn] DOCX table post-processing failed; keeping LibreOffice output: {exc}")

    return str(docx_file)


def _resolve_libreoffice_path(configured_path: str | None = None) -> str:
    """兼容旧调用方，内部转到共享的 LibreOffice 解析逻辑。"""
    return resolve_libreoffice_path(configured_path)


def convert_text_to_word_via_libreoffice(
    raw_text: str,
    output_path: str,
    html_output_path: str | None = None,
    title: str = "Document",
) -> tuple[str, str]:
    """将文本规范化为 HTML，并用 LibreOffice 转 DOCX。"""
    html_text = normalize_to_word_html(raw_text, title=title)
    html_file = html_output_path or str(Path(output_path).with_suffix(".html"))
    docx_file = convert_html_to_docx_via_libreoffice(
        html_text=html_text,
        output_path=output_path,
        html_output_path=html_file,
    )
    return html_file, docx_file


class HybridToDocxConverter:
    """
    将混合 HTML + Markdown 的 OCR 输出转换为格式化的 Word 文档。
    自己实现解析和渲染，不依赖 pandoc 或 htmldocx。
    """

    def __init__(self):
        self.doc = Document()
        self._setup_default_styles()

    def _setup_default_styles(self):
        """设置文档默认样式"""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.15

        # 设置中文字体回退
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>')
            rPr.append(rFonts)
        else:
            rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # 设置页面边距
        for section in self.doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    # ----------------------------------------------------------
    # 主转换入口
    # ----------------------------------------------------------
    def convert(self, raw_text: str, output_path: str):
        """主转换方法"""
        # Step 1: 预处理 — 统一格式
        processed = self._preprocess(raw_text)

        # Step 2: 按块解析并渲染
        self._parse_and_render(processed)

        # Step 3: 保存
        self.doc.save(output_path)
        print(f"✅ Word 文档已保存: {output_path}")

    # ----------------------------------------------------------
    # 预处理
    # ----------------------------------------------------------
    def _preprocess(self, text: str) -> str:
        """将 Markdown 语法统一转为 HTML 标签，方便后续用 BS4 解析"""
        # 去除可能的 code fence 包裹
        text = re.sub(r"^```(?:markdown|html)?\s*\n?", "", text)
        text = re.sub(r"\n?```\s*$", "", text)

        # Markdown bold **text** → <strong>text</strong>
        # 注意不要破坏已经在 HTML 标签属性里的 ** (极少见)
        text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)

        # Markdown italic *text* → <em>text</em>  (单个 * 且不在 ** 内)
        text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<em>\1</em>", text)

        # Markdown headers: # Header → <h1>Header</h1>
        text = re.sub(r"^######\s+(.+)$", r"<h6>\1</h6>", text, flags=re.MULTILINE)
        text = re.sub(r"^#####\s+(.+)$", r"<h5>\1</h5>", text, flags=re.MULTILINE)
        text = re.sub(r"^####\s+(.+)$", r"<h4>\1</h4>", text, flags=re.MULTILINE)
        text = re.sub(r"^###\s+(.+)$", r"<h3>\1</h3>", text, flags=re.MULTILINE)
        text = re.sub(r"^##\s+(.+)$", r"<h2>\1</h2>", text, flags=re.MULTILINE)
        text = re.sub(r"^#\s+(.+)$", r"<h1>\1</h1>", text, flags=re.MULTILINE)

        # Markdown ordered list: 1. item → <ol_item>item</ol_item> (自定义标记)
        text = re.sub(
            r"^(\d+)\.\s+(.+)$",
            r'<p class="ol_item" data-num="\1">\2</p>',
            text,
            flags=re.MULTILINE,
        )

        # Markdown unordered list: - item → <p class="ul_item">item</p>
        text = re.sub(
            r"^[-*+]\s+(.+)$",
            r'<p class="ul_item">\1</p>',
            text,
            flags=re.MULTILINE,
        )

        # Markdown horizontal rule --- → <hr>
        text = re.sub(r"^---+\s*$", "<hr/>", text, flags=re.MULTILINE)

        # 包装为 HTML body 以便 BS4 解析
        html = f"<body>{text}</body>"
        return html

    # ----------------------------------------------------------
    # 解析与渲染
    # ----------------------------------------------------------
    def _parse_and_render(self, html: str):
        """用 BeautifulSoup 解析 HTML 并逐块渲染到 docx"""
        soup = BeautifulSoup(html, "html.parser")
        body = soup.find("body") or soup

        for element in body.children:
            self._render_element(element)

    def _render_element(self, element):
        """递归渲染单个元素"""
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text and text != "\n":
                # 纯文本行 → 段落
                lines = text.split("\n")
                for line in lines:
                    line = line.strip()
                    if line:
                        para = self.doc.add_paragraph()
                        self._add_inline_text(para, line)
            return

        if not isinstance(element, Tag):
            return

        tag_name = element.name.lower()

        if tag_name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._render_heading(element, int(tag_name[1]))
        elif tag_name == "table":
            self._render_table(element)
        elif tag_name == "hr":
            self._render_hr()
        elif tag_name == "page_break":
            para = self.doc.add_paragraph()
            run = para.add_run()
            run.add_break(WD_BREAK.PAGE)
        elif tag_name == "br":
            self.doc.add_paragraph()  # 空行
        elif tag_name == "p":
            self._render_paragraph(element)
        elif tag_name == "div":
            self._render_div(element)
        elif tag_name in ("strong", "b", "em", "i", "span"):
            # 顶层的 inline 元素 → 创建段落
            para = self.doc.add_paragraph()
            self._render_inline(para, element)
        elif tag_name in ("ul", "ol"):
            self._render_list(element)
        else:
            # 其他标签: 递归处理子元素
            for child in element.children:
                self._render_element(child)

    # ----------------------------------------------------------
    # 标题渲染
    # ----------------------------------------------------------
    def _render_heading(self, element, level: int):
        heading = self.doc.add_heading(level=min(level, 9))
        heading.clear()
        self._render_inline_children(heading, element)
        # 调整标题样式
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    # ----------------------------------------------------------
    # 段落渲染
    # ----------------------------------------------------------
    def _render_paragraph(self, element):
        para = self.doc.add_paragraph()
        css_class = element.get("class", [])

        if "ol_item" in css_class:
            num = element.get("data-num", "1")
            run = para.add_run(f"{num}. ")
            self._render_inline_children(para, element)
            para.paragraph_format.left_indent = Cm(1.0)
        elif "ul_item" in css_class:
            run = para.add_run("• ")
            self._render_inline_children(para, element)
            para.paragraph_format.left_indent = Cm(1.0)
        else:
            self._render_inline_children(para, element)

        # 解析 style 中的 text-align
        self._apply_paragraph_style(para, element)

    def _render_div(self, element):
        """处理 div 元素"""
        align = (element.get("align") or "").lower()
        style = element.get("style", "")

        # 检查 div 内部是否包含块级元素
        has_block = any(
            isinstance(c, Tag) and c.name in ("table", "div", "p", "h1", "h2", "h3", "h4", "h5", "h6", "hr", "ul", "ol")
            for c in element.children
        )

        if has_block:
            # 包含块级子元素 → 递归处理每个子元素
            for child in element.children:
                if isinstance(child, Tag):
                    if child.name == "table":
                        # 传递父 div 的对齐信息
                        self._render_table(child, parent_align=align)
                    else:
                        self._render_element(child)
                elif isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        para = self.doc.add_paragraph()
                        self._add_inline_text(para, text)
                        if align == "right":
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif align == "center":
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # 纯 inline 内容 → 作为一个段落
            para = self.doc.add_paragraph()
            self._render_inline_children(para, element)
            if align == "right":
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == "center":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self._apply_paragraph_style(para, element)

    # ----------------------------------------------------------
    # 表格渲染 (核心)
    # ----------------------------------------------------------
    def _render_table(self, table_element, parent_align: str = ""):
        """将 HTML <table> 渲染为 Word 表格"""
        rows_data = []
        for tr in table_element.find_all("tr", recursive=False):
            # 也检查 thead/tbody 内的 tr
            cells = tr.find_all(["td", "th"], recursive=False)
            rows_data.append(cells)

        if not rows_data:
            # 可能 tr 在 thead/tbody 内
            for section in table_element.find_all(["thead", "tbody", "tfoot"], recursive=False):
                for tr in section.find_all("tr", recursive=False):
                    cells = tr.find_all(["td", "th"], recursive=False)
                    rows_data.append(cells)

        if not rows_data:
            return

        num_rows = len(rows_data)
        num_cols = max(len(row) for row in rows_data)

        if num_cols == 0:
            return

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.autofit = True

        # 解析表格级样式
        table_style = table_element.get("style", "")
        has_border_none = "border: none" in table_style or "border:none" in table_style

        # 设置表格对齐
        if parent_align == "right":
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        elif parent_align == "center":
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_idx, cells in enumerate(rows_data):
            for col_idx, cell_elem in enumerate(cells):
                if col_idx >= num_cols:
                    break
                cell = table.cell(row_idx, col_idx)

                # 清空默认段落
                cell.paragraphs[0].clear()

                # 渲染单元格内容
                self._render_cell_content(cell, cell_elem)

                # 应用单元格样式
                self._apply_cell_style(cell, cell_elem, has_border_none)

        # 如果表格整体无边框，移除所有边框
        if has_border_none:
            self._remove_table_borders(table)

    def _render_cell_content(self, cell, cell_elem):
        """渲染表格单元格的内容"""
        # 获取单元格的第一个段落
        para = cell.paragraphs[0]

        # 检查单元格内是否有块级元素
        block_elements = cell_elem.find_all(["div", "p", "br"], recursive=False)

        if not block_elements:
            # 简单 inline 内容
            self._render_inline_children(para, cell_elem)
        else:
            # 有块级内容 → 逐个处理
            first_para = True
            for child in cell_elem.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        if first_para:
                            self._add_inline_text(para, text)
                            first_para = False
                        else:
                            new_para = cell.add_paragraph()
                            self._add_inline_text(new_para, text)
                elif isinstance(child, Tag):
                    if child.name == "br":
                        if first_para:
                            first_para = False
                        else:
                            para = cell.add_paragraph()
                    elif child.name == "div":
                        target_para = para if first_para else cell.add_paragraph()
                        first_para = False
                        self._render_div_in_cell(target_para, child, cell)
                    elif child.name in ("strong", "b"):
                        target_para = para if first_para else cell.add_paragraph()
                        first_para = False
                        run = target_para.add_run(child.get_text())
                        run.bold = True
                    else:
                        target_para = para if first_para else cell.add_paragraph()
                        first_para = False
                        self._render_inline(target_para, child)

    def _render_div_in_cell(self, para, div_elem, cell):
        """在单元格内渲染 div"""
        style = div_elem.get("style", "")

        # 检查是否有上边框（签名线）
        has_border_top = "border-top" in style

        # 如果有 border-top，先添加一条线
        if has_border_top:
            self._add_bottom_border_to_paragraph(para)

        # 检查字体样式（如手写体）
        is_cursive = "cursive" in style or "Brush Script" in style
        font_size_match = re.search(r"font-size:\s*(\d+)px", style)
        font_size = int(font_size_match.group(1)) if font_size_match else None

        # 渲染 div 内容
        for child in div_elem.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    run = para.add_run(text)
                    if is_cursive:
                        run.font.name = "Segoe Script"  # 模拟手写体
                        run.italic = True
                    if font_size:
                        run.font.size = Pt(font_size * 0.75)  # px to pt 近似
            elif isinstance(child, Tag):
                if child.name == "br":
                    # 在同一个 cell 中新建段落
                    para = cell.add_paragraph()
                    if has_border_top:
                        # 后续段落不需要边框
                        pass
                elif child.name in ("strong", "b"):
                    run = para.add_run(child.get_text())
                    run.bold = True
                else:
                    self._render_inline(para, child)

    def _apply_cell_style(self, cell, cell_elem, table_border_none: bool):
        """应用单元格样式"""
        style = cell_elem.get("style", "")

        # 文字对齐
        if "text-align: left" in style or "text-align:left" in style:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif "text-align: center" in style or "text-align:center" in style:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "text-align: right" in style or "text-align:right" in style:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 垂直对齐
        if "vertical-align: middle" in style:
            cell.vertical_alignment = 1  # WD_ALIGN_VERTICAL.CENTER
        elif "vertical-align: bottom" in style:
            cell.vertical_alignment = 2  # WD_ALIGN_VERTICAL.BOTTOM

        # 宽度
        width_match = re.search(r"width:\s*(\d+)%", style)
        if width_match:
            # 按百分比设置 (基于 A4 可用宽度约 16cm)
            pct = int(width_match.group(1))
            cell.width = Cm(16.0 * pct / 100)

        # 单元格级别的边框移除
        if "border: none" in style or "border:none" in style or table_border_none:
            self._remove_cell_borders(cell)

    def _remove_table_borders(self, table):
        """移除整个表格的所有边框"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
            tbl.insert(0, tblPr)

        borders_xml = f"""
        <w:tblBorders {nsdecls('w')}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tblBorders>
        """
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(parse_xml(borders_xml))

    def _remove_cell_borders(self, cell):
        """移除单个单元格的边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders_xml = f"""
        <w:tcBorders {nsdecls('w')}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tcBorders>
        """
        existing = tcPr.find(qn("w:tcBorders"))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(parse_xml(borders_xml))

    def _add_bottom_border_to_paragraph(self, para):
        """为段落添加上边框线（模拟签名线 border-top）"""
        pPr = para._p.get_or_add_pPr()
        borders_xml = f"""
        <w:pBdr {nsdecls('w')}>
            <w:top w:val="single" w:sz="4" w:space="1" w:color="000000"/>
        </w:pBdr>
        """
        existing = pPr.find(qn("w:pBdr"))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(parse_xml(borders_xml))

    # ----------------------------------------------------------
    # 水平线
    # ----------------------------------------------------------
    def _render_hr(self):
        para = self.doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        borders_xml = f"""
        <w:pBdr {nsdecls('w')}>
            <w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>
        </w:pBdr>
        """
        pPr.append(parse_xml(borders_xml))

    # ----------------------------------------------------------
    # 列表渲染
    # ----------------------------------------------------------
    def _render_list(self, list_elem):
        is_ordered = list_elem.name == "ol"
        for idx, li in enumerate(list_elem.find_all("li", recursive=False), 1):
            para = self.doc.add_paragraph()
            prefix = f"{idx}. " if is_ordered else "• "
            para.add_run(prefix)
            self._render_inline_children(para, li)
            para.paragraph_format.left_indent = Cm(1.0)

    # ----------------------------------------------------------
    # Inline 内容渲染
    # ----------------------------------------------------------
    def _render_inline_children(self, para, element):
        """渲染元素的所有 inline 子节点到指定段落"""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                # 保留空格但去除纯换行
                text = text.replace("\n", " ")
                if text.strip() or text == " ":
                    self._add_inline_text(para, text)
            elif isinstance(child, Tag):
                self._render_inline(para, child)

    def _render_inline(self, para, element):
        """渲染单个 inline 元素"""
        tag = element.name.lower()

        if tag in ("strong", "b"):
            text = element.get_text()
            run = para.add_run(text)
            run.bold = True
        elif tag in ("em", "i"):
            text = element.get_text()
            run = para.add_run(text)
            run.italic = True
        elif tag == "u":
            text = element.get_text()
            run = para.add_run(text)
            run.underline = True
        elif tag == "br":
            para.add_run().add_break()
        elif tag == "span":
            style = element.get("style", "")
            text = element.get_text()
            run = para.add_run(text)
            self._apply_run_style(run, style)
        elif tag == "a":
            text = element.get_text()
            run = para.add_run(text)
            run.font.color.rgb = RGBColor(0, 0, 238)
            run.underline = True
        elif tag == "sup":
            text = element.get_text()
            run = para.add_run(text)
            run.font.superscript = True
        elif tag == "sub":
            text = element.get_text()
            run = para.add_run(text)
            run.font.subscript = True
        elif tag == "div":
            # inline 上下文中的 div → 换行后继续
            self._render_inline_children(para, element)
        else:
            # 未知 inline 标签 → 提取文本
            self._render_inline_children(para, element)

    def _add_inline_text(self, para, text: str):
        """添加纯文本 run"""
        if text:
            para.add_run(text)

    def _apply_run_style(self, run, style: str):
        """从 CSS style 字符串中提取并应用 run 样式"""
        if "font-weight: bold" in style or "font-weight:bold" in style:
            run.bold = True
        if "font-style: italic" in style or "font-style:italic" in style:
            run.italic = True
        if "text-decoration: underline" in style:
            run.underline = True

        # 字体大小
        size_match = re.search(r"font-size:\s*(\d+)px", style)
        if size_match:
            run.font.size = Pt(int(size_match.group(1)) * 0.75)

        # 字体颜色
        color_match = re.search(r"color:\s*#([0-9a-fA-F]{6})", style)
        if color_match:
            hex_color = color_match.group(1)
            run.font.color.rgb = RGBColor(
                int(hex_color[0:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:6], 16),
            )

    def _apply_paragraph_style(self, para, element):
        """从元素 style 属性中提取段落样式"""
        style = element.get("style", "")
        if "text-align: center" in style or "text-align:center" in style:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "text-align: right" in style or "text-align:right" in style:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif "text-align: left" in style or "text-align:left" in style:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ============================================================
# 3. 主函数
# ============================================================
def main():
    # ========== 配置 ==========
    API_KEY = os.getenv("OPENROUTER_API_KEY", "")
    FILE_PATH = r"E:\Pictures\营销用图\营业执照.jpg"
    OUTPUT_DOCX = r"C:\Users\Administrator\Desktop\ocr_output.docx"
    MODEL = "google/gemini-3-flash-preview"
    # ==========================

    # 检查文件是否存在
    if not API_KEY:
        print("鉂?鏈厤缃?OPENROUTER_API_KEY")
        sys.exit(1)

    if not os.path.exists(FILE_PATH):
        print(f"❌ 文件不存在: {FILE_PATH}")
        sys.exit(1)

    # Step 1: OCR
    raw_output = ocr_file(FILE_PATH, API_KEY, MODEL)

    # (可选) 保存原始输出以供调试
    raw_output_path = OUTPUT_DOCX.replace(".docx", "_raw.txt")
    with open(raw_output_path, "w", encoding="utf-8") as f:
        f.write(raw_output)
    print(f"📄 原始 OCR 输出已保存: {raw_output_path}")

    # Step 2: 转换为 Word
    html_output_path = OUTPUT_DOCX.replace(".docx", ".html")
    convert_text_to_word_via_libreoffice(
        raw_output,
        OUTPUT_DOCX,
        html_output_path=html_output_path,
    )

    print(f"\n🎉 全部完成！输出文件: {OUTPUT_DOCX}")


# ============================================================
# 4. 也可以单独使用转换器（跳过 OCR，直接处理已有文本）
# ============================================================
def convert_text_to_docx(raw_text: str, output_path: str):
    """
    便捷函数：直接将混合 HTML/Markdown 文本转换为 Word 文档

    Usage:
        convert_text_to_docx(your_ocr_output, "output.docx")
    """
    convert_text_to_word_via_libreoffice(raw_text, output_path)


if __name__ == "__main__":
    main()
