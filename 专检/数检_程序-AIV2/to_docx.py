#!/usr/bin/env python3
"""
LaTeX 公式 → DOCX 文件转换器（修复版）

将包含 LaTeX 公式的文本转换为 Word 文档，
公式以 Word 原生数学公式对象（OMML）呈现。

依赖安装：
    pip install python-docx latex2mathml lxml
"""

from __future__ import annotations

import re
import copy
import traceback
from pathlib import Path
from dataclasses import dataclass, field
from typing import Literal

from lxml import etree
import latex2mathml.converter
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# 1. 命名空间定义
# ============================================================

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
MATHML_NS = 'http://www.w3.org/1998/Math/MathML'


def _m(tag: str) -> str:
    """OMML 命名空间标签。"""
    return f'{{{MATH_NS}}}{tag}'


def _local(element: etree._Element) -> str:
    """获取元素的本地标签名（去除命名空间）。"""
    tag = element.tag
    if '}' in tag:
        return tag.split('}', 1)[1]
    return tag


# ============================================================
# 2. 尝试加载微软 XSLT 样式表
# ============================================================

def _load_xslt() -> etree.XSLT | None:
    """尝试加载微软 MML2OMML.XSL 文件。"""
    candidates = [
        Path(r"C:/Program Files/Microsoft Office/root/Office16/MML2OMML.XSL"),
        Path(r"C:/Program Files (x86)/Microsoft Office/root/Office16/MML2OMML.XSL"),
        Path(r"C:/Program Files/Microsoft Office/Office16/MML2OMML.XSL"),
        Path(r"C:/Program Files/Microsoft Office/root/Office15/MML2OMML.XSL"),
        Path("MML2OMML.XSL"),
        Path("mml2omml.xsl"),
    ]
    for p in candidates:
        if p.exists():
            try:
                return etree.XSLT(etree.parse(str(p)))
            except Exception:
                continue
    return None


_GLOBAL_XSLT = _load_xslt()


# ============================================================
# 3. MathML → OMML 手动转换器（完全重写，健壮版）
# ============================================================

class MathMLToOMML:
    """将 MathML 元素树转换为 OMML 元素树。"""

    def convert(self, mml_root: etree._Element) -> etree._Element:
        """入口：将 MathML <math> 根元素转为 OMML <m:oMath>。"""
        o_math = etree.Element(_m('oMath'))
        self._process_children(mml_root, o_math)
        return o_math

    def _process_children(self, mml_parent: etree._Element, omml_parent: etree._Element) -> None:
        """递归处理所有子节点，将结果追加到 omml_parent。"""
        for child in mml_parent:
            results = self._convert_node(child)
            for r in results:
                omml_parent.append(r)

    def _convert_node(self, node: etree._Element) -> list[etree._Element]:
        """
        将单个 MathML 节点转为 OMML 节点列表。
        返回列表是因为某些节点（如 mrow）会展开为多个兄弟节点。
        """
        tag = _local(node)

        match tag:
            case 'mi' | 'mn' | 'mo' | 'mtext' | 'ms':
                return [self._make_run(node)]
            case 'mrow' | 'mstyle' | 'mpadded' | 'merror' | 'mphantom':
                return self._convert_mrow(node)
            case 'mfrac':
                return [self._convert_mfrac(node)]
            case 'msqrt':
                return [self._convert_msqrt(node)]
            case 'mroot':
                return [self._convert_mroot(node)]
            case 'msup':
                return [self._convert_msup(node)]
            case 'msub':
                return [self._convert_msub(node)]
            case 'msubsup':
                return [self._convert_msubsup(node)]
            case 'munderover':
                return [self._convert_msubsup(node)]
            case 'munder':
                return [self._convert_msub(node)]
            case 'mover':
                return [self._convert_mover(node)]
            case 'mfenced':
                return [self._convert_mfenced(node)]
            case 'mtable':
                return [self._convert_mtable(node)]
            case 'mtr':
                return [self._convert_mtr(node)]
            case 'mtd':
                return [self._convert_mtd(node)]
            case 'mspace':
                return [self._make_text_run(' ')]
            case 'none':
                return []
            case _:
                # 未知标签：尝试处理子节点
                return self._convert_mrow(node)

    def _make_run(self, token_node: etree._Element) -> etree._Element:
        """将 MathML 文本节点 (mi/mn/mo/mtext) 转为 OMML <m:r>。"""
        text = self._get_text(token_node)
        tag = _local(token_node)

        r = etree.Element(_m('r'))

        # mi 多字符 → 正体样式（如 sin, cos, lim）
        if tag == 'mi' and len(text) > 1:
            rpr = etree.SubElement(r, _m('rPr'))
            sty = etree.SubElement(rpr, _m('sty'))
            sty.set(_m('val'), 'p')

        t = etree.SubElement(r, _m('t'))
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        return r

    def _make_text_run(self, text: str) -> etree._Element:
        """创建包含指定文本的 OMML run。"""
        r = etree.Element(_m('r'))
        t = etree.SubElement(r, _m('t'))
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        return r

    def _convert_mrow(self, node: etree._Element) -> list[etree._Element]:
        """<mrow> → 展开为子节点列表。"""
        results: list[etree._Element] = []
        for child in node:
            results.extend(self._convert_node(child))
        return results

    def _convert_mfrac(self, node: etree._Element) -> etree._Element:
        """<mfrac> → <m:f>（分数）。"""
        children = list(node)
        f = etree.Element(_m('f'))

        # 分数属性
        f_pr = etree.SubElement(f, _m('fPr'))
        # 默认即可

        # 分子
        num = etree.SubElement(f, _m('num'))
        if len(children) > 0:
            self._fill_element(num, children[0])

        # 分母
        den = etree.SubElement(f, _m('den'))
        if len(children) > 1:
            self._fill_element(den, children[1])

        return f

    def _convert_msqrt(self, node: etree._Element) -> etree._Element:
        """<msqrt> → <m:rad>（平方根）。"""
        rad = etree.Element(_m('rad'))

        # 隐藏次数
        rad_pr = etree.SubElement(rad, _m('radPr'))
        deg_hide = etree.SubElement(rad_pr, _m('degHide'))
        deg_hide.set(_m('val'), '1')

        # 空的 deg
        etree.SubElement(rad, _m('deg'))

        # 被开方数
        e = etree.SubElement(rad, _m('e'))
        self._process_children(node, e)

        return rad

    def _convert_mroot(self, node: etree._Element) -> etree._Element:
        """<mroot> → <m:rad>（n次根）。"""
        children = list(node)
        rad = etree.Element(_m('rad'))

        rad_pr = etree.SubElement(rad, _m('radPr'))

        # 次数
        deg = etree.SubElement(rad, _m('deg'))
        if len(children) > 1:
            self._fill_element(deg, children[1])

        # 被开方数
        e = etree.SubElement(rad, _m('e'))
        if len(children) > 0:
            self._fill_element(e, children[0])

        return rad

    def _convert_msup(self, node: etree._Element) -> etree._Element:
        """<msup> → <m:sSup>（上标）。"""
        children = list(node)
        ssup = etree.Element(_m('sSup'))

        ssup_pr = etree.SubElement(ssup, _m('sSupPr'))

        e = etree.SubElement(ssup, _m('e'))
        if len(children) > 0:
            self._fill_element(e, children[0])

        sup = etree.SubElement(ssup, _m('sup'))
        if len(children) > 1:
            self._fill_element(sup, children[1])

        return ssup

    def _convert_msub(self, node: etree._Element) -> etree._Element:
        """<msub> / <munder> → <m:sSub>（下标）。"""
        children = list(node)
        ssub = etree.Element(_m('sSub'))

        ssub_pr = etree.SubElement(ssub, _m('sSubPr'))

        e = etree.SubElement(ssub, _m('e'))
        if len(children) > 0:
            self._fill_element(e, children[0])

        sub = etree.SubElement(ssub, _m('sub'))
        if len(children) > 1:
            self._fill_element(sub, children[1])

        return ssub

    def _convert_msubsup(self, node: etree._Element) -> etree._Element:
        """<msubsup> / <munderover> → <m:sSubSup>（上下标）。"""
        children = list(node)
        sss = etree.Element(_m('sSubSup'))

        sss_pr = etree.SubElement(sss, _m('sSubSupPr'))

        e = etree.SubElement(sss, _m('e'))
        if len(children) > 0:
            self._fill_element(e, children[0])

        sub = etree.SubElement(sss, _m('sub'))
        if len(children) > 1:
            self._fill_element(sub, children[1])

        sup = etree.SubElement(sss, _m('sup'))
        if len(children) > 2:
            self._fill_element(sup, children[2])

        return sss

    def _convert_mover(self, node: etree._Element) -> etree._Element:
        """<mover> → 根据上方符号决定：帽子/横线用 acc，其他用 sSup。"""
        children = list(node)

        # 检查上方是否是装饰符号（如 ^ ¯ ~ ·）
        if len(children) > 1:
            over_text = self._get_text(children[1]).strip()
            accent_map = {
                '̂': '\u0302', '^': '̂',
                '¯': '̄', '―': '\u0304', '‾': '̅',
                '~': '̃', '˜': '̃',
                '·': '\u0307', '˙': '̇',
                '→': '\u20D7',
            }
            if over_text in accent_map:
                # 使用 OMML accent
                acc = etree.Element(_m('acc'))
                acc_pr = etree.SubElement(acc, _m('accPr'))
                char_elem = etree.SubElement(acc_pr, _m('chr'))
                char_elem.set(_m('val'), accent_map[over_text])

                e = etree.SubElement(acc, _m('e'))
                self._fill_element(e, children[0])
                return acc

        # 默认当作上标
        return self._convert_msup(node)

    def _convert_mfenced(self, node: etree._Element) -> etree._Element:
        """<mfenced> → <m:d>（定界符/括号）。"""
        open_ch = node.get('open', '(')
        close_ch = node.get('close', ')')
        separators = node.get('separators', ',')

        d = etree.Element(_m('d'))
        d_pr = etree.SubElement(d, _m('dPr'))

        beg = etree.SubElement(d_pr, _m('begChr'))
        beg.set(_m('val'), open_ch)
        end_chr = etree.SubElement(d_pr, _m('endChr'))
        end_chr.set(_m('val'), close_ch)

        # 每个子元素作为一个 <m:e>
        children = list(node)
        if children:
            for i, child in enumerate(children):
                e = etree.SubElement(d, _m('e'))
                self._fill_element(e, child)
        else:
            etree.SubElement(d, _m('e'))

        return d

    def _convert_mtable(self, node: etree._Element) -> etree._Element:
        """<mtable> → <m:m>（矩阵）。"""
        m_elem = etree.Element(_m('m'))
        m_pr = etree.SubElement(m_elem, _m('mPr'))

        for child in node:
            if _local(child) == 'mtr':
                mr = self._convert_mtr(child)
                m_elem.append(mr)

        return m_elem

    def _convert_mtr(self, node: etree._Element) -> etree._Element:
        """<mtr> → <m:mr>。"""
        mr = etree.Element(_m('mr'))
        for child in node:
            if _local(child) == 'mtd':
                e = self._convert_mtd(child)
                mr.append(e)
        return mr

    def _convert_mtd(self, node: etree._Element) -> etree._Element:
        """<mtd> → <m:e>。"""
        e = etree.Element(_m('e'))
        self._process_children(node, e)
        return e

    def _fill_element(self, omml_parent: etree._Element, mml_node: etree._Element) -> None:
        """将单个 MathML 节点的转换结果填入 OMML 父元素。"""
        results = self._convert_node(mml_node)
        for r in results:
            omml_parent.append(r)

    @staticmethod
    def _get_text(node: etree._Element) -> str:
        """递归获取节点的所有文本内容。"""
        parts: list[str] = []
        if node.text:
            parts.append(node.text)
        for child in node:
            parts.append(MathMLToOMML._get_text(child))
            if child.tail:
                parts.append(child.tail)
        return ''.join(parts)


# ============================================================
# 4. LaTeX → OMML 主转换器
# ============================================================

class LaTeXToOMML:
    """LaTeX 公式 → Word OMML 元素。"""

    def __init__(self) -> None:
        self._xslt = _GLOBAL_XSLT
        self._manual = MathMLToOMML()

    def convert(self, latex: str) -> etree._Element:
        """
        LaTeX → OMML。
        链路：LaTeX → MathML(字符串) → MathML(XML) → OMML(XML)
        """
        # 预处理 LaTeX
        latex = self._preprocess(latex)

        # 第一步：LaTeX → MathML 字符串
        try:
            mathml_str = latex2mathml.converter.convert(latex)
        except Exception as e:
            print(f"  ⚠️ latex2mathml 转换失败: {latex!r}")
            print(f"     错误: {e}")
            return self._fallback_text(latex)

        # 调试：打印 MathML
        # print(f"  [DEBUG] MathML: {mathml_str[:200]}")

        # 第二步：尝试 XSLT 方案
        if self._xslt is not None:
            try:
                return self._via_xslt(mathml_str)
            except Exception:
                pass  # 降级到手动方案

        # 第三步：手动转换方案
        try:
            return self._via_manual(mathml_str)
        except Exception as e:
            print(f"  ⚠️ 手动 OMML 转换失败: {e}")
            traceback.print_exc()
            return self._fallback_text(latex)

    @staticmethod
    def _preprocess(latex: str) -> str:
        """预处理 LaTeX 源码，修复常见兼容问题。"""
        s = latex.strip()
        # latex2mathml 不认识 \Delta，替换为 Δ（直接 Unicode）
        # 但实际上 latex2mathml 支持 \Delta，先试试
        # 处理中文/Unicode 数学符号（HTML 中可能直接写了 ∆ 而非 \Delta）
        s = s.replace('∆', r'\Delta')
        s = s.replace('−', '-')
        # 移除 \displaystyle 等
        s = re.sub(r'\\(displaystyle|textstyle|scriptstyle)\s*', '', s)
        # \left \right 后面如果跟 . 表示空定界符
        s = re.sub(r'\\left\s*\.', '', s)
        s = re.sub(r'\\right\s*\.', '', s)
        return s

    def _via_xslt(self, mathml_str: str) -> etree._Element:
        """通过 XSLT 转换。"""
        if 'xmlns' not in mathml_str:
            mathml_str = mathml_str.replace('<math>', f'<math xmlns="{MATHML_NS}">')
        tree = etree.fromstring(mathml_str.encode('utf-8'))
        result = self._xslt(tree)
        root = result.getroot()
        return self._extract_omath(root)

    def _via_manual(self, mathml_str: str) -> etree._Element:
        """通过手动转换器。"""
        # 解析 MathML XML
        try:
            mml_root = etree.fromstring(mathml_str.encode('utf-8'))
        except etree.XMLSyntaxError:
            # 尝试去除命名空间后重新解析
            cleaned = re.sub(r'\s*xmlns="[^"]*"', '', mathml_str)
            mml_root = etree.fromstring(cleaned.encode('utf-8'))

        return self._manual.convert(mml_root)

    @staticmethod
    def _extract_omath(element: etree._Element) -> etree._Element:
        """从转换结果中提取 oMath 元素。"""
        local = _local(element)
        if local == 'oMath':
            return element
        if local == 'oMathPara':
            for child in element:
                if _local(child) == 'oMath':
                    return child
        # 包装
        o = etree.Element(_m('oMath'))
        o.append(copy.deepcopy(element))
        return o

    @staticmethod
    def _fallback_text(text: str) -> etree._Element:
        """转换失败时的纯文本兜底。"""
        o_math = etree.Element(_m('oMath'))
        r = etree.SubElement(o_math, _m('r'))
        t = etree.SubElement(r, _m('t'))
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        return o_math


# ============================================================
# 5. 文本段数据结构
# ============================================================

@dataclass
class FormulaSegment:
    """文本段。"""
    content: str
    kind: Literal['text', 'inline_formula', 'display_formula']


# ============================================================
# 6. DOCX 文档生成器
# ============================================================

@dataclass
class DocxFormulaWriter:
    """将混合了 LaTeX 公式的文本写入 DOCX 文件。"""

    output_path: Path
    font_name: str = "微软雅黑"
    font_size: int = 12
    _converter: LaTeXToOMML = field(default_factory=LaTeXToOMML, repr=False)
    _doc: Document = field(default_factory=Document, repr=False)

    def add_heading(self, text: str, level: int = 1) -> None:
        """添加标题。"""
        self._doc.add_heading(text, level=level)

    def add_mixed_content(self, text: str) -> None:
        """
        添加混合内容（文本 + 公式）。
        自动识别 $$...$$ 和 $...$ 公式。
        """
        segments = self._parse_segments(text)
        current_inline: list[FormulaSegment] = []

        for seg in segments:
            if seg.kind == 'display_formula':
                if current_inline:
                    self._write_inline_paragraph(current_inline)
                    current_inline = []
                self._write_display_formula(seg.content)
            else:
                current_inline.append(seg)

        if current_inline:
            self._write_inline_paragraph(current_inline)

    def add_latex_formula(self, latex: str, display: bool = True) -> None:
        """直接添加一个 LaTeX 公式。"""
        if display:
            self._write_display_formula(latex)
        else:
            self._write_inline_paragraph([
                FormulaSegment(content=latex, kind='inline_formula')
            ])

    def add_text(self, text: str) -> None:
        """添加纯文本段落。"""
        para = self._doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(self.font_size)
        run.font.name = self.font_name

    def save(self) -> Path:
        """保存文件。"""
        self._doc.save(str(self.output_path))
        print(f"✅ 文件已保存: {self.output_path.resolve()}")
        return self.output_path

    # ---- 内部方法 ----

    @staticmethod
    def _parse_segments(text: str) -> list[FormulaSegment]:
        """解析文本为段落列表。"""
        segments: list[FormulaSegment] = []
        pattern = re.compile(r'(\$\$.*?\$\$|\$(?!\$).+?\$)', re.DOTALL)

        last_end = 0
        for match in pattern.finditer(text):
            start, end = match.span()
            if start > last_end:
                plain = text[last_end:start]
                if plain.strip():
                    segments.append(FormulaSegment(content=plain, kind='text'))

            raw = match.group(0)
            if raw.startswith('$$'):
                segments.append(FormulaSegment(content=raw[2:-2].strip(), kind='display_formula'))
            else:
                segments.append(FormulaSegment(content=raw[1:-1].strip(), kind='inline_formula'))
            last_end = end

        if last_end < len(text):
            remaining = text[last_end:]
            if remaining.strip():
                segments.append(FormulaSegment(content=remaining, kind='text'))

        return segments

    def _write_display_formula(self, latex: str) -> None:
        """块级公式（居中）。"""
        para = self._doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        try:
            omml = self._converter.convert(latex)
            para._element.append(omml)
            print(f"  ✅ 块级公式成功: {latex[:50]}...")
        except Exception as e:
            run = para.add_run(f"[公式: {latex}]")
            run.font.size = Pt(self.font_size)
            print(f"  ❌ 块级公式失败: {latex!r} → {e}")

    def _write_inline_paragraph(self, segments: list[FormulaSegment]) -> None:
        """行内混排段落。"""
        para = self._doc.add_paragraph()

        for seg in segments:
            if seg.kind == 'text':
                run = para.add_run(seg.content)
                run.font.size = Pt(self.font_size)
                run.font.name = self.font_name
            elif seg.kind == 'inline_formula':
                try:
                    omml = self._converter.convert(seg.content)
                    para._element.append(omml)
                except Exception as e:
                    run = para.add_run(f" [{seg.content}] ")
                    run.font.size = Pt(self.font_size)
                    print(f"  ❌ 行内公式失败: {seg.content!r} → {e}")


# ============================================================
# 7. HTML → DOCX
# ============================================================

def html_to_docx(html_source: str, output_path: str | Path) -> Path:
    """从 HTML 中提取文本和公式，生成 DOCX。"""
    import html as html_module

    output_path = Path(output_path)
    writer = DocxFormulaWriter(output_path=output_path)

    # 清理 HTML
    text = re.sub(r'<script[^>]*>.*?</script>', '', html_source, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)

    # 提取 body
    body_match = re.search(r'<body[^>]*>(.*?)</body>', text, re.DOTALL | re.IGNORECASE)
    if body_match:
        text = body_match.group(1)

    # 提取标题
    for m in re.finditer(r'<h(\d)[^>]*>(.*?)</h\1>', text, re.DOTALL | re.IGNORECASE):
        level = int(m.group(1))
        title = re.sub(r'<[^>]+>', '', m.group(2)).strip()
        title = html_module.unescape(title)
        writer.add_heading(title, level=min(level, 4))

    # 移除标题标签
    text = re.sub(r'<h\d[^>]*>.*?</h\d>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</(p|div|li)>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<(p|div|li)[^>]*>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = html_module.unescape(text)

    for line in text.split('\n'):
        line = line.strip()
        if line:
            writer.add_mixed_content(line)

    return writer.save()


# ============================================================
# 8. 主程序 + 调试
# ============================================================

def main() -> None:
    """生成包含数学公式的 DOCX 文件。"""

    output_file = Path("math_formulas.docx")

    print("=" * 60)
    print("  LaTeX 公式 → DOCX 转换器（修复版）")
    print("=" * 60)

    if _GLOBAL_XSLT:
        print("  📋 已加载微软 MML2OMML.XSL（最佳质量）")
    else:
        print("  📋 未找到 MML2OMML.XSL，使用内置转换器")

    # --- 先单独测试每个公式的 MathML 转换 ---
    print("\n🔍 调试：逐个测试公式转换")
    print("-" * 60)

    test_formulas = {
        "求根公式": r"\frac{-b \pm \sqrt{b^2 - 4ac}}{2a}",
        "动能定理": r"W = \frac{1}{2}mv^{2} - \frac{1}{2}mv_{0}^{2}",
        "欧拉公式": r"e^{i\pi} + 1 = 0",
        "高斯积分": r"\int_{-\infty}^{\infty} e^{-x^2} dx = \sqrt{\pi}",
        "求和公式": r"\sum_{i=1}^{n} i = \frac{n(n+1)}{2}",
    }

    converter = LaTeXToOMML()

    for name, latex in test_formulas.items():
        print(f"\n  [{name}] {latex}")
        try:
            # 测试 MathML 生成
            preprocessed = converter._preprocess(latex)
            mathml = latex2mathml.converter.convert(preprocessed)
            print(f"    MathML: {mathml[:120]}...")

            # 测试 OMML 生成
            omml = converter.convert(latex)
            omml_str = etree.tostring(omml, pretty_print=True).decode()
            child_count = len(list(omml))
            print(f"    OMML 子元素数: {child_count}")
            if child_count == 0:
                print(f"    ⚠️ OMML 为空！完整 OMML:")
                print(f"    {omml_str[:300]}")
            else:
                print(f"    ✅ 转换成功")
        except Exception as e:
            print(f"    ❌ 失败: {e}")
            traceback.print_exc()

    # --- 生成 DOCX ---
    print("\n" + "=" * 60)
    print("📄 生成 DOCX 文件")
    print("=" * 60)

    writer = DocxFormulaWriter(output_path=output_file)
    writer.add_heading("数学公式演示文档", level=1)

    writer.add_heading("一、经典公式", level=2)

    for name, latex in test_formulas.items():
        writer.add_text(f"{name}：")
        writer.add_latex_formula(latex, display=True)

    writer.add_heading("二、行内公式", level=2)

    mixed_texts = [
        r"""3图表标题
  [系列名] 系列 1
  [分类] 类别1	类别2	类别3	类别4
  [数据] 4.3	2.5	3.5	4.5
  [系列名] 系列 2
  [分类] 类别1	类别2	类别3	类别4
  [数据] 2.4	4.4	1.8	2.8
  [系列名] 系列 3
  [分类] 类别1	类别2	类别3	类别4
  [数据] 2	2	3	5"""
        r"$V_{0总}^{2}$",
        r"勾股定理：$a^2 + b^2 = c^2$，几何学的基石。",
        r"三角函数：$\sin(\theta) = \frac{opp}{hyp}$，物理中广泛应用。",
        r"质能方程 $E = mc^2$ 是爱因斯坦最著名的公式。",
    ]
    for text in mixed_texts:
        writer.add_mixed_content(text)

    writer.save()

    # --- 从 HTML 生成 ---
    html_source = r'''
    <body>
        <h2>📐 渲染效果展示</h2>
        <div>
            $$\frac{-b \pm \sqrt{b^2 - 4ac}}{2a}$$
            $$W=\frac{1}{2}mv^{2}-\frac{1}{2}mv_{0}^{2}$$
        </div>
        <p>行内公式：$\sin(\theta) = \frac{opp}{hyp}$，和文字融合。</p>
        <p>勾股定理 $a^2 + b^2 = c^2$ 非常直观。</p>
    </body>
    '''

    print(f"\n📄 从 HTML 生成 DOCX...")
    html_to_docx(html_source, "from_html_formulas.docx")

    print("\n🎉 全部完成！用 Word 打开文件查看效果。")


if __name__ == "__main__":
    main()