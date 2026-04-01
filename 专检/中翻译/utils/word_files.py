from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from typing import Any, Tuple, Union

from docx import Document


def write_word_with_timestamp(
    data: Any,
    output_dir: Union[str, os.PathLike],
    prefix: str = "文本对比结果",
    as_json_pretty: bool = True,
    tz_aware: bool = False,
) -> Tuple[str, str]:
    """
    将 data 写入 output_dir 目录下的“带时间戳”的 .docx 文件，并返回：
    (文件名, 文件完整路径)

    文件名示例：
      文本对比结果_20260206_153012.docx

    写入规则：
    - 如果 as_json_pretty=True：会把 data 当作对象，按 JSON 形式美化写入 Word
    - 否则：直接 str(data) 写入 Word

    参数：
    - data: 任意对象（dict/list/str...）
    - output_dir: 输出目录（不存在会自动创建）
    - prefix: 文件名前缀
    - as_json_pretty: 是否以 pretty JSON 写入（默认 True）
    - tz_aware: 是否使用本地时区时间（默认 False 使用系统本地；True 会用 datetime.now().astimezone()）
    """
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    now = datetime.now().astimezone() if tz_aware else datetime.now()
    ts = now.strftime("%Y%m%d_%H%M%S")

    filename = f"{prefix}_{ts}.docx"
    filepath = out_dir / filename

    doc = Document()
    doc.add_heading(prefix, level=0)

    if as_json_pretty:
        # 不引入 json 模块也可，但这里写得更直观
        import json
        text = json.dumps(data, ensure_ascii=False, indent=2)
    else:
        text = str(data)

    # 为了避免一整坨挤在同一段里：按行写入
    for line in text.splitlines():
        doc.add_paragraph(line)

    doc.save(str(filepath))
    return filename, str(filepath)


# ===== 用法示例 =====
if __name__ == "__main__":
    sample = {
            "错误编号": "13",
            "错误类型": "数值",
            "原文数值": "2",
            "译文数值": "21",
            "译文修改建议值": "2",
            "修改理由": "页脚内容翻译错误，原文页码为2，译文为21。",
            "违反的规则": "(一) 总体: 准确。",
            "原文上下文": "=== 页脚内容 === 2",
            "译文上下文": "=== 页脚内容 === 21",
            "原文位置": "页脚",
            "译文位置": "Footer",
            "替换锚点": "21",
        }


    name, path = write_word_with_timestamp(sample, r"C:\Users\Administrator\Desktop\project\llm\zhongfanyi\output_word")
    print("word文件名称:", name)
    print("word文件路径:", path)
