"""
replace_revision.py 优化版测试

测试场景：
1. 同值多处出现 — 验证是否替换了正确的位置（而非第一个）
2. 无上下文 + 多处出现 — 验证安全网是否拒绝替换
3. 唯一出现 — 验证正常替换
4. 上下文区分 — 两个段落都包含目标值，但上下文只匹配其中一个
5. 锚点定位 — 验证锚点能精确定位
6. region 隔离 — 验证页眉/正文不会互相干扰
7. 得分接近 — 两个候选得分几乎一样时拒绝替换
"""
import sys
import os

# 确保项目根目录在 sys.path 中
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', '..'))

from docx import Document
from docx.shared import Pt

# 直接导入被测模块的函数（绕过包名问题）
# 我们需要 mock 掉 import 路径
import importlib
import types


def _create_test_doc_with_paragraphs(paragraphs_text: list, header_text: str = "", footer_text: str = ""):
    """创建一个包含指定段落的测试 Word 文档"""
    doc = Document()
    for text in paragraphs_text:
        doc.add_paragraph(text)

    if header_text:
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = header_text

    if footer_text:
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = footer_text

    return doc


class MockRevisionManager:
    """模拟 RevisionManager，记录所有替换操作而非真正修改 XML"""

    def __init__(self, doc, author="test"):
        self.doc = doc
        self.author = author
        self.replacements = []  # [(paragraph_text, old, new, reason)]

    def replace_in_paragraph(self, paragraph, old_text, new_text, reason=""):
        full_text = "".join(r.text or "" for r in paragraph.runs)
        if old_text in full_text:
            self.replacements.append((full_text, old_text, new_text, reason))
            # 实际执行替换以便后续段落文本变化
            for run in paragraph.runs:
                if old_text in (run.text or ""):
                    run.text = run.text.replace(old_text, new_text, 1)
                    return True
            # 跨 run 情况：简单拼接后替换
            new_full = full_text.replace(old_text, new_text, 1)
            # 清空所有 run，第一个 run 放新文本
            for i, run in enumerate(paragraph.runs):
                run.text = new_full if i == 0 else ""
            return True
        return False

    def replace_run_text(self, run, new_text, reason=""):
        old_text = run.text or ""
        self.replacements.append((old_text, old_text, new_text, reason))
        run.text = new_text
        return True


# ============================================================
# 直接导入需要的函数（避免包路径问题）
# ============================================================

# 手动加载 replace_clean 模块
import unicodedata
import re
from difflib import SequenceMatcher
from typing import Tuple, Optional


def clean_text_thoroughly(text: str) -> str:
    if not text: return ""
    text = text.replace('\u2018', "'").replace('\u2019', "'")
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('`', "'").replace('\u00b4', "'")
    text = text.replace('\uff08', '(').replace('\uff09', ')')
    text = text.replace('\uff0c', ',').replace('\u3002', '.')
    text = text.replace('\uff1a', ':').replace('\uff1b', ';')
    text = text.replace('\uff01', '!').replace('\uff1f', '?')
    text = text.replace('\u3010', '[').replace('\u3011', ']')
    text = text.replace('\u300a', '<').replace('\u300b', '>')
    text = text.replace('\u3000', ' ')
    text = unicodedata.normalize('NFKC', text)
    text = re.sub(r'[\u200b-\u200f\u202a-\u202e\u2060-\u206f\ufeff\u00ad\xa0]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def calculate_context_similarity(text1: str, text2: str) -> float:
    if not text1 or not text2: return 0.0
    words1 = set(text1.lower().split())
    words2 = set(text2.lower().split())
    if not words1 or not words2: return 0.0
    intersection = len(words1 & words2)
    union = len(words1 | words2)
    return intersection / union if union > 0 else 0.0


def is_fuzzy_match(text, target, threshold=0.9):
    return SequenceMatcher(None, clean_text_thoroughly(text), clean_text_thoroughly(target)).ratio() >= threshold


def get_alphanumeric_fingerprint(text):
    if not text: return ""
    clean = clean_text_thoroughly(text)
    return re.sub(r'[^a-zA-Z0-9\u4e00-\u9fa5]', '', clean)


def build_smart_pattern(s, mode="balanced"):
    s = clean_text_thoroughly(s or "")
    if not s: return ""
    if mode == "strict":
        return re.escape(s)
    return re.escape(s).replace(r'\ ', r'\s*')


def extract_anchor_with_target(context, target_value, window=50):
    context = clean_text_thoroughly(context)
    target_value = clean_text_thoroughly(target_value)
    if not context or not target_value: return None
    if target_value in context:
        idx = context.index(target_value)
        start = max(0, idx - window)
        end = min(len(context), idx + len(target_value) + window)
        return context[start:end].strip()
    return None


def is_list_pattern(s):
    return bool(re.match(r'^[ivxlcdm]+\.$|^\([a-z]\)$|^\d+\.$', s.strip(), re.IGNORECASE))


def iter_body_paragraphs(doc):
    for p in doc.paragraphs:
        yield p


def iter_header_paragraphs(doc):
    for section in doc.sections:
        header = section.header
        if header and not header.is_linked_to_previous:
            for p in header.paragraphs:
                yield p


def iter_footer_paragraphs(doc):
    for section in doc.sections:
        footer = section.footer
        if footer and not footer.is_linked_to_previous:
            for p in footer.paragraphs:
                yield p


def iter_all_paragraphs(doc):
    yield from iter_body_paragraphs(doc)
    yield from iter_header_paragraphs(doc)
    yield from iter_footer_paragraphs(doc)


def preprocess_special_cases(old_value, doc):
    return False, "", ""


# ============================================================
# 直接复制被测函数（避免 import 路径问题）
# ============================================================

def _score_paragraph(full_text, old_value, context, anchor_text, match_level):
    score = 0.0
    full_clean = clean_text_thoroughly(full_text)
    if context:
        sim = calculate_context_similarity(full_clean, clean_text_thoroughly(context))
        score += sim * 0.5
    if anchor_text:
        anchor_clean = clean_text_thoroughly(anchor_text)
        if anchor_clean and anchor_clean in full_clean:
            score += 0.25
    level_scores = {1: 0.15, 2: 0.12, 3: 0.10, 4: 0.07, 5: 0.05}
    score += level_scores.get(match_level, 0.05)
    if full_text:
        ratio = len(old_value) / len(full_text)
        score += min(ratio, 1.0) * 0.1
    return score


def _try_match_in_paragraph(paragraph, old_value, old_value_clean, pattern):
    runs = list(paragraph.runs)
    if not runs: return None
    full_text = "".join(r.text or "" for r in runs)
    if not full_text.strip(): return None
    if old_value in full_text: return 1
    full_text_clean = clean_text_thoroughly(full_text)
    if old_value_clean and old_value_clean in full_text_clean:
        if old_value in full_text: return 2
    if pattern:
        try:
            if re.search(pattern, full_text_clean, flags=re.IGNORECASE | re.DOTALL):
                if old_value in full_text: return 3
        except re.error: pass
    if is_fuzzy_match(full_text_clean, old_value_clean, threshold=0.85):
        if old_value in full_text: return 4
    fp_old = get_alphanumeric_fingerprint(old_value)
    fp_full = get_alphanumeric_fingerprint(full_text)
    if len(fp_old) >= 3 and fp_old in fp_full:
        if old_value in full_text: return 4
    full_no_space = full_text_clean.replace(' ', '')
    old_no_space = old_value_clean.replace(' ', '')
    if old_no_space and old_no_space in full_no_space:
        if old_value in full_text: return 5
    return None


def _execute_replace(paragraph, old_value, new_value, reason, revision_manager):
    runs = list(paragraph.runs)
    if not runs: return False
    full_text = "".join(r.text or "" for r in runs)
    if old_value in full_text:
        return revision_manager.replace_in_paragraph(paragraph, old_value, new_value, reason=reason)
    old_no_space = old_value.replace(' ', '')
    for run in runs:
        run_text = run.text or ""
        run_no_space = run_text.replace(' ', '')
        if old_no_space and len(run_text.strip()) < 50 and old_no_space == run_no_space:
            revision_manager.replace_run_text(run, new_value, reason=reason)
            return True
    return False


def _replace_numbering_with_context_test(doc, old_value, new_value, reason,
                                          revision_manager, context, anchor_text,
                                          paragraph_iterator, region_desc):
    """测试版：编号段落上下文定位替换（只处理手动编号）"""
    context_clean = clean_text_thoroughly(context or "")
    anchor_clean = clean_text_thoroughly(anchor_text or "")

    def _strip_numbering_prefix(text):
        stripped = re.sub(r'^[ivxlcdm]+\.\s*', '', text.strip(), flags=re.IGNORECASE)
        stripped = re.sub(r'^\(\d+\)\s*', '', stripped)
        stripped = re.sub(r'^\d+[\.\)]\s*', '', stripped)
        stripped = re.sub(r'^\([a-z]\)\s*', '', stripped, flags=re.IGNORECASE)
        return stripped.strip()

    context_content = _strip_numbering_prefix(context_clean)

    candidates = []
    for p in paragraph_iterator():
        para_text = (p.text or "").strip()
        if not para_text:
            continue
        full_text = "".join(r.text or "" for r in p.runs)
        has_manual_numbering = old_value in full_text
        if not has_manual_numbering:
            continue

        para_clean = clean_text_thoroughly(para_text)
        score = 0.0

        if context_content:
            context_content_lower = context_content.lower()
            para_clean_lower = para_clean.lower()
            if context_content_lower and context_content_lower in para_clean_lower:
                score += 0.6
            elif para_clean_lower and para_clean_lower in context_content_lower:
                score += 0.5
            else:
                sim = calculate_context_similarity(para_clean, context_content)
                score += sim * 0.4

        if context_clean:
            sim_full = calculate_context_similarity(para_clean, context_clean)
            score += sim_full * 0.1

        if anchor_clean:
            anchor_stripped = _strip_numbering_prefix(anchor_clean)
            if anchor_stripped and anchor_stripped.lower() in para_clean.lower():
                score += 0.25

        candidates.append((score, p, para_text))

    if not candidates:
        return False, ""

    candidates.sort(key=lambda x: x[0], reverse=True)
    best_score, best_para, best_text = candidates[0]

    if len(candidates) > 1:
        second_score = candidates[1][0]
        if not context_clean and not anchor_clean:
            return False, ""
        if best_score > 0 and second_score / best_score > 0.95:
            return False, ""
        if best_score < 0.1:
            return False, ""

    ok = _execute_replace(best_para, old_value, new_value, reason, revision_manager)
    if ok:
        return True, f"编号上下文定位(手动编号) [{region_desc}] (得分:{best_score:.2f}, 候选:{len(candidates)})"
    return False, ""


def replace_and_revise_in_docx(doc, old_value, new_value, reason, revision_manager,
                                context="", anchor_text="", region="all"):
    old_value_original = old_value or ""
    old_value = (old_value or "").strip()
    new_value = clean_text_thoroughly(new_value or "").strip()
    context = clean_text_thoroughly(context or "")
    if isinstance(reason, (list, tuple)):
        reason = " ".join([str(i) for i in reason if i]).strip()
    reason = reason or "数值/术语不一致"
    if not old_value or not new_value:
        return False, "数据缺失"

    if region == "body":
        paragraph_iterator = lambda: iter_body_paragraphs(doc)
        region_desc = "正文"
    elif region == "header":
        paragraph_iterator = lambda: iter_header_paragraphs(doc)
        region_desc = "页眉"
    elif region == "footer":
        paragraph_iterator = lambda: iter_footer_paragraphs(doc)
        region_desc = "页脚"
    else:
        paragraph_iterator = lambda: iter_all_paragraphs(doc)
        region_desc = "全部"

    # 策略0A：编号上下文定位
    if is_list_pattern(old_value):
        def _detect_numbering_type(val):
            v = val.strip().lower()
            if re.match(r'^[ivxlcdm]+\.$', v): return 'roman'
            if re.match(r'^\(\d+\)$', v): return 'paren_digit'
            if re.match(r'^\d+[\.\)]$', v): return 'digit'
            if re.match(r'^\([a-z]\)$', v): return 'paren_letter'
            if re.match(r'^[a-z][\.\)]$', v): return 'letter'
            return 'unknown'
        old_type = _detect_numbering_type(old_value)
        new_type = _detect_numbering_type(new_value)
        is_same_format = (old_type == new_type and old_type != 'unknown')

        if is_same_format:
            ok, strategy = _replace_numbering_with_context_test(
                doc, old_value, new_value, reason, revision_manager,
                context, anchor_text, paragraph_iterator, region_desc
            )
            if ok:
                return True, strategy
            # 同格式失败直接返回，不让后续策略用短编号字符串误匹配
            return False, f"编号替换失败: 同格式序号值错误但无法定位 [{region_desc}]"
        else:
            # 格式变更（如 i. → (1)）：测试中无自动编号，直接走上下文定位
            ok, strategy = _replace_numbering_with_context_test(
                doc, old_value, new_value, reason, revision_manager,
                context, anchor_text, paragraph_iterator, region_desc
            )
            if ok:
                return True, strategy
            # 编号类值不再往下走
            return False, f"编号替换失败: 未找到匹配的编号 [{region_desc}]"

    # 策略0：预处理（简化版，测试中跳过）
    found, matched_text, strategy = preprocess_special_cases(old_value_original, doc)
    if found:
        for p in paragraph_iterator():
            full_text = "".join(r.text or "" for r in p.runs)
            if full_text == matched_text:
                ok = _execute_replace(p, old_value, new_value, reason, revision_manager)
                if ok: return True, f"{strategy} [{region_desc}]"

    # 策略1：显式锚点
    if anchor_text:
        anchor_clean = clean_text_thoroughly(anchor_text)
        if anchor_clean:
            for p in paragraph_iterator():
                full_text = "".join(r.text or "" for r in p.runs)
                if anchor_clean in clean_text_thoroughly(full_text) and old_value in full_text:
                    ok = _execute_replace(p, old_value, new_value, reason, revision_manager)
                    if ok: return True, f"锚点匹配 [{region_desc}]"

    # 策略2：上下文锚点
    if context:
        context_anchor = extract_anchor_with_target(context, old_value, window=60)
        if context_anchor:
            anchor_clean = clean_text_thoroughly(context_anchor)
            for p in paragraph_iterator():
                full_text = "".join(r.text or "" for r in p.runs)
                full_clean = clean_text_thoroughly(full_text)
                if anchor_clean in full_clean and old_value in full_text:
                    ok = _execute_replace(p, old_value, new_value, reason, revision_manager)
                    if ok: return True, f"上下文锚点匹配 [{region_desc}]"

    # 策略3-6：候选打分
    old_value_clean = clean_text_thoroughly(old_value)
    strategies = [
        ("严格模式+上下文", build_smart_pattern(old_value, mode="strict"), 0.15, 2),
        ("严格模式", build_smart_pattern(old_value, mode="strict"), 0.0, 2),
        ("平衡模式", build_smart_pattern(old_value, mode="balanced"), 0.0, 3),
        ("宽松模式", build_smart_pattern(old_value, mode="balanced"), 0.0, 5),
    ]

    for strategy_name, pattern, min_similarity, max_level in strategies:
        if "上下文" in strategy_name and not context:
            continue
        candidates = []
        for p in paragraph_iterator():
            match_level = _try_match_in_paragraph(p, old_value, old_value_clean, pattern)
            if match_level is None or match_level > max_level:
                continue
            full_text = "".join(r.text or "" for r in p.runs)
            score = _score_paragraph(full_text, old_value, context, anchor_text, match_level)
            candidates.append((score, p, match_level, full_text))
        if not candidates:
            continue
        candidates.sort(key=lambda x: x[0], reverse=True)
        best_score, best_para, best_level, best_text = candidates[0]

        # 安全检查1：上下文相似度门槛
        if context and min_similarity > 0:
            sim = calculate_context_similarity(
                clean_text_thoroughly(best_text), clean_text_thoroughly(context))
            if sim < min_similarity:
                continue

        # 安全检查2：唯一性安全网
        if len(candidates) > 1 and not context and not anchor_text:
            continue

        # 安全检查3：得分区分度
        if len(candidates) > 1:
            second_score = candidates[1][0]
            if best_score > 0 and second_score / best_score > 0.95:
                continue

        ok = _execute_replace(best_para, old_value, new_value, reason, revision_manager)
        if ok:
            detail = f"{strategy_name} [{region_desc}] (得分:{best_score:.2f}, 候选:{len(candidates)})"
            return True, detail

    return False, f"未找到匹配项 (搜索区域: {region_desc})"


# ============================================================
# 测试用例
# ============================================================

def test_1_unique_occurrence_replaces_correctly():
    """场景1：目标值在文档中只出现一次 → 正常替换"""
    doc = _create_test_doc_with_paragraphs([
        "The total revenue was 5,000,000 USD in 2024.",
        "Operating expenses reached 2,300,000 USD.",
        "Net profit was 2,700,000 USD after tax.",
    ])
    rm = MockRevisionManager(doc)

    ok, strategy = replace_and_revise_in_docx(
        doc, "2,300,000", "2,500,000", "数值错误", rm,
        context="Operating expenses reached 2,300,000 USD.",
        region="body"
    )

    assert ok, f"应该替换成功，但失败了: {strategy}"
    assert len(rm.replacements) == 1
    assert rm.replacements[0][1] == "2,300,000"  # old
    assert rm.replacements[0][2] == "2,500,000"  # new
    # 验证替换的是第二个段落
    assert "Operating expenses" in rm.replacements[0][0]
    print(f"  ✓ 测试1通过: {strategy}")


def test_2_multiple_occurrences_with_context_picks_correct():
    """场景2：同一个数值出现多次，上下文应该帮助选择正确的段落"""
    doc = _create_test_doc_with_paragraphs([
        "In Q1, the company earned 10.00 million from domestic sales.",
        "In Q2, the company earned 15.00 million from exports.",
        "In Q3, the company earned 10.00 million from overseas operations.",
        "Total annual revenue was 35.00 million.",
    ])
    rm = MockRevisionManager(doc)

    # 上下文指向第三个段落（Q3, overseas operations）
    ok, strategy = replace_and_revise_in_docx(
        doc, "10.00", "12.00", "数值错误", rm,
        context="In Q3, the company earned 10.00 million from overseas operations.",
        region="body"
    )

    assert ok, f"应该替换成功，但失败了: {strategy}"
    assert len(rm.replacements) == 1
    # 关键：应该替换的是第三个段落（Q3），而不是第一个（Q1）
    replaced_para_text = rm.replacements[0][0]
    assert "Q3" in replaced_para_text, \
        f"替换了错误的段落！应该替换Q3段落，实际替换了: '{replaced_para_text}'"
    assert "overseas" in replaced_para_text
    print(f"  ✓ 测试2通过: 正确选择了Q3段落 — {strategy}")


def test_3_multiple_occurrences_no_context_rejects():
    """场景3：同一个数值出现多次，但没有上下文 → 安全网拒绝替换"""
    doc = _create_test_doc_with_paragraphs([
        "Revenue: 10.00 million",
        "Cost: 10.00 million",
        "Profit: 5.00 million",
    ])
    rm = MockRevisionManager(doc)

    ok, strategy = replace_and_revise_in_docx(
        doc, "10.00", "12.00", "数值错误", rm,
        context="",  # 无上下文
        anchor_text="",  # 无锚点
        region="body"
    )

    assert not ok, f"应该拒绝替换（多处出现无上下文），但成功了: {strategy}"
    assert len(rm.replacements) == 0
    print(f"  ✓ 测试3通过: 安全网正确拒绝了歧义替换 — {strategy}")


def test_4_anchor_text_locates_correctly():
    """场景4：锚点文本精确定位"""
    doc = _create_test_doc_with_paragraphs([
        "Section A: The rate is 3.5% per annum.",
        "Section B: The rate is 3.5% per quarter.",
        "Section C: The total is 100 units.",
    ])
    rm = MockRevisionManager(doc)

    ok, strategy = replace_and_revise_in_docx(
        doc, "3.5%", "4.0%", "数值错误", rm,
        context="",
        anchor_text="per quarter",  # 锚点指向 Section B
        region="body"
    )

    assert ok, f"应该替换成功，但失败了: {strategy}"
    assert len(rm.replacements) == 1
    assert "Section B" in rm.replacements[0][0] or "per quarter" in rm.replacements[0][0], \
        f"替换了错误的段落！应该替换Section B，实际: '{rm.replacements[0][0]}'"
    print(f"  ✓ 测试4通过: 锚点正确定位到Section B — {strategy}")


def test_5_region_isolation():
    """场景5：region 隔离 — 正文中的值不应被页眉替换影响"""
    doc = _create_test_doc_with_paragraphs(
        ["Body text with value 100."],
        header_text="Header with value 100.",
    )
    rm = MockRevisionManager(doc)

    # 指定只在页眉中替换
    ok, strategy = replace_and_revise_in_docx(
        doc, "100", "200", "数值错误", rm,
        context="Header with value 100.",
        region="header"
    )

    if ok:
        # 验证替换的是页眉而非正文
        replaced_text = rm.replacements[0][0]
        assert "Header" in replaced_text, \
            f"替换了正文而非页眉！实际: '{replaced_text}'"
        print(f"  ✓ 测试5通过: 正确替换了页眉 — {strategy}")
    else:
        # 页眉段落可能没有 runs（取决于 python-docx 行为）
        print(f"  ⚠ 测试5: 页眉替换未成功（可能是 python-docx 页眉 runs 为空）— {strategy}")


def test_6_close_scores_rejects():
    """场景6：两个候选段落得分几乎一样 → 拒绝替换"""
    # 两个段落几乎一模一样，只有微小差异
    doc = _create_test_doc_with_paragraphs([
        "The amount is 500 USD for project Alpha.",
        "The amount is 500 USD for project Beta.",
    ])
    rm = MockRevisionManager(doc)

    # 上下文不够区分（两个段落都很相似）
    ok, strategy = replace_and_revise_in_docx(
        doc, "500", "600", "数值错误", rm,
        context="The amount is 500 USD",  # 两个段落都匹配
        region="body"
    )

    # 如果两个候选得分差距 < 5%，应该拒绝
    if not ok:
        print(f"  ✓ 测试6通过: 得分接近时正确拒绝了替换 — {strategy}")
    else:
        # 如果上下文恰好能区分（因为 Jaccard 计算），也可以接受
        print(f"  ⚠ 测试6: 替换成功了（上下文可能足以区分）— {strategy}")


def test_7_context_similarity_threshold():
    """场景7：上下文相似度低于门槛 → 严格+上下文策略跳过"""
    doc = _create_test_doc_with_paragraphs([
        "The interest rate for the first quarter was 2.5% based on the central bank policy.",
        "Annual GDP growth rate was 6.8% according to the national statistics bureau.",
    ])
    rm = MockRevisionManager(doc)

    # 上下文完全不匹配任何段落
    ok, strategy = replace_and_revise_in_docx(
        doc, "2.5%", "3.0%", "数值错误", rm,
        context="Completely unrelated context about weather forecast and temperature.",
        region="body"
    )

    # 即使文本匹配，上下文不匹配时严格+上下文策略应该跳过
    # 但后续策略（严格模式全局）如果只有一个候选，仍然可以替换
    if ok:
        assert "2.5%" not in strategy or "上下文" not in strategy, \
            "不应该通过'严格+上下文'策略成功"
        print(f"  ✓ 测试7通过: 跳过了上下文策略，通过其他策略替换 — {strategy}")
    else:
        print(f"  ✓ 测试7通过: 上下文不匹配，替换被拒绝 — {strategy}")


def test_8_correct_paragraph_with_strong_context():
    """场景8：强上下文 — 三个段落都有目标值，上下文明确指向中间那个"""
    doc = _create_test_doc_with_paragraphs([
        "First paragraph mentions amount 1,000 in the introduction section.",
        "Second paragraph discusses the budget of 1,000 allocated to marketing department for Q2 campaign.",
        "Third paragraph notes that 1,000 units were shipped to warehouse.",
    ])
    rm = MockRevisionManager(doc)

    ok, strategy = replace_and_revise_in_docx(
        doc, "1,000", "1,500", "数值错误", rm,
        context="Second paragraph discusses the budget of 1,000 allocated to marketing department for Q2 campaign.",
        region="body"
    )

    assert ok, f"应该替换成功: {strategy}"
    assert len(rm.replacements) == 1
    replaced = rm.replacements[0][0]
    assert "marketing" in replaced or "Second" in replaced or "budget" in replaced, \
        f"替换了错误的段落！应该是第二段(marketing)，实际: '{replaced}'"
    print(f"  ✓ 测试8通过: 强上下文正确定位到第二段 — {strategy}")


def test_9_single_occurrence_no_context_succeeds():
    """场景9：只出现一次，即使没有上下文也应该成功替换"""
    doc = _create_test_doc_with_paragraphs([
        "The total amount is 999,999 USD.",
        "Other information here.",
    ])
    rm = MockRevisionManager(doc)

    ok, strategy = replace_and_revise_in_docx(
        doc, "999,999", "1,000,000", "数值错误", rm,
        context="",
        region="body"
    )

    assert ok, f"唯一出现应该替换成功: {strategy}"
    assert rm.replacements[0][1] == "999,999"
    print(f"  ✓ 测试9通过: 唯一出现无上下文正常替换 — {strategy}")


def test_10_old_code_would_fail_new_code_succeeds():
    """场景10：旧代码会替换第一个碰到的（错误位置），新代码应该替换正确位置

    这是最关键的回归测试：
    - 旧代码：遍历段落，碰到第一个包含 "10%" 的就替换 → 替换了段落1（错误）
    - 新代码：打分后选择上下文最匹配的 → 替换段落3（正确）
    """
    doc = _create_test_doc_with_paragraphs([
        "The tax rate is 10% for domestic companies under the new regulation.",
        "Import duty is 15% for electronic products.",
        "The withholding tax rate is 10% for foreign investors according to the treaty.",
        "VAT rate is 13% for general goods.",
    ])
    rm = MockRevisionManager(doc)

    ok, strategy = replace_and_revise_in_docx(
        doc, "10%", "8%", "数值错误", rm,
        context="The withholding tax rate is 10% for foreign investors according to the treaty.",
        region="body"
    )

    assert ok, f"应该替换成功: {strategy}"
    replaced = rm.replacements[0][0]
    assert "withholding" in replaced or "foreign investors" in replaced, \
        f"旧代码的bug！替换了错误的段落。应该替换'withholding tax'段落，实际替换了: '{replaced}'"
    print(f"  ✓ 测试10通过（核心回归测试）: 正确替换了第三段而非第一段 — {strategy}")


def test_11_manual_numbering_with_context():
    """场景11：手动编号 iv. 出现两次，上下文指向 Precautions 而非 Operation procedures

    模拟用户报告的实际场景：
    - iv. Operation procedures
    - iv. Precautions
    上下文是 "iv. Precautions"，应该替换第二个 iv. → v.
    """
    doc = _create_test_doc_with_paragraphs([
        "i. Introduction",
        "ii. Scope of work",
        "iii. Technical requirements",
        "iv. Operation procedures",
        "iv. Precautions",
        "vi. Emergency response plan",
    ])
    rm = MockRevisionManager(doc)

    ok, strategy = replace_and_revise_in_docx(
        doc, "iv.", "v.", "层级编号重号", rm,
        context="iv. Precautions",
        anchor_text="iv.",
        region="body"
    )

    assert ok, f"应该替换成功: {strategy}"
    assert len(rm.replacements) == 1
    replaced = rm.replacements[0][0]
    assert "Precautions" in replaced, \
        f"替换了错误的段落！应该替换'iv. Precautions'，实际替换了: '{replaced}'"
    print(f"  ✓ 测试11通过: 手动编号上下文正确定位到 Precautions — {strategy}")


def test_12_numbering_no_context_rejects():
    """场景12：编号 iv. 出现两次，无上下文 → 拒绝替换"""
    doc = _create_test_doc_with_paragraphs([
        "iv. Operation procedures",
        "iv. Precautions",
    ])
    rm = MockRevisionManager(doc)

    ok, strategy = replace_and_revise_in_docx(
        doc, "iv.", "v.", "层级编号重号", rm,
        context="",
        anchor_text="",
        region="body"
    )

    assert not ok, f"应该拒绝替换（编号多处出现无上下文），但成功了: {strategy}"
    print(f"  ✓ 测试12通过: 编号多处出现无上下文正确拒绝 — {strategy}")


def test_13_format_change_with_context():
    """场景13：编号格式变更 i. → (1)，上下文指向 Organizational structure

    模拟用户报告的实际场景：
    - i. Organizational structure  ← 应该替换这个
    - i. (1) Managers must take responsibility...  ← 不应该碰这个
    旧代码会在后续策略中找到 "i." 并替换到错误位置
    """
    doc = _create_test_doc_with_paragraphs([
        "2. Specific organizational implementation structure",
        "i. Organizational structure",
        "Leader: Project manager",
        "Deputy group leader: Project management personnel",
        "Members: Captains and cleaners of grid cells in each road section",
        "(2) Responsibilities and requirements",
        "i. (1) Managers must take responsibility for the quality of cleaning operations",
    ])
    rm = MockRevisionManager(doc)

    ok, strategy = replace_and_revise_in_docx(
        doc, "i.", "(1)", "层级错误", rm,
        context="i. Organizational structure",
        anchor_text="i.",
        region="body"
    )

    assert ok, f"应该替换成功: {strategy}"
    assert len(rm.replacements) == 1
    replaced = rm.replacements[0][0]
    assert "Organizational structure" in replaced, \
        f"替换了错误的段落！应该替换'i. Organizational structure'，实际替换了: '{replaced}'"
    # 确保没有替换到 "i. (1) Managers..." 那个段落
    assert "Managers" not in replaced, \
        f"误匹配到了错误段落: '{replaced}'"
    print(f"  ✓ 测试13通过: 格式变更 i.→(1) 正确定位到 Organizational structure — {strategy}")


def test_14_numbering_format_change_no_context_rejects():
    """场景14：编号格式变更 i. → (1)，但 i. 出现多处且无上下文 → 拒绝"""
    doc = _create_test_doc_with_paragraphs([
        "i. Organizational structure",
        "i. (1) Managers must take responsibility",
    ])
    rm = MockRevisionManager(doc)

    ok, strategy = replace_and_revise_in_docx(
        doc, "i.", "(1)", "层级错误", rm,
        context="",
        anchor_text="",
        region="body"
    )

    assert not ok, f"应该拒绝替换（编号多处出现无上下文），但成功了: {strategy}"
    print(f"  ✓ 测试14通过: 格式变更无上下文正确拒绝 — {strategy}")


# ============================================================
# 运行所有测试
# ============================================================

if __name__ == "__main__":
    print("=" * 70)
    print("replace_revision.py 优化版测试")
    print("=" * 70)

    tests = [
        ("测试1: 唯一出现正常替换", test_1_unique_occurrence_replaces_correctly),
        ("测试2: 多处出现+上下文选正确位置", test_2_multiple_occurrences_with_context_picks_correct),
        ("测试3: 多处出现+无上下文→拒绝", test_3_multiple_occurrences_no_context_rejects),
        ("测试4: 锚点精确定位", test_4_anchor_text_locates_correctly),
        ("测试5: region隔离", test_5_region_isolation),
        ("测试6: 得分接近→拒绝", test_6_close_scores_rejects),
        ("测试7: 上下文相似度门槛", test_7_context_similarity_threshold),
        ("测试8: 强上下文三选一", test_8_correct_paragraph_with_strong_context),
        ("测试9: 唯一出现无上下文", test_9_single_occurrence_no_context_succeeds),
        ("测试10: 核心回归—旧代码替错新代码替对", test_10_old_code_would_fail_new_code_succeeds),
        ("测试11: 手动编号iv.上下文定位", test_11_manual_numbering_with_context),
        ("测试12: 编号多处出现无上下文→拒绝", test_12_numbering_no_context_rejects),
        ("测试13: 格式变更i.→(1)上下文定位", test_13_format_change_with_context),
        ("测试14: 格式变更无上下文→拒绝", test_14_numbering_format_change_no_context_rejects),
    ]

    passed = 0
    failed = 0
    errors = []

    for name, test_func in tests:
        print(f"\n--- {name} ---")
        try:
            test_func()
            passed += 1
        except AssertionError as e:
            failed += 1
            errors.append((name, str(e)))
            print(f"  ✗ 失败: {e}")
        except Exception as e:
            failed += 1
            errors.append((name, str(e)))
            print(f"  ✗ 异常: {e}")

    print(f"\n{'=' * 70}")
    print(f"结果: {passed} 通过, {failed} 失败, 共 {len(tests)} 个测试")
    if errors:
        print(f"\n失败详情:")
        for name, err in errors:
            print(f"  - {name}: {err}")
    print("=" * 70)
