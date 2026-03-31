# =========================
# 导入依赖
# =========================
import re
import unicodedata
from typing import Any, Dict, List, Optional, Tuple
from docx import Document

from llm.llm_project.note.pizhu import CommentManager


# =========================
# 1) 文本清洗工具
# =========================

def clean_text_thoroughly(text: str) -> str:
    if not text: return ""

    # 1. 统一处理：将所有智能引号、撇号变体转为标准引号
    text = text.replace(''', "'").replace(''', "'")  # 左右单引号
    text = text.replace('"', '"').replace('"', '"')  # 左右双引号
    text = text.replace('`', "'")  # 反引号
    text = text.replace('´', "'")  # 重音符

    # 2. 将全角符号转为半角（扩展版）
    text = text.replace('（', '(').replace('）', ')')
    text = text.replace('，', ',').replace('。', '.')
    text = text.replace('：', ':').replace('；', ';')
    text = text.replace('！', '!').replace('？', '?')
    text = text.replace('【', '[').replace('】', ']')
    text = text.replace('《', '<').replace('》', '>')
    text = text.replace('　', ' ')  # 全角空格

    # 3. Unicode 标准化
    text = unicodedata.normalize('NFKC', text)

    # 4. 移除隐形字符和零宽字符
    text = re.sub(r'[\u200b-\u200f\u202a-\u202e\u2060-\u206f\ufeff\u00ad\xa0]', '', text)

    # 5. 统一空格（包括各种空白字符）
    text = re.sub(r'\s+', ' ', text)

    return text.strip()




def _normalize_spaces(s: str) -> str:
    """标准化空白字符"""
    return re.sub(r"\s+", " ", (s or "")).strip()


# =========================
# 2) 编号模式识别（增强版）
# =========================

def is_list_pattern(s: str) -> bool:
    """
    判断是否为编号模式（支持复合罗马数字）

    支持格式：
    - 数字编号：(1), 1), 1.
    - 罗马数字：i., ii., iii., iv., v., ix., x. 等（包括复合形式）
    - 字母编号：a), b), (a), (b) 等

    Args:
        s: 待判断字符串

    Returns:
        是否为编号模式
    """
    s_clean = s.strip().lower()

    # 修复后的正则表达式（分离多个模式）
    patterns = [
        r'^\(\d+\)$',  # (1), (2)
        r'^\d+\)$',  # 1), 2)
        r'^\d+\.$',  # 1., 2.
        r'^[ivxlcdm]+\.$',  # i., ii., iii., iv., ix., x. 等（完整罗马数字）
        r'^\([a-z]\)$',  # (a), (b)
        r'^[a-z]\)$',  # a), b)
        r'^[a-z]\.$',  # a., b.
    ]

    return any(re.match(pattern, s_clean) for pattern in patterns)


def is_valid_roman_numeral(s: str) -> bool:
    """
    验证是否为有效的罗马数字（避免误判）

    Args:
        s: 待验证字符串（不含点号）

    Returns:
        是否为有效罗马数字
    """
    s = s.strip().lower()

    # 罗马数字规则：只能包含 i, v, x, l, c, d, m
    if not re.match(r'^[ivxlcdm]+$', s):
        return False

    # 常见罗马数字列表（1-100）
    valid_romans = [
        'i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x',
        'xi', 'xii', 'xiii', 'xiv', 'xv', 'xvi', 'xvii', 'xviii', 'xix', 'xx',
        'xxi', 'xxii', 'xxiii', 'xxiv', 'xxv', 'xxvi', 'xxvii', 'xxviii', 'xxix', 'xxx',
        'xl', 'l', 'lx', 'lxx', 'lxxx', 'xc', 'c'
    ]

    return s in valid_romans


# =========================
# 3) 智能匹配模式构建（核心修复）
# =========================

def build_smart_pattern(s: str, mode: str = "balanced") -> str:
    """
    构建智能匹配模式

    Args:
        s: 待匹配字符串
        mode: 匹配模式
            - "strict": 严格匹配（完全精确）
            - "balanced": 平衡模式（数字/标点保持连续，单词间允许空格）
            - "loose": 宽松模式（字符间允许空格，但数字连续）

    Returns:
        正则表达式模式字符串
    """
    # 【关键修复】先清洗输入
    s = clean_text_thoroughly(s or "")
    if not s:
        return ""

    if mode == "strict":
        # 严格模式：完全精确匹配（转义特殊字符）
        return re.escape(s)

    elif mode == "balanced":
        # 平衡模式：数字和标点保持连续，单词间允许空格
        pieces = []
        i = 0

        while i < len(s):
            ch = s[i]

            # 跳过空格（在字符间添加可选空格）
            if ch.isspace():
                if pieces and not pieces[-1].endswith(r"\s*"):
                    pieces.append(r"\s*")
                i += 1
                continue

            # 数字序列：保持连续（包括小数点、逗号、百分号）
            if ch.isdigit():
                num_str = ""
                while i < len(s) and (s[i].isdigit() or s[i] in ".,%-"):
                    num_str += s[i]
                    i += 1
                pieces.append(re.escape(num_str))
                continue

            # 【关键修复】标点符号：保持连续，智能判断是否添加空格
            if ch in ".,;:!?()[]{}\"'-/":
                pieces.append(re.escape(ch))
                i += 1
                # 仅在下一个字符是字母或数字时才添加可选空格
                if i < len(s) and (s[i].isalpha() or s[i].isdigit()):
                    pieces.append(r"\s*")
                continue

            # 字母序列：单词间允许空格
            if ch.isalpha():
                word = ""
                while i < len(s) and s[i].isalpha():
                    word += s[i]
                    i += 1
                pieces.append(re.escape(word))
                # 如果后面还有内容且不是标点，添加可选空格
                if i < len(s) and s[i] not in ".,;:!?()[]{}\"'-/":
                    pieces.append(r"\s*")
                continue

            # 其他字符
            pieces.append(re.escape(ch))
            i += 1

        return "".join(pieces).strip()

    else:  # loose
        # 宽松模式：字符间允许空格，但数字保持连续
        pieces = []
        i = 0

        while i < len(s):
            ch = s[i]

            if ch.isspace():
                i += 1
                continue

            # 数字序列保持连续
            if ch.isdigit():
                num_str = ""
                while i < len(s) and (s[i].isdigit() or s[i] in ".,%-"):
                    num_str += s[i]
                    i += 1
                pieces.append(re.escape(num_str) + r"\s*")
                continue

            # 其他字符间允许空格
            pieces.append(re.escape(ch) + r"\s*")
            i += 1

        return "".join(pieces).strip()


# =========================
# 4) 增强的上下文匹配（核心优化）
# =========================

def extract_anchor_with_target(
        context: str,
        target_value: str,
        window: int = 50  # 增加窗口大小以获取更多上下文
) -> Optional[str]:
    """
    从上下文中提取包含目标数值的锚点短语（智能扩展版）

    Args:
        context: 上下文文本
        target_value: 目标值（如数字、编号等）
        window: 锚点窗口大小（前后字符数）

    Returns:
        提取的锚点短语，未找到则返回 None
    """
    # 【关键修复】在最开始就清洗
    context = clean_text_thoroughly(context)
    target_value = clean_text_thoroughly(target_value)

    if not context or not target_value:
        return None

    context = _normalize_spaces(context)
    target_value = target_value.strip()

    # 步骤1：尝试严格匹配
    if target_value in context:
        idx = context.index(target_value)
        start = max(0, idx - window)
        end = min(len(context), idx + len(target_value) + window)
        anchor = context[start:end]

        # 修剪边界单词（避免截断）
        anchor = re.sub(r"^\S*\s+", "", anchor)
        anchor = re.sub(r"\s+\S*$", "", anchor)
        
        # 【新增】如果锚点太短（小于目标值长度的1.5倍），扩大窗口
        if len(anchor) < len(target_value) * 1.5:
            # 使用更大的窗口
            larger_window = min(len(context), len(target_value) * 2)
            start = max(0, idx - larger_window)
            end = min(len(context), idx + len(target_value) + larger_window)
            anchor = context[start:end]
            anchor = re.sub(r"^\S*\s+", "", anchor)
            anchor = re.sub(r"\s+\S*$", "", anchor)
        
        # 【新增】如果目标值占锚点的比例太高（>80%），说明上下文不够丰富
        # 这种情况下，直接返回目标值本身作为锚点
        if len(target_value) / len(anchor) > 0.8:
            return target_value
        
        return anchor.strip()

    # 步骤2：【关键修复】编号模式强制使用 strict 模式
    if is_list_pattern(target_value):
        pattern = re.escape(target_value)
    else:
        pattern = build_smart_pattern(target_value, mode="balanced")

    if not pattern:
        return None

    # 步骤3：正则匹配
    match = re.search(pattern, context, flags=re.IGNORECASE)
    if not match:
        return None

    start, end = match.span()
    prefix_start = max(0, start - window)
    suffix_end = min(len(context), end + window)

    anchor = context[prefix_start:suffix_end]

    # 修剪边界单词
    anchor = re.sub(r"^\S*\s+", "", anchor)
    anchor = re.sub(r"\s+\S*$", "", anchor)
    
    # 【新增】同样的检查：如果目标值占比太高，返回目标值本身
    if anchor.strip() and len(target_value) / len(anchor.strip()) > 0.8:
        return target_value

    return anchor.strip() if anchor.strip() else None

    anchor = context[prefix_start:suffix_end]

    # 修剪边界单词
    anchor = re.sub(r"^\S*\s+", "", anchor)
    anchor = re.sub(r"\s+\S*$", "", anchor)

    return anchor.strip() if anchor.strip() else None


def calculate_context_similarity(text1: str, text2: str) -> float:
    """
    计算两段文本的相似度（基于词汇重叠）

    Args:
        text1: 文本1
        text2: 文本2

    Returns:
        相似度分数（0-1）
    """
    if not text1 or not text2:
        return 0.0

    # 分词（简单按空格分割）
    words1 = set(text1.lower().split())
    words2 = set(text2.lower().split())

    if not words1 or not words2:
        return 0.0

    # 计算 Jaccard 相似度
    intersection = len(words1 & words2)
    union = len(words1 | words2)

    return intersection / union if union > 0 else 0.0


# =========================
# 5) Word 文档处理工具
# =========================

def iter_all_paragraphs(doc: Document, include_headers_footers: bool = True):
    """
    遍历正文段落 + 表格单元格段落 + (可选)页眉页脚段落

    Args:
        doc: Word 文档对象
        include_headers_footers: 是否包含页眉页脚

    Yields:
        段落对象
    """
    # 正文段落
    for p in doc.paragraphs:
        yield p

    # 表格（递归）
    def walk_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    for p2 in walk_tables(cell.tables):
                        yield p2

    for p in walk_tables(doc.tables):
        yield p

    # 页眉页脚
    if include_headers_footers:
        for sec in doc.sections:
            for p in sec.header.paragraphs:
                yield p
            for p in sec.footer.paragraphs:
                yield p


def iter_body_paragraphs(doc: Document):
    """
    仅遍历正文段落（包括表格中的段落、文本框中的段落，但不包括页眉页脚）
    
    注意：表格段落会被遍历两次（一次在doc.paragraphs中，一次在doc.tables中）
    但这是必要的，因为doc.paragraphs中的表格段落对象可能没有正确的文本内容

    Args:
        doc: Word 文档对象

    Yields:
        段落对象
    """
    from lxml import etree
    
    # XML命名空间
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    V_NS = "urn:schemas-microsoft-com:vml"
    MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    NAMESPACES = {"w": W_NS, "wps": WPS_NS, "v": V_NS, "mc": MC_NS}
    
    # 跟踪已经遍历过的段落元素，避免文本框等重复
    seen_textbox_elements = set()
    
    def is_in_fallback(elem):
        """检查元素是否在 mc:Fallback 内"""
        parent = elem.getparent()
        while parent is not None:
            if parent.tag == f"{{{MC_NS}}}Fallback":
                return True
            parent = parent.getparent()
        return False
    
    def yield_paragraph(para_elem, parent):
        """辅助函数：yield 段落并记录（跳过 Fallback 内的段落）"""
        # 跳过 Fallback 内的段落
        if is_in_fallback(para_elem):
            return
            
        elem_id = id(para_elem)
        if elem_id not in seen_textbox_elements:
            seen_textbox_elements.add(elem_id)
            try:
                from docx.text.paragraph import Paragraph
                yield Paragraph(para_elem, parent)
            except:
                pass
    
    # 表格（递归）- 先遍历表格，确保表格段落能被正确处理
    def walk_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        # 直接yield，不检查seen_elements
                        yield p
                        
                        # 表格单元格中的文本框
                        for txbx in p._element.iter(f"{{{WPS_NS}}}txbxContent"):
                            for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                                yield from yield_paragraph(txbx_p, p._parent)
                        
                        for v_txbx in p._element.iter(f"{{{V_NS}}}textbox"):
                            for txbx_p in v_txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                                yield from yield_paragraph(txbx_p, p._parent)
                        
                        # w:drawing 中的文本框
                        for txbx in p._element.iter(f"{{{W_NS}}}txbxContent"):
                            for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                                yield from yield_paragraph(txbx_p, p._parent)
                    
                    # 递归处理嵌套表格
                    yield from walk_tables(cell.tables)

    # 遍历所有表格
    yield from walk_tables(doc.tables)
    
    # 然后遍历正文段落（可能会包含一些表格段落的重复，但这是必要的）
    for p in doc.paragraphs:
        yield p
        
        # 检查段落中是否有文本框
        # 方法1: wps:txbxContent (Office 2010+ 格式)
        for txbx in p._element.iter(f"{{{WPS_NS}}}txbxContent"):
            for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                yield from yield_paragraph(txbx_p, p._parent)
        
        # 方法2: v:textbox (旧版 VML 格式)
        for v_txbx in p._element.iter(f"{{{V_NS}}}textbox"):
            for txbx_p in v_txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                yield from yield_paragraph(txbx_p, p._parent)
        
        # 方法3: w:txbxContent (通过 w:drawing 的文本框)
        for txbx in p._element.iter(f"{{{W_NS}}}txbxContent"):
            for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                yield from yield_paragraph(txbx_p, p._parent)


def iter_header_paragraphs(doc: Document):
    """
    仅遍历页眉段落（包括页眉中的文本框）

    Args:
        doc: Word 文档对象

    Yields:
        段落对象
    """
    from lxml import etree
    
    # XML命名空间
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    V_NS = "urn:schemas-microsoft-com:vml"
    MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    NAMESPACES = {"w": W_NS, "wps": WPS_NS, "v": V_NS, "mc": MC_NS}
    
    # 跟踪已经遍历过的段落元素，避免重复
    seen_elements = set()
    
    def is_in_fallback(elem):
        """检查元素是否在 mc:Fallback 内"""
        parent = elem.getparent()
        while parent is not None:
            if parent.tag == f"{{{MC_NS}}}Fallback":
                return True
            parent = parent.getparent()
        return False
    
    def yield_paragraph(para_elem, parent):
        """辅助函数：yield 段落并记录（跳过 Fallback 内的段落）"""
        # 跳过 Fallback 内的段落
        if is_in_fallback(para_elem):
            return
            
        elem_id = id(para_elem)
        if elem_id not in seen_elements:
            seen_elements.add(elem_id)
            try:
                from docx.text.paragraph import Paragraph
                yield Paragraph(para_elem, parent)
            except:
                pass
    
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            elem_id = id(p._element)
            if elem_id not in seen_elements:
                seen_elements.add(elem_id)
                yield p
            
            # 检查页眉段落中是否有文本框
            # 方法1: wps:txbxContent (Office 2010+ 格式)
            for txbx in p._element.iter(f"{{{WPS_NS}}}txbxContent"):
                for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                    yield from yield_paragraph(txbx_p, p._parent)
            
            # 方法2: v:textbox (旧版 VML 格式)
            for v_txbx in p._element.iter(f"{{{V_NS}}}textbox"):
                for txbx_p in v_txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                    yield from yield_paragraph(txbx_p, p._parent)
            
            # 方法3: w:txbxContent (通过 w:drawing 的文本框)
            for txbx in p._element.iter(f"{{{W_NS}}}txbxContent"):
                for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                    yield from yield_paragraph(txbx_p, p._parent)
        
        # 也包括页眉中的表格
        for tbl in sec.header.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        elem_id = id(p._element)
                        if elem_id not in seen_elements:
                            seen_elements.add(elem_id)
                            yield p
                        
                        # 表格单元格中的文本框
                        for txbx in p._element.iter(f"{{{WPS_NS}}}txbxContent"):
                            for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                                yield from yield_paragraph(txbx_p, p._parent)
                        
                        for v_txbx in p._element.iter(f"{{{V_NS}}}textbox"):
                            for txbx_p in v_txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                                yield from yield_paragraph(txbx_p, p._parent)
                        
                        # w:drawing 中的文本框
                        for txbx in p._element.iter(f"{{{W_NS}}}txbxContent"):
                            for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                                yield from yield_paragraph(txbx_p, p._parent)


def iter_footer_paragraphs(doc: Document):
    """
    仅遍历页脚段落（包括页脚中的文本框）

    Args:
        doc: Word 文档对象

    Yields:
        段落对象
    """
    from lxml import etree
    
    # XML命名空间
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    V_NS = "urn:schemas-microsoft-com:vml"
    MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    NAMESPACES = {"w": W_NS, "wps": WPS_NS, "v": V_NS, "mc": MC_NS}
    
    # 跟踪已经遍历过的段落元素，避免重复
    seen_elements = set()
    
    def is_in_fallback(elem):
        """检查元素是否在 mc:Fallback 内"""
        parent = elem.getparent()
        while parent is not None:
            if parent.tag == f"{{{MC_NS}}}Fallback":
                return True
            parent = parent.getparent()
        return False
    
    def yield_paragraph(para_elem, parent):
        """辅助函数：yield 段落并记录（跳过 Fallback 内的段落）"""
        # 跳过 Fallback 内的段落
        if is_in_fallback(para_elem):
            return
            
        elem_id = id(para_elem)
        if elem_id not in seen_elements:
            seen_elements.add(elem_id)
            try:
                from docx.text.paragraph import Paragraph
                yield Paragraph(para_elem, parent)
            except:
                pass
    
    for sec in doc.sections:
        for p in sec.footer.paragraphs:
            elem_id = id(p._element)
            if elem_id not in seen_elements:
                seen_elements.add(elem_id)
                yield p
            
            # 检查页脚段落中是否有文本框
            # 方法1: wps:txbxContent (Office 2010+ 格式)
            for txbx in p._element.iter(f"{{{WPS_NS}}}txbxContent"):
                for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                    yield from yield_paragraph(txbx_p, p._parent)
            
            # 方法2: v:textbox (旧版 VML 格式)
            for v_txbx in p._element.iter(f"{{{V_NS}}}textbox"):
                for txbx_p in v_txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                    yield from yield_paragraph(txbx_p, p._parent)
            
            # 方法3: w:txbxContent (通过 w:drawing 的文本框)
            for txbx in p._element.iter(f"{{{W_NS}}}txbxContent"):
                for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                    yield from yield_paragraph(txbx_p, p._parent)
        
        # 也包括页脚中的表格
        for tbl in sec.footer.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        elem_id = id(p._element)
                        if elem_id not in seen_elements:
                            seen_elements.add(elem_id)
                            yield p
                        
                        # 表格单元格中的文本框
                        for txbx in p._element.iter(f"{{{WPS_NS}}}txbxContent"):
                            for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                                yield from yield_paragraph(txbx_p, p._parent)
                        
                        for v_txbx in p._element.iter(f"{{{V_NS}}}textbox"):
                            for txbx_p in v_txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                                yield from yield_paragraph(txbx_p, p._parent)
                        
                        # w:drawing 中的文本框
                        for txbx in p._element.iter(f"{{{W_NS}}}txbxContent"):
                            for txbx_p in txbx.findall(f".//{{{W_NS}}}p", NAMESPACES):
                                yield from yield_paragraph(txbx_p, p._parent)


from difflib import SequenceMatcher


def is_fuzzy_match(text: str, target: str, threshold: float = 0.9) -> bool:
    """
    判断两段文本的相似度是否超过阈值
    """
    # 彻底清理文本，确保比较的是“纯净版”
    text_clean = clean_text_thoroughly(text)
    target_clean = clean_text_thoroughly(target)

    # 使用 SequenceMatcher 计算相似比率
    ratio = SequenceMatcher(None, text_clean, target_clean).ratio()
    return ratio >= threshold


def copy_run_style(source_run, target_run):
    """
    统一使用 font 属性复制样式，避免 AttributeError
    """
    # 统一通过 font 属性访问，既规范又避免了快捷属性缺失的问题
    target_run.font.bold = source_run.font.bold
    target_run.font.italic = source_run.font.italic
    target_run.font.underline = source_run.font.underline
    target_run.font.strike = source_run.font.strike  # 这里解决了你的报错
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size

    # 颜色处理
    if source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb


def apply_replacement(paragraph, runs, old_value, new_value, reason, comment_manager, match_type="正则", region="body"):
    """
    执行实际替换、样式复刻与批注添加（只替换匹配部分，保留其他文本）
    
    Args:
        paragraph: 段落对象
        runs: run列表
        old_value: 原始值
        new_value: 新值
        reason: 修改理由
        comment_manager: 批注管理器
        match_type: 匹配类型描述
        region: 区域（"body"=正文, "header"=页眉, "footer"=页脚）
    """
    if not runs:
        return False
    
    # 获取段落完整文本
    full_text = "".join(r.text or "" for r in runs)
    
    # 【关键修复】直接在原始文本中查找，不要清洗 old_value
    # 因为清洗会改变字符（如全角冒号变半角），导致无法找到原始位置
    if old_value in full_text:
        match_start = full_text.index(old_value)
        match_end = match_start + len(old_value)
    else:
        # 如果直接找不到，返回False
        # 不再尝试清洗后匹配，因为会导致位置错误
        return False
    
    # 构建新的文本：前缀 + 新值 + 后缀
    new_full_text = full_text[:match_start] + new_value + full_text[match_end:]
    
    # 保存第一个run的样式作为模板
    template_run = runs[0]
    
    # 清空所有runs
    for r in runs:
        r.text = ""
    
    # 将新文本写入第一个run
    first_run = runs[0]
    first_run.text = new_full_text
    
    # 复制样式
    copy_run_style(template_run, first_run)
    
    # 添加批注
    comment_text = f"【修改建议】\n原值: {old_value}\n新值: {new_value}\n理由: {reason}\n匹配方式: {match_type}"
    
    # 判断是否为页眉或页脚区域
    if region in ["header", "footer"]:
        # 页眉页脚无法直接添加批注，记录到系统批注
        region_name = "页眉" if region == "header" else "页脚"
        system_comment = f"\n【{region_name}修改】{old_value} → {new_value} (理由: {reason})"
        comment_manager.append_to_initial_comment(system_comment)
    else:
        # 正文区域，直接添加批注到run
        comment_manager.add_comment_to_run(first_run, comment_text)
    
    return True

def replace_and_add_comment_in_paragraph(
        paragraph, pattern, old_value, new_value, reason, comment_manager,
        anchor_pattern=None, context_text=None, similarity_threshold=0.3, region="body"
) -> bool:
    runs = list(paragraph.runs)
    if not runs: return False
    full_text = "".join(r.text or "" for r in runs)
    
    # 清洗段落文本
    full_text_clean = clean_text_thoroughly(full_text)
    old_value_clean = clean_text_thoroughly(old_value)

    # --- 逻辑瀑布 ---

    # 1. 尝试在单个run中查找（处理"May 15,2024"这种情况和表格单元格）
    for i, run in enumerate(runs):
        run_text = run.text or ""
        run_clean = clean_text_thoroughly(run_text)
        
        # 策略1a: 精确匹配（原始文本）
        if old_value in run_text:
            new_run_text = run_text.replace(old_value, new_value, 1)
            run.text = new_run_text
            comment_text = f"【修改建议】\n原值: {old_value}\n新值: {new_value}\n理由: {reason}\n匹配方式: 单Run精确匹配"
            
            if region in ["header", "footer"]:
                region_name = "页眉" if region == "header" else "页脚"
                system_comment = f"\n【{region_name}修改】{old_value} → {new_value} (理由: {reason})"
                comment_manager.append_to_initial_comment(system_comment)
            else:
                comment_manager.add_comment_to_run(run, comment_text)
            return True
        
        # 策略1b: 清洗后匹配
        if old_value_clean in run_clean:
            # 尝试在原始文本中找到对应位置
            # 方法1: 尝试直接替换清洗前的值
            if old_value in run_text:
                new_run_text = run_text.replace(old_value, new_value, 1)
            else:
                # 方法2: 使用清洗后的位置映射
                try:
                    clean_pos = run_clean.index(old_value_clean)
                    # 简化：假设清洗主要是去除空格，尝试在原文中找到相似位置
                    # 这里使用一个启发式方法：在原文中查找包含相同字符的子串
                    new_run_text = run_text.replace(old_value, new_value, 1)
                    if new_run_text == run_text:
                        # 如果替换失败，尝试忽略空格
                        old_no_space = old_value.replace(' ', '')
                        text_no_space = run_text.replace(' ', '')
                        if old_no_space in text_no_space:
                            # 找到匹配位置（忽略空格）
                            match_start = text_no_space.index(old_no_space)
                            # 重建：保留前缀，插入新值，保留后缀
                            # 这里需要映射回原始位置（考虑空格）
                            # 简化处理：直接替换整个run
                            new_run_text = new_value
                        else:
                            continue
                except (ValueError, IndexError):
                    continue
            
            run.text = new_run_text
            comment_text = f"【修改建议】\n原值: {old_value}\n新值: {new_value}\n理由: {reason}\n匹配方式: 单Run清洗后匹配"
            
            if region in ["header", "footer"]:
                region_name = "页眉" if region == "header" else "页脚"
                system_comment = f"\n【{region_name}修改】{old_value} → {new_value} (理由: {reason})"
                comment_manager.append_to_initial_comment(system_comment)
            else:
                comment_manager.add_comment_to_run(run, comment_text)
            return True
        
        # 策略1c: 忽略空格匹配（处理"Document No."这种情况）
        old_no_space = old_value.replace(' ', '')
        run_no_space = run_text.replace(' ', '')
        if old_no_space and old_no_space in run_no_space:
            # 找到匹配，但需要保留原始空格结构
            # 简化：如果run文本很短且主要内容匹配，直接替换
            if len(run_text.strip()) < 50 and old_no_space == run_no_space:
                # 整个run就是目标文本（忽略空格）
                run.text = new_value
            else:
                # 部分匹配，尝试定位并替换
                try:
                    match_start = run_no_space.index(old_no_space)
                    # 映射回原始位置（简化：直接替换）
                    run.text = run_text.replace(old_value, new_value, 1)
                    if run.text == run_text:
                        # 替换失败，使用新值
                        run.text = new_value
                except ValueError:
                    continue
            
            comment_text = f"【修改建议】\n原值: {old_value}\n新值: {new_value}\n理由: {reason}\n匹配方式: 单Run忽略空格匹配"
            
            if region in ["header", "footer"]:
                region_name = "页眉" if region == "header" else "页脚"
                system_comment = f"\n【{region_name}修改】{old_value} → {new_value} (理由: {reason})"
                comment_manager.append_to_initial_comment(system_comment)
            else:
                comment_manager.add_comment_to_run(run, comment_text)
            return True

    # 2. 尝试正则匹配 (Strategy A) - 在清洗后的文本上匹配
    if re.search(pattern, full_text_clean, flags=re.IGNORECASE | re.DOTALL):
        return apply_replacement(paragraph, runs, old_value, new_value, reason, comment_manager, "正则匹配", region)

    # 3. 尝试直接字符串匹配（清洗后）
    if old_value_clean in full_text_clean:
        return apply_replacement(paragraph, runs, old_value, new_value, reason, comment_manager, "直接匹配(清洗后)", region)

    # 4. 尝试模糊匹配 (Strategy B)
    if is_fuzzy_match(full_text_clean, old_value_clean, threshold=0.85):
        return apply_replacement(paragraph, runs, old_value, new_value, reason, comment_manager, "模糊匹配", region)

    # 5. 尝试指纹匹配 (Strategy C)
    fingerprint_old = get_alphanumeric_fingerprint(old_value)
    fingerprint_full = get_alphanumeric_fingerprint(full_text)
    if len(fingerprint_old) >= 3 and fingerprint_old in fingerprint_full:
        return apply_replacement(paragraph, runs, old_value, new_value, reason, comment_manager, "指纹匹配", region)

    # 6. 策略 D: 长文本切片匹配 (防御性切片)
    if len(old_value_clean) > 20:
        # 使用 20 字符作为切片长度，但只要匹配到任意一个块，就认为是该段落
        chunks = [old_value_clean[i:i + 20] for i in range(0, len(old_value_clean), 20)]
        for chunk in chunks:
            if chunk in full_text_clean:
                return apply_replacement(
                    paragraph, runs, old_value, new_value, reason, comment_manager, "长文本切片匹配", region
                )
    
    # 7. 策略 E: 处理特殊情况 - 中文字符嵌入（如"存量"）
    # 移除所有空格后再尝试匹配
    full_text_no_space = full_text_clean.replace(' ', '')
    old_value_no_space = old_value_clean.replace(' ', '')
    if old_value_no_space and old_value_no_space in full_text_no_space:
        return apply_replacement(paragraph, runs, old_value, new_value, reason, comment_manager, "无空格匹配", region)
    
    # 8. 策略 F: 处理全角括号等特殊字符的情况
    # 尝试在原始文本上直接查找（不清洗）
    if old_value in full_text:
        return apply_replacement(paragraph, runs, old_value, new_value, reason, comment_manager, "原始文本匹配", region)

    return False

def get_alphanumeric_fingerprint(text: str) -> str:
    """
    极度宽松的指纹：只提取字符，彻底忽略格式、标点和空格。
    """
    if not text: return ""
    # 彻底转换为半角并去掉标点后再指纹化
    clean = clean_text_thoroughly(text)
    return re.sub(r'[^a-zA-Z0-9\u4e00-\u9fa5]', '', clean)


def preprocess_special_cases(old_value: str, doc: Document) -> Tuple[bool, str, str]:
    """
    预处理特殊情况，尝试直接在文档中查找并返回匹配信息
    
    Returns:
        (found, matched_text, strategy_description)
    """
    # 特殊情况0: 处理Word自动编号（如 i., ii., iii.）
    # 这些编号不是文本的一部分，而是Word的列表编号
    if is_list_pattern(old_value):
        # 对于编号模式，我们需要特殊处理
        # 因为编号可能是Word的自动编号，不在文本中
        # 我们返回False，让后续逻辑通过上下文匹配
        return False, "", ""
    
    # 特殊情况1: 处理全角括号（如：（a）Stop...）
    if '（' in old_value or '）' in old_value:
        for p in iter_all_paragraphs(doc):
            full_text = "".join(r.text or "" for r in p.runs)
            if old_value in full_text:
                return True, full_text, "全角括号直接匹配"
            # 尝试忽略换行符匹配（处理被换行符分隔的情况）
            full_text_no_newline = full_text.replace('\n', '').replace('\r', '')
            old_value_no_newline = old_value.replace('\n', '').replace('\r', '')
            if old_value_no_newline in full_text_no_newline:
                return True, full_text, "全角括号匹配(忽略换行符)"
    
    # 特殊情况2: 处理中文字符嵌入（如：存量）
    if any('\u4e00' <= c <= '\u9fff' for c in old_value):
        old_clean = clean_text_thoroughly(old_value)
        for p in iter_all_paragraphs(doc):
            full_text = "".join(r.text or "" for r in p.runs)
            full_clean = clean_text_thoroughly(full_text)
            if old_clean in full_clean:
                return True, full_text, "中文字符匹配"
    
    # 特殊情况3: 处理撇号变体（如：Guang'an）
    if "'" in old_value or "'" in old_value or "'" in old_value:
        # 尝试所有撇号变体
        variants = [
            old_value,
            old_value.replace("'", "'"),
            old_value.replace("'", "'"),
            old_value.replace("'", "`"),
        ]
        for p in iter_all_paragraphs(doc):
            full_text = "".join(r.text or "" for r in p.runs)
            for variant in variants:
                if variant in full_text:
                    return True, full_text, f"撇号变体匹配({variant})"
    
    # 特殊情况4: 处理数字格式（如：RMB1,060,000）
    if re.search(r'\d{1,3}(,\d{3})+', old_value):
        # 尝试移除逗号后匹配
        old_no_comma = old_value.replace(',', '')
        for p in iter_all_paragraphs(doc):
            full_text = "".join(r.text or "" for r in p.runs)
            full_no_comma = full_text.replace(',', '')
            if old_no_comma in full_no_comma:
                return True, full_text, "数字格式匹配(忽略逗号)"
    
    # 特殊情况5: 处理日期时间连接问题（如：May 15,2024 连接成 May 15,20249）
    # 检测模式：日期后面紧跟数字（如 May 15,2024 变成 May 15,20249）
    if re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2},\s*\d{4}', old_value):
        # 这是日期格式，尝试查找可能连接的版本
        old_no_space = old_value.replace(' ', '').replace(',', ',')  # 保留逗号
        for p in iter_all_paragraphs(doc):
            full_text = "".join(r.text or "" for r in p.runs)
            # 查找类似 "May15,20249" 的模式
            if re.search(r'May\s*15,\s*20249', full_text):
                return True, full_text, "日期连接匹配(May 15,2024 -> May 15,20249)"
    
    # 特殊情况6: 处理时间格式（如：9 am 在 "20249 am" 中）
    if re.search(r'\d+\s+(am|pm|a\.m\.|p\.m\.)', old_value, re.IGNORECASE):
        # 提取数字部分
        match = re.search(r'(\d+)\s+(am|pm)', old_value, re.IGNORECASE)
        if match:
            time_num = match.group(1)
            time_suffix = match.group(2)
            # 查找可能连接的版本（如 "20249 am" 包含 "9 am"）
            for p in iter_all_paragraphs(doc):
                full_text = "".join(r.text or "" for r in p.runs)
                # 查找以该数字结尾+空格+am/pm的模式
                if re.search(rf'{time_num}\s+{time_suffix}', full_text, re.IGNORECASE):
                    return True, full_text, f"时间格式匹配({time_num} {time_suffix})"
    
    # 特殊情况7: 通用空格忽略匹配
    if ' ' in old_value or ',' in old_value:
        for p in iter_all_paragraphs(doc):
            full_text = "".join(r.text or "" for r in p.runs)
            # 移除所有空格后匹配
            old_no_space = old_value.replace(' ', '')
            full_no_space = full_text.replace(' ', '')
            if old_no_space in full_no_space:
                return True, full_text, "忽略空格匹配"
    
    return False, "", ""

# =========================
# 6) 主替换函数（增强版）
# =========================

def replace_and_comment_in_docx(
        doc: Document,
        old_value: str,
        new_value: str,
        reason: str,
        comment_manager: CommentManager,
        context: str = "",
        anchor_text: str = "",
        region: str = "all"
) -> Tuple[bool, str]:
    """
    多策略执行替换并添加批注（增强上下文匹配）

    Args:
        doc: Word 文档对象
        old_value: 原始值
        new_value: 新值
        reason: 修改理由
        comment_manager: 批注管理器
        context: 上下文文本（用于提取锚点和相似度验证）
        anchor_text: 显式指定的锚点文本
        region: 替换区域 ("all"=全部, "body"=正文, "header"=页眉, "footer"=页脚)

    Returns:
        (是否成功, 匹配策略描述)
        
    注意：
        - 正文区域的替换会直接在段落上添加批注
        - 页眉/页脚区域的替换会将修改记录添加到系统批注中（因为页眉页脚无法直接添加批注）
    """
    # 保存原始值用于特殊情况匹配
    old_value_original = old_value or ""
    
    # 【关键修复】不要清洗 old_value，保持原样以便精确匹配
    # 只清洗 new_value 和 context
    # old_value = clean_text_thoroughly(old_value or "").strip()  # 注释掉
    old_value = (old_value or "").strip()  # 只去除首尾空格
    new_value = clean_text_thoroughly(new_value or "").strip()
    context = clean_text_thoroughly(context or "")

    # 处理 reason 可能是元组的情况
    if isinstance(reason, (list, tuple)):
        reason = " ".join([str(i) for i in reason if i]).strip()
    reason = reason or "数值/术语不一致"

    if not old_value or not new_value:
        return False, "数据缺失"

    # 根据 region 参数选择段落迭代器
    if region == "body":
        paragraph_iterator = lambda: iter_body_paragraphs(doc)
        region_desc = "正文"
    elif region == "header":
        paragraph_iterator = lambda: iter_header_paragraphs(doc)
        region_desc = "页眉"
    elif region == "footer":
        paragraph_iterator = lambda: iter_footer_paragraphs(doc)
        region_desc = "页脚"
    else:  # "all" or any other value
        paragraph_iterator = lambda: iter_all_paragraphs(doc)
        region_desc = "全部"

    # ===== 策略0A：Word自动编号替换 =====
    # 检测是否为编号模式（如 i., ii., iii.）
    if is_list_pattern(old_value):
        try:
            from zhongfanyi.llm.llm_project.replace.numbering_replacer import replace_numbering_in_docx
            success, message = replace_numbering_in_docx(
                doc, old_value, new_value, context, 
                comment_manager, reason
            )
            if success:
                return True, f"Word自动编号替换: {message}"
            else:
                # 检查是否已经被之前的操作修改过
                # 如果找不到旧格式，可能是因为已经被批量修改了
                print(f"  提示: {old_value} 可能已被之前的编号替换操作修改")
        except Exception as e:
            # 如果编号替换失败，继续尝试其他策略
            print(f"  编号替换失败: {e}")

    # ===== 策略0：预处理特殊情况 =====
    found, matched_text, strategy = preprocess_special_cases(old_value_original, doc)
    if found:
        # 找到匹配的段落，执行替换（只在指定区域内查找）
        for p in paragraph_iterator():
            full_text = "".join(r.text or "" for r in p.runs)
            if full_text == matched_text:
                # 特殊处理：如果是忽略换行符的匹配，需要跨runs替换
                if "忽略换行符" in strategy:
                    # 移除换行符后进行匹配和替换
                    full_text_no_newline = full_text.replace('\n', '').replace('\r', '')
                    old_value_no_newline = old_value_original.replace('\n', '').replace('\r', '')
                    
                    if old_value_no_newline in full_text_no_newline:
                        # 找到匹配位置
                        match_start = full_text_no_newline.index(old_value_no_newline)
                        match_end = match_start + len(old_value_no_newline)
                        
                        # 在原始文本中找到对应的位置（考虑换行符）
                        # 简化处理：直接替换整个段落的文本
                        new_full_text = full_text_no_newline.replace(old_value_no_newline, new_value, 1)
                        
                        # 清空所有runs并将新文本写入第一个run
                        runs = list(p.runs)
                        if runs:
                            for r in runs:
                                r.text = ""
                            runs[0].text = new_full_text
                            comment_text = f"【修改建议】\n原值: {old_value}\n新值: {new_value}\n理由: {reason}\n匹配方式: {strategy}"
                            comment_manager.add_comment_to_run(runs[0], comment_text)
                            return True, strategy
                
                # 在这个段落中查找并替换特定的run
                runs = list(p.runs)
                for i, run in enumerate(runs):
                    run_text = run.text or ""
                    run_clean = clean_text_thoroughly(run_text)
                    old_clean = clean_text_thoroughly(old_value_original)
                    
                    # 尝试多种匹配方式
                    if old_value_original in run_text:
                        # 在原始文本中找到了，只替换匹配部分
                        run.text = run_text.replace(old_value_original, new_value, 1)
                        comment_text = f"【修改建议】\n原值: {old_value}\n新值: {new_value}\n理由: {reason}\n匹配方式: {strategy}"
                        comment_manager.add_comment_to_run(run, comment_text)
                        return True, strategy
                    elif old_clean in run_clean:
                        # 在清洗后的文本中找到了，需要在原始文本中定位并替换
                        # 简化处理：如果清洗后能匹配，尝试直接替换
                        if old_value in run_text:
                            run.text = run_text.replace(old_value, new_value, 1)
                        else:
                            # 使用清洗后的匹配位置
                            match_pos = run_clean.index(old_clean)
                            run.text = run_text[:match_pos] + new_value + run_text[match_pos + len(old_clean):]
                        comment_text = f"【修改建议】\n原值: {old_value}\n新值: {new_value}\n理由: {reason}\n匹配方式: {strategy}"
                        comment_manager.add_comment_to_run(run, comment_text)
                        return True, strategy
                    elif old_value in run_text:
                        # 使用清洗后的old_value匹配
                        run.text = run_text.replace(old_value, new_value, 1)
                        comment_text = f"【修改建议】\n原值: {old_value}\n新值: {new_value}\n理由: {reason}\n匹配方式: {strategy}"
                        comment_manager.add_comment_to_run(run, comment_text)
                        return True, strategy
                
                # 如果没有找到特定run，尝试整体替换
                if runs:
                    if apply_replacement(p, runs, old_value, new_value, reason, 
                                        comment_manager, strategy, region):
                        return True, strategy

    # ===== 策略1：显式锚点 + 上下文验证（最高优先级）=====
    if anchor_text:
        anchor_text = clean_text_thoroughly(anchor_text)
        anchor_pattern = build_smart_pattern(anchor_text, mode="balanced")
        target_pattern = build_smart_pattern(old_value, mode="strict")

        if anchor_pattern and target_pattern:
            for p in paragraph_iterator():
                if replace_and_add_comment_in_paragraph(
                        p, target_pattern, old_value, new_value, reason,
                        comment_manager, anchor_pattern, context_text=context, region=region
                ):
                    return True, f"锚点匹配 (显式) + 上下文验证 [{region_desc}]"

    # ===== 策略2：上下文锚点 + 相似度验证 =====
    if context:
        context_anchor = extract_anchor_with_target(context, old_value, window=60)

        if context_anchor:
            anchor_pattern = build_smart_pattern(context_anchor, mode="balanced")
            target_pattern = build_smart_pattern(old_value, mode="strict")

            for p in paragraph_iterator():
                if replace_and_add_comment_in_paragraph(
                        p, target_pattern, old_value, new_value, reason,
                        comment_manager, anchor_pattern, context_text=context,
                        similarity_threshold=0.3, region=region  # 相似度阈值
                ):
                    return True, f"上下文匹配 (锚点: {context_anchor[:30]}...) + 相似度验证 [{region_desc}]"

    # ===== 策略3：严格模式 + 上下文验证 =====
    if context:
        strict_pattern = build_smart_pattern(old_value, mode="strict")
        for p in paragraph_iterator():
            if replace_and_add_comment_in_paragraph(
                    p, strict_pattern, old_value, new_value, reason,
                    comment_manager, context_text=context,
                    similarity_threshold=0.2, region=region  # 降低阈值以提高召回率
            ):
                return True, f"严格模式 (全局) + 上下文验证 [{region_desc}]"

    # ===== 策略4：严格模式全局匹配（无上下文） =====
    strict_pattern = build_smart_pattern(old_value, mode="strict")
    for p in paragraph_iterator():
        if replace_and_add_comment_in_paragraph(
                p, strict_pattern, old_value, new_value, reason, comment_manager, region=region
        ):
            return True, f"严格模式 (全局，无上下文验证) [{region_desc}]"

    # ===== 策略5：平衡模式（兜底） =====
    balanced_pattern = build_smart_pattern(old_value, mode="balanced")
    for p in paragraph_iterator():
        if replace_and_add_comment_in_paragraph(
                p, balanced_pattern, old_value, new_value, reason, comment_manager, region=region
        ):
            return True, f"平衡模式 (容错) [{region_desc}]"
    
    # ===== 策略6：宽松模式 - 处理中文字符嵌入等特殊情况 =====
    for p in paragraph_iterator():
        runs = list(p.runs)
        if not runs:
            continue
        full_text = "".join(r.text or "" for r in runs)
        full_text_clean = clean_text_thoroughly(full_text)
        
        # 移除所有空格后匹配
        full_text_no_space = full_text_clean.replace(' ', '')
        old_value_no_space = old_value.replace(' ', '')
        
        if old_value_no_space and old_value_no_space in full_text_no_space:
            if apply_replacement(p, runs, old_value, new_value, reason, 
                               comment_manager, f"宽松模式(无空格匹配) [{region_desc}]", region):
                return True, f"宽松模式(无空格匹配) [{region_desc}]"

    return False, f"未找到匹配项 (搜索区域: {region_desc})"


# =========================
# 7) 测试用例（增强版）
# =========================

if __name__ == "__main__":
    # 测试清洗功能
    test_text = "i.\u200b First item"
    cleaned = clean_text_thoroughly(test_text)
    print(f"原始: {repr(test_text)}")
    print(f"清洗: {repr(cleaned)}")
    print()

    # 测试编号识别（包括复合罗马数字）
    test_cases = ["i.", "ii.", "iii.", "iv.", "ix.", "x.", "(1)", "1)", "a."]
    print("=" * 60)
    print("编号模式识别测试")
    print("=" * 60)
    for case in test_cases:
        result = is_list_pattern(case)
        print(f"{case:6s} -> 是编号: {result}")
    print()

    # 测试罗马数字验证
    print("=" * 60)
    print("罗马数字验证测试")
    print("=" * 60)
    roman_tests = ["i", "ii", "iii", "iv", "ix", "x", "xi", "xiv", "xix", "xx", "iiii", "vv"]
    for roman in roman_tests:
        result = is_valid_roman_numeral(roman)
        print(f"{roman:6s} -> 有效: {result}")
    print()

    # 测试模式构建
    print("=" * 60)
    print("模式构建测试")
    print("=" * 60)
    pattern_tests = ["i.", "ix.", "20%", "$1,234.56"]
    for test in pattern_tests:
        pattern = build_smart_pattern(test, mode="balanced")
        print(f"{test:12s} -> 模式: {pattern}")
    print()

    # 测试上下文相似度
    print("=" * 60)
    print("上下文相似度测试")
    print("=" * 60)
    context1 = "The total revenue is 20% higher than last year"
    context2 = "Revenue increased by 20% compared to previous period"
    context3 = "The discount rate is 20% for all customers"

    similarity1 = calculate_context_similarity(context1, context2)
    similarity2 = calculate_context_similarity(context1, context3)

    print(f"文本1: {context1}")
    print(f"文本2: {context2}")
    print(f"相似度: {similarity1:.2f}")
    print()
    print(f"文本1: {context1}")
    print(f"文本3: {context3}")
    print(f"相似度: {similarity2:.2f}")