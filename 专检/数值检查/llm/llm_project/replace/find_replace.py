import os
import shutil
from typing import List, Dict, Any, Tuple
from datetime import datetime
from docx import Document
from lxml import etree
from zipfile import ZipFile
import warnings
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from llm.llm_project.llm_check.check import Match
from llm.llm_project.parsers.body_extractor import extract_body_text
from llm.llm_project.parsers.footer_extractor import extract_footers
from llm.llm_project.parsers.header_extractor import extract_headers
from llm.utils.json_files import write_json_with_timestamp

warnings.filterwarnings("ignore")

# =========================
# 0) 基础配置
# =========================

BACKUP_DIR_NAME = "backup"

# =========================
# 1) DOCX 预加载
# =========================

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NAMESPACES = {"w": W_NS}


def get_xml_text(element) -> str:
    """从任意XML元素及其子元素中提取纯文本"""
    texts = []
    for t in element.iter(f"{{{W_NS}}}t"):
        if t.text:
            texts.append(t.text)
    return "".join(texts)



class DocContentLoader:
    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.footnotes = {}
        self.endnotes = {}
        self.comments = {}
        self.headers = []
        self.footers = []
        self._load_all()

    def _load_xml_map(self, zip_file, filename, tag_name, id_attr="id"):
        data_map = {}
        if filename not in zip_file.namelist():
            return data_map

        with zip_file.open(filename) as f:
            tree = etree.parse(f)
            for elem in tree.findall(f".//w:{tag_name}", NAMESPACES):
                eid = elem.get(f"{{{W_NS}}}{id_attr}")
                if eid in ["-1", "0"]:
                    continue
                texts = [t.text for t in elem.iter(f"{{{W_NS}}}t") if t.text]
                full_text = "".join(texts).strip()
                if full_text:
                    data_map[eid] = full_text
        return data_map

    def _load_headers_footers(self, zip_file, prefix):
        contents = []
        files = [f for f in zip_file.namelist() if f.startswith(f"word/{prefix}") and f.endswith(".xml")]
        files.sort(key=lambda x: int("".join(filter(str.isdigit, x)) or 0))
        for f_name in files:
            try:
                with zip_file.open(f_name) as f:
                    tree = etree.parse(f)
                    texts = [t.text for t in tree.iter(f"{{{W_NS}}}t") if t.text]
                    full_text = "".join(texts).strip()
                    if full_text:
                        contents.append(full_text)
            except Exception:
                pass
        return contents

    def _load_all(self):
        with ZipFile(self.doc_path, "r") as zf:
            self.footnotes = self._load_xml_map(zf, "word/footnotes.xml", "footnote")
            self.endnotes = self._load_xml_map(zf, "word/endnotes.xml", "endnote")
            self.comments = self._load_xml_map(zf, "word/comments.xml", "comment")
            self.headers = self._load_headers_footers(zf, "header")
            self.footers = self._load_headers_footers(zf, "footer")


# =========================
# 2) 错误报告解析（增强版）
# =========================

def _normalize_spaces(s: str) -> str:
    """标准化空白字符"""
    return re.sub(r"\s+", " ", (s or "")).strip()


import re
import ast
from typing import Any, Dict, List, Optional


def _unify_labels(raw: str) -> str:
    """
    统一“字段名 + 冒号”的各种断行/空格形态，减少解析歧义。
    目标：把 '译文\\n修改建议值\\n: xxx'、'译 文修改建议值: xxx' 等统一成 '译文修改建议值: xxx'
    """
    raw = raw or ""

    # 先把全局的回车/制表等统一为 \n / 空格
    raw = raw.replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")

    # 这些字段名最关键：允许中间任意空白（含换行）
    label_patterns = [
        "错误编号",
        "错误类型",
        "原文数值",
        "译文数值",
        "译文修改建议值",
        "修改理由",
        "违反的规则",
        "原文上下文",
        "译文上下文",
        "原文位置",
        "译文位置",
        "替换锚点",
    ]

    # 把“字段名”内部的空白压平，并让冒号紧跟字段名
    for lab in label_patterns:
        # 例如：译文\s*修改\s*建议\s*值\s*[:：]  ->  译文修改建议值:
        chars = list(lab)
        pat = r"".join(re.escape(c) + r"\s*" for c in chars) + r"[:：]?"
        raw = re.sub(pat, lab, raw, flags=re.IGNORECASE)

        # 统一冒号：字段名后如果紧跟的不是冒号，就补齐（仅当后面确实有内容时）
        raw = re.sub(rf"({re.escape(lab)})\s*\n?\s*[:：]\s*", rf"\1: ", raw)

    # 修复你原来对“错误编号”做的特殊替换（保留）
    raw = raw.replace("错误\n编号", "错误编号")
    raw = raw.replace("错误\n编号:", "错误编号:")
    raw = raw.replace("编号:", "错误编号:")

    return raw

def _split_merged_fields(one: Dict[str, Any]) -> Dict[str, Any]:
    """
    兜底：若译文数值里吞了“译文修改建议值: ...”，则拆分并回填两个字段。
    允许出现：'译 文修改建议值'、换行、多个空格等。
    """
    tv = (one.get("译文数值") or "").strip()
    sv = (one.get("译文修改建议值") or "").strip()

    if not tv:
        return one

    # 允许各种空白：译\s*文\s*修\s*改\s*建\s*议\s*值
    marker_pat = r"(译\s*文\s*修\s*改\s*建\s*议\s*值)\s*[:：]\s*"
    m = re.search(marker_pat, tv, flags=re.IGNORECASE)
    if not m:
        return one

    left = tv[:m.start()].strip()
    right = tv[m.end():].strip()

    if left:
        one["译文数值"] = re.sub(r"\s+", " ", left).strip()

    # 如果原本建议值为空，则用拆出来的补上；如果不为空，也可以选择保留原值（更安全）
    if right and not sv:
        one["译文修改建议值"] = re.sub(r"\s+", " ", right).strip()

    return one


import re
import ast
from typing import Any, Dict, List, Optional


def _extract_first_braced_block(raw: str) -> Optional[str]:
    """
    从 raw 中提取第一个完整的 {...} 或 [...] 字面量块（后面允许跟说明文字）。
    用于兼容你给的那种 dict/list 结构化输入。
    """
    if not raw:
        return None

    m = re.search(r"[\{\[]", raw)
    if not m:
        return None

    start = m.start()
    open_ch = raw[start]
    close_ch = "}" if open_ch == "{" else "]"

    depth = 0
    in_str = False
    str_ch = ""
    escape = False

    for i in range(start, len(raw)):
        ch = raw[i]

        if in_str:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == str_ch:
                in_str = False
            continue

        if ch in ("'", '"'):
            in_str = True
            str_ch = ch
            continue

        if ch == open_ch:
            depth += 1
        elif ch == close_ch:
            depth -= 1
            if depth == 0:
                return raw[start:i + 1]

    return None


def _unify_labels_keep_colon(raw: str) -> str:
    """
    关键：把字段名里的各种空格/换行统一掉，并且【强制保留/补齐冒号】为 '字段名:'。
    解决 Word 自动换行：'错误\\n类型:'、'译 文修改建议值\\n:' 等。
    """
    raw = raw or ""
    raw = raw.replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")

    labels = [
        "错误编号",
        "错误类型",
        "原文数值",
        "译文数值",
        "译文修改建议值",
        "修改理由",
        "违反的规则",
        "原文上下文",
        "译文上下文",
        "原文位置",
        "译文位置",
        "替换锚点",
    ]

    for lab in labels:
        pat = (
                r"(?m)"  # 多行模式
                r"(^|\n)"  # 行首
                + r"\s*"
                + r"".join(re.escape(c) + r"\s*" for c in lab)
                + r"\s*[:：]?"
        )
        raw = re.sub(pat, r"\1" + lab + ":", raw, flags=re.IGNORECASE)

    # 冒号后统一一个空格（含中文冒号）
    raw = re.sub(r"[:：]\s*", ": ", raw)
    return raw


def _split_merged_translation_fields(one: Dict[str, Any]) -> Dict[str, Any]:
    """
    兜底：如果译文数值吞了“译文修改建议值: ...”，则拆开两个字段。
    _unify_labels_keep_colon 会把各种形态统一成 '译文修改建议值:'，因此这里拆分很稳。
    """
    tv = (one.get("译文数值") or "").strip()
    sv = (one.get("译文修改建议值") or "").strip()
    if not tv:
        return one

    marker = "译文修改建议值:"
    idx = tv.find(marker)
    if idx == -1:
        return one

    left = tv[:idx].strip()
    right = tv[idx + len(marker):].strip()

    if left:
        one["译文数值"] = re.sub(r"\s+", " ", left).strip()

    # 如果原本建议值为空，用 right 补齐（更安全；不覆盖已有值）
    if right and not sv:
        one["译文修改建议值"] = re.sub(r"\s+", " ", right).strip()

    return one


def load_error_list_from_text(raw: str) -> List[Dict[str, Any]]:
    """
    解析错误报告（增强版）：
    1) 优先解析结构化 dict/list（你的示例那种）
    2) 再解析文本型报告（支持字段名断行/空格、同一行字段、冒号断行）
    3) 兜底拆分：译文数值吞掉译文修改建议值
    """
    raw = raw or ""

    # ========== 1) 结构化优先：提取第一个 {...} / [...] 后 literal_eval ==========
    block = _extract_first_braced_block(raw)
    if block:
        try:
            obj = ast.literal_eval(block)
            if isinstance(obj, dict):
                obj = [obj]
            if isinstance(obj, list):
                fixed = []
                for it in obj:
                    if isinstance(it, dict):
                        fixed.append(_split_merged_translation_fields(it))
                if fixed:
                    return fixed
        except Exception:
            pass  # 失败就走文本解析

    # ========== 2) 文本解析：统一标签（保留/补齐冒号） ==========
    cleaned = _unify_labels_keep_colon(raw)

    # 每条错误从 “错误编号: 数字” 开始
    parts = re.split(r"(?=错误编号:\s*\d+)", cleaned)
    errors: List[Dict[str, Any]] = []

    # 关键：边界不要依赖换行，直接用 “下一个字段名:” 做 lookahead
    key_map = [
        ("错误编号", r"错误编号:\s*(.+?)(?=错误类型:|$)"),
        ("错误类型", r"错误类型:\s*(.+?)(?=原文数值:|$)"),
        ("原文数值", r"原文数值:\s*(.+?)(?=译文数值:|$)"),
        ("译文数值", r"译文数值:\s*(.+?)(?=译文修改建议值:|修改理由:|$)"),
        ("译文修改建议值", r"译文修改建议值:\s*(.+?)(?=修改理由:|$)"),
        ("修改理由", r"修改理由:\s*(.+?)(?=违反的规则:|$)"),
        ("违反的规则", r"违反的规则:\s*(.+?)(?=原文上下文:|$)"),
        ("原文上下文", r"原文上下文:\s*(.+?)(?=译文上下文:|$)"),
        ("译文上下文", r"译文上下文:\s*(.+?)(?=原文位置:|$)"),
        ("原文位置", r"原文位置:\s*(.+?)(?=译文位置:|$)"),
        ("译文位置", r"译文位置:\s*(.+?)(?=替换锚点:|$)"),
        ("替换锚点", r"替换锚点:\s*(.+?)(?=错误编号:|$)"),
    ]

    for p in parts:
        if "错误编号:" not in p:
            continue

        one: Dict[str, Any] = {}
        for k, pat in key_map:
            mm = re.search(pat, p, flags=re.DOTALL | re.IGNORECASE)
            if mm:
                v = mm.group(1).strip()
                v = re.sub(r"\s+", " ", v).strip()
                one[k] = v

        # 兜底拆分合并字段
        one = _split_merged_translation_fields(one)

        if one.get("错误编号") or one.get("译文数值") or one.get("译文修改建议值"):
            errors.append(one)

    if not errors:
        raise ValueError("未在错误报告文本中找到可解析的错误列表")

    return errors




# =========================
# 3) 改进的匹配策略
# =========================

def build_smart_pattern(s: str, mode: str = "balanced") -> str:
    """
    构建智能匹配模式

    Args:
        s: 待匹配字符串
        mode: 匹配模式
            - "strict": 严格匹配（完全精确）
            - "balanced": 平衡模式（数字/标点保持连续，单词间允许空格）
            - "loose": 宽松模式（字符间允许空格，但数字连续）
    """
    s = (s or "").strip()
    if not s:
        return ""

    if mode == "strict":
        # 严格模式：完全精确匹配（转义特殊字符）
        return re.escape(s)

    elif mode == "balanced":
        # 平衡模式：数字和标点保持连续，单词间允许空格
        pieces = []
        i = 0
        while i < len(s):
            ch = s[i]

            # 跳过空格
            if ch.isspace():
                if pieces and not pieces[-1].endswith(r"\s*"):
                    pieces.append(r"\s*")
                i += 1
                continue

            # 数字序列：保持连续（包括小数点、逗号）
            if ch.isdigit():
                num_str = ""
                while i < len(s) and (s[i].isdigit() or s[i] in ".,"):
                    num_str += s[i]
                    i += 1
                pieces.append(re.escape(num_str))
                continue

            # 标点符号：保持连续
            if ch in ".,;:!?()[]{}\"'-/":
                pieces.append(re.escape(ch))
                i += 1
                continue

            # 字母序列：单词间允许空格
            if ch.isalpha():
                word = ""
                while i < len(s) and s[i].isalpha():
                    word += s[i]
                    i += 1
                pieces.append(re.escape(word))
                if i < len(s) and not s[i].isspace():
                    continue
                if i < len(s):
                    pieces.append(r"\s*")
                continue

            # 其他字符
            pieces.append(re.escape(ch))
            i += 1

        return "".join(pieces).strip()

    else:  # loose
        # 宽松模式：字符间允许空格，但数字保持连续
        pieces = []
        i = 0
        while i < len(s):
            ch = s[i]

            if ch.isspace():
                i += 1
                continue

            # 数字序列保持连续
            if ch.isdigit():
                num_str = ""
                while i < len(s) and (s[i].isdigit() or s[i] in ".,"):
                    num_str += s[i]
                    i += 1
                pieces.append(re.escape(num_str) + r"\s*")
                continue

            # 其他字符间允许空格
            pieces.append(re.escape(ch) + r"\s*")
            i += 1

        return "".join(pieces).strip()


def extract_anchor_with_target(context: str, target_value: str, window: int = 30) -> Optional[str]:
    """
    从上下文中提取包含目标数值的锚点短语（更大窗口）
    """
    if not context or not target_value:
        return None

    context = _normalize_spaces(context)
    target_value = target_value.strip()

    # 先尝试严格匹配
    if target_value in context:
        idx = context.index(target_value)
        start = max(0, idx - window)
        end = min(len(context), idx + len(target_value) + window)
        anchor = context[start:end]
        # 修剪边界单词
        anchor = re.sub(r"^\S*\s+", "", anchor)
        anchor = re.sub(r"\s+\S*$", "", anchor)
        return anchor.strip()

    # 再尝试平衡模式匹配
    pattern = build_smart_pattern(target_value, mode="balanced")
    if not pattern:
        return None

    match = re.search(pattern, context, flags=re.IGNORECASE)
    if not match:
        return None

    start, end = match.span()
    prefix_start = max(0, start - window)
    suffix_end = min(len(context), end + window)

    anchor = context[prefix_start:suffix_end]
    anchor = re.sub(r"^\S*\s+", "", anchor)
    anchor = re.sub(r"\s+\S*$", "", anchor)

    return anchor.strip() if anchor.strip() else None


# =========================
# 4) Word批注功能（核心修改 - 替换+批注版）
# =========================

class CommentManager:
    """Word批注管理器（完整XML支持）"""

    def __init__(self, doc: Document):
        self.doc = doc
        self._comment_id = 0
        self._comments_part = None
        self._init_comments_part()

    def _init_comments_part(self):
        """初始化comments.xml部分（如果不存在则创建）"""
        try:
            package = self.doc.part.package

            # 查找现有的comments part
            for rel in self.doc.part.rels.values():
                if "comments" in rel.target_ref:
                    self._comments_part = rel.target_part
                    break

            if self._comments_part:
                # 解析现有批注ID
                root = self._comments_part.element
                for comment in root.findall(f".//{{{W_NS}}}comment"):
                    cid = comment.get(f"{{{W_NS}}}id")
                    if cid and cid.isdigit():
                        self._comment_id = max(self._comment_id, int(cid))
            else:
                # 创建新的comments part
                self._create_comments_part()

            self._comment_id += 1

        except Exception as e:
            print(f"⚠️ 初始化批注部分失败: {e}")
            self._comment_id = 1

    def _create_comments_part(self):
        """创建comments.xml文件和关系"""
        try:
            from docx.opc.part import XmlPart
            from docx.opc.packuri import PackURI

            # 创建comments.xml的XML结构
            comments_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="{W_NS}" 
            xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
            xmlns:o="urn:schemas-microsoft-com:office:office"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            xmlns:w10="urn:schemas-microsoft-com:office:word"
            xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
            xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
            xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
            xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
            xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
            xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
            mc:Ignorable="w14 w15 wp14">
</w:comments>'''

            # 解析XML
            comments_element = etree.fromstring(comments_xml.encode('utf-8'))

            # 创建Part对象
            partname = PackURI('/word/comments.xml')
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'

            package = self.doc.part.package
            self._comments_part = XmlPart(partname, content_type, comments_element, package)

            # 添加关系
            self.doc.part.relate_to(self._comments_part, RT.COMMENTS)

            print("✓ 成功创建 comments.xml")

        except Exception as e:
            print(f"⚠️ 创建comments.xml失败: {e}")
            print("   将使用备用方案（批注可能无法显示）")

    def _get_next_comment_id(self) -> int:
        """获取下一个批注ID"""
        cid = self._comment_id
        self._comment_id += 1
        return cid

    def create_initial_comment(self) -> bool:
        """
        在文档开头创建一个初始化批注（确保comments.xml结构完整）
        """
        try:
            if not self.doc.paragraphs:
                print("⚠️ 文档无段落，无法创建初始批注")
                return False

            first_para = self.doc.paragraphs[0]
            if not first_para.runs:
                first_para.add_run(".")  # 添加一个占位符

            first_run = first_para.runs[0]

            # 创建初始批注
            init_text = f"【翻译校对系统】批注功能已启用 - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            success = self.add_comment_to_run(first_run, init_text, author="系统")

            if success:
                print("✓ 已创建初始化批注")

            return success

        except Exception as e:
            print(f"⚠️ 创建初始批注失败: {e}")
            return False

    def add_comment_to_run(self, run, comment_text: str, author: str = "翻译校对") -> bool:
        """
        为run添加批注

        Args:
            run: Run对象
            comment_text: 批注内容
            author: 批注作者

        Returns:
            是否成功添加批注
        """
        try:
            if not self._comments_part:
                print("⚠️ comments part不存在，无法添加批注")
                return False

            # 获取批注ID
            comment_id = self._get_next_comment_id()

            # 创建批注范围标记
            comment_range_start = OxmlElement('w:commentRangeStart')
            comment_range_start.set(qn('w:id'), str(comment_id))

            comment_range_end = OxmlElement('w:commentRangeEnd')
            comment_range_end.set(qn('w:id'), str(comment_id))

            # 创建批注引用
            comment_reference = OxmlElement('w:r')
            comment_ref_elem = OxmlElement('w:commentReference')
            comment_ref_elem.set(qn('w:id'), str(comment_id))
            comment_reference.append(comment_ref_elem)

            # 在段落中插入标记
            p_elem = run._element.getparent()
            run_index = list(p_elem).index(run._element)

            p_elem.insert(run_index, comment_range_start)
            p_elem.insert(run_index + 2, comment_range_end)
            p_elem.insert(run_index + 3, comment_reference)

            # 添加批注内容到comments.xml
            self._add_comment_to_xml(comment_id, author, comment_text)

            return True

        except Exception as e:
            print(f"    ⚠️ 添加批注失败: {e}")
            return False

    def _add_comment_to_xml(self, comment_id: int, author: str, text: str):
        """将批注内容添加到comments.xml"""
        try:
            if not self._comments_part:
                return

            root = self._comments_part.element

            # 创建批注元素
            comment = OxmlElement('w:comment')
            comment.set(qn('w:id'), str(comment_id))
            comment.set(qn('w:author'), author)
            comment.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            comment.set(qn('w:initials'), author[:2] if len(author) >= 2 else author)

            # 创建段落
            p = OxmlElement('w:p')

            # 添加段落属性（使用批注样式）
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), 'CommentText')
            pPr.append(pStyle)
            p.append(pPr)

            # 添加文本run
            r = OxmlElement('w:r')

            # 添加run属性
            rPr = OxmlElement('w:rPr')
            r.append(rPr)

            # 添加文本
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text
            r.append(t)

            p.append(r)
            comment.append(p)

            # 添加到comments根元素
            root.append(comment)

        except Exception as e:
            print(f"    ⚠️ 写入批注XML失败: {e}")


def iter_all_paragraphs(doc: Document, include_headers_footers: bool = True):
    """遍历正文段落 + 表格单元格段落 + (可选)页眉页脚段落"""
    # 正文段落
    for p in doc.paragraphs:
        yield p

    # 表格（递归）
    def walk_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    for p2 in walk_tables(cell.tables):
                        yield p2

    for p in walk_tables(doc.tables):
        yield p

    if include_headers_footers:
        for sec in doc.sections:
            for p in sec.header.paragraphs:
                yield p
            for p in sec.footer.paragraphs:
                yield p


def replace_and_add_comment_in_paragraph(
        paragraph,
        pattern: str,
        old_value: str,
        new_value: str,
        comment_manager: CommentManager,
        anchor_pattern: Optional[str] = None
) -> bool:
    """
    【核心修改】在段落内定位目标值 → 替换为新值 → 在新值上添加批注

    Args:
        paragraph: 段落对象
        pattern: 目标值的匹配模式
        old_value: 原值（待替换）
        new_value: 新值（替换后的值）
        comment_manager: 批注管理器
        anchor_pattern: 锚点模式

    Returns:
        是否成功替换并添加批注
    """
    runs = list(paragraph.runs)
    if not runs:
        return False

    full = "".join(r.text for r in runs)

    # 如果提供了锚点模式，先验证段落是否匹配锚点
    if anchor_pattern:
        if not re.search(anchor_pattern, full, flags=re.DOTALL | re.IGNORECASE):
            return False

    m = re.search(pattern, full, flags=re.DOTALL)
    if not m:
        return False

    start, end = m.span()
    if start == end:
        return False

    # 计算每个 run 在 full 中的区间
    spans = []
    pos = 0
    for r in runs:
        t = r.text or ""
        spans.append((r, pos, pos + len(t)))
        pos += len(t)

    # 找到覆盖 start/end 的 run
    first_idx = None
    last_idx = None
    for i, (_, s, e) in enumerate(spans):
        if first_idx is None and start < e and end > s:
            first_idx = i
        if start < e and end > s:
            last_idx = i

    if first_idx is None or last_idx is None:
        return False

    try:
        # ===== 步骤1：执行替换 =====

        # 如果匹配跨越多个run，需要合并处理
        if first_idx == last_idx:
            # 单个run内的替换
            run = spans[first_idx][0]
            run_start = spans[first_idx][1]
            local_start = start - run_start
            local_end = end - run_start

            old_text = run.text
            new_text = old_text[:local_start] + new_value + old_text[local_end:]
            run.text = new_text

            # 新值在该run中的位置
            new_value_start = local_start
            new_value_end = local_start + len(new_value)
            target_run = run

        else:
            # 跨多个run的替换（复杂情况）
            # 策略：清空中间run，在第一个run完成替换
            first_run = spans[first_idx][0]
            first_run_start = spans[first_idx][1]
            local_start = start - first_run_start

            # 构建替换后的文本
            prefix = first_run.text[:local_start]

            # 清空中间的run
            for i in range(first_idx + 1, last_idx + 1):
                spans[i][0].text = ""

            # 第一个run设置为：前缀 + 新值
            first_run.text = prefix + new_value
            target_run = first_run

        # ===== 步骤2：在新值上添加批注 =====

        # 批注内容：说明这是修改建议
        comment_text = f"【修改建议】\n原值: {old_value}\n新值: {new_value}\n\n请审核此修改是否正确"

        # 在替换后的run上添加批注
        success = comment_manager.add_comment_to_run(target_run, comment_text)

        if not success:
            print(f"    ⚠️ 替换成功但批注添加失败")
            return False

        return True

    except Exception as e:
        print(f"    ⚠️ 替换或批注失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def replace_and_comment_in_docx(
        doc: Document,
        old_value: str,
        new_value: str,
        comment_manager: CommentManager,
        context: str = "",
        anchor_text: str = ""
) -> Tuple[bool, str]:
    """
    在整个 doc 内为"单条错误"执行替换并添加批注（多策略，优先级优化）

    Args:
        doc: Document 对象
        old_value: 原值
        new_value: 新值
        comment_manager: 批注管理器
        context: 译文上下文
        anchor_text: 替换锚点（优先级最高）

    Returns:
        (是否成功, 策略描述)
    """
    old_value = (old_value or "").strip()
    new_value = (new_value or "").strip()
    if not old_value or not new_value:
        return False, "old/new 缺失"

    # ===== 策略1：替换锚点（最高优先级）=====
    if anchor_text:
        anchor_text = _normalize_spaces(anchor_text)
        anchor_pattern = build_smart_pattern(anchor_text, mode="balanced")
        target_pattern = build_smart_pattern(old_value, mode="strict")

        if anchor_pattern and target_pattern:
            for p in iter_all_paragraphs(doc, include_headers_footers=True):
                if replace_and_add_comment_in_paragraph(
                        p, target_pattern, old_value, new_value,
                        comment_manager, anchor_pattern
                ):
                    return True, f"替换锚点: '{anchor_text[:40]}...'"

    # ===== 策略2：译文上下文 + 严格匹配 =====
    if context:
        context_anchor = extract_anchor_with_target(context, old_value, window=40)
        if context_anchor:
            anchor_pattern = build_smart_pattern(context_anchor, mode="balanced")
            target_pattern = build_smart_pattern(old_value, mode="strict")

            if anchor_pattern and target_pattern:
                for p in iter_all_paragraphs(doc, include_headers_footers=True):
                    if replace_and_add_comment_in_paragraph(
                            p, target_pattern, old_value, new_value,
                            comment_manager, anchor_pattern
                    ):
                        return True, f"上下文精确: '{context_anchor[:40]}...'"

    # ===== 策略3：严格模式全局匹配 =====
    if len(old_value) >= 3 or re.search(r'[,.\-/]', old_value):
        strict_pattern = build_smart_pattern(old_value, mode="strict")
        if strict_pattern:
            for p in iter_all_paragraphs(doc, include_headers_footers=True):
                if replace_and_add_comment_in_paragraph(
                        p, strict_pattern, old_value, new_value, comment_manager
                ):
                    return True, "严格匹配(全局)"

    # ===== 策略4：平衡模式 =====
    balanced_pattern = build_smart_pattern(old_value, mode="balanced")
    if balanced_pattern:
        for p in iter_all_paragraphs(doc, include_headers_footers=True):
            if replace_and_add_comment_in_paragraph(
                    p, balanced_pattern, old_value, new_value, comment_manager
            ):
                return True, "平衡模式"

    # ===== 策略5：宽松模式（最后尝试）=====
    loose_pattern = build_smart_pattern(old_value, mode="loose")
    if loose_pattern:
        for p in iter_all_paragraphs(doc, include_headers_footers=True):
            if replace_and_add_comment_in_paragraph(
                    p, loose_pattern, old_value, new_value, comment_manager
            ):
                return True, "宽松模式"

    return False, "未匹配"


# =========================
# 5) 文件备份功能
# =========================

def ensure_backup_copy(src_docx_path: str) -> str:
    """
    把译文文件复制到 backup/ 下，生成不重复的新副本文件名
    """
    src_docx_path = os.path.abspath(src_docx_path)
    if not os.path.exists(src_docx_path):
        raise FileNotFoundError(f"译文文件不存在: {src_docx_path}")

    base_dir = os.path.dirname(src_docx_path)
    backup_dir = os.path.join(base_dir, BACKUP_DIR_NAME)
    os.makedirs(backup_dir, exist_ok=True)

    src_name = os.path.basename(src_docx_path)
    stem, ext = os.path.splitext(src_name)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    dst_name = f"{stem}_corrected_{timestamp}{ext}"
    dst_path = os.path.join(backup_dir, dst_name)

    shutil.copy2(src_docx_path, dst_path)
    return dst_path


# =========================
# 6) 主流程
# =========================
if __name__ == "__main__":
    # 示例文件路径
    original_path = r"C:\Users\Administrator\Desktop\project\效果\TP251117023，北京中翻译，中译英（字数2w）\原文-B251124195-Y-更新1121-附件1：中国银行股份有限公司模型风险管理办法（2025年修订）.docx"  # 请替换为原文文件路径
    translated_path = r"C:\Users\Administrator\Desktop\project\效果\TP251117023，北京中翻译，中译英（字数2w）\测试译文-清洁版-B251124195-附件1：中国银行股份有限公司模型风险管理政策（2025年修订）-.docx"  # 请替换为译文文件路径
    error_docx_path = r"C:\Users\Administrator\Desktop\project\效果\TP251107025，扬杰科技，中译英（字数3w）\中翻译\文本对比结果.docx"

    #处理页眉
    original_header_text=extract_headers(original_path)
    translated_header_text=extract_headers(translated_path)
    #处理页脚
    original_footer_text = extract_footers(original_path)
    translated_footer_text = extract_footers(translated_path)
    #处理正文(含脚注/表格/自动编号)
    original_body_text=extract_body_text(original_path)
    translated_body_text=extract_body_text(translated_path)
    print("======页眉===========")
    print(original_header_text)
    print(translated_header_text)
    print("======页脚===========")
    print(original_footer_text)
    print(translated_footer_text)
    print("======正文===========")
    print(original_body_text)
    print(translated_body_text)

    # # 实例化对象并进行对比
    matcher = Match()
    #正文对比
    print("======正在检查正文===========")
    body_result = matcher.compare_texts(original_body_text, translated_body_text)
    body_result_name, body_result_path = write_json_with_timestamp(body_result, r"C:\Users\Administrator\Desktop\project\llm\llm_project\zhengwen\output_json")
    #页眉对比
    print("======正在检查页眉===========")
    header_result = matcher.compare_texts(original_header_text, translated_header_text)
    header_result_name, header_result_path = write_json_with_timestamp(header_result, r"C:\Users\Administrator\Desktop\project\llm\llm_project\yemei\output_json")
    #页脚对比
    print("======正在检查页脚===========")
    footer_result = matcher.compare_texts(original_footer_text, translated_footer_text)
    footer_result_name, footer_result_path = write_json_with_timestamp(footer_result, r"C:\Users\Administrator\Desktop\project\llm\llm_project\yejiao\output_json")
    print("================================")

    if not os.path.exists(error_docx_path):
        raise FileNotFoundError(f"错误报告文件不存在: {error_docx_path}")

    # 1) 复制译文到 backup/
    backup_copy_path = ensure_backup_copy(translated_path)
    print(f"✅ 已复制译文副本到: {backup_copy_path}")

    # 2) 读取错误报告并解析
    print("\n正在提取正文错误报告...")
    body_err_doc = Document(body_result_path)
    body_error_text = "\n".join([p.text for p in body_err_doc.paragraphs if p.text])

    print("\n正在提取页眉错误报告...")
    header_err_doc = Document(header_result_path)
    header_error_text = "\n".join([p.text for p in header_err_doc.paragraphs if p.text])

    print("\n正在提取页脚错误报告...")
    footer_err_doc = Document(footer_result_path)
    footer_error_text = "\n".join([p.text for p in footer_err_doc.paragraphs if p.text])

    print("正在解析正文错误列表...")
    body_errors = load_error_list_from_text(body_error_text)
    print(body_errors)
    for err in body_errors:
        print(err)
    print(f"✓ 解析到 {len(body_errors)} 条错误")

    print("正在解析页眉错误列表...")
    header_errors = load_error_list_from_text(header_error_text)
    print(header_errors)
    for err in header_errors:
        print(err)
    print(f"✓ 解析到 {len(header_errors)} 条错误")

    print("正在解析页脚错误列表...")
    footer_errors = load_error_list_from_text(footer_error_text)
    print(footer_errors)
    for err in footer_errors:
        print(err)
    print(f"✓ 解析到 {len(footer_errors)} 条错误")

    # 打印解析结果（调试用）
    print("\n==================== 正文解析结果预览 ====================")
    for i, e in enumerate(body_errors, 1):
        print(f"\n错误 #{i}:")
        for k, v in e.items():
            print(f"  {k}: {v if len(str(v)) > 100 else v}")
    if len(body_errors) > 3:
        print(f"\n... 还有 {len(body_errors) - 3} 条错误")
    print("=" * 60 + "\n")

    print("\n==================== 页眉解析结果预览 ====================")
    for i, e in enumerate(header_errors, 1):
        print(f"\n错误 #{i}:")
        for k, v in e.items():
            print(f"  {k}: {v if len(str(v)) > 100 else v}")
    if len(header_errors) > 3:
        print(f"\n... 还有 {len(header_errors) - 3} 条错误")
    print("=" * 60 + "\n")

    print("\n==================== 页脚解析结果预览 ====================")
    for i, e in enumerate(footer_errors, 1):
        print(f"\n错误 #{i}:")
        for k, v in e.items():
            print(f"  {k}: {v if len(str(v)) > 100 else v}")
    if len(footer_errors) > 3:
        print(f"\n... 还有 {len(footer_errors) - 3} 条错误")
    print("=" * 60 + "\n")

    # 3) 打开副本 docx
    print("正在加载文档...")
    doc = Document(backup_copy_path)

    # 4) 创建批注管理器并初始化
    print("正在初始化批注系统...")
    comment_manager = CommentManager(doc)

    # 【关键】创建初始批注以确保 comments.xml 结构完整
    if comment_manager.create_initial_comment():
        print("✓ 批注系统初始化成功\n")
    else:
        print("⚠️ 批注系统初始化失败，但将继续尝试处理\n")

    # 5) 逐条执行替换并添加批注
    print("==================== 开始处理正文错误 ====================\n")
    body_success_count = 0
    body_fail_count = 0

    for idx, e in enumerate(body_errors, 1):
        err_id = e.get("错误编号", "?")
        err_type = e.get("错误类型", "")
        old = (e.get("译文数值") or "").strip()
        new = (e.get("译文修改建议值") or "").strip()
        trans_context = e.get("译文上下文", "") or ""
        anchor = e.get("替换锚点", "") or ""

        if not old or not new:
            print(f"[{idx}/{len(body_errors)}] [跳过] 错误 #{err_id}: old/new 缺失")
            body_fail_count += 1
            continue

        # 执行替换并添加批注(正文)
        ok, strategy = replace_and_comment_in_docx(
            doc, old, new, comment_manager,
            context=trans_context,
            anchor_text=anchor
        )

        if ok:
            print(f"[{idx}/{len(body_errors)}] [✓成功] 错误 #{err_id} ({err_type})")
            print(f"    策略: {strategy}")
            print(f"    操作: '{old}' → '{new}' (已替换并添加批注)")
            if anchor:
                print(f"    锚点: {anchor}...")
            elif trans_context:
                print(f"    上下文: {trans_context}...")
            body_success_count += 1
        else:
            print(f"[{idx}/{len(body_errors)}] [✗失败] 错误 #{err_id} ({err_type})")
            print(f"    未找到匹配: '{old}'")
            if anchor:
                print(f"    锚点: {anchor}...")
            print(f"    上下文: {trans_context if trans_context else '无'}...")
            body_fail_count += 1
        print()

    print("==================== 开始处理页眉错误 ====================\n")
    header_success_count = 0
    header_fail_count = 0

    for idx, e in enumerate(header_errors, 1):
        err_id = e.get("错误编号", "?")
        err_type = e.get("错误类型", "")
        old = (e.get("译文数值") or "").strip()
        new = (e.get("译文修改建议值") or "").strip()
        trans_context = e.get("译文上下文", "") or ""
        anchor = e.get("替换锚点", "") or ""

        if not old or not new:
            print(f"[{idx}/{len(header_errors)}] [跳过] 错误 #{err_id}: old/new 缺失")
            header_fail_count += 1
            continue

        # 执行替换并添加批注(正文)
        ok, strategy = replace_and_comment_in_docx(
            doc, old, new, comment_manager,
            context=trans_context,
            anchor_text=anchor
        )

        if ok:
            print(f"[{idx}/{len(header_errors)}] [✓成功] 错误 #{err_id} ({err_type})")
            print(f"    策略: {strategy}")
            print(f"    操作: '{old}' → '{new}' (已替换并添加批注)")
            if anchor:
                print(f"    锚点: {anchor}...")
            elif trans_context:
                print(f"    上下文: {trans_context}...")
            header_success_count += 1
        else:
            print(f"[{idx}/{len(header_errors)}] [✗失败] 错误 #{err_id} ({err_type})")
            print(f"    未找到匹配: '{old}'")
            if anchor:
                print(f"    锚点: {anchor}...")
            print(f"    上下文: {trans_context if trans_context else '无'}...")
            header_fail_count += 1
        print()

    print("==================== 开始处理页脚错误 ====================\n")
    footer_success_count = 0
    footer_fail_count = 0

    for idx, e in enumerate(footer_errors, 1):
        err_id = e.get("错误编号", "?")
        err_type = e.get("错误类型", "")
        old = (e.get("译文数值") or "").strip()
        new = (e.get("译文修改建议值") or "").strip()
        trans_context = e.get("译文上下文", "") or ""
        anchor = e.get("替换锚点", "") or ""

        if not old or not new:
            print(f"[{idx}/{len(footer_errors)}] [跳过] 错误 #{err_id}: old/new 缺失")
            footer_fail_count += 1
            continue

        # 执行替换并添加批注(正文)
        ok, strategy = replace_and_comment_in_docx(
            doc, old, new, comment_manager,
            context=trans_context,
            anchor_text=anchor
        )

        if ok:
            print(f"[{idx}/{len(footer_errors)}] [✓成功] 错误 #{err_id} ({err_type})")
            print(f"    策略: {strategy}")
            print(f"    操作: '{old}' → '{new}' (已替换并添加批注)")
            if anchor:
                print(f"    锚点: {anchor}...")
            elif trans_context:
                print(f"    上下文: {trans_context}...")
            footer_success_count += 1
        else:
            print(f"[{idx}/{len(footer_errors)}] [✗失败] 错误 #{err_id} ({err_type})")
            print(f"    未找到匹配: '{old}'")
            if anchor:
                print(f"    锚点: {anchor}...")
            print(f"    上下文: {trans_context if trans_context else '无'}...")
            footer_fail_count += 1
        print()

    # 6) 保存文档
    print("正在保存文档...")

    doc.save(backup_copy_path)

    print(f"\n==================== 正文处理完成 ====================")
    print(f"成功: {body_success_count} | 失败: {body_fail_count} | 总计: {len(body_errors)}")
    if len(body_errors) > 0:
        print(f"成功率: {body_success_count / len(body_errors) * 100:.1f}%")
    print(f"\n✅ 已保存到: {backup_copy_path}")
    print("⚠️ 原始译文文件未被修改")

    print(f"\n==================== 页眉处理完成 ====================")
    print(f"成功: {header_success_count} | 失败: {header_fail_count} | 总计: {len(header_errors)}")
    if len(header_errors) > 0:
        print(f"成功率: {header_success_count / len(header_errors) * 100:.1f}%")
    print(f"\n✅ 已保存到: {backup_copy_path}")
    print("⚠️ 原始译文文件未被修改")

    print(f"\n==================== 页脚处理完成 ====================")
    print(f"成功: {footer_success_count} | 失败: {footer_fail_count} | 总计: {len(footer_errors)}")
    if len(footer_errors) > 0:
        print(f"成功率: {footer_success_count / len(footer_errors) * 100:.1f}%")
    print(f"\n✅ 已保存到: {backup_copy_path}")
    print("⚠️ 原始译文文件未被修改")

    print(f"\n==================== 文章处理完成 ====================")
    count=len(body_errors)+len(header_errors)+len(footer_errors)
    success_count=body_success_count+header_success_count+footer_success_count
    fail_count=body_fail_count+header_fail_count+footer_fail_count
    print(f"成功: {success_count} | 失败: {fail_count} | 总计: {count}")
    if len(footer_errors) > 0:
        print(f"成功率: {success_count / len(footer_errors) * 100:.1f}%")
    print(f"\n✅ 已保存到: {backup_copy_path}")
    print("⚠️ 原始译文文件未被修改")


