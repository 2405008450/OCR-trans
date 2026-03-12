import os
import shutil
from typing import List, Dict, Any, Tuple
from datetime import datetime
from docx import Document
from lxml import etree
from zipfile import ZipFile
import warnings
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
# =========================
# 4) Word批注功能（核心修改 - 替换+批注版）
# =========================
warnings.filterwarnings("ignore")

# =========================
# 0) 基础配置
# =========================

BACKUP_DIR_NAME = "backup"

# =========================
# 1) DOCX 预加载
# =========================

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NAMESPACES = {"w": W_NS}
class CommentManager:
    """Word批注管理器（完整XML支持）"""

    def __init__(self, doc: Document):
        self.doc = doc
        self._comment_id = 0
        self._comments_part = None
        self._init_comments_part()

    def _init_comments_part(self):
        """初始化comments.xml部分（如果不存在则创建）"""
        try:
            package = self.doc.part.package

            # 查找现有的comments part
            for rel in self.doc.part.rels.values():
                if "comments" in rel.target_ref:
                    self._comments_part = rel.target_part
                    break

            if self._comments_part:
                # 解析现有批注ID
                root = self._comments_part.element
                for comment in root.findall(f".//{{{W_NS}}}comment"):
                    cid = comment.get(f"{{{W_NS}}}id")
                    if cid and cid.isdigit():
                        self._comment_id = max(self._comment_id, int(cid))
            else:
                # 创建新的comments part
                self._create_comments_part()

            self._comment_id += 1

        except Exception as e:
            print(f"初始化批注部分失败: {e}")
            self._comment_id = 1

    def _create_comments_part(self):
        """创建comments.xml文件和关系"""
        try:
            from docx.opc.part import XmlPart
            from docx.opc.packuri import PackURI

            # 创建comments.xml的XML结构
            comments_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="{W_NS}" 
            xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
            xmlns:o="urn:schemas-microsoft-com:office:office"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            xmlns:w10="urn:schemas-microsoft-com:office:word"
            xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
            xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
            xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
            xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
            xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
            xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
            mc:Ignorable="w14 w15 wp14">
</w:comments>'''

            # 解析XML
            comments_element = etree.fromstring(comments_xml.encode('utf-8'))

            # 创建Part对象
            partname = PackURI('/word/comments.xml')
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'

            package = self.doc.part.package
            self._comments_part = XmlPart(partname, content_type, comments_element, package)

            # 添加关系
            self.doc.part.relate_to(self._comments_part, RT.COMMENTS)

            print("成功创建 comments.xml")

        except Exception as e:
            print(f"创建comments.xml失败: {e}")
            print("   将使用备用方案（批注可能无法显示）")

    def _get_next_comment_id(self) -> int:
        """获取下一个批注ID"""
        cid = self._comment_id
        self._comment_id += 1
        return cid

    def create_initial_comment(self) -> bool:
        """
        在文档开头创建一个初始化批注（确保comments.xml结构完整）
        """
        try:
            if not self.doc.paragraphs:
                print("⚠️ 文档无段落，无法创建初始批注")
                return False

            first_para = self.doc.paragraphs[0]
            if not first_para.runs:
                first_para.add_run(".")  # 添加一个占位符

            first_run = first_para.runs[0]

            # 获取自动编号批注功能状态
            try:
                from llm.llm_project.replace.numbering_replacer import ENABLE_NUMBERING_COMMENTS
                if ENABLE_NUMBERING_COMMENTS:
                    numbering_status = "✓ 自动编号批注功能已开启，将为每个修改的段落添加批注"
                else:
                    numbering_status = "ℹ 自动编号批注功能已关闭"
            except Exception:
                numbering_status = "⚠ 无法获取自动编号批注功能状态"

            # 创建初始批注
            init_text = f"【翻译校对系统】\n批注功能已启用 - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n{numbering_status}"
            success = self.add_comment_to_run(first_run, init_text, author="系统")

            if success:
                print("已创建初始化批注")

            return success

        except Exception as e:
            print(f"创建初始批注失败: {e}")
            return False

    def append_to_initial_comment(self, additional_text: str) -> bool:
        """
        追加内容到初始系统批注
        
        Args:
            additional_text: 要追加的文本
            
        Returns:
            是否成功
        """
        try:
            if not self._comments_part:
                print("⚠️ 批注部分不存在")
                return False

            root = self._comments_part.element
            
            # 查找第一个批注（ID=1，系统批注）
            first_comment = root.find(f".//{{{W_NS}}}comment[@{{{W_NS}}}id='1']")
            
            if first_comment is None:
                print("未找到初始系统批注")
                return False
            
            # 查找批注中的段落
            p = first_comment.find(f".//{{{W_NS}}}p")
            if p is None:
                print("批注格式异常")
                return False
            
            # 创建新的文本run
            r = OxmlElement('w:r')
            
            # 添加换行
            br = OxmlElement('w:br')
            r.append(br)
            
            # 添加文本
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = additional_text
            r.append(t)
            
            p.append(r)
            
            return True
            
        except Exception as e:
            print(f"追加批注内容失败: {e}")
            return False

    def add_comment_to_run(self, run, comment_text: str, author: str = "翻译校对") -> bool:
        """
        【安全版】批注只包一个纯 run，不做 index 偏移假设
        """
        try:
            if not self._comments_part:
                return False

            comment_id = self._get_next_comment_id()

            # commentRangeStart
            start = OxmlElement("w:commentRangeStart")
            start.set(qn("w:id"), str(comment_id))

            # commentRangeEnd
            end = OxmlElement("w:commentRangeEnd")
            end.set(qn("w:id"), str(comment_id))

            # commentReference
            ref_run = OxmlElement("w:r")
            ref = OxmlElement("w:commentReference")
            ref.set(qn("w:id"), str(comment_id))
            ref_run.append(ref)

            r_elem = run._element
            p_elem = r_elem.getparent()

            # ✅ 关键：只在 run 前后插，不算 index 偏移
            p_elem.insert(p_elem.index(r_elem), start)
            p_elem.insert(p_elem.index(r_elem) + 1, end)
            p_elem.insert(p_elem.index(r_elem) + 2, ref_run)

            self._add_comment_to_xml(comment_id, author, comment_text)
            return True

        except Exception as e:
            print(f"添加批注失败: {e}")
            return False

    def _add_comment_to_xml(self, comment_id: int, author: str, text: str):
        """将批注内容添加到comments.xml"""
        try:
            if not self._comments_part:
                return

            root = self._comments_part.element

            # 创建批注元素
            comment = OxmlElement('w:comment')
            comment.set(qn('w:id'), str(comment_id))
            comment.set(qn('w:author'), author)
            comment.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            comment.set(qn('w:initials'), author[:2] if len(author) >= 2 else author)

            # 创建段落
            p = OxmlElement('w:p')

            # 添加段落属性（使用批注样式）
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), 'CommentText')
            pPr.append(pStyle)
            p.append(pPr)

            # 添加文本run
            r = OxmlElement('w:r')

            # 添加run属性
            rPr = OxmlElement('w:rPr')
            r.append(rPr)

            # 添加文本
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text
            r.append(t)

            p.append(r)
            comment.append(p)

            # 添加到comments根元素
            root.append(comment)

        except Exception as e:
            print(f"    写入批注XML失败: {e}")