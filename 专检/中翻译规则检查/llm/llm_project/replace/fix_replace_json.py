import os
import shutil
from typing import List, Dict, Any, Tuple
from datetime import datetime
from docx import Document
from lxml import etree
from zipfile import ZipFile
import warnings
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import json
from llm.llm_project.llm_check.check import Match
from llm.llm_project.parsers.body_extractor import extract_body_text
from llm.llm_project.parsers.footer_extractor import extract_footers
from llm.llm_project.parsers.header_extractor import extract_headers
from llm.utils.clean_json import clean_markdown_json, load_json_file
from llm.utils.json_files import write_json_with_timestamp

warnings.filterwarnings("ignore")

# =========================
# 0) 基础配置
# =========================

BACKUP_DIR_NAME = "backup"

# =========================
# 1) DOCX 预加载
# =========================

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NAMESPACES = {"w": W_NS}


def get_xml_text(element) -> str:
    """从任意XML元素及其子元素中提取纯文本"""
    texts = []
    for t in element.iter(f"{{{W_NS}}}t"):
        if t.text:
            texts.append(t.text)
    return "".join(texts)



class DocContentLoader:
    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.footnotes = {}
        self.endnotes = {}
        self.comments = {}
        self.headers = []
        self.footers = []
        self._load_all()

    def _load_xml_map(self, zip_file, filename, tag_name, id_attr="id"):
        data_map = {}
        if filename not in zip_file.namelist():
            return data_map

        with zip_file.open(filename) as f:
            tree = etree.parse(f)
            for elem in tree.findall(f".//w:{tag_name}", NAMESPACES):
                eid = elem.get(f"{{{W_NS}}}{id_attr}")
                if eid in ["-1", "0"]:
                    continue
                texts = [t.text for t in elem.iter(f"{{{W_NS}}}t") if t.text]
                full_text = "".join(texts).strip()
                if full_text:
                    data_map[eid] = full_text
        return data_map

    def _load_headers_footers(self, zip_file, prefix):
        contents = []
        files = [f for f in zip_file.namelist() if f.startswith(f"word/{prefix}") and f.endswith(".xml")]
        files.sort(key=lambda x: int("".join(filter(str.isdigit, x)) or 0))
        for f_name in files:
            try:
                with zip_file.open(f_name) as f:
                    tree = etree.parse(f)
                    texts = [t.text for t in tree.iter(f"{{{W_NS}}}t") if t.text]
                    full_text = "".join(texts).strip()
                    if full_text:
                        contents.append(full_text)
            except Exception:
                pass
        return contents

    def _load_all(self):
        with ZipFile(self.doc_path, "r") as zf:
            self.footnotes = self._load_xml_map(zf, "word/footnotes.xml", "footnote")
            self.endnotes = self._load_xml_map(zf, "word/endnotes.xml", "endnote")
            self.comments = self._load_xml_map(zf, "word/comments.xml", "comment")
            self.headers = self._load_headers_footers(zf, "header")
            self.footers = self._load_headers_footers(zf, "footer")


# =========================
# 2) 错误报告解析（增强版）
# =========================

def _normalize_spaces(s: str) -> str:
    """标准化空白字符"""
    return re.sub(r"\s+", " ", (s or "")).strip()


import re
import ast
from typing import Any, Dict, List, Optional


def _unify_labels(raw: str) -> str:
    """
    统一“字段名 + 冒号”的各种断行/空格形态，减少解析歧义。
    目标：把 '译文\\n修改建议值\\n: xxx'、'译 文修改建议值: xxx' 等统一成 '译文修改建议值: xxx'
    """
    raw = raw or ""

    # 先把全局的回车/制表等统一为 \n / 空格
    raw = raw.replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")

    # 这些字段名最关键：允许中间任意空白（含换行）
    label_patterns = [
        "错误编号",
        "错误类型",
        "原文数值",
        "译文数值",
        "译文修改建议值",
        "修改理由",
        "违反的规则",
        "原文上下文",
        "译文上下文",
        "原文位置",
        "译文位置",
        "替换锚点",
    ]

    # 把“字段名”内部的空白压平，并让冒号紧跟字段名
    for lab in label_patterns:
        # 例如：译文\s*修改\s*建议\s*值\s*[:：]  ->  译文修改建议值:
        chars = list(lab)
        pat = r"".join(re.escape(c) + r"\s*" for c in chars) + r"[:：]?"
        raw = re.sub(pat, lab, raw, flags=re.IGNORECASE)

        # 统一冒号：字段名后如果紧跟的不是冒号，就补齐（仅当后面确实有内容时）
        raw = re.sub(rf"({re.escape(lab)})\s*\n?\s*[:：]\s*", rf"\1: ", raw)

    # 修复你原来对“错误编号”做的特殊替换（保留）
    raw = raw.replace("错误\n编号", "错误编号")
    raw = raw.replace("错误\n编号:", "错误编号:")
    raw = raw.replace("编号:", "错误编号:")

    return raw

def _split_merged_fields(one: Dict[str, Any]) -> Dict[str, Any]:
    """
    兜底：若译文数值里吞了“译文修改建议值: ...”，则拆分并回填两个字段。
    允许出现：'译 文修改建议值'、换行、多个空格等。
    """
    tv = (one.get("译文数值") or "").strip()
    sv = (one.get("译文修改建议值") or "").strip()

    if not tv:
        return one

    # 允许各种空白：译\s*文\s*修\s*改\s*建\s*议\s*值
    marker_pat = r"(译\s*文\s*修\s*改\s*建\s*议\s*值)\s*[:：]\s*"
    m = re.search(marker_pat, tv, flags=re.IGNORECASE)
    if not m:
        return one

    left = tv[:m.start()].strip()
    right = tv[m.end():].strip()

    if left:
        one["译文数值"] = re.sub(r"\s+", " ", left).strip()

    # 如果原本建议值为空，则用拆出来的补上；如果不为空，也可以选择保留原值（更安全）
    if right and not sv:
        one["译文修改建议值"] = re.sub(r"\s+", " ", right).strip()

    return one


import re
import ast
from typing import Any, Dict, List, Optional


def _extract_first_braced_block(raw: str) -> Optional[str]:
    """
    从 raw 中提取第一个完整的 {...} 或 [...] 字面量块（后面允许跟说明文字）。
    用于兼容你给的那种 dict/list 结构化输入。
    """
    if not raw:
        return None

    m = re.search(r"[\{\[]", raw)
    if not m:
        return None

    start = m.start()
    open_ch = raw[start]
    close_ch = "}" if open_ch == "{" else "]"

    depth = 0
    in_str = False
    str_ch = ""
    escape = False

    for i in range(start, len(raw)):
        ch = raw[i]

        if in_str:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == str_ch:
                in_str = False
            continue

        if ch in ("'", '"'):
            in_str = True
            str_ch = ch
            continue

        if ch == open_ch:
            depth += 1
        elif ch == close_ch:
            depth -= 1
            if depth == 0:
                return raw[start:i + 1]

    return None


def _unify_labels_keep_colon(raw: str) -> str:
    """
    关键：把字段名里的各种空格/换行统一掉，并且【强制保留/补齐冒号】为 '字段名:'。
    解决 Word 自动换行：'错误\\n类型:'、'译 文修改建议值\\n:' 等。
    """
    raw = raw or ""
    raw = raw.replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")

    labels = [
        "错误编号",
        "错误类型",
        "原文数值",
        "译文数值",
        "译文修改建议值",
        "修改理由",
        "违反的规则",
        "原文上下文",
        "译文上下文",
        "原文位置",
        "译文位置",
        "替换锚点",
    ]

    for lab in labels:
        pat = (
                r"(?m)"  # 多行模式
                r"(^|\n)"  # 行首
                + r"\s*"
                + r"".join(re.escape(c) + r"\s*" for c in lab)
                + r"\s*[:：]?"
        )
        raw = re.sub(pat, r"\1" + lab + ":", raw, flags=re.IGNORECASE)

    # 冒号后统一一个空格（含中文冒号）
    raw = re.sub(r"[:：]\s*", ": ", raw)
    return raw


def _split_merged_translation_fields(one: Dict[str, Any]) -> Dict[str, Any]:
    """
    兜底：如果译文数值吞了“译文修改建议值: ...”，则拆开两个字段。
    _unify_labels_keep_colon 会把各种形态统一成 '译文修改建议值:'，因此这里拆分很稳。
    """
    tv = (one.get("译文数值") or "").strip()
    sv = (one.get("译文修改建议值") or "").strip()
    if not tv:
        return one

    marker = "译文修改建议值:"
    idx = tv.find(marker)
    if idx == -1:
        return one

    left = tv[:idx].strip()
    right = tv[idx + len(marker):].strip()

    if left:
        one["译文数值"] = re.sub(r"\s+", " ", left).strip()

    # 如果原本建议值为空，用 right 补齐（更安全；不覆盖已有值）
    if right and not sv:
        one["译文修改建议值"] = re.sub(r"\s+", " ", right).strip()

    return one


def load_error_list_from_text(raw: str) -> List[Dict[str, Any]]:
    """
    解析错误报告（增强版）：
    1) 优先解析结构化 dict/list（你的示例那种）
    2) 再解析文本型报告（支持字段名断行/空格、同一行字段、冒号断行）
    3) 兜底拆分：译文数值吞掉译文修改建议值
    """
    raw = raw or ""

    # ========== 1) 结构化优先：提取第一个 {...} / [...] 后 literal_eval ==========
    block = _extract_first_braced_block(raw)
    if block:
        try:
            obj = ast.literal_eval(block)
            if isinstance(obj, dict):
                obj = [obj]
            if isinstance(obj, list):
                fixed = []
                for it in obj:
                    if isinstance(it, dict):
                        fixed.append(_split_merged_translation_fields(it))
                if fixed:
                    return fixed
        except Exception:
            pass  # 失败就走文本解析

    # ========== 2) 文本解析：统一标签（保留/补齐冒号） ==========
    cleaned = _unify_labels_keep_colon(raw)

    # 每条错误从 “错误编号: 数字” 开始
    parts = re.split(r"(?=错误编号:\s*\d+)", cleaned)
    errors: List[Dict[str, Any]] = []

    # 关键：边界不要依赖换行，直接用 “下一个字段名:” 做 lookahead
    key_map = [
        ("错误编号", r"错误编号:\s*(.+?)(?=错误类型:|$)"),
        ("错误类型", r"错误类型:\s*(.+?)(?=原文数值:|$)"),
        ("原文数值", r"原文数值:\s*(.+?)(?=译文数值:|$)"),
        ("译文数值", r"译文数值:\s*(.+?)(?=译文修改建议值:|修改理由:|$)"),
        ("译文修改建议值", r"译文修改建议值:\s*(.+?)(?=修改理由:|$)"),
        ("修改理由", r"修改理由:\s*(.+?)(?=违反的规则:|$)"),
        ("违反的规则", r"违反的规则:\s*(.+?)(?=原文上下文:|$)"),
        ("原文上下文", r"原文上下文:\s*(.+?)(?=译文上下文:|$)"),
        ("译文上下文", r"译文上下文:\s*(.+?)(?=原文位置:|$)"),
        ("原文位置", r"原文位置:\s*(.+?)(?=译文位置:|$)"),
        ("译文位置", r"译文位置:\s*(.+?)(?=替换锚点:|$)"),
        ("替换锚点", r"替换锚点:\s*(.+?)(?=错误编号:|$)"),
    ]

    for p in parts:
        if "错误编号:" not in p:
            continue

        one: Dict[str, Any] = {}
        for k, pat in key_map:
            mm = re.search(pat, p, flags=re.DOTALL | re.IGNORECASE)
            if mm:
                v = mm.group(1).strip()
                v = re.sub(r"\s+", " ", v).strip()
                one[k] = v

        # 兜底拆分合并字段
        one = _split_merged_translation_fields(one)

        if one.get("错误编号") or one.get("译文数值") or one.get("译文修改建议值"):
            errors.append(one)

    if not errors:
        print("未在错误报告文本中找到可解析的错误列表")

    return errors




# =========================
# 3) 改进的匹配策略
# =========================

def build_smart_pattern(s: str, mode: str = "balanced") -> str:
    """
    构建智能匹配模式

    Args:
        s: 待匹配字符串
        mode: 匹配模式
            - "strict": 严格匹配（完全精确）
            - "balanced": 平衡模式（数字/标点保持连续，单词间允许空格）
            - "loose": 宽松模式（字符间允许空格，但数字连续）
    """
    s = (s or "").strip()
    if not s:
        return ""

    if mode == "strict":
        # 严格模式：完全精确匹配（转义特殊字符）
        return re.escape(s)

    elif mode == "balanced":
        # 平衡模式：数字和标点保持连续，单词间允许空格
        pieces = []
        i = 0
        while i < len(s):
            ch = s[i]

            # 跳过空格
            if ch.isspace():
                if pieces and not pieces[-1].endswith(r"\s*"):
                    pieces.append(r"\s*")
                i += 1
                continue

            # 数字序列：保持连续（包括小数点、逗号）
            if ch.isdigit():
                num_str = ""
                while i < len(s) and (s[i].isdigit() or s[i] in ".,"):
                    num_str += s[i]
                    i += 1
                pieces.append(re.escape(num_str))
                continue

            # 标点符号：保持连续
            if ch in ".,;:!?()[]{}\"'-/":
                pieces.append(re.escape(ch))
                i += 1
                continue

            # 字母序列：单词间允许空格
            if ch.isalpha():
                word = ""
                while i < len(s) and s[i].isalpha():
                    word += s[i]
                    i += 1
                pieces.append(re.escape(word))
                if i < len(s) and not s[i].isspace():
                    continue
                if i < len(s):
                    pieces.append(r"\s*")
                continue

            # 其他字符
            pieces.append(re.escape(ch))
            i += 1

        return "".join(pieces).strip()

    else:  # loose
        # 宽松模式：字符间允许空格，但数字保持连续
        pieces = []
        i = 0
        while i < len(s):
            ch = s[i]

            if ch.isspace():
                i += 1
                continue

            # 数字序列保持连续
            if ch.isdigit():
                num_str = ""
                while i < len(s) and (s[i].isdigit() or s[i] in ".,"):
                    num_str += s[i]
                    i += 1
                pieces.append(re.escape(num_str) + r"\s*")
                continue

            # 其他字符间允许空格
            pieces.append(re.escape(ch) + r"\s*")
            i += 1

        return "".join(pieces).strip()


def extract_anchor_with_target(context: str, target_value: str, window: int = 30) -> Optional[str]:
    """
    从上下文中提取包含目标数值的锚点短语（更大窗口）
    """
    if not context or not target_value:
        return None

    context = _normalize_spaces(context)
    target_value = target_value.strip()

    # 先尝试严格匹配
    if target_value in context:
        idx = context.index(target_value)
        start = max(0, idx - window)
        end = min(len(context), idx + len(target_value) + window)
        anchor = context[start:end]
        # 修剪边界单词
        anchor = re.sub(r"^\S*\s+", "", anchor)
        anchor = re.sub(r"\s+\S*$", "", anchor)
        return anchor.strip()

    # 再尝试平衡模式匹配
    pattern = build_smart_pattern(target_value, mode="balanced")
    if not pattern:
        return None

    match = re.search(pattern, context, flags=re.IGNORECASE)
    if not match:
        return None

    start, end = match.span()
    prefix_start = max(0, start - window)
    suffix_end = min(len(context), end + window)

    anchor = context[prefix_start:suffix_end]
    anchor = re.sub(r"^\S*\s+", "", anchor)
    anchor = re.sub(r"\s+\S*$", "", anchor)

    return anchor.strip() if anchor.strip() else None


# =========================
# 4) Word批注功能（核心修改 - 替换+批注版）
# =========================

class CommentManager:
    """Word批注管理器（完整XML支持）"""

    def __init__(self, doc: Document):
        self.doc = doc
        self._comment_id = 0
        self._comments_part = None
        self._init_comments_part()

    def _init_comments_part(self):
        """初始化comments.xml部分（如果不存在则创建）"""
        try:
            package = self.doc.part.package

            # 查找现有的comments part
            for rel in self.doc.part.rels.values():
                if "comments" in rel.target_ref:
                    self._comments_part = rel.target_part
                    break

            if self._comments_part:
                # 解析现有批注ID
                root = self._comments_part.element
                for comment in root.findall(f".//{{{W_NS}}}comment"):
                    cid = comment.get(f"{{{W_NS}}}id")
                    if cid and cid.isdigit():
                        self._comment_id = max(self._comment_id, int(cid))
            else:
                # 创建新的comments part
                self._create_comments_part()

            self._comment_id += 1

        except Exception as e:
            print(f"⚠️ 初始化批注部分失败: {e}")
            self._comment_id = 1

    def _create_comments_part(self):
        """创建comments.xml文件和关系"""
        try:
            from docx.opc.part import XmlPart
            from docx.opc.packuri import PackURI

            # 创建comments.xml的XML结构
            comments_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="{W_NS}" 
            xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
            xmlns:o="urn:schemas-microsoft-com:office:office"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            xmlns:w10="urn:schemas-microsoft-com:office:word"
            xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
            xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
            xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
            xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
            xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
            xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
            mc:Ignorable="w14 w15 wp14">
</w:comments>'''

            # 解析XML
            comments_element = etree.fromstring(comments_xml.encode('utf-8'))

            # 创建Part对象
            partname = PackURI('/word/comments.xml')
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'

            package = self.doc.part.package
            self._comments_part = XmlPart(partname, content_type, comments_element, package)

            # 添加关系
            self.doc.part.relate_to(self._comments_part, RT.COMMENTS)

            print("✓ 成功创建 comments.xml")

        except Exception as e:
            print(f"⚠️ 创建comments.xml失败: {e}")
            print("   将使用备用方案（批注可能无法显示）")

    def _get_next_comment_id(self) -> int:
        """获取下一个批注ID"""
        cid = self._comment_id
        self._comment_id += 1
        return cid

    def create_initial_comment(self) -> bool:
        """
        在文档开头创建一个初始化批注（确保comments.xml结构完整）
        """
        try:
            if not self.doc.paragraphs:
                print("⚠️ 文档无段落，无法创建初始批注")
                return False

            first_para = self.doc.paragraphs[0]
            if not first_para.runs:
                first_para.add_run(".")  # 添加一个占位符

            first_run = first_para.runs[0]

            # 创建初始批注
            init_text = f"【翻译校对系统】批注功能已启用 - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            success = self.add_comment_to_run(first_run, init_text, author="系统")

            if success:
                print("✓ 已创建初始化批注")

            return success

        except Exception as e:
            print(f"⚠️ 创建初始批注失败: {e}")
            return False

    def add_comment_to_run(self, run, comment_text: str, author: str = "翻译校对") -> bool:
        """
        【安全版】批注只包一个纯 run，不做 index 偏移假设
        """
        try:
            if not self._comments_part:
                return False

            comment_id = self._get_next_comment_id()

            # commentRangeStart
            start = OxmlElement("w:commentRangeStart")
            start.set(qn("w:id"), str(comment_id))

            # commentRangeEnd
            end = OxmlElement("w:commentRangeEnd")
            end.set(qn("w:id"), str(comment_id))

            # commentReference
            ref_run = OxmlElement("w:r")
            ref = OxmlElement("w:commentReference")
            ref.set(qn("w:id"), str(comment_id))
            ref_run.append(ref)

            r_elem = run._element
            p_elem = r_elem.getparent()

            # ✅ 关键：只在 run 前后插，不算 index 偏移
            p_elem.insert(p_elem.index(r_elem), start)
            p_elem.insert(p_elem.index(r_elem) + 1, end)
            p_elem.insert(p_elem.index(r_elem) + 2, ref_run)

            self._add_comment_to_xml(comment_id, author, comment_text)
            return True

        except Exception as e:
            print(f"⚠️ 添加批注失败: {e}")
            return False

    def _add_comment_to_xml(self, comment_id: int, author: str, text: str):
        """将批注内容添加到comments.xml"""
        try:
            if not self._comments_part:
                return

            root = self._comments_part.element

            # 创建批注元素
            comment = OxmlElement('w:comment')
            comment.set(qn('w:id'), str(comment_id))
            comment.set(qn('w:author'), author)
            comment.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            comment.set(qn('w:initials'), author[:2] if len(author) >= 2 else author)

            # 创建段落
            p = OxmlElement('w:p')

            # 添加段落属性（使用批注样式）
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), 'CommentText')
            pPr.append(pStyle)
            p.append(pPr)

            # 添加文本run
            r = OxmlElement('w:r')

            # 添加run属性
            rPr = OxmlElement('w:rPr')
            r.append(rPr)

            # 添加文本
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text
            r.append(t)

            p.append(r)
            comment.append(p)

            # 添加到comments根元素
            root.append(comment)

        except Exception as e:
            print(f"    ⚠️ 写入批注XML失败: {e}")


def iter_all_paragraphs(doc: Document, include_headers_footers: bool = True):
    """遍历正文段落 + 表格单元格段落 + (可选)页眉页脚段落"""
    # 正文段落
    for p in doc.paragraphs:
        yield p

    # 表格（递归）
    def walk_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    for p2 in walk_tables(cell.tables):
                        yield p2

    for p in walk_tables(doc.tables):
        yield p

    if include_headers_footers:
        for sec in doc.sections:
            for p in sec.header.paragraphs:
                yield p
            for p in sec.footer.paragraphs:
                yield p


def replace_and_add_comment_in_paragraph(
        paragraph,
        pattern: str,
        old_value: str,
        new_value: str,
        reason: str,
        comment_manager,
        anchor_pattern: str = None
) -> bool:
    """
    【100% 原位替换版】
    - 只替换命中内容
    - 不改变句子其余部分
    - 不改变 run 顺序
    - 批注只作用于 new_value
    """

    runs = list(paragraph.runs)
    if not runs:
        return False

    full_text = "".join(r.text or "" for r in runs)

    # 锚点校验（不通过直接跳过）
    if anchor_pattern:
        if not re.search(anchor_pattern, full_text, flags=re.IGNORECASE | re.DOTALL):
            return False

    m = re.search(pattern, full_text, flags=re.IGNORECASE | re.DOTALL)
    if not m:
        return False

    start, end = m.span()
    if start == end:
        return False

    # ===== 1. 建立 run → 全文位置映射 =====
    spans = []
    cursor = 0
    for r in runs:
        text = r.text or ""
        spans.append((r, cursor, cursor + len(text)))
        cursor += len(text)

    # 找到命中的 run 区间
    hit = [(r, s, e) for r, s, e in spans if start < e and end > s]
    if not hit:
        return False

    first_run, fs, fe = hit[0]
    last_run, ls, le = hit[-1]

    # ===== 2. 计算前后文本 =====
    prefix = first_run.text[: max(0, start - fs)]
    suffix = last_run.text[max(0, end - ls):]

    # ===== 3. 清空命中 run =====
    for r, _, _ in hit:
        r.text = ""

    # ===== 4. 重建 run（关键：在原位置插入）=====
    parent = first_run._element.getparent()
    insert_pos = parent.index(first_run._element)

    # prefix
    if prefix:
        prefix_run = paragraph.add_run(prefix)
        parent.remove(prefix_run._element)
        parent.insert(insert_pos, prefix_run._element)
        insert_pos += 1

    # new value（核心替换）
    new_run = paragraph.add_run(new_value)
    parent.remove(new_run._element)
    parent.insert(insert_pos, new_run._element)
    insert_pos += 1

    # 复制样式（保持原格式）
    new_run.bold = first_run.bold
    new_run.italic = first_run.italic
    new_run.font.name = first_run.font.name
    new_run.font.size = first_run.font.size

    # suffix
    if suffix:
        suffix_run = paragraph.add_run(suffix)
        parent.remove(suffix_run._element)
        parent.insert(insert_pos, suffix_run._element)

    # ===== 5. 批注只加在 new_run 上 =====
    comment_text = (
        f"【修改建议】\n"
        f"原值: {old_value}\n"
        f"新值: {new_value}\n"
        f"修改理由: {reason}"
    )

    return comment_manager.add_comment_to_run(new_run, comment_text)


def replace_and_comment_in_docx(
        doc: Document,
        old_value: str,
        new_value: str,
        reason: str,
        comment_manager: CommentManager,
        context: str = "",
        anchor_text: str = ""
) -> Tuple[bool, str]:
    """
    多策略执行替换并添加批注
    """
    old_value = (old_value or "").strip()
    new_value = (new_value or "").strip()
    # 这里的 reason 可能是元组 (来自 main.py)，做一层强转处理
    if isinstance(reason, (list, tuple)):
        reason = " ".join([str(i) for i in reason if i]).strip()
    reason = reason or "数值/术语不一致"

    if not old_value or not new_value:
        return False, "数据缺失"

    # ===== 策略1：替换锚点（最高优先级）=====
    if anchor_text:
        anchor_pattern = build_smart_pattern(anchor_text, mode="balanced")
        target_pattern = build_smart_pattern(old_value, mode="strict")
        if anchor_pattern and target_pattern:
            for p in iter_all_paragraphs(doc):
                if replace_and_add_comment_in_paragraph(
                        p, target_pattern, old_value, new_value, reason, comment_manager, anchor_pattern
                ):
                    return True, f"锚点匹配"

    # ===== 策略2：上下文锚点 =====
    if context:
        context_anchor = extract_anchor_with_target(context, old_value, window=40)
        if context_anchor:
            anchor_pattern = build_smart_pattern(context_anchor, mode="balanced")
            target_pattern = build_smart_pattern(old_value, mode="strict")
            for p in iter_all_paragraphs(doc):
                if replace_and_add_comment_in_paragraph(
                        p, target_pattern, old_value, new_value, reason, comment_manager, anchor_pattern
                ):
                    return True, f"上下文匹配"

    # ===== 策略3：严格模式全局匹配 =====
    strict_pattern = build_smart_pattern(old_value, mode="strict")
    for p in iter_all_paragraphs(doc):
        if replace_and_add_comment_in_paragraph(
                p, strict_pattern, old_value, new_value, reason, comment_manager
        ):
            return True, "严格模式"

    # ===== 策略4：平衡模式 =====
    balanced_pattern = build_smart_pattern(old_value, mode="balanced")
    for p in iter_all_paragraphs(doc):
        if replace_and_add_comment_in_paragraph(
                p, balanced_pattern, old_value, new_value, reason, comment_manager
        ):
            return True, "平衡模式"

    return False, "未找到匹配项"


# =========================
# 5) 文件备份功能
# =========================

def ensure_backup_copy(src_docx_path: str) -> str:
    """
    把译文文件复制到 backup/ 下，生成不重复的新副本文件名
    """
    src_docx_path = os.path.abspath(src_docx_path)
    if not os.path.exists(src_docx_path):
        raise FileNotFoundError(f"译文文件不存在: {src_docx_path}")

    base_dir = os.path.dirname(src_docx_path)
    backup_dir = os.path.join(base_dir, BACKUP_DIR_NAME)
    os.makedirs(backup_dir, exist_ok=True)

    src_name = os.path.basename(src_docx_path)
    stem, ext = os.path.splitext(src_name)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    dst_name = f"{stem}_corrected_{timestamp}{ext}"
    dst_path = os.path.join(backup_dir, dst_name)

    shutil.copy2(src_docx_path, dst_path)
    return dst_path


# =========================
# 6) 主流程
# =========================
if __name__ == "__main__":
    # 示例文件路径
    original_path = r"C:\Users\Administrator\Desktop\project\llm\文本对比结果\加粗斜体测试\原文-加粗斜体.docx"  # 请替换为原文文件路径
    translated_path = r"C:\Users\Administrator\Desktop\project\llm\文本对比结果\加粗斜体测试\译文-加粗斜体.docx"  # 请替换为译文文件路径
    #处理页眉
    original_header_text=extract_headers(original_path)
    translated_header_text=extract_headers(translated_path)
    #处理页脚
    original_footer_text = extract_footers(original_path)
    translated_footer_text = extract_footers(translated_path)
    #处理正文(含脚注/表格/自动编号)
    original_body_text=extract_body_text(original_path)
    translated_body_text=extract_body_text(translated_path)
    print("======页眉===========")
    print(original_header_text)
    print(translated_header_text)
    print("======页脚===========")
    print(original_footer_text)
    print(translated_footer_text)
    print("======正文===========")
    print(original_body_text)
    print(translated_body_text)

    #实例化对象并进行对比
    matcher = Match()
    #正文对比
    print("======正在检查正文===========")
    if original_body_text and translated_body_text:
        # 两个值都不为空，正常执行比较
        body_result = matcher.compare_texts(original_body_text, translated_body_text)
    else:
        # 任意一个为空，生成空结果
        body_result = {}  # 或者 body_result = []，根据你的 write_json_with_timestamp 函数期望的格式
        print("原文或译文为空，检查结果为空")

    body_result_name, body_result_path = write_json_with_timestamp(
        body_result,
        r"C:\Users\Administrator\Desktop\project\llm\llm_project\zhengwen\output_json"
    )
    # body_result = matcher.compare_texts(original_body_text, translated_body_text)
    # body_result_name, body_result_path = write_json_with_timestamp(body_result,r"C:\Users\Administrator\Desktop\project\llm\llm_project\zhengwen\output_json")
    # #页眉对比
    print("======正在检查页眉===========")
    if original_header_text and translated_header_text:
        # 两个值都不为空，正常执行比较
        header_result = matcher.compare_texts(original_header_text, translated_header_text)
    else:
        # 任意一个为空，生成空结果
        header_result = {}
        print("原文或译文为空，检查结果为空")

    header_result_name, header_result_path = write_json_with_timestamp(
        header_result,
        r"C:\Users\Administrator\Desktop\project\llm\llm_project\yemei\output_json"
    )
    # header_result = matcher.compare_texts(original_header_text, translated_header_text)
    # header_result_name, header_result_path = write_json_with_timestamp(header_result, r"C:\Users\Administrator\Desktop\project\llm\llm_project\yemei\output_json")
    # #页脚对比
    print("======正在检查页脚===========")
    if original_footer_text and translated_footer_text:
        # 两个值都不为空，正常执行比较
        footer_result = matcher.compare_texts(original_footer_text, translated_footer_text)
    else:
        # 任意一个为空，生成空结果
        footer_result = {}
        print("原文或译文为空，检查结果为空")

    footer_result_name, footer_result_path = write_json_with_timestamp(
        footer_result,
        r"C:\Users\Administrator\Desktop\project\llm\llm_project\yejiao\output_json"
    )
    # footer_result = matcher.compare_texts(original_footer_text, translated_footer_text)
    # footer_result_name, footer_result_path = write_json_with_timestamp(footer_result, r"C:\Users\Administrator\Desktop\project\llm\llm_project\yejiao\output_json")

    print("================================")
    # if not os.path.exists(error_docx_path):
    #     raise FileNotFoundError(f"错误报告文件不存在: {error_docx_path}")
    # body_result_path=r"C:\Users\Administrator\Desktop\project\llm\文本对比结果\加粗斜体测试\AI检查结果.json"
    # header_result_path=r"C:\Users\Administrator\Desktop\project\llm\llm_project\yemei\output_json\文本对比结果_20260208_144950.json"
    # footer_result_path=r"C:\Users\Administrator\Desktop\project\llm\llm_project\yejiao\output_json\文本对比结果_20260208_145004.json"

    # 1) 复制译文到 backup/
    backup_copy_path = ensure_backup_copy(translated_path)
    print(f"✅ 已复制译文副本到: {backup_copy_path}")

    # 2) 读取错误报告并解析
    print("\n正在提取解析正文错误报告...")
    body_errors = load_json_file(body_result_path)
    print("正文错误报告",body_errors)
    for err in body_errors:
        print(err)
    print("正文错误解析个数：",len(body_errors))

    print("\n正在提取解析页眉错误报告...")
    header_errors = load_json_file(header_result_path)
    print("页眉错误报告", header_errors)
    print("页眉错误解析个数：", len(header_errors))

    print("\n正在提取解析页脚错误报告...")
    footer_errors = load_json_file(footer_result_path)
    print("页脚错误报告", footer_errors)
    print("页脚错误解析个数：", len(footer_errors))


    # 3) 打开副本 docx
    print("正在加载文档...")
    doc = Document(backup_copy_path)

    # 4) 创建批注管理器并初始化
    print("正在初始化批注系统...")
    comment_manager = CommentManager(doc)

    # 【关键】创建初始批注以确保 comments.xml 结构完整
    if comment_manager.create_initial_comment():
        print("✓ 批注系统初始化成功\n")
    else:
        print("⚠️ 批注系统初始化失败，但将继续尝试处理\n")

    # 5) 逐条执行替换并添加批注
    print("==================== 开始处理正文错误 ====================\n")
    body_success_count = 0
    body_fail_count = 0

    for idx, e in enumerate(body_errors, 1):
        err_id = e.get("错误编号", "?")
        err_type = e.get("错误类型", "")
        old = (e.get("译文数值") or "").strip()
        new = (e.get("译文修改建议值") or "").strip()
        reason = (e.get("修改理由"),"")
        trans_context = e.get("译文上下文", "") or ""
        anchor = e.get("替换锚点", "") or ""

        if not old or not new:
            print(f"[{idx}/{len(body_errors)}] [跳过] 错误 #{err_id}: old/new 缺失")
            body_fail_count += 1
            continue

        # 执行替换并添加批注(正文)
        ok, strategy = replace_and_comment_in_docx(
            doc, old, new, reason,comment_manager,
            context=trans_context,
            anchor_text=anchor
        )

        if ok:
            print(f"[{idx}/{len(body_errors)}] [✓成功] 错误 #{err_id} ({err_type})")
            print(f"    策略: {strategy}")
            print(f"    修改理由: {reason}")
            print(f"    操作: '{old}' → '{new}' (已替换并添加批注)")
            if anchor:
                print(f"    锚点: {anchor}...")
            elif trans_context:
                print(f"    上下文: {trans_context}...")
            body_success_count += 1
        else:
            print(f"[{idx}/{len(body_errors)}] [✗失败] 错误 #{err_id} ({err_type})")
            print(f"    未找到匹配: '{old}'")
            if anchor:
                print(f"    锚点: {anchor}...")
            print(f"    上下文: {trans_context if trans_context else '无'}...")
            body_fail_count += 1
        print()

    print(f"\n==================== 正文处理完成 ====================")
    print(f"成功: {body_success_count} | 失败: {body_fail_count} | 总计: {len(body_errors)}")
    if len(body_errors) > 0:
        print(f"成功率: {body_success_count / len(body_errors) * 100:.1f}%")
    print(f"\n✅ 已保存到: {backup_copy_path}")
    print("⚠️ 原始译文文件未被修改")

    print("==================== 开始处理页眉错误 ====================\n")
    header_success_count = 0
    header_fail_count = 0

    for idx, e in enumerate(header_errors, 1):
        err_id = e.get("错误编号", "?")
        err_type = e.get("错误类型", "")
        old = (e.get("译文数值") or "").strip()
        new = (e.get("译文修改建议值") or "").strip()
        reason = (e.get("修改理由"), "")
        trans_context = e.get("译文上下文", "") or ""
        anchor = e.get("替换锚点", "") or ""

        if not old or not new:
            print(f"[{idx}/{len(header_errors)}] [跳过] 错误 #{err_id}: old/new 缺失")
            header_fail_count += 1
            continue

        # 执行替换并添加批注(正文)
        ok, strategy = replace_and_comment_in_docx(
            doc, old, new, reason,comment_manager,
            context=trans_context,
            anchor_text=anchor
        )

        if ok:
            print(f"[{idx}/{len(header_errors)}] [✓成功] 错误 #{err_id} ({err_type})")
            print(f"    修改理由: {reason}")
            print(f"    策略: {strategy}")
            print(f"    操作: '{old}' → '{new}' (已替换并添加批注)")
            if anchor:
                print(f"    锚点: {anchor}...")
            elif trans_context:
                print(f"    上下文: {trans_context}...")
            header_success_count += 1
        else:
            print(f"[{idx}/{len(header_errors)}] [✗失败] 错误 #{err_id} ({err_type})")
            print(f"    未找到匹配: '{old}'")
            if anchor:
                print(f"    锚点: {anchor}...")
            print(f"    上下文: {trans_context if trans_context else '无'}...")
            header_fail_count += 1
        print()
    print(f"\n==================== 页眉处理完成 ====================")
    print(f"成功: {header_success_count} | 失败: {header_fail_count} | 总计: {len(header_errors)}")
    if len(header_errors) > 0:
        print(f"成功率: {header_success_count / len(header_errors) * 100:.1f}%")
    print(f"\n✅ 已保存到: {backup_copy_path}")
    print("⚠️ 原始译文文件未被修改")

    print("==================== 开始处理页脚错误 ====================\n")
    footer_success_count = 0
    footer_fail_count = 0

    for idx, e in enumerate(footer_errors, 1):
        err_id = e.get("错误编号", "?")
        err_type = e.get("错误类型", "")
        old = (e.get("译文数值") or "").strip()
        new = (e.get("译文修改建议值") or "").strip()
        reason = (e.get("修改理由"), "")
        trans_context = e.get("译文上下文", "") or ""
        anchor = e.get("替换锚点", "") or ""

        if not old or not new:
            print(f"[{idx}/{len(footer_errors)}] [跳过] 错误 #{err_id}: old/new 缺失")
            footer_fail_count += 1
            continue

        # 执行替换并添加批注(正文)
        ok, strategy = replace_and_comment_in_docx(
            doc, old, new, reason,comment_manager,
            context=trans_context,
            anchor_text=anchor
        )

        if ok:
            print(f"[{idx}/{len(footer_errors)}] [✓成功] 错误 #{err_id} ({err_type})")
            print(f"    修改理由: {reason}")
            print(f"    策略: {strategy}")
            print(f"    操作: '{old}' → '{new}' (已替换并添加批注)")
            if anchor:
                print(f"    锚点: {anchor}...")
            elif trans_context:
                print(f"    上下文: {trans_context}...")
            footer_success_count += 1
        else:
            print(f"[{idx}/{len(footer_errors)}] [✗失败] 错误 #{err_id} ({err_type})")
            print(f"    未找到匹配: '{old}'")
            if anchor:
                print(f"    锚点: {anchor}...")
            print(f"    上下文: {trans_context if trans_context else '无'}...")
            footer_fail_count += 1
        print()
    print(f"\n==================== 页脚处理完成 ====================")
    print(f"成功: {footer_success_count} | 失败: {footer_fail_count} | 总计: {len(footer_errors)}")
    if len(footer_errors) > 0:
        print(f"成功率: {footer_success_count / len(footer_errors) * 100:.1f}%")
    print(f"\n✅ 已保存到: {backup_copy_path}")
    print("⚠️ 原始译文文件未被修改")

    # 6) 保存文档
    print("正在保存文档...")
    doc.save(backup_copy_path)

    print(f"\n==================== 文章处理完成 ====================")
    # 汇总计算
    total_count = len(body_errors) + len(header_errors) + len(footer_errors)
    total_success = body_success_count + header_success_count + footer_success_count
    total_fail = body_fail_count + header_fail_count + footer_fail_count
    # 计算被跳过的数量（因为数值缺失导致无法处理的项）
    total_skipped = total_count - total_success - total_fail

    # 有效分母 = 总计 - 跳过
    effective_total = total_success + total_fail

    print(f"\n" + "=" * 50)
    print(f"📊 最终处理报告 (已保存至副本)")
    print(f"总计数据: {total_count} 条")
    print(f"状态分布: [✓成功: {total_success}] [✗失败: {total_fail}] [⏩跳过: {total_skipped}]")

    if effective_total > 0:
        success_rate = (total_success / effective_total) * 100
        print(f"有效成功率: {success_rate:.1f}%  (计算公式: 成功 / (成功+失败))")
    else:
        print("有效成功率: 0% (无有效可处理数据)")

    print(f"保存路径: {backup_copy_path}")
    print("=" * 50)
    print(f"\n✅ 已保存到: {backup_copy_path}")
    print("⚠️ 原始译文文件未被修改")


